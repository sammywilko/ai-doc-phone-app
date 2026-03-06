"""
Documentary Production App - Flask backend with Firestore and Vertex AI
"""
import os
import re
import json
import hashlib
import threading
import base64
import uuid
from datetime import datetime
from urllib.parse import urlparse

from dotenv import load_dotenv
load_dotenv()  # Load .env file for local development

import requests
from flask import Flask, render_template, request, jsonify, Response, send_from_directory
from flask_cors import CORS
from google.cloud import firestore, storage
from weasyprint import HTML
import vertexai
from vertexai.generative_models import GenerativeModel, Part, Tool, grounding
from google.cloud.aiplatform_v1beta1 import Tool as GapicTool

app = Flask(__name__, static_folder='static', static_url_path='/static')
CORS(app, origins=[
    "http://localhost:3000",
    "http://localhost:5173",
    r"https://.*\.run\.app",
])

# Configuration
PROJECT_ID = os.environ.get("GCP_PROJECT_ID", "your-project-id")
LOCATION = os.environ.get("GCP_LOCATION", "us-central1")
MODEL_NAME = os.environ.get("MODEL_NAME", "gemini-2.0-flash-001")

# Per-agent model config — override via env vars, script_writer defaults to Pro
AGENT_MODELS = {
    'research_specialist': os.environ.get('MODEL_RESEARCH', MODEL_NAME),
    'archive_specialist':  os.environ.get('MODEL_ARCHIVE',  MODEL_NAME),
    'interview_producer':  os.environ.get('MODEL_INTERVIEW', MODEL_NAME),
    'script_writer':       os.environ.get('MODEL_SCRIPT',   'gemini-3.1-pro-preview'),
    'fact_checker':        os.environ.get('MODEL_FACTCHECK', MODEL_NAME),
}
STORAGE_BUCKET = os.environ.get("STORAGE_BUCKET", f"{PROJECT_ID}-doc-assets")
MAX_URLS_PER_QUERY = 10
DOWNLOAD_TIMEOUT = 30

# App version and environment
APP_VERSION = os.environ.get("APP_VERSION", "1.0.0")
APP_ENV = os.environ.get("APP_ENV", "prod")

# Initialize Vertex AI
vertexai.init(project=PROJECT_ID, location=LOCATION)
model = GenerativeModel(MODEL_NAME)

# Initialize Firestore
db = firestore.Client(project=PROJECT_ID)

# Initialize Cloud Storage
storage_client = storage.Client(project=PROJECT_ID)

# Collection prefix based on environment (dev uses separate collections)
COLLECTION_PREFIX = "dev_" if APP_ENV == "dev" else ""

# Collection names (prefixed by environment)
COLLECTIONS = {
    'projects': f'{COLLECTION_PREFIX}doc_projects',
    'episodes': f'{COLLECTION_PREFIX}doc_episodes',
    'series': f'{COLLECTION_PREFIX}doc_series',
    'research': f'{COLLECTION_PREFIX}doc_research',
    'interviews': f'{COLLECTION_PREFIX}doc_interviews',
    'shots': f'{COLLECTION_PREFIX}doc_shots',
    'assets': f'{COLLECTION_PREFIX}doc_assets',
    'scripts': f'{COLLECTION_PREFIX}doc_scripts',
    'feedback': f'{COLLECTION_PREFIX}doc_feedback',
    # Production Factory collections
    'research_documents': f'{COLLECTION_PREFIX}doc_research_documents',
    'archive_logs': f'{COLLECTION_PREFIX}doc_archive_logs',
    'interview_transcripts': f'{COLLECTION_PREFIX}doc_interview_transcripts',
    'script_versions': f'{COLLECTION_PREFIX}doc_script_versions',
    'compliance_items': f'{COLLECTION_PREFIX}doc_compliance_items',
    'agent_tasks': f'{COLLECTION_PREFIX}doc_agent_tasks',
    'users': f'{COLLECTION_PREFIX}doc_users',
    'youtube_clips': f'{COLLECTION_PREFIX}doc_youtube_clips',
    'youtube_batches': f'{COLLECTION_PREFIX}doc_youtube_batches',
    'story_cards': f'{COLLECTION_PREFIX}doc_story_cards',
    'episode_arrangements': f'{COLLECTION_PREFIX}doc_episode_arrangements',
    'universal_config': f'{COLLECTION_PREFIX}doc_universal_config',
    'series_config':    f'{COLLECTION_PREFIX}doc_series_config',
    'style_references': f'{COLLECTION_PREFIX}doc_style_references',
    'style_batches':    f'{COLLECTION_PREFIX}doc_style_batches',
    'notifications': f'{COLLECTION_PREFIX}doc_notifications',
    'messages':      f'{COLLECTION_PREFIX}doc_messages',
    'golden_scripts': f'{COLLECTION_PREFIX}doc_golden_scripts',
    'beat_sheets': f'{COLLECTION_PREFIX}doc_beat_sheets',
}

DEFAULT_UNIVERSAL_RULES = """DOCUMENTARY CRAFT RULES — ALL SERIES (Thomas's Layer)

SENTENCE STRUCTURE
- Shorter sentences. Maximum 20 words for voice-over.
- No alliteration. Ever.
- Active voice only.

STORY STRUCTURE
- Every segment needs a problem, complication, and resolution.
- Open with a question the audience wants answered.
- Never summarise at the end — let the story close itself.

SOURCE STANDARDS
- No Wikipedia as a primary citation. Follow to Wikipedia's cited source.
- Every major claim needs a named source.
- Flag any claim that could require legal review.

TONE
- Grounded, intelligent, never sensational.
- Respect the subject matter and the audience.
- Jargon must be explained within one sentence of first use.
"""

STYLE_LAB_PASS1_PROMPT = """You are a professional documentary editor and story structure analyst.

Watch this video carefully and extract a detailed BEAT SHEET — the structural blueprint of the edit.

For each distinct story beat or scene change, record:
- A timecode (MM:SS or HH:MM:SS)
- Which act it belongs to (Act 1 / Act 2 / Act 3 etc.)
- The beat type (cold_open, hook, exposition, rising_action, climax, resolution, cliffhanger, transition, montage, interview_segment, archive_sequence, b_roll)
- A one-line description of what happens in that beat
- Whether the pacing shifts at this point (faster, slower, pause, unchanged)

Also identify:
- Overall scene count
- Average scene length in seconds (estimate)
- All major pacing shifts with timecodes and direction (accelerating / decelerating / pause / cut)

Return ONLY valid JSON in this exact schema:
{
  "acts": [
    {
      "act_number": 1,
      "act_title": "string",
      "start_timecode": "00:00",
      "end_timecode": "02:30",
      "beats": [
        {
          "timecode": "00:00",
          "beat_type": "cold_open",
          "description": "string",
          "pacing": "unchanged"
        }
      ]
    }
  ],
  "scene_count": 0,
  "avg_scene_length_seconds": 0,
  "total_duration_seconds": 0,
  "pacing_shifts": [
    {
      "timecode": "00:00",
      "direction": "accelerating",
      "description": "string"
    }
  ]
}"""

STYLE_LAB_PASS2_PROMPT = """You are an expert documentary style analyst. You have already extracted the beat sheet for this video (provided below as context). Now perform a deep analysis across five style pillars.

BEAT SHEET CONTEXT:
{beat_sheet_json}

Watch the video again with the beat sheet in mind and analyze these five pillars in depth:

1. NARRATIVE PACING — How does the story breathe? Measure the rhythm of information delivery, the balance of fast vs slow sections, how tension builds and releases, and the overall tempo.

2. VISUAL LANGUAGE — What is the visual grammar? Analyze shot composition (wide/medium/close ratios), camera movement patterns, colour grading mood, lighting style, depth of field choices, and any signature framing techniques.

3. EDITING RHYTHMS — How are cuts used as a storytelling tool? Analyze average cut length, transition types (hard cut, dissolve, L-cut, J-cut), montage patterns, intercutting strategies, and how edit pace maps to emotional beats.

4. TONE & ATMOSPHERE — What feeling does the piece evoke? Analyze the music/sound design approach, narration style and register, emotional arc, use of silence, tension/release patterns, and overall mood palette.

5. REVERSE ENGINEERING — If you had to recreate this exact style, what are the concrete production instructions? Provide a shot list template, edit decision list rules, sound design palette, pacing formula, and any signature techniques that define this style.

Return ONLY valid JSON in this exact schema:
{
  "narrative_pacing": {
    "summary": "2-3 sentence overview",
    "tempo": "fast|moderate|slow|variable",
    "information_density": "high|medium|low",
    "tension_pattern": "string describing build/release cycle",
    "act_balance": "string describing time allocation across acts",
    "key_metrics": {
      "avg_beat_duration_seconds": 0,
      "fastest_sequence_timecode": "00:00",
      "slowest_sequence_timecode": "00:00",
      "beats_per_minute": 0
    }
  },
  "visual_language": {
    "summary": "2-3 sentence overview",
    "dominant_shot_type": "wide|medium|close|mixed",
    "camera_movement": "static|handheld|gimbal|dolly|mixed",
    "color_palette": "warm|cool|neutral|desaturated|high_contrast",
    "lighting_style": "natural|studio|dramatic|mixed",
    "signature_techniques": ["string"],
    "key_metrics": {
      "wide_shot_percentage": 0,
      "medium_shot_percentage": 0,
      "close_up_percentage": 0,
      "moving_shot_percentage": 0
    }
  },
  "editing_rhythms": {
    "summary": "2-3 sentence overview",
    "avg_cut_length_seconds": 0,
    "dominant_transition": "hard_cut|dissolve|l_cut|j_cut|mixed",
    "montage_frequency": "frequent|occasional|rare|none",
    "intercutting_style": "string",
    "pace_mapping": "string describing how edit pace follows emotional beats",
    "key_metrics": {
      "total_cuts": 0,
      "fastest_cut_seconds": 0,
      "slowest_cut_seconds": 0,
      "cuts_per_minute": 0
    }
  },
  "tone_atmosphere": {
    "summary": "2-3 sentence overview",
    "overall_mood": "string",
    "music_approach": "score|source|minimal|none|mixed",
    "narration_style": "authoritative|conversational|poetic|absent",
    "emotional_arc": "string",
    "silence_usage": "frequent|occasional|rare|none",
    "key_metrics": {
      "music_percentage": 0,
      "silence_percentage": 0,
      "narration_percentage": 0,
      "interview_percentage": 0
    }
  },
  "reverse_engineering": {
    "summary": "2-3 sentence overview of how to recreate this style",
    "shot_list_template": ["string — each is a shot instruction"],
    "edit_rules": ["string — each is an editing rule"],
    "sound_design_palette": ["string — each is a sound/music instruction"],
    "pacing_formula": "string — the rhythm recipe",
    "signature_techniques": ["string — defining characteristics to replicate"]
  }
}"""

STYLE_LAB_PASS3_PROMPT = """You are an expert documentary style analyst. You have already extracted a beat sheet for this video (provided below). Now perform a deep analysis of the CONTENT LAYER — three pillars that examine how the opening, interviews, and archive footage function as storytelling tools.

BEAT SHEET CONTEXT:
{beat_sheet_json}

Watch the video again with the beat sheet in mind and analyze these three pillars:

1. OPENING CONTRACT — Forensic breakdown of the first 90-120 seconds. What is the deal the show makes with the viewer? Identify the hook type, the implicit promise, the planted question, and every element used in the opening (archive, interview, narration, music stings, title cards). Provide a timecoded breakdown of the opening sequence.

2. INTERVIEW GRAMMAR — How are talking heads deployed as a storytelling tool? Identify the role interviews play (driver, supporter, or absent), average soundbite length, range of soundbite lengths, frequency of interviews, framing style, how interviewees are introduced, the narration-to-interview ratio, and key patterns in how interviews are used. Provide timecoded examples.

3. ARCHIVE INTEGRATION PLAYBOOK — How is archive footage used as storytelling? What percentage is archive vs filmed vs graphics? What is the dominant usage pattern (illustration, evidence, emotional punctuation, montage texture)? Detail each usage category, identify archive sources, treatment style, and opportunities where AI-generated content could fill gaps. Provide timecoded examples.

Return ONLY valid JSON in this exact schema:
{{
  "opening_contract": {{
    "summary": "2-3 sentences",
    "hook_type": "question|shock|mystery|character|spectacle|cold_open|montage",
    "hook_description": "What grabs the viewer",
    "implicit_promise": "The deal — stay and I'll show you X",
    "planted_question": "The central question seeded in the opening",
    "opening_duration_seconds": 0,
    "elements_used": ["archive", "interview", "narration", "music_sting", "title_card"],
    "grab_technique": "e.g. mid-action cold open, provocative claim",
    "timecoded_breakdown": [
      {{ "timecode": "00:00", "element": "string", "purpose": "hook|context|promise|question" }}
    ],
    "effectiveness_notes": "What makes this opening work"
  }},
  "interview_grammar": {{
    "summary": "2-3 sentences",
    "interview_role": "driver|supporter|absent",
    "avg_soundbite_length_seconds": 0,
    "soundbite_range": {{ "shortest_seconds": 0, "longest_seconds": 0 }},
    "interview_frequency": "constant|frequent|moderate|sparse|none",
    "framing_style": "full_frame|cutaway_heavy|mixed|over_shoulder",
    "introduction_method": "lower_third|narrated|self_identified|none",
    "narration_vs_interview_ratio": "60/40 narration-led",
    "interview_to_narration_balance": "interview_dominant|narration_dominant|balanced",
    "interviewee_count": 0,
    "key_patterns": ["soundbites used as act transitions"],
    "timecoded_examples": [
      {{ "timecode": "00:00", "interviewee": "string", "soundbite_length_seconds": 0, "usage": "driving_narrative|supporting_point|emotional_beat|expert_authority|transition" }}
    ]
  }},
  "archive_integration": {{
    "summary": "2-3 sentences",
    "archive_ratio_percentage": 0,
    "filmed_ratio_percentage": 0,
    "graphics_ratio_percentage": 0,
    "dominant_usage_pattern": "illustration|evidence|emotional_punctuation|montage_texture|mixed",
    "usage_categories": {{
      "illustration": {{ "frequency": "frequent|occasional|rare|none", "description": "string" }},
      "evidence": {{ "frequency": "string", "description": "string" }},
      "emotional_punctuation": {{ "frequency": "string", "description": "string" }},
      "montage_texture": {{ "frequency": "string", "description": "string" }}
    }},
    "archive_sources_detected": ["news broadcast", "home video"],
    "archive_treatment": "raw|graded|slowed|mixed|effects_applied",
    "ai_replacement_opportunities": ["gaps where AI-generated content could fill"],
    "timecoded_examples": [
      {{ "timecode": "00:00", "category": "illustration|evidence|emotional_punctuation|montage_texture", "description": "string", "duration_seconds": 0 }}
    ]
  }}
}}"""

STYLE_LAB_PASS4_PROMPT = """You are an expert documentary style analyst. You have already extracted a beat sheet and analysed multiple style pillars for this video (provided below). Now perform a deep analysis of the PRODUCTION LAYER — four pillars that examine the sound design, narrative engine, ad break structure, and a comparative fingerprint.

BEAT SHEET CONTEXT:
{beat_sheet_json}

PRIOR PILLAR ANALYSIS:
{pillars_json}

PASS 3 ANALYSIS:
{pass3_json}

AD BREAK CONFIG:
{ad_break_config}

Watch the video again and analyze these four pillars:

1. SOUND & SCORE ARCHITECTURE — Deep sound design layering. Identify the layers: atmosphere beds, stingers/hits, score themes, silence usage, and foley. What genre, tempo, and mood does the music convey? When does music enter and exit relative to story beats? What is the score-to-silence ratio? Provide timecoded key moments and production notes for recreating this sound design.

2. STORY ENGINE — The narrative machine that keeps viewers watching. What is the primary engine (mystery, character, countdown, revelation, journey)? Is there a secondary engine? How does the engine manifest in the edit? What drives the viewer from scene to scene? Map the question chain — the sequence of questions/tensions that maintain engagement. What techniques sustain tension? Map engines to acts.

3. AD BREAK ENGINEERING — For broadcast shows with ad breaks. If the ad break config indicates no ad breaks, note N/A but still analyze any natural break points. How are cliffhangers constructed before breaks? What is the bump structure (pre-break tease, post-break recap)? How does intensity escalate across successive breaks? Provide timecoded break analysis.

4. COMPARATIVE FINGERPRINT — Express the style as a mashup of known documentary references. This runs last because it synthesises all prior analysis. Create a fingerprint statement (e.g. "The pacing of Sunderland 'Til I Die meets the interview style of The Jinx with the sound design of Planet Earth"). Provide dimension-by-dimension comparisons, a pitch deck line, and an editor shorthand.

Return ONLY valid JSON in this exact schema:
{{
  "sound_score_architecture": {{
    "summary": "2-3 sentences",
    "layers": {{
      "atmos_beds": {{ "presence": "constant|frequent|sparse|none", "description": "string" }},
      "stingers_hits": {{ "frequency": "frequent|occasional|rare|none", "description": "string" }},
      "score_themes": {{ "count": 0, "recurrence_pattern": "string", "description": "string" }},
      "silence": {{ "usage": "frequent|occasional|rare|none", "purpose": "string" }},
      "foley": {{ "presence": "prominent|subtle|none", "description": "string" }}
    }},
    "music_genre": "orchestral|electronic|acoustic|hybrid",
    "music_tempo": "slow|moderate|fast|variable",
    "music_mood": "tension|wonder|melancholy|triumphant",
    "music_entry_exit_pattern": "when music enters/exits relative to story beats",
    "score_to_silence_ratio": "80/20 scored",
    "key_moments": [
      {{ "timecode": "00:00", "sound_element": "score|stinger|silence|atmos|foley", "description": "string" }}
    ],
    "production_notes": "Instructions for recreating this sound design"
  }},
  "story_engine": {{
    "summary": "2-3 sentences",
    "primary_engine": "mystery|character|countdown|revelation|journey",
    "secondary_engine": "mystery|character|countdown|revelation|journey|none",
    "engine_description": "How the engine manifests in the edit",
    "viewer_drive_mechanism": "What drives the viewer from scene to scene",
    "question_chain": ["sequence of questions/tensions that maintain engagement"],
    "tension_sustain_techniques": ["withholding key reveal until act 3", "intercutting timelines"],
    "act_engine_mapping": [
      {{ "act": 1, "engine_mode": "string", "primary_tension": "string" }}
    ],
    "engagement_hooks": ["specific narrative hooks used to retain viewers"]
  }},
  "ad_break_engineering": {{
    "summary": "2-3 sentences (or N/A if no ad breaks)",
    "has_ad_breaks": true,
    "break_count": 0,
    "cliffhanger_construction": {{
      "pattern": "How cliffhangers are built before breaks",
      "techniques": ["unanswered question", "visual tease", "mid-sentence cut"]
    }},
    "bump_structure": {{
      "pre_break_tease": "10-30s before the break",
      "post_break_recap": "How the show re-hooks after the break",
      "recap_duration_seconds": 0
    }},
    "escalation_pattern": "How intensity escalates across successive breaks",
    "re_hook_techniques": ["methods to recapture attention after break"],
    "breaks": [
      {{ "break_number": 1, "timecode": "00:00", "cliffhanger_type": "question|reveal|danger|emotional|mystery", "cliffhanger_description": "string", "re_hook_description": "string" }}
    ]
  }},
  "comparative_fingerprint": {{
    "summary": "2-3 sentences",
    "fingerprint_statement": "The pacing of X meets the interview style of Y with the sound design of Z",
    "reference_comparisons": [
      {{ "dimension": "pacing|interview_style|sound_design|archive_use|visual_language|tone|editing|narrative_engine", "reference_show": "string", "similarity_description": "string" }}
    ],
    "pitch_deck_line": "Single sentence for a pitch deck",
    "editor_shorthand": "Brief reference an editor would understand"
  }}
}}"""

VALID_ROLES = [
    'exec_producer', 'series_producer', 'production_manager', 'line_producer',
    'producer', 'ap', 'researcher', 'editor', 'archive_producer', 'legal',
    'archivist', 'ai_director', 'ai_generator',
]

SEED_USERS = [
    # NASA
    {'username': 'Mark Carter',         'role': 'series_producer',  'avatar': '',  'bio': 'Series Producer — NASA Uncovered',               'customInstructions': ''},
    {'username': 'Rachel Bruce',        'role': 'archive_producer', 'avatar': '',  'bio': 'Archive Producer — NASA Uncovered',              'customInstructions': ''},
    {'username': 'Daniel Ward',         'role': 'researcher',       'avatar': '',  'bio': 'Researcher — NASA Uncovered',                    'customInstructions': ''},
    # Abandoned Places: Uncovered
    {'username': 'Dominic Hill',        'role': 'series_producer',  'avatar': '',  'bio': 'Series Producer — Abandoned Places: Uncovered',  'customInstructions': ''},
    {'username': 'Jake Williams',       'role': 'archive_producer', 'avatar': '',  'bio': 'Archive Producer — Abandoned Places: Uncovered', 'customInstructions': ''},
    {'username': 'Steven Hartley',      'role': 'researcher',       'avatar': '',  'bio': 'Researcher — Abandoned Places: Uncovered',       'customInstructions': ''},
    # Superstructures Uncovered
    {'username': 'Al Blane',            'role': 'series_producer',  'avatar': '',  'bio': 'Series Producer — Superstructures Uncovered',    'customInstructions': ''},
    {'username': 'Simon Yeoman Taylor', 'role': 'archive_producer', 'avatar': '',  'bio': 'Archive Producer — Superstructures Uncovered',   'customInstructions': ''},
    {'username': 'Lily Roberts',        'role': 'researcher',       'avatar': '',  'bio': 'Researcher — Superstructures Uncovered',         'customInstructions': ''},
    # AI Team
    {'username': 'Sam Wilkinson',       'role': 'ai_director',      'avatar': '',  'bio': 'Series AI Director',                 'customInstructions': ''},
    {'username': 'Ivan Gould',          'role': 'ai_director',      'avatar': '',  'bio': 'AI Director',                        'customInstructions': ''},
    {'username': 'Tony',                'role': 'ai_director',      'avatar': '',  'bio': 'AI Director',                        'customInstructions': ''},
    # Executive & Production
    {'username': 'Thomas Viner',        'role': 'exec_producer',    'avatar': '',  'bio': 'Executive Producer',                 'customInstructions': ''},
    {'username': 'Sarah Dellow',        'role': 'line_producer',    'avatar': '',  'bio': 'Line Producer',                      'customInstructions': ''},
]

# ============== Production Factory Constants ==============

# Episode workflow phases
EPISODE_PHASES = {
    'research': {'order': 1, 'name': 'Research', 'description': 'Deep research and fact gathering'},
    'archive': {'order': 2, 'name': 'Archive', 'description': 'Archive collection and processing'},
    'script': {'order': 3, 'name': 'Script', 'description': 'Script generation and refinement'},
    'voiceover': {'order': 4, 'name': 'Voiceover', 'description': 'VO generation and QC'},
    'assembly': {'order': 5, 'name': 'Assembly', 'description': 'Quickture final assembly'},
}

# Phase statuses
PHASE_STATUSES = ['pending', 'in_progress', 'review', 'approved', 'rejected']

# Script version types
SCRIPT_VERSION_TYPES = ['V1_initial', 'V2_producer_review', 'V3_interview_additions', 'V4_locked']

# AI Agent types for script generation
AGENT_TYPES = {
    'research_specialist': {
        'name': 'Research Specialist',
        'role': 'Foundation & Fact Accuracy',
        'responsibilities': ['Verify timeline accuracy', 'Ensure technical explanations', 'Flag claims requiring corroboration', 'Suggest narrative structure']
    },
    'archive_specialist': {
        'name': 'Archive Specialist',
        'role': 'Visual Storytelling',
        'responsibilities': ['Match archive to script moments', 'Identify visual sequences', 'Flag missing footage', 'Suggest pacing']
    },
    'interview_producer': {
        'name': 'Interview Producer',
        'role': 'Human Voices & Emotional Beats',
        'responsibilities': ['Extract best soundbites', 'Identify emotional peaks', 'Match interview content to structure', 'Suggest questions for gaps']
    },
    'script_writer': {
        'name': 'Script Writer',
        'role': 'Narrative Construction',
        'responsibilities': ['Build voiceover narrative', 'Structure story arc', 'Write to broadcast standards', 'Match tone to series bible']
    },
    'fact_checker': {
        'name': 'Fact Checker',
        'role': 'Verification & Compliance',
        'responsibilities': ['Cross-reference every claim', 'Flag legal review items', 'Verify dates and names', 'Generate source citations']
    }
}


# ============== Helper Functions ==============

def doc_to_dict(doc):
    """Convert Firestore document to dict with id."""
    if doc.exists:
        data = doc.to_dict()
        data['id'] = doc.id
        return data
    return None


def get_all_docs(collection_name, project_id=None):
    """Get all documents from a collection, optionally filtered by project."""
    collection = db.collection(COLLECTIONS[collection_name])
    if project_id:
        # Try camelCase first, then snake_case, merge & deduplicate
        camel_docs = {doc.id: doc_to_dict(doc) for doc in collection.where('projectId', '==', project_id).stream()}
        snake_docs = {doc.id: doc_to_dict(doc) for doc in collection.where('project_id', '==', project_id).stream()}
        merged = {**snake_docs, **camel_docs}
        return list(merged.values())
    else:
        docs = collection.stream()
    return [doc_to_dict(doc) for doc in docs]


def get_doc(collection_name, doc_id):
    """Get a single document by ID."""
    doc = db.collection(COLLECTIONS[collection_name]).document(doc_id).get()
    return doc_to_dict(doc)


def create_doc(collection_name, data):
    """Create a new document."""
    data['createdAt'] = datetime.utcnow().isoformat()
    data['updatedAt'] = datetime.utcnow().isoformat()
    doc_ref = db.collection(COLLECTIONS[collection_name]).document()
    doc_ref.set(data)
    data['id'] = doc_ref.id
    return data


def update_doc(collection_name, doc_id, data):
    """Update an existing document."""
    data['updatedAt'] = datetime.utcnow().isoformat()
    db.collection(COLLECTIONS[collection_name]).document(doc_id).update(data)
    return get_doc(collection_name, doc_id)


def delete_doc(collection_name, doc_id):
    """Delete a document."""
    db.collection(COLLECTIONS[collection_name]).document(doc_id).delete()
    return True


# ============== Production Factory Helper Functions ==============

def initialize_episode_workflow(episode_id):
    """Initialize workflow phases for a new episode."""
    workflow = {
        'currentPhase': 'research',
        'phases': {}
    }
    for phase_key, phase_info in EPISODE_PHASES.items():
        workflow['phases'][phase_key] = {
            'status': 'pending' if phase_info['order'] > 1 else 'in_progress',
            'startedAt': None,
            'completedAt': None,
            'reviewNotes': [],
            'assignedTo': None
        }
    # Set first phase as in_progress
    workflow['phases']['research']['startedAt'] = datetime.utcnow().isoformat()
    return workflow


def update_episode_phase(episode_id, phase, status, notes=None):
    """Update an episode's workflow phase status."""
    episode = get_doc('episodes', episode_id)
    if not episode:
        return None

    workflow = episode.get('workflow', initialize_episode_workflow(episode_id))

    if phase not in workflow['phases']:
        return None

    workflow['phases'][phase]['status'] = status

    if status == 'in_progress' and not workflow['phases'][phase]['startedAt']:
        workflow['phases'][phase]['startedAt'] = datetime.utcnow().isoformat()
    elif status == 'approved':
        workflow['phases'][phase]['completedAt'] = datetime.utcnow().isoformat()
        # Move to next phase
        phase_order = EPISODE_PHASES[phase]['order']
        for next_phase, info in EPISODE_PHASES.items():
            if info['order'] == phase_order + 1:
                workflow['currentPhase'] = next_phase
                workflow['phases'][next_phase]['status'] = 'in_progress'
                workflow['phases'][next_phase]['startedAt'] = datetime.utcnow().isoformat()
                break

    if notes:
        workflow['phases'][phase]['reviewNotes'].append({
            'text': notes,
            'timestamp': datetime.utcnow().isoformat()
        })

    return update_doc('episodes', episode_id, {'workflow': workflow})


def get_episode_workflow_status(episode_id):
    """Get detailed workflow status for an episode."""
    episode = get_doc('episodes', episode_id)
    if not episode:
        return None

    workflow = episode.get('workflow', {})
    current_phase = workflow.get('currentPhase', 'research')

    # Calculate overall progress
    completed_phases = sum(
        1 for p in workflow.get('phases', {}).values()
        if p.get('status') == 'approved'
    )
    total_phases = len(EPISODE_PHASES)

    return {
        'episodeId': episode_id,
        'currentPhase': current_phase,
        'currentPhaseName': EPISODE_PHASES.get(current_phase, {}).get('name', 'Unknown'),
        'progress': (completed_phases / total_phases) * 100,
        'completedPhases': completed_phases,
        'totalPhases': total_phases,
        'phases': workflow.get('phases', {}),
        'phaseDefinitions': EPISODE_PHASES
    }


def create_episode_with_buckets(data):
    """Create a new episode with initialized workflow and empty buckets."""
    # Initialize workflow
    data['workflow'] = initialize_episode_workflow(None)

    # Initialize buckets (these are logical containers, actual documents stored separately)
    data['researchBucket'] = {
        'uploadedDocuments': [],
        'agentOutputs': [],
        'factCheckSources': [],
        'notebookLMSource': None
    }
    data['archiveBucket'] = {
        'quicktureLogs': [],
        'nasaMetadata': [],
        'interviewTranscripts': [],
        'referenceFootage': []
    }
    data['scriptWorkspace'] = {
        'referenceTemplate': None,
        'currentVersion': None,
        'versionHistory': [],
        'producerFeedback': []
    }
    data['compliancePackage'] = {
        'exifMetadataLogs': [],
        'sourceCitations': [],
        'archiveLicenses': [],
        'legalSignoff': None
    }

    # Episode brief fields
    data['brief'] = data.get('brief', {
        'summary': '',
        'storyBeats': [],
        'targetInterviewees': [],
        'archiveRequirements': [],
        'uniqueAngle': ''
    })

    return create_doc('episodes', data)


def get_docs_by_episode(collection_name, episode_id):
    """Get all documents from a collection filtered by episode."""
    collection = db.collection(COLLECTIONS[collection_name])
    docs = collection.where('episodeId', '==', episode_id).stream()
    return [doc_to_dict(doc) for doc in docs]


def get_docs_by_series(collection_name, series_id):
    """Get all documents from a collection filtered by series."""
    collection = db.collection(COLLECTIONS[collection_name])
    docs = collection.where('seriesId', '==', series_id).stream()
    return [doc_to_dict(doc) for doc in docs]


def create_agent_task(episode_id, agent_type, task_type, input_data):
    """Create a new AI agent task."""
    task_data = {
        'episodeId': episode_id,
        'agentType': agent_type,
        'agentInfo': AGENT_TYPES.get(agent_type, {}),
        'taskType': task_type,
        'status': 'pending',
        'inputData': input_data,
        'outputData': None,
        'startedAt': None,
        'completedAt': None,
        'error': None
    }
    return create_doc('agent_tasks', task_data)


def update_agent_task(task_id, status, output_data=None, error=None):
    """Update an agent task status."""
    update_data = {'status': status}

    if status == 'in_progress':
        update_data['startedAt'] = datetime.utcnow().isoformat()
    elif status in ['completed', 'failed']:
        update_data['completedAt'] = datetime.utcnow().isoformat()

    if output_data:
        update_data['outputData'] = output_data
    if error:
        update_data['error'] = error

    return update_doc('agent_tasks', task_id, update_data)


def get_project_dashboard_stats(project_id):
    """Get dashboard statistics for a project."""
    # Get all series for project
    series_list = get_all_docs('series', project_id)

    # Get all episodes for project
    all_episodes = get_all_docs('episodes', project_id)

    # Calculate phase statistics
    phase_stats = {phase: {'pending': 0, 'in_progress': 0, 'review': 0, 'approved': 0}
                   for phase in EPISODE_PHASES.keys()}

    for episode in all_episodes:
        workflow = episode.get('workflow', {})
        current_phase = workflow.get('currentPhase', 'research')
        phases = workflow.get('phases', {})

        for phase_key, phase_data in phases.items():
            status = phase_data.get('status', 'pending')
            if status in phase_stats[phase_key]:
                phase_stats[phase_key][status] += 1

    # Episodes by series
    episodes_by_series = {}
    for series in series_list:
        series_episodes = [e for e in all_episodes if e.get('seriesId') == series['id']]
        episodes_by_series[series['id']] = {
            'seriesName': series.get('title', 'Unknown'),
            'totalEpisodes': len(series_episodes),
            'completedEpisodes': sum(1 for e in series_episodes
                                     if e.get('workflow', {}).get('currentPhase') == 'assembly'
                                     and e.get('workflow', {}).get('phases', {}).get('assembly', {}).get('status') == 'approved')
        }

    return {
        'projectId': project_id,
        'totalSeries': len(series_list),
        'totalEpisodes': len(all_episodes),
        'phaseStats': phase_stats,
        'episodesBySeries': episodes_by_series,
        'bottlenecks': [
            {'phase': phase, 'count': stats['review']}
            for phase, stats in phase_stats.items()
            if stats['review'] > 0
        ]
    }


# ============== AI Functions ==============

def clean_ai_response(text):
    """Strip markdown code-block wrappers (```json ... ```) from AI output."""
    cleaned = text.strip()
    # Remove ```json or ``` prefix
    if cleaned.startswith("```json"):
        cleaned = cleaned[len("```json"):]
    elif cleaned.startswith("```"):
        cleaned = cleaned[3:]
    # Remove trailing ```
    if cleaned.endswith("```"):
        cleaned = cleaned[:-3]
    return cleaned.strip()


def generate_ai_response(prompt, system_prompt="", model_name=None, generation_config=None):
    """Generate AI response using Vertex AI. Pass model_name to use a specific model for this call."""
    try:
        full_prompt = f"{system_prompt}\n\n{prompt}" if system_prompt else prompt
        active_model = GenerativeModel(model_name) if model_name else model
        config = generation_config or {}
        response = active_model.generate_content(full_prompt, generation_config=config)
        return response.text
    except Exception as e:
        return f"AI error: {str(e)}"


def generate_grounded_research(prompt, system_prompt=""):
    """Placeholder - AI research functionality disabled in this build."""
    return {
        'text': 'AI Research functionality is not available in this build.',
        'sources': []
    }


# ============== Source Document Functions ==============

def extract_urls(text):
    """Extract URLs from text, limited to MAX_URLS_PER_QUERY."""
    url_pattern = r'https?://[^\s<>\[\]()"\']+'
    urls = re.findall(url_pattern, text)
    cleaned_urls = []
    for url in urls:
        url = url.rstrip('.,;:!?)')
        if url and len(url) > 10:
            cleaned_urls.append(url)
    unique_urls = list(dict.fromkeys(cleaned_urls))
    return unique_urls[:MAX_URLS_PER_QUERY]


def validate_url(url, timeout=3):
    """Check if a URL is accessible (returns True if reachable)."""
    try:
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
        }
        response = requests.head(url, headers=headers, timeout=timeout, allow_redirects=True)
        return response.status_code < 400
    except:
        # Try GET if HEAD fails (some servers don't support HEAD)
        try:
            response = requests.get(url, headers=headers, timeout=timeout, stream=True)
            response.close()
            return response.status_code < 400
        except:
            return False


def filter_valid_urls(urls, max_to_check=5):
    """Filter URLs to only include valid, accessible ones (limited to prevent timeout)."""
    valid_urls = []
    checked = 0
    for url in urls:
        if checked >= max_to_check:
            break
        if len(valid_urls) >= 3:  # Stop once we have enough valid URLs
            break
        if validate_url(url):
            valid_urls.append(url)
            print(f"✓ Valid URL: {url[:60]}...")
        else:
            print(f"✗ Invalid URL: {url[:60]}...")
        checked += 1
    return valid_urls


def convert_to_pdf(html_content, url):
    """Convert HTML content to PDF."""
    try:
        parsed = urlparse(url)
        base_url = f"{parsed.scheme}://{parsed.netloc}"
        if '<base' not in html_content.lower():
            html_content = html_content.replace(
                '<head>',
                f'<head><base href="{base_url}">',
                1
            )
        html = HTML(string=html_content, base_url=base_url)
        pdf_bytes = html.write_pdf()
        return pdf_bytes
    except Exception as e:
        print(f"PDF conversion error for {url}: {e}")
        return None


def download_and_store(url, bucket_name, project_id, research_id):
    """Download a URL, convert to PDF, and store in GCS bucket."""
    try:
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
        }
        response = requests.get(url, headers=headers, timeout=DOWNLOAD_TIMEOUT, allow_redirects=True)
        response.raise_for_status()

        content_type = response.headers.get('Content-Type', '').split(';')[0].strip()

        parsed_url = urlparse(url)
        path = parsed_url.path.strip('/')
        if path:
            base_filename = path.split('/')[-1]
            base_filename = re.sub(r'[^\w\-.]', '_', base_filename)
        else:
            base_filename = parsed_url.netloc.replace('.', '_')

        if '.' in base_filename:
            base_filename = base_filename.rsplit('.', 1)[0]

        title = base_filename.replace('_', ' ').title()
        if 'html' in content_type:
            title_match = re.search(r'<title[^>]*>([^<]+)</title>', response.text, re.IGNORECASE)
            if title_match:
                title = title_match.group(1).strip()

        url_hash = hashlib.md5(url.encode()).hexdigest()[:8]
        bucket = storage_client.bucket(bucket_name)

        result = {"url": url, "title": title, "status": "success"}

        if content_type == 'application/pdf':
            blob_path = f"{project_id}/{url_hash}_{base_filename}.pdf"
            blob = bucket.blob(blob_path)
            blob.upload_from_string(response.content, content_type='application/pdf')
            result["gcsPath"] = blob_path
            result["size_bytes"] = len(response.content)
            result["filename"] = f"{base_filename}.pdf"

        elif 'html' in content_type or 'text' in content_type:
            pdf_bytes = convert_to_pdf(response.text, url)
            if pdf_bytes:
                blob_path = f"{project_id}/{url_hash}_{base_filename}.pdf"
                blob = bucket.blob(blob_path)
                blob.upload_from_string(pdf_bytes, content_type='application/pdf')
                result["gcsPath"] = blob_path
                result["size_bytes"] = len(pdf_bytes)
                result["filename"] = f"{base_filename}.pdf"
            else:
                blob_path = f"{project_id}/{url_hash}_{base_filename}.html"
                blob = bucket.blob(blob_path)
                blob.upload_from_string(response.content, content_type='text/html')
                result["gcsPath"] = blob_path
                result["size_bytes"] = len(response.content)
                result["filename"] = f"{base_filename}.html"
        else:
            ext = content_type.split('/')[-1] if '/' in content_type else 'bin'
            blob_path = f"{project_id}/{url_hash}_{base_filename}.{ext}"
            blob = bucket.blob(blob_path)
            blob.upload_from_string(response.content, content_type=content_type)
            result["gcsPath"] = blob_path
            result["size_bytes"] = len(response.content)
            result["filename"] = f"{base_filename}.{ext}"

        if result.get("gcsPath"):
            create_source_document_asset(project_id, research_id, result)

        return result

    except Exception as e:
        return {"url": url, "status": "error", "error": str(e)}


def create_source_document_asset(project_id, research_id, doc_result):
    """Create a Firestore asset entry for a downloaded source document."""
    try:
        asset_data = {
            "projectId": project_id,
            "researchId": research_id,
            "title": doc_result.get("title", "Untitled Document"),
            "type": "Document",
            "source": doc_result.get("url", ""),
            "gcsPath": doc_result.get("gcsPath", ""),
            "status": "Acquired",
            "isSourceDocument": True,
            "sizeBytes": doc_result.get("size_bytes", 0),
            "filename": doc_result.get("filename", ""),
            "createdAt": datetime.utcnow().isoformat(),
            "updatedAt": datetime.utcnow().isoformat()
        }
        doc_ref = db.collection(COLLECTIONS['assets']).document()
        doc_ref.set(asset_data)
        print(f"Created source document asset: {doc_result.get('title')}")
    except Exception as e:
        print(f"Error creating asset: {e}")


def process_source_documents_async(urls, bucket_name, project_id, research_id):
    """Background thread to download and process source documents."""
    print(f"Starting download of {len(urls)} sources for research {research_id}")
    success_count = 0
    error_count = 0
    for url in urls:
        try:
            print(f"Downloading: {url[:80]}...")
            result = download_and_store(url, bucket_name, project_id, research_id)
            if result.get("status") == "error":
                print(f"Failed to download {url}: {result.get('error')}")
                error_count += 1
            else:
                print(f"Successfully downloaded: {result.get('title', url)}")
                success_count += 1
        except Exception as e:
            print(f"Error processing {url}: {e}")
            error_count += 1
    print(f"Download complete: {success_count} success, {error_count} errors")


def ensure_bucket_exists(bucket_name):
    """Create bucket if it doesn't exist."""
    try:
        bucket = storage_client.bucket(bucket_name)
        if not bucket.exists():
            bucket = storage_client.create_bucket(bucket_name, location=LOCATION)
        return True
    except Exception as e:
        print(f"Bucket error: {e}")
        return False


# ============== Routes ==============

@app.route("/")
def index():
    """Serve the main app interface from static folder."""
    return send_from_directory('static', 'index.html')


@app.route("/health")
def health():
    """Health check endpoint."""
    return jsonify({"status": "healthy"})


# ============== Project Routes ==============

@app.route("/api/projects", methods=["GET"])
def get_projects():
    """Get all projects."""
    projects = get_all_docs('projects')
    return jsonify(projects)


@app.route("/api/projects", methods=["POST"])
def create_project():
    """Create a new project."""
    data = request.get_json()
    project = create_doc('projects', data)
    return jsonify(project), 201


@app.route("/api/projects/<project_id>", methods=["GET"])
def get_project(project_id):
    """Get a single project."""
    project = get_doc('projects', project_id)
    if project:
        return jsonify(project)
    return jsonify({"error": "Project not found"}), 404


@app.route("/api/projects/<project_id>", methods=["PUT"])
def update_project(project_id):
    """Update a project."""
    data = request.get_json()
    project = update_doc('projects', project_id, data)
    return jsonify(project)


@app.route("/api/projects/<project_id>", methods=["DELETE"])
def delete_project(project_id):
    """Delete a project and all related data."""
    # Delete all related data first
    for collection in ['episodes', 'series', 'research', 'interviews', 'shots', 'assets', 'scripts']:
        docs = db.collection(COLLECTIONS[collection]).where('projectId', '==', project_id).stream()
        for doc in docs:
            doc.reference.delete()

    # Delete the project itself
    delete_doc('projects', project_id)
    return jsonify({"success": True})


# ============== Episode Routes ==============

@app.route("/api/projects/<project_id>/episodes", methods=["GET"])
def get_episodes(project_id):
    """Get all episodes for a project."""
    episodes = get_all_docs('episodes', project_id)
    episodes.sort(key=lambda e: e.get('order', e.get('episodeNumber', 0)))
    return jsonify(episodes)


@app.route("/api/episodes", methods=["POST"])
def create_episode():
    """Create a new episode."""
    data = request.get_json()
    episode = create_doc('episodes', data)
    return jsonify(episode), 201


@app.route("/api/episodes/<episode_id>", methods=["PUT"])
def update_episode(episode_id):
    """Update an episode."""
    data = request.get_json()
    episode = update_doc('episodes', episode_id, data)
    return jsonify(episode)


@app.route("/api/episodes/<episode_id>", methods=["DELETE"])
def delete_episode(episode_id):
    """Delete an episode."""
    delete_doc('episodes', episode_id)
    return jsonify({"success": True})


# ============== Series Routes ==============

@app.route("/api/projects/<project_id>/series", methods=["GET"])
def get_series(project_id):
    """Get all series for a project."""
    series = get_all_docs('series', project_id)
    # Sort by order field
    series.sort(key=lambda s: s.get('order', 0))
    return jsonify(series)


@app.route("/api/series", methods=["POST"])
def create_series():
    """Create a new series."""
    data = request.get_json()
    series = create_doc('series', data)
    return jsonify(series), 201


@app.route("/api/series/<series_id>", methods=["PUT"])
def update_series(series_id):
    """Update a series."""
    data = request.get_json()
    series = update_doc('series', series_id, data)
    return jsonify(series)


@app.route("/api/series/<series_id>", methods=["DELETE"])
def delete_series(series_id):
    """Delete a series and ungroup its episodes."""
    # Remove seriesId from all episodes in this series
    episodes = db.collection(COLLECTIONS['episodes']).where('seriesId', '==', series_id).stream()
    for ep in episodes:
        ep.reference.update({'seriesId': None, 'updatedAt': datetime.utcnow().isoformat()})

    delete_doc('series', series_id)
    return jsonify({"success": True})


# ============== Research Routes ==============

@app.route("/api/projects/<project_id>/research", methods=["GET"])
def get_research(project_id):
    """Get all research for a project."""
    research = get_all_docs('research', project_id)
    return jsonify(research)


@app.route("/api/research", methods=["POST"])
def create_research():
    """Create a new research item, or run AI research if topic+engine are present."""
    data = request.get_json()

    # Detect AI research request from frontend's summarizeResearch()
    # Frontend sends {topic, engine, systemInstruction} and expects
    # {summary, key_facts, expert_suggestions, urls} back
    if data.get('topic') and data.get('engine'):
        topic = data['topic']
        system_instruction = data.get('systemInstruction', '')

        prompt = f"""Research the following topic for a documentary production.

Topic: {topic}

Provide comprehensive research and return ONLY a valid JSON object (no markdown fences) with:
- "summary": string (2-3 detailed paragraphs of background research)
- "key_facts": array of strings (5-8 key facts with source references)
- "expert_suggestions": array of strings (3-5 expert names with their role/title to interview)
- "urls": array of strings (relevant source URLs, use real credible sources)"""

        system_prompt = system_instruction or "You are a documentary research specialist. Always respond with valid JSON only, no markdown code fences."

        raw_result = generate_ai_response(prompt, system_prompt)
        cleaned = clean_ai_response(raw_result)

        try:
            result = json.loads(cleaned)
        except (json.JSONDecodeError, ValueError):
            # If AI didn't return valid JSON, wrap the text response
            result = {
                "summary": cleaned,
                "key_facts": [],
                "expert_suggestions": [],
                "urls": []
            }

        return jsonify(result)

    # Standard CRUD: create a research document in Firestore
    research = create_doc('research', data)
    return jsonify(research), 201


@app.route("/api/research/<research_id>", methods=["PUT"])
def update_research(research_id):
    """Update a research item."""
    data = request.get_json()
    research = update_doc('research', research_id, data)
    return jsonify(research)


@app.route("/api/research/<research_id>", methods=["DELETE"])
def delete_research(research_id):
    """Delete a research item."""
    delete_doc('research', research_id)
    return jsonify({"success": True})


# ============== Interview Routes ==============

@app.route("/api/projects/<project_id>/interviews", methods=["GET"])
def get_interviews(project_id):
    """Get all interviews for a project."""
    interviews = get_all_docs('interviews', project_id)
    return jsonify(interviews)


@app.route("/api/interviews", methods=["POST"])
def create_interview():
    """Create a new interview."""
    data = request.get_json()
    interview = create_doc('interviews', data)
    return jsonify(interview), 201


@app.route("/api/interviews/<interview_id>", methods=["PUT"])
def update_interview(interview_id):
    """Update an interview."""
    data = request.get_json()
    interview = update_doc('interviews', interview_id, data)
    return jsonify(interview)


@app.route("/api/interviews/<interview_id>", methods=["DELETE"])
def delete_interview(interview_id):
    """Delete an interview."""
    delete_doc('interviews', interview_id)
    return jsonify({"success": True})


# ============== Shot Routes ==============

@app.route("/api/projects/<project_id>/shots", methods=["GET"])
def get_shots(project_id):
    """Get all shots for a project."""
    shots = get_all_docs('shots', project_id)
    return jsonify(shots)


@app.route("/api/shots", methods=["POST"])
def create_shot():
    """Create a new shot."""
    data = request.get_json()
    if 'projectId' in data:
        data['project_id'] = data['projectId']
    elif 'project_id' in data:
        data['projectId'] = data['project_id']
    if 'episodeId' in data:
        data['episode_id'] = data['episodeId']
    elif 'episode_id' in data:
        data['episodeId'] = data['episode_id']
    shot = create_doc('shots', data)
    return jsonify(shot), 201


@app.route("/api/shots/<shot_id>", methods=["PUT"])
def update_shot(shot_id):
    """Update a shot."""
    data = request.get_json()
    if 'projectId' in data:
        data['project_id'] = data['projectId']
    elif 'project_id' in data:
        data['projectId'] = data['project_id']
    if 'episodeId' in data:
        data['episode_id'] = data['episodeId']
    elif 'episode_id' in data:
        data['episodeId'] = data['episode_id']
    shot = update_doc('shots', shot_id, data)
    return jsonify(shot)


@app.route("/api/shots/<shot_id>", methods=["DELETE"])
def delete_shot(shot_id):
    """Delete a shot."""
    delete_doc('shots', shot_id)
    return jsonify({"success": True})


# ============== Asset Routes ==============

@app.route("/api/projects/<project_id>/assets", methods=["GET"])
def get_assets(project_id):
    """Get all assets for a project."""
    assets = get_all_docs('assets', project_id)
    return jsonify(assets)


@app.route("/api/assets", methods=["POST"])
def create_asset():
    """Create a new asset."""
    data = request.get_json()
    asset = create_doc('assets', data)
    return jsonify(asset), 201


@app.route("/api/assets/<asset_id>", methods=["PUT"])
def update_asset(asset_id):
    """Update an asset."""
    data = request.get_json()
    asset = update_doc('assets', asset_id, data)
    return jsonify(asset)


@app.route("/api/assets/<asset_id>", methods=["DELETE"])
def delete_asset(asset_id):
    """Delete an asset and its GCS file if it exists."""
    # Get asset first to check for GCS file
    asset = get_doc('assets', asset_id)
    if asset and asset.get('gcsPath'):
        try:
            bucket = storage_client.bucket(STORAGE_BUCKET)
            blob = bucket.blob(asset['gcsPath'])
            if blob.exists():
                blob.delete()
        except Exception as e:
            print(f"Error deleting GCS file: {e}")

    delete_doc('assets', asset_id)
    return jsonify({"success": True})


@app.route("/api/assets/upload", methods=["POST"])
def upload_asset_file():
    """Upload a file for an asset and create/update the asset record."""
    if 'file' not in request.files:
        return jsonify({"error": "No file provided"}), 400

    file = request.files['file']
    if file.filename == '':
        return jsonify({"error": "No file selected"}), 400

    # Get form data
    project_id = request.form.get('projectId')
    asset_id = request.form.get('assetId')  # Optional - for updating existing asset
    episode_id = request.form.get('episodeId')  # Optional - for episode research documents
    series_id = request.form.get('seriesId')  # Optional - for series research documents
    is_research_document = request.form.get('isResearchDocument', 'false').lower() == 'true'
    title = request.form.get('title', file.filename)
    asset_type = request.form.get('type', 'Document')
    status = request.form.get('status', 'Acquired')
    notes = request.form.get('notes', '')

    if not project_id:
        return jsonify({"error": "Project ID is required"}), 400

    try:
        # Read file content
        file_content = file.read()
        file_size = len(file_content)

        # Determine content type
        content_type = file.content_type or 'application/octet-stream'

        # Generate safe filename
        original_filename = file.filename
        safe_filename = re.sub(r'[^\w\-.]', '_', original_filename)

        # Generate unique path
        file_hash = hashlib.md5(file_content).hexdigest()[:8]
        blob_path = f"assets/{project_id}/{file_hash}_{safe_filename}"

        # Upload to GCS
        ensure_bucket_exists(STORAGE_BUCKET)
        bucket = storage_client.bucket(STORAGE_BUCKET)
        blob = bucket.blob(blob_path)
        blob.upload_from_string(file_content, content_type=content_type)

        print(f"Uploaded asset file: {blob_path} ({file_size} bytes)")

        # Create or update asset document
        asset_data = {
            "projectId": project_id,
            "title": title,
            "type": asset_type,
            "status": status,
            "notes": notes,
            "gcsPath": blob_path,
            "filename": original_filename,
            "mimeType": content_type,
            "sizeBytes": file_size,
            "hasFile": True,
            "isResearchDocument": is_research_document,
            "updatedAt": datetime.utcnow().isoformat()
        }

        # Add optional entity associations for research documents
        if episode_id:
            asset_data["episodeId"] = episode_id
        if series_id:
            asset_data["seriesId"] = series_id

        if asset_id:
            # Update existing asset
            # First delete old file if exists
            existing = get_doc('assets', asset_id)
            if existing and existing.get('gcsPath') and existing['gcsPath'] != blob_path:
                try:
                    old_blob = bucket.blob(existing['gcsPath'])
                    if old_blob.exists():
                        old_blob.delete()
                except:
                    pass

            asset = update_doc('assets', asset_id, asset_data)
        else:
            # Create new asset
            asset_data['createdAt'] = datetime.utcnow().isoformat()
            asset = create_doc('assets', asset_data)

        return jsonify({
            "success": True,
            "asset": asset,
            "gcsPath": blob_path,
            "filename": original_filename,
            "size": file_size
        }), 201

    except Exception as e:
        print(f"Error uploading asset file: {e}")
        return jsonify({"error": str(e)}), 500


@app.route("/api/assets/<asset_id>/file", methods=["GET"])
def get_asset_file(asset_id):
    """Download an asset's file using streaming for all sizes."""
    from flask import stream_with_context
    import sys

    print(f"Download request for asset: {asset_id}", file=sys.stderr)

    asset = get_doc('assets', asset_id)
    if not asset:
        print(f"Asset not found: {asset_id}", file=sys.stderr)
        return jsonify({"error": "Asset not found"}), 404

    if not asset.get('gcsPath'):
        print(f"Asset has no gcsPath: {asset_id}", file=sys.stderr)
        return jsonify({"error": "Asset has no file"}), 404

    gcs_path = asset.get('gcsPath')
    print(f"Attempting to access GCS path: {gcs_path}", file=sys.stderr)

    try:
        bucket = storage_client.bucket(STORAGE_BUCKET)
        blob = bucket.blob(gcs_path)

        print(f"Checking if blob exists...", file=sys.stderr)
        if not blob.exists():
            print(f"Blob does not exist: {gcs_path}", file=sys.stderr)
            return jsonify({"error": "File not found in storage"}), 404

        content_type = asset.get('mimeType', 'application/octet-stream')
        filename = asset.get('filename', 'download')

        # Get file size
        print(f"Reloading blob metadata...", file=sys.stderr)
        blob.reload()
        file_size = blob.size

        print(f"Streaming asset file: {gcs_path} ({file_size} bytes)", file=sys.stderr)

        # Stream using GCS range requests - this works within Cloud Run's response limits
        # Each chunk is fetched separately, avoiding memory issues
        chunk_size = 5 * 1024 * 1024  # 5MB chunks for faster streaming

        def generate():
            """Generator that yields file chunks using range requests."""
            offset = 0
            chunk_num = 0
            while offset < file_size:
                end = min(offset + chunk_size, file_size)
                try:
                    # GCS range requests are inclusive on both ends
                    chunk = blob.download_as_bytes(start=offset, end=end - 1)
                    chunk_num += 1
                    if chunk_num % 10 == 0:  # Log every 10 chunks
                        print(f"Chunk {chunk_num}: {offset}-{end} of {file_size} ({100*end/file_size:.1f}%)")
                    yield chunk
                    offset = end
                except Exception as e:
                    print(f"Error downloading chunk at offset {offset}: {e}")
                    raise

            print(f"Stream complete: {chunk_num} chunks, {file_size} bytes total")

        # Use chunked transfer encoding for streaming (no Content-Length)
        # This allows Cloud Run to stream without buffering the entire response
        response = Response(
            stream_with_context(generate()),
            mimetype=content_type,
        )
        response.headers['Content-Disposition'] = f'attachment; filename="{filename}"'
        response.headers['X-Content-Length'] = str(file_size)  # Hint for client progress
        return response

    except Exception as e:
        print(f"Error downloading asset file: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500


# Chunked upload for large asset files (>32MB)
CHUNK_SIZE = 10 * 1024 * 1024  # 10MB chunks


@app.route("/api/assets/upload/init", methods=["POST"])
def init_asset_chunked_upload():
    """Initialize a chunked upload session for large asset files."""
    data = request.get_json()
    filename = data.get('filename', 'upload')
    content_type = data.get('contentType', 'application/octet-stream')
    file_size = data.get('fileSize', 0)
    total_chunks = data.get('totalChunks', 1)
    project_id = data.get('projectId')
    title = data.get('title', filename)
    asset_type = data.get('type', 'Document')
    status = data.get('status', 'Acquired')
    notes = data.get('notes', '')

    if not project_id:
        return jsonify({"error": "Project ID is required"}), 400

    # Generate unique upload ID
    upload_id = hashlib.md5(f"{filename}{project_id}{datetime.utcnow().isoformat()}".encode()).hexdigest()[:16]

    # Generate safe filename and blob path
    safe_filename = re.sub(r'[^\w\-.]', '_', filename)
    blob_path = f"assets/{project_id}/{upload_id}_{safe_filename}"

    print(f"Initialized chunked upload: {upload_id} for {filename} ({file_size} bytes, {total_chunks} chunks)")

    return jsonify({
        "uploadId": upload_id,
        "blobPath": blob_path,
        "totalChunks": total_chunks,
        "projectId": project_id,
        "filename": filename,
        "contentType": content_type,
        "title": title,
        "type": asset_type,
        "status": status,
        "notes": notes
    })


@app.route("/api/assets/upload/chunk/<upload_id>", methods=["POST"])
def upload_asset_chunk(upload_id):
    """Upload a chunk of a large asset file."""
    chunk_index = int(request.form.get('chunkIndex', 0))
    total_chunks = int(request.form.get('totalChunks', 1))
    blob_path = request.form.get('blobPath', '')
    content_type = request.form.get('contentType', 'application/octet-stream')

    if 'chunk' not in request.files:
        return jsonify({"error": "No chunk data"}), 400

    chunk = request.files['chunk']
    chunk_data = chunk.read()

    try:
        ensure_bucket_exists(STORAGE_BUCKET)
        bucket = storage_client.bucket(STORAGE_BUCKET)

        # Store chunk temporarily
        chunk_blob_name = f"uploads/chunks/{upload_id}/chunk_{chunk_index:04d}"
        chunk_blob = bucket.blob(chunk_blob_name)
        chunk_blob.upload_from_string(chunk_data, content_type='application/octet-stream')

        print(f"Uploaded chunk {chunk_index + 1}/{total_chunks} for {upload_id} ({len(chunk_data)} bytes)")

        return jsonify({
            "status": "uploaded",
            "chunkIndex": chunk_index,
            "totalChunks": total_chunks,
            "bytesUploaded": len(chunk_data)
        })

    except Exception as e:
        print(f"Chunk upload error: {e}")
        return jsonify({"error": f"Chunk upload failed: {str(e)}"}), 500


@app.route("/api/assets/upload/complete/<upload_id>", methods=["POST"])
def complete_asset_chunked_upload(upload_id):
    """Complete a chunked upload by combining chunks and creating the asset."""
    data = request.get_json()
    blob_path = data.get('blobPath', '')
    content_type = data.get('contentType', 'application/octet-stream')
    project_id = data.get('projectId')
    filename = data.get('filename', 'upload')
    title = data.get('title', filename)
    asset_type = data.get('type', 'Document')
    status = data.get('status', 'Acquired')
    notes = data.get('notes', '')
    asset_id = data.get('assetId')  # Optional - for updating existing asset

    if not project_id:
        return jsonify({"error": "Project ID is required"}), 400

    try:
        bucket = storage_client.bucket(STORAGE_BUCKET)

        # List all chunks
        chunk_blobs = list(bucket.list_blobs(prefix=f"uploads/chunks/{upload_id}/"))
        chunk_blobs.sort(key=lambda b: b.name)

        if not chunk_blobs:
            return jsonify({"error": "No chunks found for this upload"}), 400

        print(f"Combining {len(chunk_blobs)} chunks for {upload_id}")

        # Combine all chunk data
        combined_data = b''
        for cb in chunk_blobs:
            combined_data += cb.download_as_bytes()

        file_size = len(combined_data)

        # Upload combined file to final location
        final_blob = bucket.blob(blob_path)
        final_blob.upload_from_string(combined_data, content_type=content_type)

        print(f"Combined file uploaded: {blob_path} ({file_size} bytes)")

        # Clean up chunks
        for cb in chunk_blobs:
            cb.delete()

        # Create or update asset document
        asset_data = {
            "projectId": project_id,
            "title": title,
            "type": asset_type,
            "status": status,
            "notes": notes,
            "gcsPath": blob_path,
            "filename": filename,
            "mimeType": content_type,
            "sizeBytes": file_size,
            "hasFile": True,
            "updatedAt": datetime.utcnow().isoformat()
        }

        if asset_id:
            # Update existing asset
            existing = get_doc('assets', asset_id)
            if existing and existing.get('gcsPath') and existing['gcsPath'] != blob_path:
                try:
                    old_blob = bucket.blob(existing['gcsPath'])
                    if old_blob.exists():
                        old_blob.delete()
                except:
                    pass
            asset = update_doc('assets', asset_id, asset_data)
        else:
            asset_data['createdAt'] = datetime.utcnow().isoformat()
            asset = create_doc('assets', asset_data)

        return jsonify({
            "success": True,
            "asset": asset,
            "gcsPath": blob_path,
            "filename": filename,
            "size": file_size
        }), 201

    except Exception as e:
        print(f"Error completing chunked upload: {e}")
        # Try to clean up chunks on error
        try:
            bucket = storage_client.bucket(STORAGE_BUCKET)
            for cb in bucket.list_blobs(prefix=f"uploads/chunks/{upload_id}/"):
                cb.delete()
        except:
            pass
        return jsonify({"error": str(e)}), 500


# ============== Research Documents Query Routes ==============

@app.route("/api/episodes/<episode_id>/research-documents", methods=["GET"])
def get_episode_research_documents(episode_id):
    """Get all research documents for an episode."""
    try:
        docs_ref = db.collection(COLLECTIONS['assets']).where(
            'episodeId', '==', episode_id
        ).where(
            'isResearchDocument', '==', True
        )
        documents = []
        for doc in docs_ref.stream():
            doc_data = doc.to_dict()
            doc_data['id'] = doc.id
            documents.append(doc_data)
        return jsonify(documents)
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/series/<series_id>/research-documents", methods=["GET"])
def get_series_research_documents(series_id):
    """Get all research documents for a series."""
    try:
        docs_ref = db.collection(COLLECTIONS['assets']).where(
            'seriesId', '==', series_id
        ).where(
            'isResearchDocument', '==', True
        )
        documents = []
        for doc in docs_ref.stream():
            doc_data = doc.to_dict()
            doc_data['id'] = doc.id
            documents.append(doc_data)
        return jsonify(documents)
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/projects/<project_id>/research-documents", methods=["GET"])
def get_project_research_documents(project_id):
    """Get all research documents for a project (project-level only, not episode/series)."""
    try:
        # Get research documents that are project-level (no episodeId or seriesId)
        docs_ref = db.collection(COLLECTIONS['assets']).where(
            'projectId', '==', project_id
        ).where(
            'isResearchDocument', '==', True
        )
        documents = []
        for doc in docs_ref.stream():
            doc_data = doc.to_dict()
            # Filter to only project-level (no episode or series association)
            if not doc_data.get('episodeId') and not doc_data.get('seriesId'):
                doc_data['id'] = doc.id
                documents.append(doc_data)
        return jsonify(documents)
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/projects/<project_id>/all-research-documents", methods=["GET"])
def get_all_project_research_documents(project_id):
    """Get all research documents for a project (including episode and series docs)."""
    try:
        docs_ref = db.collection(COLLECTIONS['assets']).where(
            'projectId', '==', project_id
        ).where(
            'isResearchDocument', '==', True
        )
        documents = []
        for doc in docs_ref.stream():
            doc_data = doc.to_dict()
            doc_data['id'] = doc.id
            documents.append(doc_data)
        return jsonify(documents)
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/projects/<project_id>/assets/clear-sources", methods=["DELETE"])
def clear_source_documents(project_id):
    """Delete all source documents for a project."""
    try:
        # Get all source documents
        docs_ref = db.collection(COLLECTIONS['assets']).where(
            'projectId', '==', project_id
        ).where(
            'isSourceDocument', '==', True
        )

        deleted_count = 0
        bucket = storage_client.bucket(STORAGE_BUCKET)

        for doc in docs_ref.stream():
            doc_data = doc.to_dict()
            # Delete GCS file if exists
            if doc_data.get('gcsPath'):
                try:
                    blob = bucket.blob(doc_data['gcsPath'])
                    if blob.exists():
                        blob.delete()
                except Exception as e:
                    print(f"Error deleting GCS file {doc_data['gcsPath']}: {e}")

            # Delete Firestore document
            doc.reference.delete()
            deleted_count += 1

        return jsonify({"success": True, "deleted": deleted_count})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/projects/<project_id>/assets/download-sources", methods=["POST"])
def download_additional_sources(project_id):
    """Download source documents from a list of URLs."""
    data = request.get_json()
    urls = data.get('urls', [])
    research_id = data.get('researchId', datetime.utcnow().strftime("%Y%m%d_%H%M%S"))

    if not urls:
        return jsonify({"error": "No URLs provided"}), 400

    ensure_bucket_exists(STORAGE_BUCKET)
    results = []
    for url in urls[:5]:  # Max 5 per request
        try:
            result = download_and_store(url, STORAGE_BUCKET, project_id, research_id)
            results.append({
                "url": url,
                "status": "completed" if result.get("status") == "success" else "error",
                "title": result.get("title", ""),
                "filename": result.get("filename", ""),
                "error": result.get("error")
            })
        except Exception as e:
            results.append({"url": url, "status": "error", "error": str(e)})

    return jsonify({"results": results})


@app.route("/api/projects/<project_id>/assets/download-all", methods=["GET"])
def download_all_source_documents(project_id):
    """Download all source documents as a ZIP file."""
    import zipfile
    import io

    try:
        # Get all source documents
        docs_ref = db.collection(COLLECTIONS['assets']).where(
            'projectId', '==', project_id
        ).where(
            'isSourceDocument', '==', True
        )

        # Create ZIP in memory
        zip_buffer = io.BytesIO()
        bucket = storage_client.bucket(STORAGE_BUCKET)

        with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
            for doc in docs_ref.stream():
                doc_data = doc.to_dict()
                if doc_data.get('gcsPath'):
                    try:
                        blob = bucket.blob(doc_data['gcsPath'])
                        if blob.exists():
                            content = blob.download_as_bytes()
                            filename = doc_data.get('filename') or doc_data['gcsPath'].split('/')[-1]
                            zip_file.writestr(filename, content)
                    except Exception as e:
                        print(f"Error adding {doc_data['gcsPath']} to ZIP: {e}")

        zip_buffer.seek(0)

        return Response(
            zip_buffer.getvalue(),
            mimetype='application/zip',
            headers={'Content-Disposition': f'attachment; filename="source-documents-{project_id[:8]}.zip"'}
        )
    except Exception as e:
        return jsonify({"error": str(e)}), 500


# ============== Script Routes ==============

@app.route("/api/projects/<project_id>/scripts", methods=["GET"])
def get_scripts(project_id):
    """Get all scripts for a project."""
    scripts = get_all_docs('scripts', project_id)
    return jsonify(scripts)


@app.route("/api/scripts", methods=["POST"])
def create_script():
    """Create a new script."""
    data = request.get_json()
    if 'projectId' in data:
        data['project_id'] = data['projectId']
    elif 'project_id' in data:
        data['projectId'] = data['project_id']
    if 'episodeId' in data:
        data['episode_id'] = data['episodeId']
    elif 'episode_id' in data:
        data['episodeId'] = data['episode_id']
    script = create_doc('scripts', data)
    return jsonify(script), 201


@app.route("/api/scripts/<script_id>", methods=["PUT"])
def update_script(script_id):
    """Update a script."""
    data = request.get_json()
    if 'projectId' in data:
        data['project_id'] = data['projectId']
    elif 'project_id' in data:
        data['projectId'] = data['project_id']
    if 'episodeId' in data:
        data['episode_id'] = data['episodeId']
    elif 'episode_id' in data:
        data['episodeId'] = data['episode_id']
    script = update_doc('scripts', script_id, data)
    return jsonify(script)


@app.route("/api/scripts/<script_id>", methods=["DELETE"])
def delete_script(script_id):
    """Delete a script."""
    delete_doc('scripts', script_id)
    return jsonify({"success": True})


@app.route("/api/projects/<project_id>/seed-scripts", methods=["POST"])
def seed_scripts(project_id):
    """Seed parsed script segments into a project. Idempotent — skips segments that already have scripts.

    Each segment becomes a script doc with parts > scenes > beats structure.
    Segments are independent and can be combined into episodes later.
    """
    try:
        data = request.get_json() or {}
        segments = data.get('segments', [])

        if not segments:
            return jsonify({"error": "No segments provided"}), 400

        # Get existing scripts for this project
        existing_scripts = get_all_docs('scripts', project_id)
        existing_segment_nums = {s.get('segment_number') for s in existing_scripts if s.get('segment_number')}

        created = 0
        skipped = 0

        for seg in segments:
            seg_num = seg.get('segment_number')
            if seg_num in existing_segment_nums:
                skipped += 1
                continue

            script_doc = {
                'projectId': project_id,
                'project_id': project_id,
                'segment_number': seg_num,
                'title': seg.get('title', f'Segment {seg_num}'),
                'series': seg.get('series', 1),
                'version': 1,
                'status': 'draft',
                'source_file': seg.get('source_file', ''),
                'parts': seg.get('parts', []),
                'raw_text': seg.get('raw_text', ''),
                'stats': seg.get('stats', {}),
            }
            create_doc('scripts', script_doc)
            created += 1

        return jsonify({
            "success": True,
            "created": created,
            "skipped": skipped,
            "total_segments": len(segments),
        })

    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/scripts/project/<project_id>/current", methods=["GET"])
def get_current_script(project_id):
    """Get the most recent script for a project."""
    scripts = get_all_docs('scripts', project_id)
    if not scripts:
        return jsonify({"error": "No scripts found for this project"}), 404
    # Return the most recently updated script
    scripts.sort(key=lambda s: s.get('updatedAt', s.get('createdAt', '')), reverse=True)
    return jsonify(scripts[0])


# ============== Feedback Routes ==============

@app.route("/api/feedback", methods=["POST"])
def submit_feedback():
    """Submit user feedback with optional screenshot."""
    data = request.get_json()

    feedback_text = data.get('text', '')
    feedback_type = data.get('type', 'general')
    screenshot_data = data.get('screenshot')
    project_id = data.get('projectId')
    project_title = data.get('projectTitle')
    current_tab = data.get('currentTab')
    user_agent = data.get('userAgent', '')
    screen_size = data.get('screenSize', '')
    timestamp = data.get('timestamp', datetime.utcnow().isoformat())

    if not feedback_text:
        return jsonify({"error": "Feedback text is required"}), 400

    # Save screenshot to GCS if provided
    screenshot_path = None
    screenshot_base64 = None
    if screenshot_data and screenshot_data.startswith('data:image'):
        try:
            # Extract base64 data
            header, base64_data = screenshot_data.split(',', 1)
            screenshot_base64 = base64_data
            image_data = base64.b64decode(base64_data)

            # Generate unique filename
            feedback_id = str(uuid.uuid4())[:8]
            screenshot_filename = f"feedback/{feedback_id}_screenshot.jpg"

            # Upload to GCS
            ensure_bucket_exists(STORAGE_BUCKET)
            bucket = storage_client.bucket(STORAGE_BUCKET)
            blob = bucket.blob(screenshot_filename)
            blob.upload_from_string(image_data, content_type='image/jpeg')

            screenshot_path = screenshot_filename
            print(f"Feedback screenshot saved: {screenshot_filename}")

        except Exception as e:
            print(f"Error saving feedback screenshot: {e}")
            # Continue without screenshot

    # Create feedback document
    feedback_doc = {
        'text': feedback_text,
        'type': feedback_type,
        'name': data.get('name', 'Anonymous'),
        'version': data.get('version', ''),
        'url': data.get('url', ''),
        'screenshotPath': screenshot_path,
        'projectId': project_id,
        'projectTitle': project_title,
        'currentTab': current_tab,
        'userAgent': user_agent,
        'screenSize': screen_size,
        'status': 'new',
        'response': '',
        'createdAt': timestamp,
        'updatedAt': timestamp
    }

    # Save to Firestore
    doc_ref = db.collection(COLLECTIONS['feedback']).document()
    doc_ref.set(feedback_doc)
    feedback_doc['id'] = doc_ref.id

    print(f"Feedback submitted: [{feedback_type}] {feedback_text[:50]}...")

    # --- Send email notification via SendGrid (best-effort) ---
    sendgrid_api_key = os.environ.get('SENDGRID_API_KEY', '')
    sender = os.environ.get('EMAIL_SENDER', 'aim-studio@arrowmedia.com')

    if sendgrid_api_key:
        try:
            recipients = [
                {"email": "sam.wilkinson@arrowmedia.com"},
                {"email": "arrowmedia@netlinux.co.uk"}
            ]
            subject = f"[AiM Feedback] {feedback_type}: {feedback_text[:60]}{'...' if len(feedback_text) > 60 else ''}"

            screenshot_note = '<p style="color:#888;font-size:12px;"><em>Screenshot attached</em></p>' if screenshot_base64 else ''

            html_body = f"""
            <div style="font-family:Arial,sans-serif;max-width:700px;margin:0 auto;">
                <div style="background:#1a1a1a;padding:20px;border-radius:8px 8px 0 0;">
                    <h1 style="color:#e53e3e;margin:0;font-size:20px;">AiM Feedback Report</h1>
                    <p style="color:#a0a0a0;margin:4px 0 0;font-size:13px;">Type: {feedback_type.upper()}</p>
                </div>
                <div style="padding:20px;background:#ffffff;border:1px solid #ddd;border-top:none;border-radius:0 0 8px 8px;">
                    <p><strong>From:</strong> {feedback_doc.get('name', 'Anonymous')}</p>
                    <p><strong>Description:</strong></p>
                    <div style="background:#f7f7f7;padding:12px;border-radius:6px;margin:8px 0;">
                        {feedback_text}
                    </div>
                    {f'<p><strong>Project:</strong> {project_title}</p>' if project_title else ''}
                    {f'<p><strong>Phase:</strong> {current_tab}</p>' if current_tab else ''}
                    {screenshot_note}
                    <hr style="border:none;border-top:1px solid #eee;margin:16px 0;">
                    <p style="color:#888;font-size:11px;">
                        {timestamp} &middot; {screen_size} &middot; Feedback ID: {doc_ref.id}
                    </p>
                    <p style="color:#aaa;font-size:10px;">{user_agent[:120]}</p>
                </div>
            </div>"""

            sg_payload = {
                "personalizations": [{"to": recipients, "subject": subject}],
                "from": {"email": sender, "name": "AiM Documentary Studio"},
                "content": [{"type": "text/html", "value": html_body}]
            }

            if screenshot_base64:
                sg_payload["attachments"] = [{
                    "content": screenshot_base64,
                    "type": "image/jpeg",
                    "filename": "screenshot.jpg",
                    "disposition": "attachment"
                }]

            sg_response = requests.post(
                "https://api.sendgrid.com/v3/mail/send",
                json=sg_payload,
                headers={
                    "Authorization": f"Bearer {sendgrid_api_key}",
                    "Content-Type": "application/json"
                },
                timeout=10
            )

            if sg_response.status_code in (200, 201, 202):
                print(f"Feedback email sent to {len(recipients)} recipients")
            else:
                print(f"[WARN] SendGrid returned {sg_response.status_code}: {sg_response.text}")

        except Exception as e:
            print(f"[WARN] Feedback email failed (feedback still saved): {e}")
    else:
        print("[INFO] SendGrid not configured — skipping feedback email")

    return jsonify({
        "success": True,
        "id": doc_ref.id,
        "feedbackId": doc_ref.id,
        "message": "Feedback received"
    }), 201


@app.route("/api/feedback", methods=["GET"])
def get_all_feedback():
    """Get all feedback (admin view)."""
    docs = db.collection(COLLECTIONS['feedback']).order_by(
        'createdAt', direction=firestore.Query.DESCENDING
    ).limit(100).stream()

    feedback_list = []
    for doc in docs:
        fb = doc.to_dict()
        fb['id'] = doc.id
        feedback_list.append(fb)

    return jsonify(feedback_list)


@app.route("/api/feedback/<feedback_id>", methods=["PUT"])
def update_feedback_status(feedback_id):
    """Update a feedback entry (status and admin response)."""
    try:
        data = request.get_json()
        update_data = {
            "updatedAt": datetime.utcnow().isoformat()
        }
        if 'status' in data:
            update_data['status'] = data['status']
        if 'response' in data:
            update_data['response'] = data['response']

        doc_ref = db.collection(COLLECTIONS['feedback']).document(feedback_id)
        doc_ref.update(update_data)
        return jsonify({"success": True})
    except Exception as e:
        print(f"[ERROR] Failed to update feedback: {e}")
        return jsonify({"error": str(e)}), 500


@app.route("/api/projects/<project_id>/blueprint-content", methods=["GET"])
def get_blueprint_content(project_id):
    """Get the blueprint markdown content for a project."""
    project = get_doc('projects', project_id)
    if not project:
        return jsonify({"error": "Project not found"}), 404

    # Check if project has blueprint content stored
    blueprint_content = project.get('blueprintContent', '')

    if not blueprint_content and project.get('blueprintFile'):
        # Try to extract content from stored document
        try:
            file_path = project['blueprintFile'].get('path', '')
            if file_path:
                bucket = storage_client.bucket(STORAGE_BUCKET)
                blob = bucket.blob(file_path)
                if blob.exists():
                    content = blob.download_as_text()
                    # If it's HTML (from PDF), try to extract text
                    if content.startswith('<!DOCTYPE') or content.startswith('<html'):
                        blueprint_content = "Blueprint document available for download"
                    else:
                        blueprint_content = content
        except Exception as e:
            print(f"Error loading blueprint content: {e}")

    return jsonify({"content": blueprint_content})


@app.route("/api/ai/generate-script", methods=["POST"])
def ai_generate_script():
    """Generate a Quickture-compatible documentary script for an episode."""
    data = request.get_json()
    episode_id = data.get('episodeId', '')
    episode_title = data.get('episodeTitle', '')
    episode_description = data.get('episodeDescription', '')
    project_title = data.get('projectTitle', '')
    project_description = data.get('projectDescription', '')
    project_style = data.get('projectStyle', '')
    project_id = data.get('projectId', '')
    duration = data.get('duration', '45 minutes')

    # Get research for this episode if available
    research_content = ""
    if episode_id:
        research_docs = db.collection(COLLECTIONS['research']).where(
            'episodeId', '==', episode_id
        ).limit(1).stream()
        for doc in research_docs:
            research_content = doc.to_dict().get('content', '')[:5000]
            break

    system_prompt = f"""You are a professional documentary script writer creating scripts optimized for AI-assisted editing tools like Quickture.

## PROJECT CONTEXT
- Project: {project_title}
- Style: {project_style or 'Documentary'}
- Description: {project_description}

## QUICKTURE-COMPATIBLE SCRIPT FORMAT

Create a detailed documentary script that includes:

1. **HEADER SECTION**
   - Episode title and number
   - Target duration
   - Style/tone notes for editors

2. **SCENE BREAKDOWN** (use this format for each scene):
   ```
   SCENE [NUMBER]: [SCENE TITLE]
   Duration: [estimated time]
   Location: [setting]
   Mood: [emotional tone]

   VISUAL:
   [Description of what we see - B-roll, interviews, graphics]

   AUDIO:
   [Narration, interview soundbites, ambient sound, music cues]

   NARRATION:
   "[Exact narration text if any]"

   INTERVIEW BITES:
   - [Subject name]: "[Key quote or topic to cover]"

   B-ROLL NEEDED:
   - [List of specific shots needed]

   GRAPHICS/TEXT:
   - [Any lower thirds, titles, or info graphics]

   TRANSITION:
   [How this scene connects to the next]
   ```

3. **SHOT LIST SUMMARY**
   - Numbered list of all shots with descriptions
   - Technical notes (wide/close, handheld/tripod, etc.)

4. **INTERVIEW GUIDE**
   - Questions for each subject
   - Key points to cover

5. **MUSIC/SOUND DESIGN NOTES**
   - Mood suggestions per scene
   - Transition audio cues

## REQUIREMENTS
- Be specific and actionable for editors
- Include timing estimates for pacing
- Mark emotional beats and story arc moments
- Note any archival footage or graphics needed
- Keep narration concise and documentary-style
"""

    research_section = f"\n\nRESEARCH AVAILABLE:\n{research_content}" if research_content else ""

    prompt = f"""Create a detailed, Quickture-compatible documentary script for:

Episode: {episode_title}
Description: {episode_description}
Target Duration: {duration}
{research_section}

Generate a comprehensive production script with scene breakdowns, shot lists, narration, and interview guides."""

    result = generate_ai_response(prompt, system_prompt)

    # Save the script
    script_data = {
        'projectId': project_id,
        'episodeId': episode_id,
        'title': f"Script: {episode_title}",
        'content': result,
        'format': 'quickture',
        'duration': duration,
        'status': 'Draft'
    }

    if project_id:
        saved_script = create_doc('scripts', script_data)
        return jsonify({
            "script": result,
            "saved": True,
            "scriptId": saved_script['id']
        })

    return jsonify({"script": result, "saved": False})


# ============== AI Routes ==============

def get_research_document_contents(episode_id=None, series_id=None, project_id=None):
    """Fetch and read contents of research documents for context."""
    documents_context = []

    try:
        # Get episode research documents
        if episode_id:
            docs = db.collection(COLLECTIONS['assets']).where('episodeId', '==', episode_id).where('isResearchDocument', '==', True).stream()
            for doc in docs:
                data = doc.to_dict()
                content = read_document_content(data.get('gcsPath'), data.get('mimeType', ''))
                if content:
                    documents_context.append({
                        'source': f"Episode Document: {data.get('title', data.get('filename', 'Unknown'))}",
                        'content': content
                    })

        # Get series research documents
        if series_id:
            docs = db.collection(COLLECTIONS['assets']).where('seriesId', '==', series_id).where('isResearchDocument', '==', True).stream()
            for doc in docs:
                data = doc.to_dict()
                content = read_document_content(data.get('gcsPath'), data.get('mimeType', ''))
                if content:
                    documents_context.append({
                        'source': f"Series Document: {data.get('title', data.get('filename', 'Unknown'))}",
                        'content': content
                    })

        # Get project-level research documents
        if project_id:
            docs = db.collection(COLLECTIONS['assets']).where('projectId', '==', project_id).where('isResearchDocument', '==', True).stream()
            for doc in docs:
                data = doc.to_dict()
                # Only include project-level docs (not linked to episode/series)
                if not data.get('episodeId') and not data.get('seriesId'):
                    content = read_document_content(data.get('gcsPath'), data.get('mimeType', ''))
                    if content:
                        documents_context.append({
                            'source': f"Project Document: {data.get('title', data.get('filename', 'Unknown'))}",
                            'content': content
                        })
    except Exception as e:
        print(f"[ERROR] Error fetching research documents: {e}")

    return documents_context


def read_document_content(gcs_path, mime_type=''):
    """Read content from a document in GCS."""
    if not gcs_path:
        return None

    try:
        bucket = storage_client.bucket(STORAGE_BUCKET)
        blob = bucket.blob(gcs_path)
        content = blob.download_as_bytes()

        # For text-based files, decode to string
        if mime_type.startswith('text/') or gcs_path.endswith(('.txt', '.md', '.csv')):
            return content.decode('utf-8', errors='ignore')[:10000]  # Limit to 10k chars

        # Extract text from PDF documents
        if gcs_path.endswith('.pdf'):
            try:
                import io
                from PyPDF2 import PdfReader
                reader = PdfReader(io.BytesIO(content))
                pages_text = []
                for page in reader.pages:
                    text = page.extract_text()
                    if text:
                        pages_text.append(text)
                extracted = "\n\n".join(pages_text)
                return extracted[:10000] if extracted else None  # Limit to 10k chars
            except Exception as pdf_err:
                print(f"[ERROR] PDF extraction failed for {gcs_path}: {pdf_err}")
                return None

        return None
    except Exception as e:
        print(f"[ERROR] Error reading document {gcs_path}: {e}")
        return None


@app.route("/api/ai/simple-research", methods=["POST"])
def ai_simple_research():
    """AI research query augmented with uploaded research documents from episode and series."""
    data = request.get_json()
    title = data.get('title', '')
    description = data.get('description', '')
    user_query = data.get('query', '')  # User's custom research prompt
    episode_id = data.get('episodeId', '')
    series_id = data.get('seriesId', '')
    project_id = data.get('projectId', '')
    save_research = data.get('save', True)

    print(f"[DEBUG] simple-research called: title={title}, query={user_query[:50] if user_query else 'None'}..., episodeId={episode_id}, seriesId={series_id}")

    # Fetch research documents for context
    research_docs = get_research_document_contents(
        episode_id=episode_id,
        series_id=series_id,
        project_id=project_id
    )

    print(f"[DEBUG] Found {len(research_docs)} research documents for context")

    # Build context from research documents
    context_section = ""
    if research_docs:
        context_section = "\n\n## Reference Documents\n\nThe following research documents have been uploaded and should be used as context:\n\n"
        for doc in research_docs:
            context_section += f"### {doc['source']}\n{doc['content']}\n\n"

    # Use user's query if provided, otherwise fall back to title/description
    research_query = user_query if user_query else f"Research background information for the documentary episode titled '{title}': {description}"

    prompt = f"""You are researching for a documentary episode.

Episode: {title}
{f'Description: {description}' if description else ''}
{context_section}

## Research Request

{research_query}

## Instructions

Based on {'the reference documents above and ' if research_docs else ''}the research request, provide:
- Key facts and background information
- Relevant sources and references (with URLs where possible)
- Interview suggestions (people to talk to)
- Visual/archive material recommendations

{f'IMPORTANT: Incorporate and build upon the information from the {len(research_docs)} provided reference document(s). Reference specific details from them where relevant.' if research_docs else ''}

Format your response with:
- Clear sections with headers
- Bullet points for key facts
- Include real, clickable URLs to credible sources (news sites, Wikipedia, .gov, .edu, .org sites)
- Mark each source with its URL in markdown link format: [Source Name](URL)"""

    system_prompt = """You are a documentary research assistant. Provide comprehensive background research with real source links. Always format URLs as markdown links that can be clicked. Focus on factual, verifiable information from credible sources. When reference documents are provided, incorporate their information and expand upon it."""

    raw_result = generate_ai_response(prompt, system_prompt)
    result = clean_ai_response(raw_result)

    print(f"[DEBUG] AI response length: {len(result)} chars")

    response_data = {
        "result": result,
        "title": title,
        "query": research_query,
        "saved": False,
        "documentsUsed": len(research_docs)
    }

    # Save research to episode if episodeId provided
    if save_research and episode_id and project_id:
        try:
            print(f"[DEBUG] Saving research to episode {episode_id}")
            # Update the episode with the research content
            episode_ref = db.collection(COLLECTIONS['episodes']).document(episode_id)
            episode_ref.update({
                'research': result,
                'researchGeneratedAt': datetime.utcnow().isoformat(),
                'updatedAt': datetime.utcnow().isoformat()
            })
            response_data['saved'] = True
            response_data['episodeId'] = episode_id
            print(f"[DEBUG] Research saved successfully to episode {episode_id}")
        except Exception as e:
            print(f"[ERROR] Failed to save research: {e}")
            response_data['saveError'] = str(e)

    return jsonify(response_data)


@app.route("/api/episodes/<episode_id>/research", methods=["GET"])
def get_episode_research(episode_id):
    """Get saved research for an episode."""
    print(f"[DEBUG] Getting research for episode {episode_id}")
    try:
        episode = get_doc('episodes', episode_id)
        if not episode:
            print(f"[DEBUG] Episode {episode_id} not found")
            return jsonify({"error": "Episode not found"}), 404

        research = episode.get('research', '')
        generated_at = episode.get('researchGeneratedAt', '')

        print(f"[DEBUG] Found research: {len(research)} chars, generated at: {generated_at}")

        return jsonify({
            "research": research,
            "generatedAt": generated_at,
            "episodeId": episode_id,
            "episodeTitle": episode.get('title', '')
        })
    except Exception as e:
        print(f"[ERROR] Failed to get research: {e}")
        return jsonify({"error": str(e)}), 500


@app.route("/api/episodes/<episode_id>/research", methods=["DELETE"])
def delete_episode_research(episode_id):
    """Delete saved research for an episode."""
    print(f"[DEBUG] Deleting research for episode {episode_id}")
    try:
        episode_ref = db.collection(COLLECTIONS['episodes']).document(episode_id)
        episode_ref.update({
            'research': '',
            'researchGeneratedAt': '',
            'updatedAt': datetime.utcnow().isoformat()
        })
        print(f"[DEBUG] Research deleted for episode {episode_id}")
        return jsonify({"success": True})
    except Exception as e:
        print(f"[ERROR] Failed to delete research: {e}")
        return jsonify({"error": str(e)}), 500


@app.route("/api/episodes/<episode_id>/research", methods=["PUT"])
def save_episode_research(episode_id):
    """Save research to an episode and extract links as reference assets."""
    print(f"[DEBUG] Saving research for episode {episode_id}")
    try:
        data = request.get_json()
        research = data.get('research', '')

        if not research:
            return jsonify({"error": "No research content provided"}), 400

        # Get episode to find projectId
        episode_ref = db.collection(COLLECTIONS['episodes']).document(episode_id)
        episode_doc = episode_ref.get()
        if not episode_doc.exists:
            return jsonify({"error": "Episode not found"}), 404

        episode_data = episode_doc.to_dict()
        project_id = episode_data.get('projectId')
        episode_title = episode_data.get('title', 'Unknown Episode')

        # Save research to episode
        episode_ref.update({
            'research': research,
            'researchGeneratedAt': datetime.utcnow().isoformat(),
            'updatedAt': datetime.utcnow().isoformat()
        })
        print(f"[DEBUG] Research saved for episode {episode_id}, length: {len(research)}")

        # Extract markdown links and create assets as reference links
        links_created = 0
        markdown_links = []

        if project_id:
            # Find all markdown links: [text](url)
            markdown_links = re.findall(r'\[([^\]]+)\]\((https?://[^)]+)\)', research)
            print(f"[DEBUG] Found {len(markdown_links)} links in research")

            for link_text, link_url in markdown_links:
                try:
                    # Check if asset with this URL already exists for this project
                    existing = db.collection(COLLECTIONS['assets']).where(
                        'projectId', '==', project_id
                    ).where(
                        'source', '==', link_url
                    ).limit(1).get()

                    if len(list(existing)) > 0:
                        print(f"[DEBUG] Asset already exists for URL: {link_url[:50]}...")
                        continue

                    # Create asset for this link as a reference (link only, no download)
                    asset_data = {
                        "projectId": project_id,
                        "episodeId": episode_id,
                        "title": link_text[:100],
                        "type": "Reference",
                        "source": link_url,
                        "status": "Identified",
                        "isSourceDocument": False,
                        "isResearchLink": True,
                        "sourceEpisode": episode_title,
                        "notes": f"Extracted from research for: {episode_title}",
                        "createdAt": datetime.utcnow().isoformat(),
                        "updatedAt": datetime.utcnow().isoformat()
                    }
                    doc_ref = db.collection(COLLECTIONS['assets']).document()
                    doc_ref.set(asset_data)
                    links_created += 1
                    print(f"[DEBUG] Created asset: {link_text[:50]}...")
                except Exception as link_error:
                    print(f"[ERROR] Failed to create asset for {link_url}: {link_error}")

        print(f"[DEBUG] Created {links_created} new reference assets")
        return jsonify({
            "success": True,
            "episodeId": episode_id,
            "linksExtracted": len(markdown_links),
            "assetsCreated": links_created
        })
    except Exception as e:
        print(f"[ERROR] Failed to save research: {e}")
        return jsonify({"error": str(e)}), 500


@app.route("/api/ai/interview-questions", methods=["POST"])
def ai_interview_questions():
    """Generate interview questions."""
    data = request.get_json()
    subject = data.get('subject', '')
    role = data.get('role', '')
    context = data.get('context', '')
    project_title = data.get('projectTitle', '')

    system_prompt = """You are a documentary interview specialist. Generate thoughtful, open-ended questions that elicit detailed stories and insights. Focus on emotional moments, specific details, and unique perspectives."""

    prompt = f"""Generate 8-10 interview questions for {subject}, who is/was a {role}, for a documentary about "{project_title}".
Focus on: {context}.
Make questions specific, open-ended, and designed to get compelling stories."""

    result = generate_ai_response(prompt, system_prompt,
                                   generation_config={"max_output_tokens": 8192, "temperature": 0.7})
    return jsonify({"result": result})


@app.route("/api/ai/script-outline", methods=["POST"])
def ai_script_outline():
    """Generate script outline."""
    data = request.get_json()
    title = data.get('title', '')
    topic = data.get('topic', '')
    duration = data.get('duration', '45 minutes')
    project_title = data.get('projectTitle', '')

    system_prompt = """You are a documentary scriptwriter. Create detailed episode outlines with clear acts, narrative arcs, and visual storytelling elements."""

    prompt = f"""Create a detailed outline for a documentary episode titled "{title}" ({duration}) about: {topic}.
This is for the documentary "{project_title}".
Include acts, key beats, narrative arc, and suggested visuals."""

    result = generate_ai_response(prompt, system_prompt,
                                   generation_config={"max_output_tokens": 16384, "temperature": 0.7})
    return jsonify({"result": result})


@app.route("/api/ai/shot-ideas", methods=["POST"])
def ai_shot_ideas():
    """Generate shot ideas."""
    data = request.get_json()
    scene = data.get('scene', '')
    project_title = data.get('projectTitle', '')

    system_prompt = """You are a documentary cinematographer. Suggest creative, visually compelling shot ideas with specific camera movements, angles, and equipment."""

    prompt = f"""Suggest 5-7 creative shots for: {scene}.
This is for "{project_title}".
Include camera angles, movements, equipment needed, and why each shot would be compelling."""

    result = generate_ai_response(prompt, system_prompt,
                                   generation_config={"max_output_tokens": 8192, "temperature": 0.7})
    return jsonify({"result": result})


@app.route("/api/ai/expand-topic", methods=["POST"])
def ai_expand_topic():
    """Expand on a topic."""
    data = request.get_json()
    topic = data.get('topic', '')
    project_title = data.get('projectTitle', '')

    system_prompt = """You are a documentary story consultant. Help explore topics by suggesting angles, themes, narrative approaches, and key elements to investigate."""

    prompt = f"""Explore potential angles and approaches for covering "{topic}" in the documentary "{project_title}".
Suggest themes, storylines, key questions to answer, and unique perspectives."""

    result = generate_ai_response(prompt, system_prompt,
                                   generation_config={"max_output_tokens": 8192, "temperature": 0.7})
    return jsonify({"result": result})


@app.route("/api/ai/generate-topics", methods=["POST"])
def ai_generate_topics():
    """Generate episode topics from project title and description."""
    data = request.get_json()
    title = data.get('title', '')
    description = data.get('description', '')
    style = data.get('style', '')
    num_topics = data.get('numTopics', 5)

    system_prompt = """You are a documentary series planner. Generate compelling episode topics that would make a cohesive documentary series.

IMPORTANT: Respond ONLY with a JSON array of episode objects. No markdown, no explanation, just valid JSON.

Each episode object must have:
- "title": A compelling episode title (max 60 chars)
- "description": Brief description of what this episode covers (max 150 chars)
- "order": Episode number (1, 2, 3, etc.)

Example response format:
[
  {"title": "Episode Title Here", "description": "What this episode covers", "order": 1},
  {"title": "Another Episode", "description": "Description of content", "order": 2}
]"""

    style_instruction = f"\nStyle/Approach: {style}\nEnsure all episodes match this documentary style." if style else ""

    prompt = f"""Create {num_topics} episode topics for a documentary series:

Title: {title}
Description: {description}{style_instruction}

Generate episode topics that:
1. Cover the subject comprehensively
2. Have a logical narrative flow
3. Each could stand alone but contribute to the whole
4. Are engaging and specific, not generic

Return ONLY the JSON array, no other text."""

    result = generate_ai_response(prompt, system_prompt,
                                   generation_config={"max_output_tokens": 8192, "temperature": 0.7})

    # Parse the JSON response
    try:
        # Clean up the response - remove markdown code blocks if present
        cleaned = result.strip()
        if cleaned.startswith('```'):
            # Remove markdown code fence
            lines = cleaned.split('\n')
            cleaned = '\n'.join(lines[1:-1] if lines[-1].startswith('```') else lines[1:])

        import json
        topics = json.loads(cleaned)
        return jsonify({"topics": topics})
    except Exception as e:
        # If parsing fails, return the raw result for debugging
        return jsonify({"topics": [], "raw": result, "error": str(e)})


@app.route("/api/upload/init", methods=["POST"])
def init_chunked_upload():
    """Initialize a chunked upload session for large files."""
    data = request.get_json()
    filename = data.get('filename', 'upload')
    content_type = data.get('contentType', 'application/octet-stream')
    file_size = data.get('fileSize', 0)
    total_chunks = data.get('totalChunks', 1)

    # Generate unique upload ID and blob name
    upload_id = hashlib.md5(f"{filename}{datetime.utcnow().isoformat()}".encode()).hexdigest()[:16]
    ext = filename.rsplit('.', 1)[-1].lower() if '.' in filename else ''
    blob_name = f"uploads/{upload_id}.{ext}"

    return jsonify({
        "uploadId": upload_id,
        "gcsUri": f"gs://{STORAGE_BUCKET}/{blob_name}",
        "blobPath": blob_name,
        "totalChunks": total_chunks
    })


@app.route("/api/upload/chunk/<upload_id>", methods=["POST"])
def upload_chunk(upload_id):
    """Upload a chunk of a large file."""
    chunk_index = int(request.form.get('chunkIndex', 0))
    total_chunks = int(request.form.get('totalChunks', 1))
    blob_path = request.form.get('blobPath', '')
    content_type = request.form.get('contentType', 'application/octet-stream')

    if 'chunk' not in request.files:
        return jsonify({"error": "No chunk data"}), 400

    chunk = request.files['chunk']
    chunk_data = chunk.read()

    try:
        ensure_bucket_exists(STORAGE_BUCKET)
        bucket = storage_client.bucket(STORAGE_BUCKET)

        # Store chunk temporarily
        chunk_blob_name = f"uploads/chunks/{upload_id}/chunk_{chunk_index:04d}"
        chunk_blob = bucket.blob(chunk_blob_name)
        chunk_blob.upload_from_string(chunk_data, content_type='application/octet-stream')

        # If this is the last chunk, combine all chunks
        if chunk_index == total_chunks - 1:
            # List all chunks
            chunk_blobs = list(bucket.list_blobs(prefix=f"uploads/chunks/{upload_id}/"))
            chunk_blobs.sort(key=lambda b: b.name)

            # Combine chunks into final file
            final_blob = bucket.blob(blob_path)

            # Use compose for efficiency (works for up to 32 components)
            if len(chunk_blobs) <= 32:
                final_blob.compose(chunk_blobs)
            else:
                # For more than 32 chunks, we need to compose in stages
                # First, combine all chunk data
                combined_data = b''
                for cb in chunk_blobs:
                    combined_data += cb.download_as_bytes()
                final_blob.upload_from_string(combined_data, content_type=content_type)

            # Clean up chunks
            for cb in chunk_blobs:
                cb.delete()

            return jsonify({
                "status": "complete",
                "gcsUri": f"gs://{STORAGE_BUCKET}/{blob_path}",
                "chunkIndex": chunk_index,
                "totalChunks": total_chunks
            })

        return jsonify({
            "status": "uploaded",
            "chunkIndex": chunk_index,
            "totalChunks": total_chunks
        })

    except Exception as e:
        return jsonify({"error": f"Chunk upload failed: {str(e)}"}), 500


@app.route("/api/ai/analyze-blueprint", methods=["POST"])
def ai_analyze_blueprint():
    """Analyze an uploaded document or video to extract project blueprint.

    Supports two modes:
    1. Direct file upload (for small files < 32MB)
    2. GCS URI (for large files uploaded via signed URL)
    """
    import json
    import tempfile
    import mimetypes
    from vertexai.generative_models import Part

    # Check for GCS URI (for large files uploaded directly to GCS)
    gcs_uri = None
    if request.is_json:
        data = request.get_json()
        gcs_uri = data.get('gcsUri')
        num_episodes = int(data.get('numEpisodes', 5))
        filename = data.get('filename', 'video.mp4')
    else:
        num_episodes = int(request.form.get('numEpisodes', 5))
        gcs_uri = request.form.get('gcsUri')
        filename = request.form.get('filename', '')

    # Mode 1: GCS URI provided (large file already in GCS)
    if gcs_uri:
        ext = filename.rsplit('.', 1)[-1].lower() if '.' in filename else 'mp4'
        mime_type = mimetypes.guess_type(filename)[0] or 'video/mp4'
        file_content = None  # No content to read, using GCS
        print(f"Analyzing from GCS: {gcs_uri}")

    # Mode 2: Direct file upload (small files)
    elif 'file' in request.files:
        file = request.files['file']
        if file.filename == '':
            return jsonify({"error": "No file selected"}), 400

        filename = file.filename
        ext = filename.rsplit('.', 1)[-1].lower() if '.' in filename else ''
        mime_type = mimetypes.guess_type(filename)[0] or 'application/octet-stream'
        file_content = file.read()
    else:
        return jsonify({"error": "No file uploaded and no gcsUri provided"}), 400

    system_prompt = """You are a documentary production analyst. Analyze the provided content and create a comprehensive project blueprint document.

IMPORTANT: Respond ONLY with a JSON object. No markdown, no explanation, just valid JSON.

The JSON must have:
- "title": A compelling project title (max 80 chars)
- "description": A comprehensive description of what this documentary should cover (max 500 chars)
- "style": The documentary style/approach that fits best (e.g., "investigative journalism", "observational", "personal narrative", "educational", "cinematic")
- "episodes": An array of episode objects, each with "title", "description", and "order"
- "blueprintDocument": A detailed markdown document (1500-2500 words) that serves as the project blueprint, including:
  * Executive Summary
  * Project Overview and Goals
  * Target Audience
  * Visual Style and Tone
  * Key Themes to Explore
  * Production Approach
  * Episode Breakdown with descriptions
  * Editing Approach (detailed section covering):
    - Pacing and rhythm guidelines
    - Transition styles between scenes/segments
    - Use of B-roll and cutaways
    - Music and sound design direction
    - Graphics and text overlay style
    - Color grading/look recommendations
    - Interview editing approach (jump cuts vs continuous, etc.)
    - Narrative structure and story arc editing
  * Potential Challenges and Considerations

Example response format:
{
  "title": "Documentary Title",
  "description": "What this documentary is about...",
  "style": "investigative journalism",
  "episodes": [
    {"title": "Episode 1 Title", "description": "What this episode covers", "order": 1},
    {"title": "Episode 2 Title", "description": "What this episode covers", "order": 2}
  ],
  "blueprintDocument": "# Project Blueprint\\n\\n## Executive Summary\\n..."
}"""

    try:
        # Determine if it's a video or document
        is_video = mime_type.startswith('video/') or ext in ['mp4', 'mov', 'avi', 'mkv', 'webm']
        is_document = ext in ['pdf', 'txt', 'doc', 'docx', 'md'] or mime_type.startswith('text/')

        # Variable to store blueprint file info
        blueprint_file = None
        source_filename = filename

        if is_video:
            # For videos, use GCS URI for Gemini analysis
            bucket = storage_client.bucket(STORAGE_BUCKET)
            temp_blob = None

            if gcs_uri:
                # Large file already in GCS - use provided URI
                video_uri = gcs_uri
                print(f"Using existing GCS file: {video_uri}")
            else:
                # Small file - upload temporarily to GCS
                file_hash = hashlib.md5(file_content).hexdigest()
                temp_blob_name = f"temp_blueprints/{file_hash}.{ext}"
                temp_blob = bucket.blob(temp_blob_name)
                temp_blob.upload_from_string(file_content, content_type=mime_type)
                video_uri = f"gs://{STORAGE_BUCKET}/{temp_blob_name}"
                print(f"Uploaded temp file: {video_uri}")

            prompt = f"""Analyze this video and create a comprehensive documentary project blueprint.
The video is a reference, sample, or outline for a documentary project.

Based on what you see and hear in the video, create:
1. A compelling project title
2. A comprehensive description
3. The appropriate documentary style
4. {num_episodes} episode topics
5. A detailed blueprint document (1000-2000 words) covering all aspects of the project

Return ONLY the JSON object as specified."""

            # Use multimodal model with video
            video_part = Part.from_uri(video_uri, mime_type=mime_type)
            response = model.generate_content([system_prompt, video_part, prompt])
            result = response.text

            # Delete the temporary video file after analysis (only if we uploaded it)
            if temp_blob:
                def cleanup_video():
                    import time
                    time.sleep(30)
                    try:
                        temp_blob.delete()
                    except:
                        pass
                threading.Thread(target=cleanup_video, daemon=True).start()

        elif is_document:
            # For documents, analyze and generate a blueprint document
            if ext == 'pdf':
                # Use Gemini to analyze PDF
                doc_part = Part.from_data(file_content, mime_type='application/pdf')
                prompt = f"""Analyze this PDF document and create a comprehensive documentary project blueprint.
The document contains information about a documentary project or subject matter.

Based on the content, create:
1. A compelling project title
2. A comprehensive description
3. The appropriate documentary style
4. {num_episodes} episode topics
5. A detailed blueprint document (1000-2000 words) covering all aspects of the project

Return ONLY the JSON object as specified."""

                response = model.generate_content([system_prompt, doc_part, prompt])
                result = response.text
            else:
                # For text files, decode and send as text
                try:
                    text_content = file_content.decode('utf-8')
                except:
                    text_content = file_content.decode('latin-1')

                prompt = f"""Analyze this document and create a comprehensive documentary project blueprint.

DOCUMENT CONTENT:
{text_content[:50000]}

Based on the content, create:
1. A compelling project title
2. A comprehensive description
3. The appropriate documentary style
4. {num_episodes} episode topics
5. A detailed blueprint document (1000-2000 words) covering all aspects of the project

Return ONLY the JSON object as specified."""

                result = generate_ai_response(prompt, system_prompt)
        else:
            return jsonify({"error": f"Unsupported file type: {ext}"}), 400

        # Parse the JSON response
        cleaned = result.strip()
        if cleaned.startswith('```'):
            lines = cleaned.split('\n')
            cleaned = '\n'.join(lines[1:-1] if lines[-1].startswith('```') else lines[1:])

        # Clean control characters that break JSON parsing
        import re
        # Remove all ASCII control characters except newline, tab, carriage return
        cleaned = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]', '', cleaned)
        # Also remove extended control characters (0x80-0x9f)
        cleaned = re.sub(r'[\x80-\x9f]', '', cleaned)

        # Fix unescaped newlines inside JSON string values
        # This is a common issue when AI generates long text content
        def fix_json_strings(json_str):
            """Fix newlines inside JSON string values by properly escaping them."""
            result = []
            in_string = False
            escape_next = False

            for char in json_str:
                if escape_next:
                    result.append(char)
                    escape_next = False
                elif char == '\\':
                    result.append(char)
                    escape_next = True
                elif char == '"':
                    result.append(char)
                    in_string = not in_string
                elif char == '\n' and in_string:
                    # Newline inside string - escape it
                    result.append('\\n')
                elif char == '\r' and in_string:
                    # Carriage return inside string - escape it
                    result.append('\\r')
                elif char == '\t' and in_string:
                    # Tab inside string - escape it
                    result.append('\\t')
                else:
                    result.append(char)

            return ''.join(result)

        cleaned = fix_json_strings(cleaned)

        blueprint = json.loads(cleaned)

        # Save the blueprint document to GCS
        blueprint_doc_content = blueprint.get('blueprintDocument', '')
        if not blueprint_doc_content:
            # Generate a basic document if not provided
            blueprint_doc_content = f"""# {blueprint.get('title', 'Documentary Project')}

## Project Overview
{blueprint.get('description', '')}

## Documentary Style
{blueprint.get('style', 'Documentary')}

## Episodes
"""
            for ep in blueprint.get('episodes', []):
                blueprint_doc_content += f"\n### Episode {ep.get('order', '')}: {ep.get('title', '')}\n{ep.get('description', '')}\n"

        # Convert markdown to PDF
        import markdown

        # Convert markdown to HTML
        html_content = markdown.markdown(blueprint_doc_content, extensions=['tables', 'fenced_code'])

        # Create styled HTML document
        styled_html = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <style>
        body {{
            font-family: 'Helvetica Neue', Arial, sans-serif;
            line-height: 1.6;
            max-width: 800px;
            margin: 0 auto;
            padding: 40px;
            color: #333;
        }}
        h1 {{
            color: #1a1a1a;
            border-bottom: 2px solid #2563eb;
            padding-bottom: 10px;
            margin-top: 0;
        }}
        h2 {{
            color: #2563eb;
            margin-top: 30px;
            border-bottom: 1px solid #e5e7eb;
            padding-bottom: 8px;
        }}
        h3 {{
            color: #4b5563;
            margin-top: 20px;
        }}
        p {{
            margin-bottom: 12px;
        }}
        ul, ol {{
            margin-bottom: 16px;
            padding-left: 24px;
        }}
        li {{
            margin-bottom: 6px;
        }}
        strong {{
            color: #1a1a1a;
        }}
        .header {{
            text-align: center;
            margin-bottom: 40px;
            padding-bottom: 20px;
            border-bottom: 3px solid #2563eb;
        }}
        .header h1 {{
            border: none;
            margin-bottom: 10px;
        }}
        .header p {{
            color: #6b7280;
            font-style: italic;
        }}
    </style>
</head>
<body>
    <div class="header">
        <h1>{blueprint.get('title', 'Documentary Blueprint')}</h1>
        <p>Project Blueprint Document</p>
    </div>
    {html_content}
</body>
</html>"""

        # Convert HTML to PDF using WeasyPrint
        pdf_content = HTML(string=styled_html).write_pdf()

        # Save PDF to GCS
        bucket = storage_client.bucket(STORAGE_BUCKET)
        doc_hash = hashlib.md5(blueprint_doc_content.encode()).hexdigest()
        doc_blob_name = f"blueprints/{doc_hash}_blueprint.pdf"
        doc_blob = bucket.blob(doc_blob_name)
        doc_blob.upload_from_string(pdf_content, content_type='application/pdf')

        # Create blueprint file info for the document
        blueprint["blueprintFile"] = {
            "path": doc_blob_name,
            "filename": f"{blueprint.get('title', 'Blueprint')[:50]}_blueprint.pdf",
            "mimeType": "application/pdf",
            "size": len(pdf_content),
            "type": "document",
            "sourceFile": source_filename if is_video else None
        }

        # Include the markdown content for inline display
        blueprint["blueprintContent"] = blueprint_doc_content

        return jsonify({"blueprint": blueprint})

    except Exception as e:
        return jsonify({"error": str(e)}), 500


# ============== Document Serving Routes ==============

@app.route("/api/document/<path:blob_path>")
def get_document(blob_path):
    """Serve a document from GCS (inline viewing)."""
    try:
        bucket = storage_client.bucket(STORAGE_BUCKET)
        blob = bucket.blob(blob_path)

        if not blob.exists():
            return jsonify({"error": "Document not found"}), 404

        content = blob.download_as_bytes()
        content_type = blob.content_type or 'application/octet-stream'

        return Response(
            content,
            mimetype=content_type,
            headers={'Content-Disposition': f'inline; filename="{blob_path.split("/")[-1]}"'}
        )
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/download/<path:blob_path>")
def download_document(blob_path):
    """Download a document from GCS (attachment)."""
    try:
        bucket = storage_client.bucket(STORAGE_BUCKET)
        blob = bucket.blob(blob_path)

        if not blob.exists():
            return jsonify({"error": "Document not found"}), 404

        content = blob.download_as_bytes()
        content_type = blob.content_type or 'application/octet-stream'
        filename = blob_path.split("/")[-1]

        return Response(
            content,
            mimetype=content_type,
            headers={'Content-Disposition': f'attachment; filename="{filename}"'}
        )
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/projects/<project_id>/source-documents", methods=["GET"])
def get_source_documents(project_id):
    """Get all source documents for a project."""
    try:
        docs_ref = db.collection(COLLECTIONS['assets']).where(
            'projectId', '==', project_id
        ).where(
            'isSourceDocument', '==', True
        )
        documents = []
        for doc in docs_ref.stream():
            doc_data = doc.to_dict()
            doc_data['id'] = doc.id
            documents.append(doc_data)
        return jsonify(documents)
    except Exception as e:
        return jsonify({"error": str(e)}), 500


# ============== Initialize Sample Data ==============

@app.route("/api/init-sample-data", methods=["POST"])
def init_sample_data():
    """Initialize sample data for testing."""
    # Check if project already exists
    existing = list(db.collection(COLLECTIONS['projects']).limit(1).stream())
    if existing:
        return jsonify({"message": "Data already exists", "projectId": existing[0].id})

    # Create sample project
    project_data = {
        'title': 'Apollo 11: Journey to the Moon',
        'description': 'A comprehensive documentary series exploring the historic first moon landing',
        'status': 'In Production'
    }
    project = create_doc('projects', project_data)
    project_id = project['id']

    # Create sample episodes
    episodes = [
        {'projectId': project_id, 'title': 'Episode 1: The Race Begins', 'description': 'Cold War context and the space race', 'status': 'Research', 'duration': '45 min'},
        {'projectId': project_id, 'title': 'Episode 2: Preparation', 'description': 'Training and technical development', 'status': 'Planning', 'duration': '45 min'},
        {'projectId': project_id, 'title': 'Episode 3: Launch', 'description': 'The Saturn V launch and journey to the moon', 'status': 'Planning', 'duration': '45 min'},
        {'projectId': project_id, 'title': 'Episode 4: One Small Step', 'description': 'The lunar landing and moonwalk', 'status': 'Planning', 'duration': '45 min'}
    ]
    for ep in episodes:
        create_doc('episodes', ep)

    # Create sample research
    research_items = [
        {'projectId': project_id, 'title': 'NASA Archives Access', 'content': 'Contact: Dr. Sarah Mitchell at NASA History Office. Available footage includes mission control audio, cockpit recordings, and technical schematics.', 'category': 'Archive'},
        {'projectId': project_id, 'title': 'Cold War Context Research', 'content': 'Key sources: "The Right Stuff" by Tom Wolfe, Smithsonian Air & Space Museum archives, Kennedy Space Center historical records.', 'category': 'Background'}
    ]
    for r in research_items:
        create_doc('research', r)

    # Create sample interviews
    interviews = [
        {'projectId': project_id, 'subject': 'Buzz Aldrin', 'role': 'Lunar Module Pilot', 'status': 'Confirmed', 'questions': 'What were your thoughts during descent?\nDescribe the moment of landing.\nWhat did the lunar surface feel like?', 'notes': 'Available for 2-hour interview in Los Angeles'},
        {'projectId': project_id, 'subject': 'Gene Kranz', 'role': 'Flight Director', 'status': 'Requested', 'questions': 'Describe mission control during landing.\nWhat was the most critical moment?\nHow did the team prepare?', 'notes': 'Contact through NASA public affairs'}
    ]
    for i in interviews:
        create_doc('interviews', i)

    # Create sample shots
    shots = [
        {'projectId': project_id, 'description': 'Kennedy Space Center launch pads - modern day', 'location': 'KSC, Florida', 'equipment': '4K drone, cinema camera', 'status': 'Scheduled', 'shootDate': '2026-03-15'},
        {'projectId': project_id, 'description': 'Saturn V rocket at Space Center Houston', 'location': 'Houston, TX', 'equipment': 'Gimbal, 4K camera', 'status': 'Pending'}
    ]
    for s in shots:
        create_doc('shots', s)

    # Create sample assets
    assets = [
        {'projectId': project_id, 'title': 'Original Mission Control Audio', 'type': 'Audio', 'source': 'NASA Archives', 'status': 'Acquired', 'notes': 'Full 8-day mission audio, needs editing'},
        {'projectId': project_id, 'title': 'Apollo 11 Launch Footage', 'type': 'Video', 'source': 'NASA/CBS News Archives', 'status': 'Licensing', 'notes': 'Multiple camera angles, 16mm film transfer'},
        {'projectId': project_id, 'title': 'Lunar Surface Photos', 'type': 'Image', 'source': 'NASA/Hasselblad', 'status': 'Acquired', 'notes': 'High-res scans of original photos'}
    ]
    for a in assets:
        create_doc('assets', a)

    # Create sample script
    script = {
        'projectId': project_id,
        'title': 'Episode 1 Outline',
        'content': '''ACT 1: COLD WAR CONTEXT
- Sputnik shock (1957)
- Kennedy's moon speech (1961)
- Early Mercury/Gemini programs

ACT 2: THE APOLLO PROGRAM
- Tragedy of Apollo 1
- Technical challenges
- Team assembly

ACT 3: APOLLO 11 CREW
- Armstrong, Aldrin, Collins selection
- Training montage
- Mission objectives'''
    }
    create_doc('scripts', script)

    return jsonify({"message": "Sample data created", "projectId": project_id})


# ============== Production Factory Routes ==============

# --- Project Dashboard ---
@app.route("/api/projects/<project_id>/dashboard", methods=["GET"])
def get_project_dashboard(project_id):
    """Get dashboard statistics for a project."""
    try:
        stats = get_project_dashboard_stats(project_id)
        return jsonify(stats)
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/projects/<project_id>/workflow-overview", methods=["GET"])
def get_project_workflow_overview(project_id):
    """Get workflow overview for all episodes in a project."""
    try:
        episodes = get_all_docs('episodes', project_id)
        overview = []
        for episode in episodes:
            workflow_status = get_episode_workflow_status(episode['id'])
            if workflow_status:
                overview.append({
                    'episodeId': episode['id'],
                    'episodeTitle': episode.get('title', 'Untitled'),
                    'seriesId': episode.get('seriesId'),
                    **workflow_status
                })
        return jsonify(overview)
    except Exception as e:
        return jsonify({"error": str(e)}), 500


# --- Episode Workflow ---
@app.route("/api/episodes/<episode_id>/workflow", methods=["GET"])
def get_episode_workflow(episode_id):
    """Get workflow status for an episode."""
    try:
        status = get_episode_workflow_status(episode_id)
        if not status:
            return jsonify({"error": "Episode not found"}), 404
        return jsonify(status)
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/episodes/<episode_id>/workflow/phase", methods=["PUT"])
def update_workflow_phase(episode_id):
    """Update an episode's workflow phase status."""
    try:
        data = request.get_json()
        phase = data.get('phase')
        status = data.get('status')
        notes = data.get('notes')

        if not phase or not status:
            return jsonify({"error": "Phase and status are required"}), 400

        if phase not in EPISODE_PHASES:
            return jsonify({"error": f"Invalid phase. Must be one of: {list(EPISODE_PHASES.keys())}"}), 400

        if status not in PHASE_STATUSES:
            return jsonify({"error": f"Invalid status. Must be one of: {PHASE_STATUSES}"}), 400

        result = update_episode_phase(episode_id, phase, status, notes)
        if not result:
            return jsonify({"error": "Episode not found or update failed"}), 404

        return jsonify(result)
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/episodes/factory", methods=["POST"])
def create_episode_factory():
    """Create a new episode with Production Factory structure."""
    try:
        data = request.get_json()
        episode = create_episode_with_buckets(data)
        return jsonify(episode), 201
    except Exception as e:
        return jsonify({"error": str(e)}), 500


# --- Research Documents (Research Bucket) ---
@app.route("/api/episodes/<episode_id>/research-bucket", methods=["GET"])
def get_episode_research_bucket(episode_id):
    """Get all research documents from the research bucket for an episode."""
    try:
        docs = get_docs_by_episode('research_documents', episode_id)
        return jsonify(docs)
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/research-documents", methods=["POST"])
def create_research_document():
    """Create a new research document."""
    try:
        data = request.get_json()
        # Validate required fields
        if not data.get('episodeId'):
            return jsonify({"error": "episodeId is required"}), 400

        # Set document type
        data['documentType'] = data.get('documentType', 'uploaded')  # uploaded, agent_output, fact_check, notebooklm
        data['confidenceLevel'] = data.get('confidenceLevel', 'unverified')  # verified, probable, requires_confirmation

        doc = create_doc('research_documents', data)
        return jsonify(doc), 201
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/research-documents/<doc_id>", methods=["PUT"])
def update_research_document(doc_id):
    """Update a research document."""
    try:
        data = request.get_json()
        doc = update_doc('research_documents', doc_id, data)
        return jsonify(doc)
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/research-documents/<doc_id>", methods=["DELETE"])
def delete_research_document(doc_id):
    """Delete a research document."""
    delete_doc('research_documents', doc_id)
    return jsonify({"success": True})


@app.route("/api/research-documents/upload", methods=["POST"])
def upload_research_document():
    """Upload a file as a research document."""
    if 'file' not in request.files:
        return jsonify({"error": "No file provided"}), 400

    file = request.files['file']
    if file.filename == '':
        return jsonify({"error": "No file selected"}), 400

    # Get form data
    episode_id = request.form.get('episodeId')
    project_id = request.form.get('projectId')
    title = request.form.get('title', file.filename)
    document_type = request.form.get('documentType', 'uploaded')
    confidence_level = request.form.get('confidenceLevel', 'requires_confirmation')
    content = request.form.get('content', '')
    source_url = request.form.get('sourceUrl', '')

    if not episode_id:
        return jsonify({"error": "episodeId is required"}), 400

    try:
        # Read file content
        file_content = file.read()
        file_size = len(file_content)

        # Determine content type
        content_type = file.content_type or 'application/octet-stream'

        # Generate safe filename
        original_filename = file.filename
        safe_filename = re.sub(r'[^\w\-.]', '_', original_filename)

        # Generate unique path
        file_hash = hashlib.md5(file_content).hexdigest()[:8]
        blob_path = f"research/{episode_id}/{file_hash}_{safe_filename}"

        # Upload to GCS
        ensure_bucket_exists(STORAGE_BUCKET)
        bucket = storage_client.bucket(STORAGE_BUCKET)
        blob = bucket.blob(blob_path)
        blob.upload_from_string(file_content, content_type=content_type)

        print(f"Uploaded research document: {blob_path} ({file_size} bytes)")

        # Determine file type for categorization
        file_ext = original_filename.rsplit('.', 1)[-1].lower() if '.' in original_filename else ''
        file_category = 'document'
        if file_ext in ['mp4', 'mov', 'avi', 'mkv', 'webm']:
            file_category = 'video'
        elif file_ext in ['mp3', 'wav', 'aiff', 'aac', 'm4a']:
            file_category = 'audio'
        elif file_ext in ['jpg', 'jpeg', 'png', 'gif', 'tiff', 'bmp', 'webp']:
            file_category = 'image'
        elif file_ext in ['pdf', 'doc', 'docx', 'txt', 'rtf', 'md']:
            file_category = 'document'

        # Create research document entry with file info
        doc_data = {
            'episodeId': episode_id,
            'projectId': project_id,
            'title': title,
            'content': content or f'Uploaded file: {original_filename}',
            'documentType': document_type,
            'confidenceLevel': confidence_level,
            'sourceUrl': source_url,
            'fileCategory': file_category,
            'gcsPath': blob_path,
            'filename': original_filename,
            'mimeType': content_type,
            'sizeBytes': file_size,
            'hasFile': True,
            'uploadedAt': datetime.utcnow().isoformat()
        }

        doc = create_doc('research_documents', doc_data)

        return jsonify({
            "success": True,
            "researchDocument": doc,
            "gcsPath": blob_path,
            "filename": original_filename,
            "size": file_size,
            "category": file_category
        }), 201

    except Exception as e:
        print(f"Error uploading research document: {e}")
        return jsonify({"error": str(e)}), 500


# --- Archive Logs ---
@app.route("/api/episodes/<episode_id>/archive-logs", methods=["GET"])
def get_episode_archive_logs(episode_id):
    """Get all archive logs for an episode."""
    try:
        logs = get_docs_by_episode('archive_logs', episode_id)
        return jsonify(logs)
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/archive-logs", methods=["POST"])
def create_archive_log():
    """Create a new archive log entry."""
    try:
        data = request.get_json()
        if not data.get('episodeId'):
            return jsonify({"error": "episodeId is required"}), 400

        # Archive log structure matching Quickture export
        data['source'] = data.get('source', 'manual')  # quickture, nasa_api, getty, manual
        data['clipCount'] = data.get('clipCount', 0)

        log = create_doc('archive_logs', data)
        return jsonify(log), 201
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/archive-logs/import-csv", methods=["POST"])
def import_archive_csv():
    """Import archive log from Quickture CSV export."""
    try:
        data = request.get_json()
        episode_id = data.get('episodeId')
        csv_content = data.get('csvContent')
        source = data.get('source', 'quickture')

        if not episode_id or not csv_content:
            return jsonify({"error": "episodeId and csvContent are required"}), 400

        # Parse CSV content (expecting: Filename, Timecode_In, Timecode_Out, Description, Keywords, Technical_Notes, Getty_ID)
        import csv
        import io

        reader = csv.DictReader(io.StringIO(csv_content))
        clips = []
        for row in reader:
            clips.append({
                'filename': row.get('Filename', ''),
                'timecodeIn': row.get('Timecode_In', ''),
                'timecodeOut': row.get('Timecode_Out', ''),
                'description': row.get('Description', ''),
                'keywords': row.get('Keywords', '').split(',') if row.get('Keywords') else [],
                'technicalNotes': row.get('Technical_Notes', ''),
                'gettyId': row.get('Getty_ID', ''),
                'nasaId': row.get('NASA_ID', ''),
            })

        # Create archive log entry
        log_data = {
            'episodeId': episode_id,
            'source': source,
            'clips': clips,
            'clipCount': len(clips),
            'importedAt': datetime.utcnow().isoformat()
        }

        log = create_doc('archive_logs', log_data)
        return jsonify(log), 201
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/archive-logs/<log_id>", methods=["DELETE"])
def delete_archive_log(log_id):
    """Delete an archive log."""
    delete_doc('archive_logs', log_id)
    return jsonify({"success": True})


@app.route("/api/archive-logs/upload", methods=["POST"])
def upload_archive_file():
    """Upload a file to the archive bucket and create an archive log entry."""
    if 'file' not in request.files:
        return jsonify({"error": "No file provided"}), 400

    file = request.files['file']
    if file.filename == '':
        return jsonify({"error": "No file selected"}), 400

    # Get form data
    episode_id = request.form.get('episodeId')
    project_id = request.form.get('projectId')
    source_type = request.form.get('sourceType', 'upload')  # upload, nasa, getty, pond5
    description = request.form.get('description', '')
    keywords = request.form.get('keywords', '')
    technical_notes = request.form.get('technicalNotes', '')

    if not episode_id:
        return jsonify({"error": "episodeId is required"}), 400

    try:
        # Read file content
        file_content = file.read()
        file_size = len(file_content)

        # Determine content type
        content_type = file.content_type or 'application/octet-stream'

        # Generate safe filename
        original_filename = file.filename
        safe_filename = re.sub(r'[^\w\-.]', '_', original_filename)

        # Generate unique path
        file_hash = hashlib.md5(file_content).hexdigest()[:8]
        blob_path = f"archive/{episode_id}/{file_hash}_{safe_filename}"

        # Upload to GCS
        ensure_bucket_exists(STORAGE_BUCKET)
        bucket = storage_client.bucket(STORAGE_BUCKET)
        blob = bucket.blob(blob_path)
        blob.upload_from_string(file_content, content_type=content_type)

        print(f"Uploaded archive file: {blob_path} ({file_size} bytes)")

        # Determine file type for categorization
        file_ext = original_filename.rsplit('.', 1)[-1].lower() if '.' in original_filename else ''
        file_category = 'document'
        if file_ext in ['mp4', 'mov', 'avi', 'mkv', 'webm', 'mxf', 'prores']:
            file_category = 'video'
        elif file_ext in ['mp3', 'wav', 'aiff', 'aac', 'm4a']:
            file_category = 'audio'
        elif file_ext in ['jpg', 'jpeg', 'png', 'gif', 'tiff', 'bmp', 'webp']:
            file_category = 'image'
        elif file_ext in ['pdf', 'doc', 'docx', 'txt', 'rtf']:
            file_category = 'document'

        # Create archive log entry with file info
        log_data = {
            'episodeId': episode_id,
            'projectId': project_id,
            'source': source_type,
            'sourceName': original_filename,
            'fileCategory': file_category,
            'clips': [{
                'filename': original_filename,
                'description': description,
                'keywords': [k.strip() for k in keywords.split(',') if k.strip()] if keywords else [],
                'technicalNotes': technical_notes,
                'gcsPath': blob_path,
                'mimeType': content_type,
                'sizeBytes': file_size
            }],
            'clipCount': 1,
            'gcsPath': blob_path,
            'filename': original_filename,
            'mimeType': content_type,
            'sizeBytes': file_size,
            'hasFile': True,
            'uploadedAt': datetime.utcnow().isoformat()
        }

        log = create_doc('archive_logs', log_data)

        return jsonify({
            "success": True,
            "archiveLog": log,
            "gcsPath": blob_path,
            "filename": original_filename,
            "size": file_size,
            "category": file_category
        }), 201

    except Exception as e:
        print(f"Error uploading archive file: {e}")
        return jsonify({"error": str(e)}), 500


# --- Interview Transcripts ---
@app.route("/api/episodes/<episode_id>/transcripts", methods=["GET"])
def get_episode_transcripts(episode_id):
    """Get all interview transcripts for an episode."""
    try:
        transcripts = get_docs_by_episode('interview_transcripts', episode_id)
        return jsonify(transcripts)
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/transcripts", methods=["POST"])
def create_transcript():
    """Create a new interview transcript."""
    try:
        data = request.get_json()
        if not data.get('episodeId'):
            return jsonify({"error": "episodeId is required"}), 400

        # Transcript structure
        data['speakerIdentified'] = data.get('speakerIdentified', False)
        data['timecodesAligned'] = data.get('timecodesAligned', False)
        data['segments'] = data.get('segments', [])  # [{timecode, speaker, text, metadata}]

        transcript = create_doc('interview_transcripts', data)
        return jsonify(transcript), 201
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/transcripts/<transcript_id>", methods=["PUT"])
def update_transcript(transcript_id):
    """Update an interview transcript."""
    try:
        data = request.get_json()
        transcript = update_doc('interview_transcripts', transcript_id, data)
        return jsonify(transcript)
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/transcripts/<transcript_id>", methods=["DELETE"])
def delete_transcript(transcript_id):
    """Delete an interview transcript."""
    delete_doc('interview_transcripts', transcript_id)
    return jsonify({"success": True})


# --- Script Versions ---
@app.route("/api/episodes/<episode_id>/script-versions", methods=["GET"])
def get_episode_script_versions(episode_id):
    """Get all script versions for an episode."""
    try:
        versions = get_docs_by_episode('script_versions', episode_id)
        # Sort by version number
        versions.sort(key=lambda x: x.get('versionNumber', 0))
        return jsonify(versions)
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/script-versions", methods=["POST"])
def create_script_version():
    """Create a new script version."""
    try:
        data = request.get_json()
        if not data.get('episodeId'):
            return jsonify({"error": "episodeId is required"}), 400

        # Get current version count to auto-increment
        existing = get_docs_by_episode('script_versions', data['episodeId'])
        data['versionNumber'] = len(existing) + 1
        data['versionType'] = data.get('versionType', SCRIPT_VERSION_TYPES[min(len(existing), 3)])
        data['isLocked'] = data.get('isLocked', False)
        data['segments'] = data.get('segments', [])  # [{segmentNumber, title, voiceover, archiveRefs, interviewRefs, genAiVisuals}]

        version = create_doc('script_versions', data)

        # Update episode's script workspace
        episode = get_doc('episodes', data['episodeId'])
        if episode:
            script_workspace = episode.get('scriptWorkspace', {})
            script_workspace['currentVersion'] = version['id']
            if 'versionHistory' not in script_workspace:
                script_workspace['versionHistory'] = []
            script_workspace['versionHistory'].append({
                'versionId': version['id'],
                'versionNumber': data['versionNumber'],
                'versionType': data['versionType'],
                'createdAt': version['createdAt']
            })
            update_doc('episodes', data['episodeId'], {'scriptWorkspace': script_workspace})

        return jsonify(version), 201
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/script-versions/<version_id>", methods=["PUT"])
def update_script_version(version_id):
    """Update a script version."""
    try:
        data = request.get_json()
        version = update_doc('script_versions', version_id, data)
        return jsonify(version)
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/script-versions/<version_id>/lock", methods=["POST"])
def lock_script_version(version_id):
    """Lock a script version (V4 final)."""
    try:
        version = update_doc('script_versions', version_id, {
            'isLocked': True,
            'lockedAt': datetime.utcnow().isoformat(),
            'versionType': 'V4_locked'
        })
        return jsonify(version)
    except Exception as e:
        return jsonify({"error": str(e)}), 500


# --- Compliance Items ---
@app.route("/api/episodes/<episode_id>/compliance", methods=["GET"])
def get_episode_compliance(episode_id):
    """Get all compliance items for an episode."""
    try:
        items = get_docs_by_episode('compliance_items', episode_id)
        return jsonify(items)
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/compliance", methods=["POST"])
def create_compliance_item():
    """Create a new compliance item."""
    try:
        data = request.get_json()
        if not data.get('episodeId'):
            return jsonify({"error": "episodeId is required"}), 400

        # Compliance item types (supports both legacy and new frontend categories):
        #   Legacy: source_citation, archive_license, exif_metadata, legal_signoff
        #   Frontend: archive_clearance, fact_check, contributor_release, compliance_signoff
        valid_types = [
            'source_citation', 'archive_license', 'exif_metadata', 'legal_signoff',
            'archive_clearance', 'fact_check', 'contributor_release', 'compliance_signoff',
        ]
        item_type = data.get('itemType', 'source_citation')
        if item_type not in valid_types:
            return jsonify({"error": f"Invalid itemType. Must be one of: {valid_types}"}), 400
        data['itemType'] = item_type
        data['status'] = data.get('status', 'pending')  # pending, verified, flagged, cleared, confirmed

        item = create_doc('compliance_items', data)
        return jsonify(item), 201
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/compliance/<item_id>", methods=["PUT"])
def update_compliance_item(item_id):
    """Update a compliance item."""
    try:
        data = request.get_json()
        item = update_doc('compliance_items', item_id, data)
        return jsonify(item)
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/episodes/<episode_id>/compliance/export", methods=["GET"])
def export_compliance_package(episode_id):
    """Export compliance package for an episode."""
    try:
        items = get_docs_by_episode('compliance_items', episode_id)
        episode = get_doc('episodes', episode_id)

        # Group by type (supports both legacy and new frontend categories)
        package = {
            'episodeId': episode_id,
            'episodeTitle': episode.get('title', 'Unknown') if episode else 'Unknown',
            'exportedAt': datetime.utcnow().isoformat(),
            # Legacy categories
            'sourceCitations': [i for i in items if i.get('itemType') == 'source_citation'],
            'archiveLicenses': [i for i in items if i.get('itemType') == 'archive_license'],
            'exifMetadata': [i for i in items if i.get('itemType') == 'exif_metadata'],
            'legalSignoffs': [i for i in items if i.get('itemType') == 'legal_signoff'],
            # Frontend categories (Legal & Compliance 4-tab model)
            'archiveClearances': [i for i in items if i.get('itemType') == 'archive_clearance'],
            'factChecks': [i for i in items if i.get('itemType') == 'fact_check'],
            'contributorReleases': [i for i in items if i.get('itemType') == 'contributor_release'],
            'complianceSignoffs': [i for i in items if i.get('itemType') == 'compliance_signoff'],
            'summary': {
                'totalItems': len(items),
                'verified': sum(1 for i in items if i.get('status') in ('verified', 'cleared', 'confirmed')),
                'pending': sum(1 for i in items if i.get('status') == 'pending'),
                'flagged': sum(1 for i in items if i.get('status') == 'flagged')
            }
        }

        return jsonify(package)
    except Exception as e:
        return jsonify({"error": str(e)}), 500


# --- Agent Tasks ---
@app.route("/api/episodes/<episode_id>/agent-tasks", methods=["GET"])
def get_episode_agent_tasks(episode_id):
    """Get all agent tasks for an episode."""
    try:
        tasks = get_docs_by_episode('agent_tasks', episode_id)
        return jsonify(tasks)
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/agent-tasks", methods=["POST"])
def create_agent_task_route():
    """Create a new agent task."""
    try:
        data = request.get_json()
        episode_id = data.get('episodeId')
        agent_type = data.get('agentType')
        task_type = data.get('taskType')
        input_data = data.get('inputData', {})

        if not episode_id or not agent_type:
            return jsonify({"error": "episodeId and agentType are required"}), 400

        if agent_type not in AGENT_TYPES:
            return jsonify({"error": f"Invalid agent type. Must be one of: {list(AGENT_TYPES.keys())}"}), 400

        task = create_agent_task(episode_id, agent_type, task_type, input_data)
        return jsonify(task), 201
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/agent-tasks/<task_id>", methods=["PUT"])
def update_agent_task_route(task_id):
    """Update an agent task."""
    try:
        data = request.get_json()
        status = data.get('status')
        output_data = data.get('outputData')
        error = data.get('error')

        task = update_agent_task(task_id, status, output_data, error)
        return jsonify(task)
    except Exception as e:
        return jsonify({"error": str(e)}), 500


# --- AI Agent Swarm for Script Generation ---
@app.route("/api/ai/script-swarm", methods=["POST"])
def ai_script_swarm():
    """Execute script generation using multi-agent swarm architecture."""
    try:
        data = request.get_json()
        episode_id = data.get('episodeId')
        segment_number = data.get('segmentNumber')  # Optional: generate specific segment only

        if not episode_id:
            return jsonify({"error": "episodeId is required"}), 400

        # Get episode data
        episode = get_doc('episodes', episode_id)
        if not episode:
            return jsonify({"error": "Episode not found"}), 404

        # Get research documents
        research_docs = get_docs_by_episode('research_documents', episode_id)

        # Get archive logs
        archive_logs = get_docs_by_episode('archive_logs', episode_id)

        # Get interview transcripts
        transcripts = get_docs_by_episode('interview_transcripts', episode_id)

        # Get series bible if available
        series_id = episode.get('seriesId')
        series = get_doc('series', series_id) if series_id else None

        # Load universal config (Thomas's craft rules)
        universal_doc = db.collection(COLLECTIONS['universal_config']).document('global').get()
        universal_rules = universal_doc.to_dict().get('rules', DEFAULT_UNIVERSAL_RULES) if universal_doc.exists else DEFAULT_UNIVERSAL_RULES

        # Load series config for format-aware script generation
        series_config_doc = db.collection(COLLECTIONS['series_config']).document(series_id).get() if series_id else None
        series_config = series_config_doc.to_dict() if series_config_doc and series_config_doc.exists else {}
        script_format = series_config.get('workflow', {}).get('scriptFormat', '2-column')
        script_columns = series_config.get('workflow', {}).get('scriptColumns', ['visuals', 'script'])
        fact_check_mode = series_config.get('workflow', {}).get('factCheckMode', 'at_generation')
        commentary_format = series_config.get('scriptStyle', {}).get('commentaryFormat', 'CAPS')

        # Build context for agents
        research_context = "\n\n".join([
            f"**{doc.get('title', 'Untitled')}**\n{doc.get('content', '')}"
            for doc in research_docs
        ])

        archive_context = "\n".join([
            f"- {clip.get('description', '')} [{clip.get('filename', '')}] TC: {clip.get('timecodeIn', '')}-{clip.get('timecodeOut', '')}"
            for log in archive_logs
            for clip in log.get('clips', [])
        ])

        interview_context = "\n\n".join([
            f"**{t.get('speakerName', 'Unknown Speaker')}**\n" +
            "\n".join([f"[{s.get('timecode', '')}] {s.get('text', '')}" for s in t.get('segments', [])])
            for t in transcripts
        ])

        series_bible = series.get('seriesBible', '') if series else ''
        episode_brief = episode.get('brief', {})

        # Create agent tasks
        tasks_created = []

        # Agent 1: Research Specialist - Analyze research and suggest structure
        research_task = create_agent_task(episode_id, 'research_specialist', 'analyze_research', {
            'episodeBrief': episode_brief,
            'researchContext': research_context[:10000]  # Limit context size
        })
        tasks_created.append(research_task)

        # Agent 2: Archive Specialist - Match archive to story beats
        archive_task = create_agent_task(episode_id, 'archive_specialist', 'match_archive', {
            'episodeBrief': episode_brief,
            'archiveContext': archive_context[:10000]
        })
        tasks_created.append(archive_task)

        # Agent 3: Interview Producer - Extract soundbites
        interview_task = create_agent_task(episode_id, 'interview_producer', 'extract_soundbites', {
            'episodeBrief': episode_brief,
            'interviewContext': interview_context[:10000]
        })
        tasks_created.append(interview_task)

        # Execute agents (in production, this would be async)
        # For now, execute synchronously
        results = {}

        for task in tasks_created:
            agent_type = task['agentType']
            agent_info = AGENT_TYPES[agent_type]
            input_data = task['inputData']

            # Build agent-specific prompt
            system_prompt = f"""=== UNIVERSAL DOCUMENTARY CRAFT RULES (HIGHEST PRIORITY) ===
{universal_rules}

=== AGENT ROLE ===
You are the {agent_info['name']}, specialized in {agent_info['role']}.

Your responsibilities:
{chr(10).join(f'- {r}' for r in agent_info['responsibilities'])}

Series Guidelines: {series_bible[:2000] if series_bible else 'No series bible available'}

Episode Brief:
- Summary: {episode_brief.get('summary', 'Not provided')}
- Story Beats: {', '.join(episode_brief.get('storyBeats', []))}
- Unique Angle: {episode_brief.get('uniqueAngle', 'Not specified')}"""

            if agent_type == 'research_specialist':
                prompt = f"""Analyze the following research and provide:
1. A verified timeline of key events
2. Technical concepts that need explanation
3. Claims that require interview corroboration
4. Suggested 5-7 segment narrative structure

Research Documents:
{input_data.get('researchContext', 'No research available')}"""

            elif agent_type == 'archive_specialist':
                prompt = f"""Review the available archive footage and provide:
1. Key visual sequences that support the story
2. Footage gaps that need B-roll or Gen AI visuals
3. Suggested archive clips for each story beat
4. Pacing recommendations based on available material

Archive Log:
{input_data.get('archiveContext', 'No archive available')}"""

            elif agent_type == 'interview_producer':
                prompt = f"""Analyze the interview transcripts and provide:
1. Top 10 soundbites ranked by emotional impact
2. Soundbites matched to potential story segments
3. Gaps where additional interviews are needed
4. Suggestions for follow-up questions

Interview Transcripts:
{input_data.get('interviewContext', 'No interviews available')}"""

            # Execute AI call — use the per-agent model config
            agent_model = AGENT_MODELS.get(agent_type, MODEL_NAME)
            update_agent_task(task['id'], 'in_progress')
            try:
                response = generate_ai_response(prompt, system_prompt, model_name=agent_model)
                results[agent_type] = response
                update_agent_task(task['id'], 'completed', {'analysis': response})
            except Exception as e:
                update_agent_task(task['id'], 'failed', error=str(e))
                results[agent_type] = f"Error: {str(e)}"

        # Agent 4: Script Writer - Synthesize all inputs into script
        # Build format-specific instructions
        if script_format == '3-column':
            format_instructions = """Create the script in 3-COLUMN FORMAT:
BOX NO | VISUALS | AUDIO
1      | [Camera direction or archive description] | [NARRATION IN CAPS]
2      | [Archive clip reference or GEN AI VISUAL] | "[Interview sync in lowercase]"
Use sequential box numbers for each row."""
        else:
            format_instructions = """Create the script in 2-COLUMN FORMAT. For each segment:
SEGMENT N: [TITLE IN CAPS]

VISUALS: [Camera direction, archive reference, or GEN AI VISUAL description]
SCRIPT:  [NARRATION IN CAPS if commentary. "Interview sync in lowercase." [SOT: Name - Quote]]"""

        script_writer_prompt = f"""Based on the analysis from the specialist agents, create a documentary script.

=== UNIVERSAL CRAFT RULES (OBEY THESE ABSOLUTELY) ===
{universal_rules}

RESEARCH SPECIALIST ANALYSIS:
{results.get('research_specialist', 'Not available')}

ARCHIVE SPECIALIST RECOMMENDATIONS:
{results.get('archive_specialist', 'Not available')}

INTERVIEW PRODUCER SOUNDBITES:
{results.get('interview_producer', 'Not available')}

{format_instructions}

Create 5-7 segments. Commentary format: {commentary_format}. Write in broadcast documentary style. Target 45 minutes total."""

        script_task = create_agent_task(episode_id, 'script_writer', 'generate_script', {
            'agentInputs': results
        })
        update_agent_task(script_task['id'], 'in_progress')

        try:
            script_response = generate_ai_response(script_writer_prompt,
                f"You are the Script Writer agent. Build voiceover narrative with story arc (setup, complication, resolution). Write to broadcast documentary standards. {series_bible[:1000] if series_bible else ''}",
                model_name=AGENT_MODELS['script_writer'],
                generation_config={"max_output_tokens": 65536, "temperature": 0.7})
            update_agent_task(script_task['id'], 'completed', {'script': script_response})

            # Agent 5: Fact Checker - Verify claims
            fact_check_prompt = f"""Review this script and verify all factual claims:

{script_response}

For each major claim, provide:
1. The claim text
2. Source verification (from research documents)
3. Confidence level (verified/probable/requires_confirmation)
4. Any legal concerns to flag"""

            fact_task = create_agent_task(episode_id, 'fact_checker', 'verify_script', {
                'script': script_response
            })
            update_agent_task(fact_task['id'], 'in_progress')

            fact_response = generate_ai_response(fact_check_prompt,
                "You are the Fact Checker agent. Cross-reference every major claim. Flag statements requiring legal review. Generate source citation log.",
                model_name=AGENT_MODELS['fact_checker'])
            update_agent_task(fact_task['id'], 'completed', {'verification': fact_response})

            # Create script version
            script_version = create_doc('script_versions', {
                'episodeId': episode_id,
                'versionNumber': 1,
                'versionType': 'V1_initial',
                'content': script_response,
                'factCheck': fact_response,
                'agentOutputs': {
                    'research_specialist': results.get('research_specialist'),
                    'archive_specialist': results.get('archive_specialist'),
                    'interview_producer': results.get('interview_producer'),
                    'script_writer': script_response,
                    'fact_checker': fact_response
                },
                'isLocked': False
            })

            return jsonify({
                'success': True,
                'scriptVersionId': script_version['id'],
                'script': script_response,
                'factCheck': fact_response,
                'agentOutputs': results
            })

        except Exception as e:
            return jsonify({"error": str(e)}), 500

    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/ai/fact-check-revision", methods=["POST"])
def fact_check_revision():
    """Run fact checker on a specific script version."""
    try:
        data = request.get_json()
        script_id = data.get('scriptId')
        version_number = data.get('versionNumber', 1)

        if not script_id:
            return jsonify({"error": "scriptId is required"}), 400

        # Load script version
        version_docs = db.collection(COLLECTIONS['script_versions'])\
            .where('scriptId', '==', script_id)\
            .where('versionNumber', '==', version_number)\
            .limit(1).stream()
        version_list = list(version_docs)
        if not version_list:
            return jsonify({"error": "Script version not found"}), 404

        version_doc = version_list[0]
        version_data = version_doc.to_dict()
        script_content = version_data.get('content', '')

        fact_check_prompt = f"""Review this documentary script and verify all factual claims:

{script_content}

For each major claim, provide:
1. The claim text
2. Source verification status (verified/probable/requires_confirmation)
3. Any legal concerns to flag
4. Suggested corrections where needed"""

        universal_doc = db.collection(COLLECTIONS['universal_config']).document('global').get()
        universal_rules = universal_doc.to_dict().get('rules', DEFAULT_UNIVERSAL_RULES) if universal_doc.exists else DEFAULT_UNIVERSAL_RULES

        system_prompt = f"""=== UNIVERSAL DOCUMENTARY CRAFT RULES ===
{universal_rules}

You are the Fact Checker agent. Cross-reference every major claim. Flag statements requiring legal review. Generate source citation log."""

        fact_response = generate_ai_response(fact_check_prompt, system_prompt,
                                              model_name=AGENT_MODELS.get('fact_checker', MODEL_NAME))

        # Update the version doc with new fact check
        version_doc.reference.update({'factCheck': fact_response, 'factCheckedAt': datetime.utcnow().isoformat()})

        return jsonify({"success": True, "factCheck": fact_response})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/ai/generate-iv-brief", methods=["POST"])
def generate_iv_brief():
    """Generate evolved Script Brief (IV Brief) for contributors."""
    try:
        data = request.get_json()
        episode_id = data.get('episodeId')
        script_id = data.get('scriptId')
        contributor_ids = data.get('contributorIds', [])

        if not episode_id:
            return jsonify({"error": "episodeId is required"}), 400

        # Load episode + script
        episode = get_doc('episodes', episode_id)
        if not episode:
            return jsonify({"error": "Episode not found"}), 404

        # Load current script version
        script_content = ""
        if script_id:
            script_doc = get_doc('scripts', script_id)
            if script_doc:
                script_content = str(script_doc.get('parts', ''))

        episode_brief = episode.get('brief', {})

        universal_doc = db.collection(COLLECTIONS['universal_config']).document('global').get()
        universal_rules = universal_doc.to_dict().get('rules', DEFAULT_UNIVERSAL_RULES) if universal_doc.exists else DEFAULT_UNIVERSAL_RULES

        results = {}

        # For each contributor (or one general brief if no contributors specified)
        targets = contributor_ids if contributor_ids else ['general']
        for contributor_id in targets:
            # Load contributor info if available
            contributor_info = ""
            if contributor_id != 'general':
                interviews = db.collection(COLLECTIONS['interviews'])\
                    .where('id', '==', contributor_id).limit(1).stream()
                interview_list = list(interviews)
                if interview_list:
                    iv_data = interview_list[0].to_dict()
                    contributor_info = f"Contributor: {iv_data.get('topic', 'Unknown')} — {iv_data.get('scene_context', '')}"

            brief_prompt = f"""Generate an evolved Script Brief (IV Brief) for a documentary interview contributor.

Episode Summary: {episode_brief.get('summary', 'Not provided')}
Story Beats: {', '.join(episode_brief.get('storyBeats', []))}
{contributor_info}

Script Context:
{script_content[:3000] if script_content else 'Script not yet generated.'}

Create an IV Brief that:
1. Lists the specific moments in the script where this contributor's testimony is needed
2. Provides the narrative context for each moment (what came before, what should follow)
3. Suggests 3-5 questions designed to elicit those specific soundbites
4. Highlights any technical or factual points the contributor should confirm
5. Notes the ideal emotional tone for each answer

Format clearly with sections: SCRIPT MOMENTS | QUESTIONS | TECHNICAL POINTS | TONE NOTES"""

            system_prompt = f"""=== UNIVERSAL DOCUMENTARY CRAFT RULES ===
{universal_rules}

You are an experienced documentary Interview Producer. Generate precise, production-ready interview briefs that guide contributors to deliver exactly the soundbites the script needs."""

            brief_text = generate_ai_response(brief_prompt, system_prompt)

            # Save to interviews collection
            iv_brief_doc = create_doc('interviews', {
                'episodeId': episode_id,
                'briefType': 'iv_brief',
                'contributorId': contributor_id,
                'briefText': brief_text,
                'generatedAt': datetime.utcnow().isoformat(),
                'status': 'generated'
            })

            results[contributor_id] = {
                'briefId': iv_brief_doc['id'],
                'brief_text': brief_text,
            }

        return jsonify({"success": True, "briefs": results})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/import/research-batch", methods=["POST"])
def import_research_batch():
    """Batch import research documents from external agents."""
    try:
        data = request.get_json()
        docs = data if isinstance(data, list) else data.get('docs', [])

        created = []
        updated_episodes = set()

        for doc_data in docs:
            title = doc_data.get('title', 'Untitled')
            content = doc_data.get('content', '')
            episode_id = doc_data.get('episodeId')
            series_id = doc_data.get('seriesId')
            source_type = doc_data.get('sourceType', 'external_agent')

            # Create research document
            new_doc = create_doc('research_documents', {
                'title': title,
                'content': content,
                'episodeId': episode_id,
                'seriesId': series_id,
                'source': 'external_agent',
                'documentType': source_type,
                'confidenceLevel': 'requires_confirmation',
            })
            created.append(new_doc['id'])

            # Parse ## Sequence / ## Beat headers → update episode storyBeats
            if episode_id:
                beat_lines = [
                    line.lstrip('#').strip()
                    for line in content.split('\n')
                    if line.startswith('## Sequence') or line.startswith('## Beat')
                ]
                if beat_lines:
                    episode = get_doc('episodes', episode_id)
                    if episode:
                        existing_brief = episode.get('brief', {})
                        existing_beats = existing_brief.get('storyBeats', [])
                        merged_beats = list(dict.fromkeys(existing_beats + beat_lines))
                        db.collection(COLLECTIONS['episodes']).document(episode_id).update({
                            'brief.storyBeats': merged_beats
                        })
                        updated_episodes.add(episode_id)

        return jsonify({
            "success": True,
            "created": len(created),
            "updated_episodes": list(updated_episodes),
            "document_ids": created
        })
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/ai/research-agent", methods=["POST"])
def ai_research_agent():
    """Execute deep research for an episode."""
    try:
        data = request.get_json()
        episode_id = data.get('episodeId')
        episode_brief = data.get('brief', {})

        if not episode_id:
            return jsonify({"error": "episodeId is required"}), 400

        # Create research agent task
        task = create_agent_task(episode_id, 'research_specialist', 'deep_research', {
            'brief': episode_brief
        })

        # Build research prompt
        prompt = f"""Based on this episode brief, generate a comprehensive research package:

Episode Summary: {episode_brief.get('summary', 'Not provided')}
Story Beats: {', '.join(episode_brief.get('storyBeats', []))}
Target Interviewees: {', '.join(episode_brief.get('targetInterviewees', []))}
Archive Requirements: {', '.join(episode_brief.get('archiveRequirements', []))}
Unique Angle: {episode_brief.get('uniqueAngle', 'Not specified')}

Generate:
1. 10-15 specific research questions to investigate
2. Timeline of key events requiring verification
3. Technical concepts requiring explanation
4. Potential interview subjects with their expertise areas
5. Archive footage categories needed
6. List of sources to consult (academic papers, documentaries, books, official records)

Format each section clearly with headers."""

        system_prompt = """You are a documentary research specialist. Generate thorough, factually-grounded research questions and identify key sources.
Focus on verifiable facts and primary sources. Identify potential contradictions or controversies that need investigation."""

        update_agent_task(task['id'], 'in_progress')

        try:
            raw_response = generate_ai_response(prompt, system_prompt)
            response = clean_ai_response(raw_response)

            # Create research document from output
            research_doc = create_doc('research_documents', {
                'episodeId': episode_id,
                'title': 'AI Research Package',
                'content': response,
                'documentType': 'agent_output',
                'confidenceLevel': 'requires_confirmation',
                'agentTaskId': task['id']
            })

            update_agent_task(task['id'], 'completed', {
                'researchPackage': response,
                'documentId': research_doc['id']
            })

            return jsonify({
                'success': True,
                'taskId': task['id'],
                'documentId': research_doc['id'],
                'researchPackage': response
            })

        except Exception as e:
            update_agent_task(task['id'], 'failed', error=str(e))
            return jsonify({"error": str(e)}), 500

    except Exception as e:
        return jsonify({"error": str(e)}), 500


# --- Series-Level Knowledge Base ---
@app.route("/api/series/<series_id>/knowledge-base", methods=["GET"])
def get_series_knowledge_base(series_id):
    """Get series-level knowledge base documents."""
    try:
        series = get_doc('series', series_id)
        if not series:
            return jsonify({"error": "Series not found"}), 404

        knowledge_base = series.get('knowledgeBase', {
            'seriesBible': None,
            'brandGuidelines': None,
            'editorialTone': None,
            'archiveCredentials': [],
            'complianceChecklist': None,
            'researchDocuments': []
        })

        return jsonify(knowledge_base)
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/series/<series_id>/knowledge-base", methods=["PUT"])
def update_series_knowledge_base(series_id):
    """Update series-level knowledge base."""
    try:
        data = request.get_json()
        series = get_doc('series', series_id)
        if not series:
            return jsonify({"error": "Series not found"}), 404

        knowledge_base = series.get('knowledgeBase', {})
        knowledge_base.update(data)

        updated = update_doc('series', series_id, {'knowledgeBase': knowledge_base})
        return jsonify(updated)
    except Exception as e:
        return jsonify({"error": str(e)}), 500


# --- Bulk Operations ---
@app.route("/api/series/<series_id>/episodes/bulk-create", methods=["POST"])
def bulk_create_episodes(series_id):
    """Bulk create episodes for a series."""
    try:
        data = request.get_json()
        episodes_data = data.get('episodes', [])
        project_id = data.get('projectId')

        if not episodes_data:
            return jsonify({"error": "episodes array is required"}), 400

        created = []
        for ep_data in episodes_data:
            ep_data['seriesId'] = series_id
            ep_data['projectId'] = project_id
            episode = create_episode_with_buckets(ep_data)
            created.append(episode)

        return jsonify({
            'success': True,
            'created': len(created),
            'episodes': created
        }), 201
    except Exception as e:
        return jsonify({"error": str(e)}), 500


# --- Voiceover Generation ---
@app.route("/api/ai/generate-voiceover", methods=["POST"])
def ai_generate_voiceover():
    """Generate voiceover audio from script (placeholder for 11 Labs integration)."""
    try:
        data = request.get_json()
        episode_id = data.get('episodeId')
        script_version_id = data.get('scriptVersionId')
        voice_profile = data.get('voiceProfile', 'default')

        if not episode_id or not script_version_id:
            return jsonify({"error": "episodeId and scriptVersionId are required"}), 400

        # Get script version
        script_version = get_doc('script_versions', script_version_id)
        if not script_version:
            return jsonify({"error": "Script version not found"}), 404

        # Extract VO sections from script
        script_content = script_version.get('content', '')

        # Parse [VOICEOVER] sections
        vo_sections = []
        import re
        vo_pattern = r'\[VOICEOVER\](.*?)(?=\[ARCHIVE|\[INTERVIEW|\[GEN AI|\[VOICEOVER\]|SEGMENT|\Z)'
        matches = re.findall(vo_pattern, script_content, re.DOTALL | re.IGNORECASE)

        for i, match in enumerate(matches):
            vo_text = match.strip().strip('"\'')
            if vo_text:
                vo_sections.append({
                    'segmentIndex': i + 1,
                    'text': vo_text,
                    'estimatedDuration': len(vo_text.split()) / 2.5,  # ~150 words per minute
                    'status': 'pending'
                })

        # Create voiceover task record
        vo_task = create_doc('agent_tasks', {
            'episodeId': episode_id,
            'agentType': 'voiceover_generator',
            'taskType': 'generate_voiceover',
            'status': 'pending',
            'inputData': {
                'scriptVersionId': script_version_id,
                'voiceProfile': voice_profile,
                'sections': vo_sections
            },
            'outputData': None
        })

        # NOTE: In production, this would integrate with 11 Labs API
        # For now, return the parsed VO sections
        return jsonify({
            'success': True,
            'taskId': vo_task['id'],
            'voiceProfile': voice_profile,
            'sections': vo_sections,
            'totalSections': len(vo_sections),
            'estimatedTotalDuration': sum(s['estimatedDuration'] for s in vo_sections),
            'message': 'Voiceover generation queued. 11 Labs integration pending.'
        })

    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/episodes/<episode_id>/voiceover-status", methods=["GET"])
def get_voiceover_status(episode_id):
    """Get voiceover generation status for an episode."""
    try:
        tasks = get_docs_by_episode('agent_tasks', episode_id)
        vo_tasks = [t for t in tasks if t.get('agentType') == 'voiceover_generator']

        return jsonify({
            'episodeId': episode_id,
            'tasks': vo_tasks,
            'hasVoiceover': len(vo_tasks) > 0,
            'latestStatus': vo_tasks[-1].get('status') if vo_tasks else None
        })
    except Exception as e:
        return jsonify({"error": str(e)}), 500


# --- Interview Transcription ---
@app.route("/api/ai/transcribe-interview", methods=["POST"])
def ai_transcribe_interview():
    """Transcribe interview audio using Gemini (placeholder)."""
    try:
        data = request.get_json()
        episode_id = data.get('episodeId')
        interview_id = data.get('interviewId')
        audio_url = data.get('audioUrl')
        speaker_name = data.get('speakerName', 'Unknown Speaker')

        if not episode_id:
            return jsonify({"error": "episodeId is required"}), 400

        # Create transcription task
        task = create_agent_task(episode_id, 'interview_producer', 'transcribe_interview', {
            'interviewId': interview_id,
            'audioUrl': audio_url,
            'speakerName': speaker_name
        })

        # NOTE: In production, this would use Gemini's audio capabilities
        # For demo, create a placeholder transcript structure
        transcript_data = {
            'episodeId': episode_id,
            'interviewId': interview_id,
            'speakerName': speaker_name,
            'speakerIdentified': True,
            'timecodesAligned': False,
            'segments': [],
            'status': 'pending',
            'agentTaskId': task['id'],
            'sourceAudioUrl': audio_url
        }

        transcript = create_doc('interview_transcripts', transcript_data)

        return jsonify({
            'success': True,
            'taskId': task['id'],
            'transcriptId': transcript['id'],
            'message': 'Transcription queued. Gemini audio processing integration pending.'
        })

    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/transcribe", methods=["POST"])
def transcribe_file():
    """Transcribe an audio/video file using Gemini's multimodal capabilities."""
    import json as json_module
    try:
        if 'file' not in request.files:
            return jsonify({"error": "No file provided"}), 400

        file = request.files['file']
        if not file.filename:
            return jsonify({"error": "Empty filename"}), 400

        # Read file bytes
        file_bytes = file.read()
        max_size = 25 * 1024 * 1024  # 25MB
        if len(file_bytes) > max_size:
            return jsonify({"error": f"File too large. Maximum size is 25MB."}), 400

        # Determine MIME type
        mime_type = file.content_type or 'audio/mpeg'
        language = request.form.get('language', 'en')
        speaker_diarization = request.form.get('speaker_diarization', 'false') == 'true'
        word_timestamps = request.form.get('word_timestamps', 'false') == 'true'

        # Build transcription prompt
        diarization_instruction = 'Identify different speakers and label them (Speaker 1, Speaker 2, etc.).' if speaker_diarization else ''
        speaker_field = '"speaker": "Speaker 1",' if speaker_diarization else ''
        prompt = f"""Transcribe the following audio/video file accurately.
Language: {language}
{diarization_instruction}

Return your response as valid JSON with this exact structure:
{{
  "text": "The full transcription text...",
  "segments": [
    {{
      "id": "seg-1",
      "start_time": 0.0,
      "end_time": 5.2,
      "text": "Segment text...",
      {speaker_field}
      "confidence": 0.95
    }}
  ],
  "language": "{language}"
}}

Provide accurate timestamps for each segment. Each segment should be a natural sentence or phrase.
Return ONLY the JSON object, no markdown formatting."""

        # Use Gemini multimodal with the audio/video
        from vertexai.generative_models import Part
        audio_part = Part.from_data(data=file_bytes, mime_type=mime_type)

        response = model.generate_content(
            [audio_part, prompt],
            generation_config={"temperature": 0.1, "max_output_tokens": 8192}
        )

        # Parse the AI response
        response_text = response.text.strip()
        # Strip markdown code fences if present
        if response_text.startswith("```"):
            response_text = response_text.split("\n", 1)[1] if "\n" in response_text else response_text[3:]
        if response_text.endswith("```"):
            response_text = response_text[:-3].strip()
        if response_text.startswith("json"):
            response_text = response_text[4:].strip()

        result = json_module.loads(response_text)

        return jsonify({
            "text": result.get("text", ""),
            "full_text": result.get("text", ""),
            "segments": result.get("segments", []),
            "language": result.get("language", language),
            "status": "complete"
        })

    except json_module.JSONDecodeError:
        # If Gemini returns non-JSON, return the raw text as a single segment
        raw_text = response_text if 'response_text' in locals() else ""
        lang = language if 'language' in locals() else "en"
        return jsonify({
            "text": raw_text,
            "full_text": raw_text,
            "segments": [],
            "language": lang,
            "status": "complete"
        })
    except Exception as e:
        print(f"[ERROR] Transcription failed: {e}")
        return jsonify({"error": str(e)}), 500


@app.route("/api/transcripts/<transcript_id>/process", methods=["POST"])
def process_transcript(transcript_id):
    """Process transcript to extract soundbites and metadata."""
    try:
        transcript = get_doc('interview_transcripts', transcript_id)
        if not transcript:
            return jsonify({"error": "Transcript not found"}), 404

        # Get full transcript text
        full_text = "\n".join([
            f"[{s.get('timecode', '')}] {s.get('text', '')}"
            for s in transcript.get('segments', [])
        ])

        if not full_text:
            return jsonify({"error": "No transcript content to process"}), 400

        # Use AI to extract soundbites and add metadata
        prompt = f"""Analyze this interview transcript and identify:
1. Top 10 most powerful/emotional soundbites (with timecodes if available)
2. Key factual statements that could support documentary narrative
3. Emotional peaks suitable for voiceover bed
4. Any claims that require fact-checking

Interview with: {transcript.get('speakerName', 'Unknown')}

Transcript:
{full_text[:50000]}

Format each soundbite with:
- Timecode (if available)
- Quote text
- Metadata tag (emotional/factual/technical/anecdote)
- Suggested use in documentary"""

        response = generate_ai_response(prompt,
            "You are an interview producer specialist. Extract the most compelling soundbites for documentary use.")

        # Update transcript with processed data
        update_doc('interview_transcripts', transcript_id, {
            'processed': True,
            'processedAt': datetime.utcnow().isoformat(),
            'aiAnalysis': response
        })

        return jsonify({
            'success': True,
            'transcriptId': transcript_id,
            'analysis': response
        })

    except Exception as e:
        return jsonify({"error": str(e)}), 500


# --- NASA API Integration ---
@app.route("/api/nasa/search", methods=["POST"])
def nasa_search():
    """Search NASA archives for footage and images."""
    try:
        data = request.get_json()
        query = data.get('query', '')
        media_type = data.get('mediaType', 'video')  # video, image, audio
        year_start = data.get('yearStart')
        year_end = data.get('yearEnd')

        if not query:
            return jsonify({"error": "query is required"}), 400

        # NASA Images API
        nasa_api_url = "https://images-api.nasa.gov/search"
        params = {
            'q': query,
            'media_type': media_type,
        }
        if year_start:
            params['year_start'] = year_start
        if year_end:
            params['year_end'] = year_end

        response = requests.get(nasa_api_url, params=params, timeout=30)

        if response.status_code != 200:
            return jsonify({"error": f"NASA API error: {response.status_code}"}), 500

        nasa_data = response.json()
        items = nasa_data.get('collection', {}).get('items', [])

        # Format results
        results = []
        for item in items[:50]:  # Limit to 50 results
            item_data = item.get('data', [{}])[0]
            links = item.get('links', [])
            preview_url = next((l.get('href') for l in links if l.get('rel') == 'preview'), None)

            results.append({
                'nasaId': item_data.get('nasa_id'),
                'title': item_data.get('title'),
                'description': item_data.get('description', '')[:500],
                'dateCreated': item_data.get('date_created'),
                'mediaType': item_data.get('media_type'),
                'keywords': item_data.get('keywords', []),
                'previewUrl': preview_url,
                'center': item_data.get('center')
            })

        return jsonify({
            'success': True,
            'query': query,
            'totalResults': nasa_data.get('collection', {}).get('metadata', {}).get('total_hits', 0),
            'results': results
        })

    except requests.Timeout:
        return jsonify({"error": "NASA API timeout"}), 504
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/nasa/import-to-archive", methods=["POST"])
def nasa_import_to_archive():
    """Import NASA search results to episode archive bucket."""
    try:
        data = request.get_json()
        episode_id = data.get('episodeId')
        nasa_items = data.get('items', [])

        if not episode_id or not nasa_items:
            return jsonify({"error": "episodeId and items are required"}), 400

        # Create archive log from NASA items
        clips = []
        for item in nasa_items:
            clips.append({
                'filename': item.get('nasaId', ''),
                'timecodeIn': '',
                'timecodeOut': '',
                'description': item.get('title', ''),
                'keywords': item.get('keywords', []),
                'technicalNotes': f"NASA Center: {item.get('center', 'Unknown')}",
                'nasaId': item.get('nasaId'),
                'previewUrl': item.get('previewUrl'),
                'dateCreated': item.get('dateCreated')
            })

        log_data = {
            'episodeId': episode_id,
            'source': 'nasa_api',
            'clips': clips,
            'clipCount': len(clips),
            'importedAt': datetime.utcnow().isoformat()
        }

        log = create_doc('archive_logs', log_data)

        return jsonify({
            'success': True,
            'logId': log['id'],
            'importedCount': len(clips)
        })

    except Exception as e:
        return jsonify({"error": str(e)}), 500


# --- Episode Brief Management ---
@app.route("/api/episodes/<episode_id>/brief", methods=["GET"])
def get_episode_brief(episode_id):
    """Get episode brief."""
    try:
        episode = get_doc('episodes', episode_id)
        if not episode:
            return jsonify({"error": "Episode not found"}), 404

        return jsonify(episode.get('brief', {
            'summary': '',
            'storyBeats': [],
            'targetInterviewees': [],
            'archiveRequirements': [],
            'uniqueAngle': ''
        }))
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/episodes/<episode_id>/brief", methods=["PUT"])
def update_episode_brief(episode_id):
    """Update episode brief."""
    try:
        data = request.get_json()
        episode = get_doc('episodes', episode_id)
        if not episode:
            return jsonify({"error": "Episode not found"}), 404

        brief = episode.get('brief', {})
        brief.update(data)

        updated = update_doc('episodes', episode_id, {'brief': brief})
        return jsonify(updated.get('brief', {}))
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/episodes/<episode_id>/approve-phase", methods=["POST"])
def approve_episode_phase(episode_id):
    """Approve current phase and advance to next."""
    try:
        data = request.get_json()
        notes = data.get('notes', '')

        episode = get_doc('episodes', episode_id)
        if not episode:
            return jsonify({"error": "Episode not found"}), 404

        workflow = episode.get('workflow', {})
        current_phase = workflow.get('currentPhase', 'research')

        # Update phase to approved
        result = update_episode_phase(episode_id, current_phase, 'approved', notes)

        return jsonify({
            'success': True,
            'previousPhase': current_phase,
            'newPhase': result.get('workflow', {}).get('currentPhase'),
            'episode': result
        })
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/episodes/<episode_id>/request-revision", methods=["POST"])
def request_episode_revision(episode_id):
    """Request revision for current phase."""
    try:
        data = request.get_json()
        notes = data.get('notes', '')

        if not notes:
            return jsonify({"error": "Revision notes are required"}), 400

        episode = get_doc('episodes', episode_id)
        if not episode:
            return jsonify({"error": "Episode not found"}), 404

        workflow = episode.get('workflow', {})
        current_phase = workflow.get('currentPhase', 'research')

        # Update phase to rejected with notes
        result = update_episode_phase(episode_id, current_phase, 'rejected', notes)

        return jsonify({
            'success': True,
            'phase': current_phase,
            'status': 'rejected',
            'notes': notes,
            'episode': result
        })
    except Exception as e:
        return jsonify({"error": str(e)}), 500


# ============== Documentary Studio Compatibility Endpoints ==============
# Thin shim endpoints matching the request/response shapes the React frontend expects.

@app.route("/api/health")
def api_health():
    """Health check for the Documentary Studio frontend."""
    return jsonify({
        "status": "ok",
        "platform": "Vertex AI (Flask)",
        "project": PROJECT_ID,
        "location": LOCATION,
        "timestamp": datetime.utcnow().isoformat() + "Z",
        "version": APP_VERSION,
        "env": APP_ENV
    })


@app.route("/api/search-archive", methods=["POST"])
def api_search_archive():
    """Archive search via AI prompt → JSON array of clips with real thumbnails."""
    try:
        data = request.get_json()
        query = data.get("query", "")
        source = data.get("source", "general")

        prompt = f"""You are an archive search specialist. Search the "{source}" archive for documentary footage related to: "{query}".

Based on your knowledge of what would be available in {source}'s archive, provide realistic archive footage results.

Return a JSON array of archive clips. Each clip must have:
- title: string
- duration_seconds: number (30-300)
- visual_description: string
- archive_source: string
- category: string (news/documentary/raw_footage/interview/b-roll)
- year_range: string
- quality: string (HD/SD/4K/Film)
- search_term: string (a 2-4 word image search term for finding a real thumbnail of this clip)

Provide 5-8 relevant results. Return ONLY the JSON array."""

        response_text = generate_ai_response(prompt,
                                            generation_config={"max_output_tokens": 4096, "temperature": 0.3})
        try:
            import json
            result = json.loads(response_text.strip().removeprefix("```json").removesuffix("```").strip())
        except (json.JSONDecodeError, ValueError):
            result = []

        # Fetch real thumbnails from NASA Images API for each result
        for clip in result:
            search_term = clip.pop("search_term", clip.get("title", query))
            try:
                nasa_resp = requests.get(
                    "https://images-api.nasa.gov/search",
                    params={"q": search_term, "media_type": "image", "page_size": 1},
                    timeout=5
                )
                if nasa_resp.status_code == 200:
                    items = nasa_resp.json().get("collection", {}).get("items", [])
                    if items and items[0].get("links"):
                        clip["thumbnail_url"] = items[0]["links"][0].get("href", "")
            except Exception:
                pass
            # If NASA API didn't return an image, try with the main query
            if not clip.get("thumbnail_url"):
                try:
                    nasa_resp2 = requests.get(
                        "https://images-api.nasa.gov/search",
                        params={"q": query, "media_type": "image", "page_size": 1},
                        timeout=5
                    )
                    if nasa_resp2.status_code == 200:
                        items2 = nasa_resp2.json().get("collection", {}).get("items", [])
                        if items2 and items2[0].get("links"):
                            clip["thumbnail_url"] = items2[0]["links"][0].get("href", "")
                except Exception:
                    pass

        return jsonify(result)
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/nasa/search", methods=["GET"])
def api_nasa_search():
    """Search NASA Images API and return real image/video results with thumbnails."""
    query = request.args.get("q", "")
    media_type = request.args.get("media_type", "image")
    page = request.args.get("page", "1")
    page_size = request.args.get("page_size", "20")
    if not query:
        return jsonify({"items": [], "total": 0})
    try:
        resp = requests.get(
            "https://images-api.nasa.gov/search",
            params={
                "q": query,
                "media_type": media_type,
                "page": page,
                "page_size": page_size,
            },
            timeout=10,
        )
        if resp.status_code != 200:
            return jsonify({"error": f"NASA API returned {resp.status_code}"}), 502

        collection = resp.json().get("collection", {})
        items = []
        for item in collection.get("items", []):
            data = item.get("data", [{}])[0]
            links = item.get("links", [])
            thumbnail = links[0].get("href", "") if links else ""
            nasa_id = data.get("nasa_id", "")
            items.append({
                "nasa_id": nasa_id,
                "title": data.get("title", ""),
                "description": (data.get("description", "") or "")[:300],
                "date_created": data.get("date_created", ""),
                "media_type": data.get("media_type", "image"),
                "center": data.get("center", ""),
                "keywords": (data.get("keywords", []) or [])[:5],
                "thumbnail_url": thumbnail,
                "href": item.get("href", ""),
            })

        total_hits = collection.get("metadata", {}).get("total_hits", len(items))
        return jsonify({"items": items, "total": total_hits})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/analyze-clip", methods=["POST"])
def api_analyze_clip():
    """Clip visual analysis → {visual_description, mood, quality_score}."""
    try:
        data = request.get_json()
        clip_title = data.get("clipTitle", "")

        prompt = f"""Provide visual analysis for: "{clip_title}".
Return ONLY valid JSON with these keys:
- visual_description: string
- mood: string
- quality_score: number (0-100)"""

        response_text = generate_ai_response(prompt,
                                            generation_config={"max_output_tokens": 4096, "temperature": 0.3})
        try:
            import json
            result = json.loads(response_text.strip().removeprefix("```json").removesuffix("```").strip())
        except (json.JSONDecodeError, ValueError):
            result = {"visual_description": response_text, "mood": "unknown", "quality_score": 50}
        return jsonify(result)
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/find-experts", methods=["POST"])
def api_find_experts():
    """Expert discovery → JSON array of experts."""
    try:
        data = request.get_json()
        topic = data.get("topic", "")
        expertise_domain = data.get("expertise_domain", "")
        context = data.get("context", "")
        role = data.get("role", "")

        domain_hint = f'\nExpertise domain: "{expertise_domain}"' if expertise_domain else ""
        context_hint = f"\nDocumentary context: {context}" if context else ""
        role_hint = f"\nThe expert's role in the documentary: {role}" if role else ""

        prompt = f"""Find 3 real-world experts on: "{topic}".{domain_hint}{context_hint}{role_hint}

IMPORTANT: Find experts whose primary academic or professional expertise directly relates to "{topic}". Do NOT return AI, technology, or computer science researchers unless the topic is specifically about AI or technology.

Return ONLY a JSON array where each object has:
- name: string (real person with verifiable credentials)
- title: string (their professional title)
- affiliation: string (university, institution, or organisation)
- expertise_area: string (their specific domain of expertise)
- relevance: string (why they are relevant to this topic)
- relevance_score: number (0.0-1.0)"""

        response_text = generate_ai_response(prompt,
                                            generation_config={"max_output_tokens": 8192, "temperature": 0.5})
        try:
            result = json.loads(clean_ai_response(response_text))
        except (json.JSONDecodeError, ValueError):
            result = []
        return jsonify(result)
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/refine-beat", methods=["POST"])
def api_refine_beat():
    """Script beat rewrite → {refined: text}."""
    try:
        data = request.get_json()
        current_content = data.get("currentContent", "")
        instruction = data.get("instruction", "")
        context = data.get("context", "")

        prompt = f"""You are a script doctor. Rewrite this beat: "{current_content}"
Instruction: "{instruction}"
{f'Context: {context}' if context else ''}
Return ONLY the rewritten text, no markdown or explanations."""

        response_text = generate_ai_response(prompt,
                                            model_name=AGENT_MODELS['script_writer'],
                                            generation_config={"max_output_tokens": 8192, "temperature": 0.7})
        return jsonify({"refined": response_text.strip()})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/index-source", methods=["POST"])
def api_index_source():
    """Source indexing (URL/text/youtube) → analysis JSON."""
    try:
        data = request.get_json()
        source_type = data.get("type", "")
        url = data.get("url", "")
        content = data.get("content", "")
        title = data.get("title", "")

        if source_type in ("url", "youtube") and url:
            kind = "YouTube video URL" if source_type == "youtube" else "web URL"
            prompt = f"""Analyze this {kind}: "{url}"
Provide a comprehensive analysis of what this content likely contains.
Return ONLY valid JSON with:
- title: string (inferred title)
- summary: string (200-300 words)
- key_topics: array of strings (5-8 topics)
- key_facts: array of strings (5-8 facts)
- content_type: string (article/research/news/government/video/etc)
- suggested_questions: array of 3 research questions this source could answer"""

        elif source_type == "text" and content:
            prompt = f"""Analyze this research document/text:

"{content[:15000]}"

Return ONLY valid JSON with:
- title: string (infer a title if not obvious)
- summary: string (200-300 words)
- key_topics: array of strings (5-8 topics)
- key_facts: array of strings (5-8 facts)
- content_type: string (research/transcript/notes/article/etc)
- suggested_questions: array of 3 research questions this text could answer"""
        else:
            return jsonify({"error": "Invalid source type or missing content"}), 400

        response_text = generate_ai_response(prompt)
        try:
            import json
            analysis = json.loads(response_text.strip().removeprefix("```json").removesuffix("```").strip())
        except (json.JSONDecodeError, ValueError):
            analysis = {"title": title or url or "Unknown", "summary": response_text}

        return jsonify({
            "status": "indexed",
            "title": analysis.get("title", title or url or "Text Document"),
            "summary": analysis.get("summary", ""),
            "key_topics": analysis.get("key_topics", []),
            "key_facts": analysis.get("key_facts", []),
            "content_type": analysis.get("content_type", source_type),
            "suggested_questions": analysis.get("suggested_questions", [])
        })
    except Exception as e:
        return jsonify({"error": str(e), "status": "error"}), 500


@app.route("/api/query-sources", methods=["POST"])
def api_query_sources():
    """Cross-source research → {response, key_facts, source_citations, ...}."""
    try:
        data = request.get_json()
        query = data.get("query", "")
        sources = data.get("sources", [])
        engine = data.get("engine", "google_deep_research")

        source_context = "\n\n---\n\n".join(
            f"[Source {i+1}: {s.get('title', 'Untitled')}]\n{s.get('summary', '')}\nKey Facts: {'; '.join(s.get('key_facts', []))}"
            for i, s in enumerate(sources)
        )

        if engine == "google_deep_research":
            prompt = f"""[DEEP RESEARCH MODE] You are a senior documentary researcher with access to the following indexed sources:

{source_context}

Research Question: "{query}"

Conduct a thorough multi-step analysis:
1. Cross-reference information across all sources
2. Identify corroborating evidence and contradictions
3. Note any gaps in the available information
4. Synthesize findings into a comprehensive response

Return ONLY valid JSON with:
- response: string (detailed 300-500 word answer synthesizing all sources)
- key_facts: array of strings (8-12 specific facts)
- source_citations: array of objects with source_title and relevant_info
- confidence_level: string (high/medium/low)
- follow_up_questions: array of 3 questions for deeper research
- contradictions: array of conflicting information found
- gaps: array of information gaps"""
        else:
            prompt = f"""You are a documentary researcher. Based on these sources:

{source_context}

Question: "{query}"

Return ONLY valid JSON with:
- response: string (comprehensive answer)
- key_facts: array of strings
- source_citations: array of objects with source_title and relevant_info
- follow_up_questions: array of 2-3 questions"""

        response_text = generate_ai_response(prompt,
                                            generation_config={"max_output_tokens": 8192, "temperature": 0.5})
        try:
            import json
            result = json.loads(response_text.strip().removeprefix("```json").removesuffix("```").strip())
        except (json.JSONDecodeError, ValueError):
            result = {"response": response_text, "key_facts": [], "source_citations": []}
        return jsonify(result)
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/chat", methods=["POST"])
def api_chat():
    """General chat with history → {response: text}."""
    try:
        data = request.get_json()
        message = data.get("message", "")
        system_instruction = data.get("systemInstruction", "")
        history = data.get("history", [])

        # Build conversation context from history
        context_parts = []
        for msg in history:
            role = msg.get("role", "user")
            parts = msg.get("parts", [])
            text = " ".join(p.get("text", "") for p in parts)
            if text:
                context_parts.append(f"{'User' if role == 'user' else 'Assistant'}: {text}")

        conversation_context = "\n".join(context_parts)
        full_prompt = f"""{f'Previous conversation:{chr(10)}{conversation_context}{chr(10)}{chr(10)}' if conversation_context else ''}User: {message}"""

        response_text = generate_ai_response(full_prompt, system_prompt=system_instruction,
                                            generation_config={"max_output_tokens": 4096, "temperature": 0.7})
        return jsonify({"response": response_text})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


# ============== User Profile Endpoints ==============

@app.route("/api/users", methods=["GET"])
def get_users():
    """Get all user profiles."""
    users = get_all_docs('users')
    return jsonify(users)


@app.route("/api/users", methods=["POST"])
def create_user():
    """Create a new user profile."""
    data = request.get_json()
    if not data.get('username'):
        return jsonify({"error": "username is required"}), 400
    # Check for duplicate username
    existing = db.collection(COLLECTIONS['users']).where('username', '==', data['username']).limit(1).stream()
    if list(existing):
        return jsonify({"error": f"User '{data['username']}' already exists"}), 409
    role = data.get('role', 'researcher')
    if role not in VALID_ROLES:
        return jsonify({"error": f"Invalid role. Must be one of: {VALID_ROLES}"}), 400
    data['role'] = role
    if not data.get('avatar'):
        data['avatar'] = f"https://api.dicebear.com/7.x/avataaars/svg?seed={data['username']}"
    user = create_doc('users', data)
    return jsonify(user), 201


@app.route("/api/users/<user_id>", methods=["PUT"])
def update_user(user_id):
    """Update a user profile."""
    data = request.get_json()
    user = update_doc('users', user_id, data)
    return jsonify(user)


@app.route("/api/users/<user_id>", methods=["DELETE"])
def delete_user(user_id):
    """Delete a user profile."""
    delete_doc('users', user_id)
    return jsonify({"success": True})


@app.route("/api/users/seed", methods=["POST"])
def seed_users():
    """Seed initial user profiles (idempotent — only creates if collection is empty)."""
    existing = list(db.collection(COLLECTIONS['users']).limit(1).stream())
    if existing:
        return jsonify({"message": "Users already exist"})

    created = []
    for user_data in SEED_USERS:
        user = create_doc('users', dict(user_data))
        created.append(user)

    return jsonify(created), 201


# ============== Universal Config (Thomas's Layer) ==============

@app.route("/api/universal-config", methods=["GET"])
def get_universal_config():
    """Get universal craft rules — Thomas's Layer, applied to all agents."""
    try:
        doc = db.collection(COLLECTIONS['universal_config']).document('global').get()
        if not doc.exists:
            return jsonify({"id": "global", "rules": DEFAULT_UNIVERSAL_RULES, "updatedAt": None})
        return jsonify({**doc.to_dict(), "id": "global"})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/universal-config", methods=["PUT"])
def update_universal_config():
    """Save universal craft rules."""
    try:
        data = request.get_json()
        db.collection(COLLECTIONS['universal_config']).document('global').set({
            **data, "updatedAt": datetime.utcnow().isoformat()
        }, merge=True)
        return jsonify({"success": True})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


# ============== Series Config ==============

@app.route("/api/series-config", methods=["GET"])
def list_series_configs():
    """List all series configs."""
    try:
        docs = db.collection(COLLECTIONS['series_config']).stream()
        configs = [{**d.to_dict(), "id": d.id} for d in docs]
        return jsonify(configs)
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/series-config", methods=["POST"])
def create_series_config():
    """Create a new series config."""
    try:
        data = request.get_json()
        series_id = data.get('id') or data.get('seriesId')
        if not series_id:
            return jsonify({"error": "id or seriesId is required"}), 400
        db.collection(COLLECTIONS['series_config']).document(series_id).set({
            **data, "updatedAt": datetime.utcnow().isoformat()
        })
        return jsonify({**data, "id": series_id}), 201
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/series-config/<series_id>", methods=["GET"])
def get_series_config(series_id):
    """Get series config by ID."""
    try:
        doc = db.collection(COLLECTIONS['series_config']).document(series_id).get()
        if not doc.exists:
            return jsonify({"error": "Series config not found"}), 404
        return jsonify({**doc.to_dict(), "id": series_id})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/series-config/<series_id>", methods=["PUT"])
def update_series_config(series_id):
    """Update series config."""
    try:
        data = request.get_json()
        db.collection(COLLECTIONS['series_config']).document(series_id).set({
            **data, "updatedAt": datetime.utcnow().isoformat()
        }, merge=True)
        return jsonify({**data, "id": series_id})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/series-config/<series_id>", methods=["DELETE"])
def delete_series_config(series_id):
    """Delete series config."""
    try:
        db.collection(COLLECTIONS['series_config']).document(series_id).delete()
        return jsonify({"success": True})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/series-config/seed", methods=["POST"])
def seed_series_configs():
    """Seed the three default series configs."""
    try:
        configs = [
            {
                "id": "nasa-files",
                "name": "NASA Uncovered",
                "seriesProducer": "Mark Carter",
                "workflow": {
                    "scriptFormat": "2-column",
                    "scriptColumns": ["visuals", "script"],
                    "factCheckMode": "at_generation",
                    "interviewBriefType": "question_based",
                    "researchEngine": "perplexity",
                },
                "scriptStyle": {
                    "commentaryFormat": "CAPS",
                    "syncFormat": "lowercase_quotes",
                },
                "sourceStandards": {
                    "allowWikipedia": False,
                    "followWikiSources": True,
                    "minimumAuthority": "primary",
                }
            },
            {
                "id": "superstructures",
                "name": "Superstructures Uncovered",
                "seriesProducer": "Al Blane",
                "workflow": {
                    "scriptFormat": "3-column",
                    "scriptColumns": ["box", "visuals", "audio"],
                    "factCheckMode": "rolling",
                    "interviewBriefType": "iv_brief_evolved",
                    "researchEngine": "perplexity",
                },
                "scriptStyle": {
                    "commentaryFormat": "CAPS",
                    "syncFormat": "lowercase_quotes",
                },
                "sourceStandards": {
                    "allowWikipedia": False,
                    "followWikiSources": True,
                    "minimumAuthority": "primary",
                }
            },
            {
                "id": "abandoned",
                "name": "Abandoned Places: Uncovered",
                "seriesProducer": "Dominic Hill",
                "workflow": {
                    "scriptFormat": "2-column",
                    "scriptColumns": ["visuals", "script"],
                    "factCheckMode": "at_generation",
                    "interviewBriefType": "question_based",
                    "researchEngine": "perplexity",
                },
                "scriptStyle": {
                    "commentaryFormat": "CAPS",
                    "syncFormat": "lowercase_quotes",
                    "presenterLed": True,
                },
                "sourceStandards": {
                    "allowWikipedia": False,
                    "followWikiSources": True,
                    "minimumAuthority": "primary",
                }
            },
        ]
        for config in configs:
            series_id = config.pop("id")
            db.collection(COLLECTIONS['series_config']).document(series_id).set({
                **config, "id": series_id, "updatedAt": datetime.utcnow().isoformat()
            })
        return jsonify({"success": True, "seeded": 3})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


# ============== Email Briefs ==============


@app.route("/api/email-briefs", methods=["POST"])
def email_briefs():
    """Email B-Roll Briefs to the production team via SendGrid HTTP API."""
    try:
        data = request.get_json()
        project_title = data.get('project_title', 'Untitled')
        recipients = data.get('recipients', [])
        briefs = data.get('briefs', [])

        if not recipients or not briefs:
            return jsonify({"error": "recipients and briefs are required"}), 400

        # Build HTML email body
        brief_rows = ""
        for i, brief in enumerate(briefs, 1):
            prompts_html = ""
            for p in brief.get('prompts', []):
                target = p.get('target', 'image').upper()
                badge_color = '#e53e3e' if target == 'VIDEO' else '#4299e1'
                camera_line = ""
                if target == 'VIDEO':
                    cam = p.get('camera_motion', '')
                    dur = p.get('duration_guidance', '')
                    camera_line = f'<p style="margin:0;color:#888;font-size:12px;">Camera: {cam} | Duration: {dur}</p>'
                prompts_html += f"""
                <div style="margin:8px 0;padding:10px;background:#f7f7f7;border-radius:6px;">
                    <span style="display:inline-block;padding:2px 8px;border-radius:4px;background:{badge_color};color:white;font-size:11px;font-weight:bold;">{target}</span>
                    <p style="margin:6px 0 2px;color:#333;">{p.get('prompt', '')}</p>
                    {camera_line}
                    <p style="margin:0;color:#888;font-size:12px;">Aspect: {p.get('aspect_ratio', '16:9')}</p>
                </div>"""

            brief_rows += f"""
            <div style="margin:16px 0;padding:16px;border:1px solid #ddd;border-radius:8px;">
                <h3 style="margin:0 0 8px;color:#e53e3e;">Brief {i}: {brief.get('beat_id', '')}</h3>
                <p><strong>Gap:</strong> {brief.get('gap_description', '')}</p>
                <p><strong>Narrative Context:</strong> {brief.get('narrative_context', '')}</p>
                <p><strong>Style Notes:</strong> {brief.get('style_notes', '')}</p>
                <h4 style="margin:12px 0 4px;">Prompts</h4>
                {prompts_html}
            </div>"""

        html = f"""
        <div style="font-family:Arial,sans-serif;max-width:700px;margin:0 auto;">
            <div style="background:#1a1a1a;padding:20px;border-radius:8px 8px 0 0;">
                <h1 style="color:#e53e3e;margin:0;">AiM B-Roll Briefs</h1>
                <p style="color:#a0a0a0;margin:4px 0 0;">Project: {project_title} | {len(briefs)} brief(s)</p>
            </div>
            <div style="padding:20px;background:#ffffff;border:1px solid #ddd;border-top:none;border-radius:0 0 8px 8px;">
                {brief_rows}
                <hr style="border:none;border-top:1px solid #eee;margin:20px 0;">
                <p style="color:#888;font-size:12px;">Generated by AiM Documentary Studio</p>
            </div>
        </div>"""

        # Send via SendGrid HTTP API (Cloud Run blocks SMTP ports)
        sendgrid_api_key = os.environ.get('SENDGRID_API_KEY', '')
        sender = os.environ.get('EMAIL_SENDER', 'aim-studio@arrowmedia.com')

        if not sendgrid_api_key:
            # No SendGrid configured — store in Firestore for manual retrieval
            email_record = create_doc('feedback', {
                'type': 'email_brief',
                'projectTitle': project_title,
                'recipients': recipients,
                'briefCount': len(briefs),
                'html': html,
                'status': 'queued_no_sendgrid',
            })
            return jsonify({
                "success": True,
                "message": f"{len(briefs)} brief(s) queued (SendGrid not configured — stored for manual sending)",
                "email_id": email_record.get('id'),
                "email_configured": False
            })

        # Build SendGrid v3 API payload
        to_list = [{"email": r} for r in recipients]
        subject = f'[AiM] B-Roll Briefs: {project_title} ({len(briefs)} briefs)'

        sg_payload = {
            "personalizations": [{"to": to_list, "subject": subject}],
            "from": {"email": sender, "name": "AiM Documentary Studio"},
            "content": [{"type": "text/html", "value": html}]
        }

        sg_response = requests.post(
            "https://api.sendgrid.com/v3/mail/send",
            json=sg_payload,
            headers={
                "Authorization": f"Bearer {sendgrid_api_key}",
                "Content-Type": "application/json"
            }
        )

        if sg_response.status_code not in (200, 201, 202):
            print(f"[ERROR] SendGrid response: {sg_response.status_code} {sg_response.text}")
            return jsonify({
                "error": f"SendGrid returned {sg_response.status_code}: {sg_response.text}"
            }), 502

        return jsonify({
            "success": True,
            "message": f"{len(briefs)} brief(s) emailed to {', '.join(recipients)}",
            "email_configured": True
        })

    except Exception as e:
        print(f"[ERROR] Email briefs failed: {e}")
        return jsonify({"error": str(e)}), 500


# ============== Planning Persistence Endpoints ==============


@app.route("/api/projects/<project_id>/brief", methods=["GET"])
def get_project_brief(project_id):
    """Get the development brief for a project."""
    try:
        project = get_doc('projects', project_id)
        if not project:
            return jsonify({"error": "Project not found"}), 404
        return jsonify(project.get('brief', {}))
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/projects/<project_id>/brief", methods=["POST"])
def save_project_brief(project_id):
    """Save or update the development brief for a project."""
    try:
        data = request.get_json()
        project = get_doc('projects', project_id)
        if not project:
            return jsonify({"error": "Project not found"}), 404
        updates = {'brief': data}
        # Sync target_duration_minutes to project level if provided in brief
        if data.get('episode_duration_minutes'):
            updates['target_duration_minutes'] = data['episode_duration_minutes']
        update_doc('projects', project_id, updates)
        return jsonify({"success": True, "brief": data})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/projects/<project_id>/bible", methods=["GET"])
def get_project_bible(project_id):
    """Get the series bible for a project."""
    try:
        project = get_doc('projects', project_id)
        if not project:
            return jsonify({"error": "Project not found"}), 404
        return jsonify(project.get('bible', {}))
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/projects/<project_id>/bible", methods=["POST"])
def save_project_bible(project_id):
    """Save or update the series bible for a project."""
    try:
        data = request.get_json()
        project = get_doc('projects', project_id)
        if not project:
            return jsonify({"error": "Project not found"}), 404
        update_doc('projects', project_id, {'bible': data})
        return jsonify({"success": True, "bible": data})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/episodes/<episode_id>/structure", methods=["GET"])
def get_episode_structure(episode_id):
    """Get the episode structure (parts and segments)."""
    try:
        episode = get_doc('episodes', episode_id)
        if not episode:
            return jsonify({"error": "Episode not found"}), 404
        return jsonify(episode.get('structure', []))
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/episodes/<episode_id>/structure", methods=["POST"])
def save_episode_structure(episode_id):
    """Save or update the episode structure (parts and segments)."""
    try:
        data = request.get_json()
        episode = get_doc('episodes', episode_id)
        if not episode:
            return jsonify({"error": "Episode not found"}), 404
        update_doc('episodes', episode_id, {'structure': data})
        return jsonify({"success": True, "structure": data})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


# ============== Missing Frontend Endpoints ==============


@app.route("/api/analyze-document", methods=["POST"])
def api_analyze_document():
    """Analyze uploaded document content → structured summary."""
    try:
        data = request.get_json()
        content = data.get("content", "")[:15000]
        file_name = data.get("fileName", "unknown")
        file_type = data.get("fileType", "text/plain")

        if "csv" in file_type.lower() or file_name.lower().endswith(".csv"):
            prompt = f"""Analyze this CSV data from file "{file_name}":

"{content}"

Return ONLY valid JSON with:
- title: string (descriptive title for the data)
- summary: string (200-300 word overview of what the data contains)
- key_topics: array of strings (5-8 topics covered)
- key_facts: array of strings (5-8 notable facts or data points)
- timeline_events: array of objects with date and description (extract any chronological events found in the data)"""
        else:
            prompt = f"""Analyze this document "{file_name}" (type: {file_type}):

"{content}"

Return ONLY valid JSON with:
- title: string (descriptive title)
- summary: string (200-300 word overview)
- key_topics: array of strings (5-8 topics)
- key_facts: array of strings (5-8 notable facts)
- timeline_events: array (empty array for non-temporal documents)"""

        response_text = generate_ai_response(prompt)
        try:
            import json
            analysis = json.loads(response_text.strip().removeprefix("```json").removesuffix("```").strip())
        except (json.JSONDecodeError, ValueError):
            analysis = {"title": file_name, "summary": response_text, "key_topics": [], "key_facts": [], "timeline_events": []}

        return jsonify({
            "status": "analyzed",
            "title": analysis.get("title", file_name),
            "summary": analysis.get("summary", ""),
            "key_topics": analysis.get("key_topics", []),
            "key_facts": analysis.get("key_facts", []),
            "timeline_events": analysis.get("timeline_events", [])
        })
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/series-structure", methods=["POST"])
def api_series_structure():
    """Generate documentary series structure from premise."""
    try:
        data = request.get_json()
        premise = data.get("premise", "")
        episodes_count = data.get("episodesCount", 3)

        prompt = f"""You are a senior documentary series architect. Design a {episodes_count}-episode documentary series based on this premise:

"{premise}"

Return ONLY valid JSON with:
- episodes: array of {episodes_count} objects, each with:
  - title: string (compelling episode title)
  - research_focus: string (what research is needed for this episode)
  - suggested_engine: string (one of: google_deep_research, academic_search, investigative)
- themes: array of strings (3-5 overarching themes that connect the episodes)"""

        response_text = generate_ai_response(prompt)
        try:
            import json
            result = json.loads(response_text.strip().removeprefix("```json").removesuffix("```").strip())
        except (json.JSONDecodeError, ValueError):
            result = {"episodes": [], "themes": []}
        return jsonify(result)
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/generate-script", methods=["POST"])
def api_generate_script():
    """Multi-agent script generation for a documentary."""
    try:
        data = request.get_json()
        title = data.get("title", "")
        description = data.get("description", "")
        project_id = data.get("projectId", "")
        # Look up project's target duration if not explicitly provided
        default_duration = 60
        if project_id:
            try:
                project = get_doc('projects', project_id)
                if project:
                    default_duration = project.get('target_duration_minutes', 60)
            except Exception:
                pass
        duration = data.get("duration", default_duration)
        reference_style = data.get("referenceStyle", "cinematic")
        research_context = data.get("researchContext", [])
        archive_context = data.get("archiveContext", [])
        series_id = data.get("seriesId", "")

        # Load universal rules (Thomas's craft rules)
        universal_rules = DEFAULT_UNIVERSAL_RULES
        try:
            universal_doc = db.collection(COLLECTIONS['universal_config']).document('global').get()
            if universal_doc.exists:
                universal_rules = universal_doc.to_dict().get('rules', DEFAULT_UNIVERSAL_RULES)
        except Exception:
            pass

        # Load series config for format-aware generation
        series_config = {}
        if series_id:
            try:
                series_config_doc = db.collection(COLLECTIONS['series_config']).document(series_id).get()
                if series_config_doc.exists:
                    series_config = series_config_doc.to_dict()
            except Exception:
                pass

        commentary_format = series_config.get('scriptStyle', {}).get('commentaryFormat', 'CAPS')
        script_format = series_config.get('workflow', {}).get('scriptFormat', '2-column')
        fact_check_mode = series_config.get('workflow', {}).get('factCheckMode', 'at_generation')

        # Extract Style DNA directives if available
        style_dna_block = ""
        style_dna = series_config.get('styleDna', {})
        if style_dna and style_dna.get('script_directives'):
            sd = style_dna['script_directives']
            style_dna_block = f"""
=== STYLE DNA DIRECTIVES (From Reference Analysis) ===
Pacing: {sd.get('pacing_instruction', 'N/A')}
Visual: {sd.get('visual_instruction', 'N/A')}
Editing: {sd.get('editing_instruction', 'N/A')}
Tone: {sd.get('tone_instruction', 'N/A')}
Production: {sd.get('production_instruction', 'N/A')}
Opening: {sd.get('opening_instruction', 'N/A')}
Interviews: {sd.get('interview_instruction', 'N/A')}
Archive Use: {sd.get('archive_instruction', 'N/A')}
Sound & Score: {sd.get('sound_score_instruction', 'N/A')}
Story Engine: {sd.get('story_engine_instruction', 'N/A')}
Ad Breaks: {sd.get('ad_break_instruction', 'N/A')}
Comparative Style: {sd.get('comparative_reference', 'N/A')}
"""

        # --- Server-side research enrichment ---
        # Fetch research documents from Firestore even if frontend didn't send researchContext
        episode_id = data.get("episodeId", "")
        server_research_items = []

        # 1. Fetch from research_documents collection (structured research)
        if episode_id:
            try:
                research_docs = get_docs_by_episode('research_documents', episode_id)
                for rdoc in research_docs:
                    title_str = rdoc.get('title', 'Untitled')
                    summary_str = rdoc.get('summary', rdoc.get('content', ''))[:2000]
                    key_facts = rdoc.get('key_facts', [])
                    facts_str = "; ".join(key_facts) if key_facts else ""
                    entry = f"- {title_str}: {summary_str}"
                    if facts_str:
                        entry += f"\n  Key facts: {facts_str}"
                    server_research_items.append(entry)
            except Exception:
                pass

        # 2. Fetch saved episode research (from /api/ai/simple-research)
        if episode_id:
            try:
                ep_doc = get_doc('episodes', episode_id)
                if ep_doc and ep_doc.get('research'):
                    server_research_items.append(f"- Episode Research Brief:\n{ep_doc['research'][:3000]}")
            except Exception:
                pass

        # 3. Fetch uploaded research documents from assets collection
        if episode_id or series_id or data.get("projectId"):
            try:
                doc_contents = get_research_document_contents(
                    episode_id=episode_id,
                    series_id=series_id,
                    project_id=data.get("projectId", "")
                )
                for dc in doc_contents:
                    server_research_items.append(f"- {dc['source']}:\n{dc['content'][:2000]}")
            except Exception:
                pass

        # Merge frontend-provided context with server-side research
        all_research = []
        if research_context:
            all_research.extend(
                f"- {r}" if isinstance(r, str) else f"- {r.get('title', '')}: {r.get('summary', '')}"
                for r in research_context
            )
        if server_research_items:
            all_research.extend(server_research_items)

        research_summary = "\n".join(all_research) if all_research else "No research context provided."

        archive_summary = "\n".join(
            f"- {a}" if isinstance(a, str) else f"- {a.get('title', '')}: {a.get('visual_description', '')}"
            for a in archive_context
        ) if archive_context else "No archive footage context provided."

        prompt = f"""=== UNIVERSAL DOCUMENTARY CRAFT RULES (HIGHEST PRIORITY) ===
{universal_rules}

=== GENERATION DIRECTIVES ===
Commentary Format: {commentary_format}
Script Format: {script_format}
Fact Check Mode: {fact_check_mode}
{style_dna_block}
You are a team of documentary scriptwriters. Generate a detailed script for:

Title: "{title}"
Description: "{description}"
Target Duration: {duration} minutes
Style Reference: {reference_style}

Research Context:
{research_summary}

Available Archive Footage:
{archive_summary}

Return ONLY a valid JSON array where each element is an act/segment with:
- title: string (act/segment title)
- scenes: array of scene objects, each with:
  - title: string (scene title)
  - beats: array of beat objects, each with:
    - type: string (MUST be one of: voice_over, expert, archive, ai_visual)
    - content: string (the actual script text or description)
    - speaker: string (narrator name, interviewee, or empty)
    - topic: string (for expert beats: the interview topic)
    - duration_seconds: number (estimated duration for this beat)

Use these beat types:
- voice_over: for narration, voice-over text
- expert: for interview soundbites, expert commentary
- archive: for B-roll footage, archival clips, stock footage descriptions
- ai_visual: for AI-generated visuals, transitions, title cards, graphics

STRUCTURAL REQUIREMENTS:
- You MUST fill the full {duration} minutes. This is a hard requirement.
- For a {duration}-minute documentary, use this structure:
  - 6-8 acts/parts (each ~{duration // 6}-{duration // 4} minutes)
  - 3-5 scenes per act
  - 4-8 beats per scene
  - Each voice_over beat: 30-90 seconds of narration text
  - Each archive/ai_visual beat: 10-30 seconds
  - Each expert beat: 15-45 seconds
- The sum of all duration_seconds MUST total approximately {duration * 60} seconds ({duration} minutes).
- Write FULL narration text for every voice_over beat — do NOT use placeholders."""

        response_text = generate_ai_response(
            prompt,
            model_name=AGENT_MODELS['script_writer'],
            generation_config={"max_output_tokens": 65536, "temperature": 0.7}
        )
        try:
            import json
            result = json.loads(response_text.strip().removeprefix("```json").removesuffix("```").strip())
        except (json.JSONDecodeError, ValueError):
            result = []
        return jsonify(result)
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/interview-plan", methods=["POST"])
def api_interview_plan():
    """Generate interview strategy for a scene/topic."""
    try:
        data = request.get_json()
        scene_context = data.get("sceneContext", "")
        topic = data.get("topic", "")

        prompt = f"""You are an expert documentary interview strategist. Create an interview plan for:

Scene Context: "{scene_context}"
Topic: "{topic}"

Return ONLY valid JSON with:
- ideal_soundbite: string (the ideal 1-2 sentence soundbite you want the interviewee to deliver)
- questions: array of strings (8-12 interview questions, ordered from warm-up to probing, designed to naturally elicit the ideal soundbite)"""

        response_text = generate_ai_response(prompt)
        try:
            import json
            result = json.loads(response_text.strip().removeprefix("```json").removesuffix("```").strip())
        except (json.JSONDecodeError, ValueError):
            result = {"ideal_soundbite": "", "questions": []}
        return jsonify(result)
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/generate-broll", methods=["POST"])
def api_generate_broll():
    """Generate a structured B-roll visual brief (Veo not yet integrated for video generation)."""
    try:
        data = request.get_json()
        prompt = data.get("prompt", "")
        duration = data.get("duration_seconds", 15)
        style = data.get("style", "documentary")

        ai_prompt = f"""You are a documentary cinematographer creating a B-roll visual brief.

Based on this description: "{prompt}"
Target duration: {duration} seconds
Style: {style}

Return ONLY valid JSON with:
- shots: array of objects, each with:
  - description: string (detailed shot description including subject, framing, movement)
  - duration_seconds: number (estimated duration for this shot)
  - camera_movement: string (static/pan/tilt/dolly/drone/handheld/tracking)
  - framing: string (wide/medium/close-up/extreme-close-up/aerial)
  - lighting: string (natural/golden-hour/low-key/high-key/dramatic)
- colour_palette: array of 3-5 hex colour strings that define the visual mood
- camera_notes: string (overall cinematography direction)
- mood: string (the emotional tone of the sequence)
- music_suggestion: string (style/tempo of accompanying music)
- status: string (always "visual_brief" — video generation not yet available)"""

        response_text = generate_ai_response(ai_prompt,
                                             generation_config={"max_output_tokens": 4096, "temperature": 0.7})
        try:
            result = json.loads(response_text.strip().removeprefix("```json").removesuffix("```").strip())
        except (json.JSONDecodeError, ValueError):
            result = {
                "shots": [{"description": prompt, "duration_seconds": duration, "camera_movement": "static", "framing": "wide", "lighting": "natural"}],
                "colour_palette": ["#1a1a2e", "#16213e", "#0f3460"],
                "camera_notes": response_text,
                "mood": "documentary",
                "music_suggestion": "ambient underscore",
                "status": "visual_brief"
            }
        result["status"] = "visual_brief"
        return jsonify(result)
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/gcs/serve", methods=["GET"])
def api_gcs_serve():
    """Proxy a file from GCS so the frontend can display uploaded images/videos."""
    gcs_path = request.args.get("path", "")
    if not gcs_path:
        return jsonify({"error": "Missing path parameter"}), 400
    try:
        bucket = storage_client.bucket(STORAGE_BUCKET)
        blob = bucket.blob(gcs_path)
        if not blob.exists():
            return jsonify({"error": "File not found"}), 404
        content = blob.download_as_bytes()
        content_type = blob.content_type or "application/octet-stream"
        from flask import Response
        return Response(content, content_type=content_type, headers={
            "Cache-Control": "public, max-age=86400"
        })
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/gcs/buckets", methods=["GET"])
def api_gcs_buckets():
    """List GCS buckets with metadata."""
    try:
        buckets = []
        for bucket in storage_client.list_buckets():
            blobs = list(bucket.list_blobs(max_results=1000))
            file_count = len(blobs)
            size_bytes = sum(b.size or 0 for b in blobs)
            buckets.append({
                "name": bucket.name,
                "region": bucket.location or "us-central1",
                "storageClass": bucket.storage_class or "STANDARD",
                "fileCount": file_count,
                "sizeGb": round(size_bytes / (1024 ** 3), 2)
            })
        return jsonify(buckets)
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/vertex/models", methods=["GET"])
def api_vertex_models():
    """List live Cloud Run services as infrastructure models."""
    try:
        import google.auth
        import google.auth.transport.requests as gauth_requests

        credentials, project = google.auth.default()
        auth_req = gauth_requests.Request()
        credentials.refresh(auth_req)

        # Query Cloud Run Admin API for real service data
        cr_resp = requests.get(
            f"https://run.googleapis.com/v2/projects/{PROJECT_ID}/locations/{LOCATION}/services",
            headers={"Authorization": f"Bearer {credentials.token}"},
            timeout=10
        )
        if cr_resp.status_code != 200:
            raise Exception(f"Cloud Run API returned {cr_resp.status_code}")

        services = cr_resp.json().get("services", [])
        models = []
        for svc in services:
            name = svc.get("name", "").split("/")[-1]
            conditions = svc.get("conditions", [])
            ready = any(c.get("type") == "Ready" and c.get("state") == "CONDITION_SUCCEEDED" for c in conditions)
            latest_rev = svc.get("latestReadyRevision", "").split("/")[-1]
            uri = svc.get("uri", "")

            # Health-check the service to get real latency
            latency_ms = 0
            status = "active" if ready else "error"
            try:
                t0 = datetime.now()
                hc = requests.get(uri, timeout=5, allow_redirects=True)
                latency_ms = int((datetime.now() - t0).total_seconds() * 1000)
                if hc.status_code >= 500:
                    status = "error"
            except Exception:
                status = "error"
                latency_ms = 9999

            models.append({
                "id": name,
                "name": name,
                "version": latest_rev[-12:] if latest_rev else "unknown",
                "status": status,
                "latencyMs": latency_ms,
                "callsPerMin": 0,
                "url": uri
            })
        return jsonify(models)
    except Exception as e:
        # Fallback to curated list if Cloud Run API fails
        models = [
            {"id": "gemini-2.0-flash", "name": "Gemini 2.0 Flash", "version": "v2.0", "status": "active", "latencyMs": 450, "callsPerMin": 60},
            {"id": "gemini-2.5-pro", "name": "Gemini 2.5 Pro", "version": "v2.5", "status": "active", "latencyMs": 1200, "callsPerMin": 30},
        ]
        return jsonify(models)


@app.route("/api/cloud/stats", methods=["GET"])
def api_cloud_stats():
    """Live infrastructure stats from Firestore and GCS."""
    try:
        stat_collections = {
            "projects": COLLECTIONS.get("projects", "doc_projects"),
            "series": COLLECTIONS.get("series", "doc_series"),
            "episodes": COLLECTIONS.get("episodes", "doc_episodes"),
            "scripts": COLLECTIONS.get("scripts", "doc_scripts"),
            "research": COLLECTIONS.get("research", "doc_research"),
            "assets": COLLECTIONS.get("assets", "doc_assets"),
            "interviews": COLLECTIONS.get("interviews", "doc_interviews"),
            "shots": COLLECTIONS.get("shots", "doc_shots"),
            "users": "users",
        }
        counts = {}
        for label, full_name in stat_collections.items():
            docs = db.collection(full_name).limit(1000).stream()
            counts[label] = sum(1 for _ in docs)

        # GCS quick summary (just bucket count + total size of primary bucket)
        bucket_count = sum(1 for _ in storage_client.list_buckets())
        primary_size_bytes = 0
        try:
            primary_bucket = storage_client.bucket(STORAGE_BUCKET)
            blobs = list(primary_bucket.list_blobs(max_results=500))
            primary_size_bytes = sum(b.size or 0 for b in blobs)
        except Exception:
            pass

        return jsonify({
            "firestore": counts,
            "totalDocuments": sum(counts.values()),
            "gcsBucketCount": bucket_count,
            "primaryBucketSizeGb": round(primary_size_bytes / (1024 ** 3), 2),
            "primaryBucketFiles": len(blobs) if 'blobs' in dir() else 0,
            "region": LOCATION,
            "project": PROJECT_ID
        })
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/services/status")
def api_services_status():
    """Returns which external services have API keys configured."""
    services = {
        "heygen": {
            "configured": bool(HEYGEN_API_KEY),
            "note": "Avatar video generation"
        },
        "runway": {
            "configured": bool(RUNWAY_API_KEY),
            "model": "Gen-3 Alpha",
            "note": "AI video generation"
        },
        "elevenlabs": {
            "configured": True,
            "note": "Voice synthesis (per-user keys in Firestore)"
        },
        "pexels": {
            "configured": bool(PEXELS_API_KEY),
            "note": "Stock footage search"
        },
        "pixabay": {
            "configured": bool(PIXABAY_API_KEY),
            "note": "Stock footage search"
        },
        "nasa": {
            "configured": True,
            "note": "NASA Images API (no key required)"
        },
        "sendgrid": {
            "configured": bool(os.environ.get("SENDGRID_API_KEY", "")),
            "note": "Email delivery"
        },
        "design_agent": {
            "configured": bool(CC_DESIGN_AGENT_KEY),
            "url": CC_DESIGN_AGENT_URL,
            "note": "Channel Changers design service"
        },
        "quickture": {
            "configured": bool(os.environ.get("QUICKTURE_API_URL", "")),
            "url": os.environ.get("QUICKTURE_API_URL", ""),
            "note": "Avid archive ingest"
        },
        "vertex_ai": {
            "configured": True,
            "model": MODEL_NAME,
            "note": "Primary AI engine"
        }
    }
    return jsonify({"services": services})


@app.route("/api/vertex/deploy", methods=["POST"])
def api_vertex_deploy():
    """Placeholder for model deployment version bump."""
    try:
        data = request.get_json()
        model_id = data.get("modelId", "")
        current_version = data.get("currentVersion", "v1.0")

        # Increment version: v2.0 → v2.1, v3.5 → v3.6
        parts = current_version.lstrip("v").split(".")
        major = parts[0] if parts else "1"
        minor = int(parts[1]) + 1 if len(parts) > 1 else 1
        new_version = f"v{major}.{minor}"

        return jsonify({"newVersion": new_version})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/gcs/purge-cache", methods=["POST"])
def api_gcs_purge_cache():
    """Placeholder for GCS cache purge."""
    try:
        return jsonify({"success": True})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


# ============== ElevenLabs Proxy Routes ==============

@app.route("/api/elevenlabs/voices", methods=["POST"])
def proxy_elevenlabs_voices():
    """Proxy ElevenLabs voice list through backend so API key stays server-side."""
    try:
        data = request.get_json()
        user_id = data.get("userId")
        if not user_id:
            return jsonify({"error": "userId required"}), 400
        user_doc = get_doc('users', user_id)
        if not user_doc or not user_doc.get('elevenLabsApiKey'):
            return jsonify({"voices": []}), 200
        api_key = user_doc['elevenLabsApiKey']
        resp = requests.get(
            "https://api.elevenlabs.io/v1/voices",
            headers={"xi-api-key": api_key, "Content-Type": "application/json"},
            timeout=10
        )
        if not resp.ok:
            return jsonify({"voices": []}), 200
        voices_data = resp.json()
        voices = [
            {
                "id": v.get("voice_id"),
                "name": v.get("name"),
                "category": v.get("category", "Generated"),
                "provider": "elevenlabs",
                "preview_url": v.get("preview_url")
            }
            for v in voices_data.get("voices", [])
        ]
        return jsonify({"voices": voices})
    except Exception as e:
        print(f"ElevenLabs voices proxy error: {e}")
        return jsonify({"voices": []}), 200


@app.route("/api/elevenlabs/generate", methods=["POST"])
def proxy_elevenlabs_generate():
    """Proxy ElevenLabs TTS through backend so API key stays server-side.
    Falls back to Google Cloud TTS for demo voices or when no key is configured."""
    try:
        data = request.get_json()
        user_id = data.get("userId")
        voice_id = data.get("voiceId")
        text = data.get("text", "")
        settings = data.get("settings", {})
        if not user_id or not voice_id:
            return jsonify({"error": "userId and voiceId required"}), 400

        # Strip HTML/SSML tags from text for clean TTS input
        import re, uuid
        clean_text = re.sub(r'<[^>]*>', '', text).strip()
        # Strip ElevenLabs intonation tags like [slowly], [pause], etc.
        clean_text = re.sub(r'\[(?:slowly|loudly|whisper|emotional|pause|excited|serious|laugh)\]', '', clean_text, flags=re.IGNORECASE).strip()
        if not clean_text:
            return jsonify({"error": "No text to synthesize"}), 400

        user_doc = get_doc('users', user_id)
        use_elevenlabs = (
            user_doc
            and user_doc.get('elevenLabsApiKey')
            and not voice_id.startswith('demo-')
        )

        if use_elevenlabs:
            # Use ElevenLabs API
            api_key = user_doc['elevenLabsApiKey']
            resp = requests.post(
                f"https://api.elevenlabs.io/v1/text-to-speech/{voice_id}",
                headers={"xi-api-key": api_key, "Content-Type": "application/json"},
                json={
                    "text": clean_text,
                    "model_id": "eleven_multilingual_v2",
                    "voice_settings": {
                        "stability": settings.get("stability", 0.5),
                        "similarity_boost": settings.get("similarity_boost", 0.75),
                        "style": settings.get("style", 0.0),
                        "use_speaker_boost": settings.get("use_speaker_boost", True)
                    }
                },
                timeout=30
            )
            if not resp.ok:
                error_detail = resp.json() if resp.headers.get('content-type', '').startswith('application/json') else {}
                return jsonify({"error": error_detail.get("detail", {}).get("message", "Generation failed")}), 500
            audio_content = resp.content
            content_type = "audio/mpeg"
        else:
            # Fallback: Google Cloud Text-to-Speech
            from google.cloud import texttospeech
            tts_client = texttospeech.TextToSpeechClient()

            # Map demo voice IDs to Google Cloud TTS voice names
            gcp_voice_map = {
                'demo-rachel': ('en-US-Studio-O', texttospeech.SsmlVoiceGender.FEMALE),
                'demo-drew': ('en-US-Studio-M', texttospeech.SsmlVoiceGender.MALE),
                'demo-clyde': ('en-US-Studio-Q', texttospeech.SsmlVoiceGender.MALE),
                'demo-domi': ('en-US-Neural2-F', texttospeech.SsmlVoiceGender.FEMALE),
                'demo-bella': ('en-US-Neural2-C', texttospeech.SsmlVoiceGender.FEMALE),
                'demo-antoni': ('en-US-Neural2-D', texttospeech.SsmlVoiceGender.MALE),
            }
            voice_name, gender = gcp_voice_map.get(voice_id, ('en-US-Studio-M', texttospeech.SsmlVoiceGender.MALE))

            synthesis_input = texttospeech.SynthesisInput(text=clean_text)
            voice_params = texttospeech.VoiceSelectionParams(
                language_code="en-US",
                name=voice_name,
                ssml_gender=gender
            )
            audio_config = texttospeech.AudioConfig(
                audio_encoding=texttospeech.AudioEncoding.MP3,
                speaking_rate=0.95,
                pitch=0.0,
            )
            tts_resp = tts_client.synthesize_speech(
                input=synthesis_input, voice=voice_params, audio_config=audio_config
            )
            audio_content = tts_resp.audio_content
            content_type = "audio/mpeg"

        # Return audio as base64 data URI (avoids GCS permission issues for TTS clips)
        import base64
        audio_b64 = base64.b64encode(audio_content).decode("utf-8")
        audio_url = f"data:{content_type};base64,{audio_b64}"
        return jsonify({"audioUrl": audio_url})
    except Exception as e:
        print(f"ElevenLabs/TTS generate proxy error: {e}")
        return jsonify({"error": str(e)}), 500


# ============== Proxy Endpoints (server-side API keys) ==============

# Environment variables for external API keys (set in Cloud Run env or .env)
PEXELS_API_KEY = os.environ.get("PEXELS_API_KEY", "")
PIXABAY_API_KEY = os.environ.get("PIXABAY_API_KEY", "")
HEYGEN_API_KEY = os.environ.get("HEYGEN_API_KEY", "")
RUNWAY_API_KEY = os.environ.get("RUNWAY_API_KEY", "")
CC_DESIGN_AGENT_URL = os.environ.get("CC_DESIGN_AGENT_URL", "https://www.channelchangers.app")
CC_DESIGN_AGENT_KEY = os.environ.get("CC_DESIGN_AGENT_KEY", "")


@app.route("/api/proxy/pexels/search", methods=["GET"])
def proxy_pexels_search():
    """Proxy Pexels video search — API key stays server-side."""
    if not PEXELS_API_KEY:
        return jsonify({"error": "Pexels API key not configured"}), 501
    query = request.args.get("query", "")
    page = request.args.get("page", "1")
    per_page = request.args.get("per_page", "20")
    try:
        resp = requests.get(
            f"https://api.pexels.com/videos/search?query={query}&page={page}&per_page={per_page}",
            headers={"Authorization": PEXELS_API_KEY},
            timeout=15,
        )
        return jsonify(resp.json()), resp.status_code
    except Exception as e:
        return jsonify({"error": str(e)}), 502


@app.route("/api/proxy/pixabay/search", methods=["GET"])
def proxy_pixabay_search():
    """Proxy Pixabay video search — API key stays server-side."""
    if not PIXABAY_API_KEY:
        return jsonify({"error": "Pixabay API key not configured"}), 501
    query = request.args.get("query", "")
    page = request.args.get("page", "1")
    per_page = request.args.get("per_page", "20")
    try:
        resp = requests.get(
            f"https://pixabay.com/api/videos/?key={PIXABAY_API_KEY}&q={query}&page={page}&per_page={per_page}",
            timeout=15,
        )
        return jsonify(resp.json()), resp.status_code
    except Exception as e:
        return jsonify({"error": str(e)}), 502


@app.route("/api/proxy/heygen/status", methods=["GET"])
def proxy_heygen_status():
    """Check if HeyGen is configured."""
    return jsonify({"configured": bool(HEYGEN_API_KEY)})


@app.route("/api/proxy/heygen/avatars", methods=["GET"])
def proxy_heygen_avatars():
    """List HeyGen avatars."""
    if not HEYGEN_API_KEY:
        return jsonify({"error": "HeyGen API key not configured"}), 501
    try:
        resp = requests.get(
            "https://api.heygen.com/v2/avatars",
            headers={"X-Api-Key": HEYGEN_API_KEY},
            timeout=15,
        )
        data = resp.json()
        return jsonify({"avatars": data.get("data", {}).get("avatars", [])}), resp.status_code
    except Exception as e:
        return jsonify({"error": str(e)}), 502


@app.route("/api/proxy/heygen/generate", methods=["POST"])
def proxy_heygen_generate():
    """Generate HeyGen avatar video."""
    if not HEYGEN_API_KEY:
        return jsonify({"error": "HeyGen API key not configured"}), 501
    data = request.get_json()
    avatar_id = data.get("avatarId")
    audio_url = data.get("audioUrl")
    background_url = data.get("backgroundUrl")

    body = {
        "video_inputs": [{
            "character": {"type": "avatar", "avatar_id": avatar_id, "avatar_style": "normal"},
            "voice": {"type": "audio", "audio_url": audio_url},
            **({"background": {"type": "image", "value": background_url}} if background_url
               else {"background": {"type": "color", "value": "#0a0a0a"}}),
        }],
        "dimension": {"width": 1920, "height": 1080},
    }
    try:
        resp = requests.post(
            "https://api.heygen.com/v2/video/generate",
            json=body,
            headers={"X-Api-Key": HEYGEN_API_KEY, "Content-Type": "application/json"},
            timeout=30,
        )
        result = resp.json()
        return jsonify({"video_id": result.get("data", {}).get("video_id", "")}), resp.status_code
    except Exception as e:
        return jsonify({"error": str(e)}), 502


@app.route("/api/proxy/heygen/status/<video_id>", methods=["GET"])
def proxy_heygen_video_status(video_id):
    """Poll HeyGen video status."""
    if not HEYGEN_API_KEY:
        return jsonify({"error": "HeyGen API key not configured"}), 501
    try:
        resp = requests.get(
            f"https://api.heygen.com/v1/video_status.get?video_id={video_id}",
            headers={"X-Api-Key": HEYGEN_API_KEY},
            timeout=15,
        )
        data = resp.json().get("data", {})
        return jsonify({
            "status": data.get("status", "failed"),
            "video_url": data.get("video_url"),
            "error": data.get("error", {}).get("message") if isinstance(data.get("error"), dict) else None,
            "duration": data.get("duration"),
        }), resp.status_code
    except Exception as e:
        return jsonify({"error": str(e)}), 502


@app.route("/api/proxy/runway/status", methods=["GET"])
def proxy_runway_status():
    """Check if Runway is configured."""
    return jsonify({"configured": bool(RUNWAY_API_KEY)})


@app.route("/api/proxy/runway/generate", methods=["POST"])
def proxy_runway_generate():
    """Generate Runway video."""
    if not RUNWAY_API_KEY:
        return jsonify({"error": "Runway API key not configured"}), 501
    data = request.get_json()
    body = {
        "model": "gen4_turbo",
        "promptText": data.get("prompt", ""),
        "duration": data.get("duration", 5),
        "ratio": data.get("aspectRatio", "16:9"),
        "watermark": data.get("watermark", False),
    }
    if data.get("imageUrl"):
        body["promptImage"] = data["imageUrl"]
    try:
        resp = requests.post(
            "https://api.dev.runwayml.com/v1/image_to_video",
            json=body,
            headers={
                "Authorization": f"Bearer {RUNWAY_API_KEY}",
                "X-Runway-Version": "2024-11-06",
                "Content-Type": "application/json",
            },
            timeout=30,
        )
        return jsonify(resp.json()), resp.status_code
    except Exception as e:
        return jsonify({"error": str(e)}), 502


@app.route("/api/proxy/runway/tasks/<task_id>", methods=["GET"])
def proxy_runway_task_status(task_id):
    """Poll Runway task status."""
    if not RUNWAY_API_KEY:
        return jsonify({"error": "Runway API key not configured"}), 501
    try:
        resp = requests.get(
            f"https://api.dev.runwayml.com/v1/tasks/{task_id}",
            headers={
                "Authorization": f"Bearer {RUNWAY_API_KEY}",
                "X-Runway-Version": "2024-11-06",
            },
            timeout=15,
        )
        return jsonify(resp.json()), resp.status_code
    except Exception as e:
        return jsonify({"error": str(e)}), 502


@app.route("/api/proxy/design-agent/generate", methods=["POST"])
def proxy_design_agent():
    """Proxy CC Design Agent requests — API key stays server-side."""
    if not CC_DESIGN_AGENT_KEY:
        return jsonify({"error": "CC Design Agent API key not configured"}), 501
    data = request.get_json()
    try:
        resp = requests.post(
            f"{CC_DESIGN_AGENT_URL}/api/design-agent/generate",
            json={
                "prompt": data.get("prompt", ""),
                "aspectRatio": data.get("aspect_ratio", "16:9"),
                "enhancePrompt": True,
                "count": 1,
                "saveToProject": False,
            },
            headers={
                "Authorization": f"Bearer {CC_DESIGN_AGENT_KEY}",
                "Content-Type": "application/json",
            },
            timeout=60,
        )
        return jsonify(resp.json()), resp.status_code
    except Exception as e:
        return jsonify({"error": str(e)}), 502


@app.route("/api/proxy/quickture", methods=["POST"])
def proxy_quickture():
    """Proxy QuickTure requests — URL configured server-side."""
    quickture_url = os.environ.get("QUICKTURE_API_URL", "")
    if not quickture_url:
        return jsonify({"error": "QuickTure API URL not configured"}), 501
    data = request.get_json()
    try:
        resp = requests.post(quickture_url, json=data, timeout=30)
        return jsonify(resp.json()), resp.status_code
    except Exception as e:
        return jsonify({"error": str(e)}), 502


# ============== YouTube Analyzer ==============

KEVIN_CHANNEL_URL = "https://www.youtube.com/@theabandonedproject"

# Kevin Lacey's full video catalog mapped to episode/story slots.
# youtube_url = None until the user finds and adds the actual video URL.
# channel_search_url = pre-built search URL on Kevin's channel page.
KEVIN_LACEY_CATALOG = [
    # CALIFORNIA
    {"title": "ABANDONED CALIFORNIA RAILROAD // Goat Canyon Trestle Bridge", "location_name": "Goat Canyon Trestle Bridge", "state_country": "Carrizo Gorge, California", "episode": 6, "slot": "A", "kevin_views": 31000, "youtube_url": None},
    {"title": "Trona, California Ghost Town: Exploring Abandoned Houses", "location_name": "Trona", "state_country": "San Bernardino County, California", "episode": 3, "slot": "C", "kevin_views": 26000, "youtube_url": None},
    {"title": "Amboy California: Route 66 Ghost town", "location_name": "Amboy, Route 66", "state_country": "San Bernardino County, California", "episode": 12, "slot": "B", "kevin_views": 13000, "youtube_url": None},
    {"title": "This California Neighborhood was Covered in Sand", "location_name": "Newberry Springs Buried Town", "state_country": "San Bernardino County, California", "episode": 15, "slot": "B", "kevin_views": 11000, "youtube_url": None},
    {"title": "Exploring an Abandoned Nazi Camp in LA // Murphy's Ranch", "location_name": "Murphy's Ranch (Nazi Camp)", "state_country": "Pacific Palisades, Los Angeles", "episode": 17, "slot": "B", "kevin_views": 7500, "youtube_url": None},
    {"title": "ABANDONED Pennsylvania Bayless Paper Mill and Austin Dam", "location_name": "Austin Dam Collapse", "state_country": "Austin, Pennsylvania", "episode": 1, "slot": "A", "kevin_views": 7000, "youtube_url": None},
    {"title": "Abandoned Military Base Hospital in California", "location_name": "George AFB Hospital", "state_country": "Victorville, California", "episode": 1, "slot": "C", "kevin_views": 6500, "youtube_url": None},
    {"title": "Inside Pablo Escobar Abandoned Mansion in Colombia", "location_name": "Pablo Escobar Mansion", "state_country": "Medellín, Colombia", "episode": 20, "slot": "C", "kevin_views": 5400, "youtube_url": None},
    {"title": "WE FOUND CHARLES MANSON'S TRUCK IN A CALIFORNIA GHOST TOWN", "location_name": "Ballarat Ghost Town / Barker Ranch", "state_country": "Death Valley, California", "episode": 1, "slot": "B", "kevin_views": 5500, "youtube_url": None},
    {"title": "ABANDONED MILITARY BUNKERS // Fort MacArthur", "location_name": "Fort MacArthur", "state_country": "San Pedro, California", "episode": 4, "slot": "C", "kevin_views": 4700, "youtube_url": None},
    {"title": "Exploring ATOLIA - Mojave Desert Mine", "location_name": "Atolia Mine", "state_country": "Mojave Desert, California", "episode": 9, "slot": "C", "kevin_views": 4300, "youtube_url": None},
    {"title": "Exploring an Abandoned Military Base in California", "location_name": "George AFB (Main Base)", "state_country": "Victorville, California", "episode": 1, "slot": "A", "kevin_views": 4100, "youtube_url": None},
    {"title": "Abandoned California City: Bombay Beach", "location_name": "Bombay Beach / Salton Sea", "state_country": "Imperial County, California", "episode": 3, "slot": "A", "kevin_views": 3900, "youtube_url": None},
    {"title": "Rinconada Quicksilver Mine – Massive Abandoned Site", "location_name": "Rinconada Quicksilver Mine", "state_country": "San Benito County, California", "episode": 14, "slot": "C", "kevin_views": 3000, "youtube_url": None},
    {"title": "EXPLORING UNDERGROUND- LOS ANGELES MISSILE SILO", "location_name": "LA-88 Nike Missile Silo", "state_country": "Chatsworth, Los Angeles", "episode": 4, "slot": "B", "kevin_views": 3200, "youtube_url": None},
    {"title": "ABANDONED SAN FRANCISCO - Gold Gate Battery Spencer/Wagner", "location_name": "Golden Gate Batteries (Spencer/Wagner)", "state_country": "Marin County, California", "episode": 4, "slot": "C", "kevin_views": 2400, "youtube_url": None},
    {"title": "Eagle Mountain SOLD for $22.5 MILLION", "location_name": "Eagle Mountain (Kaiser Steel Plant)", "state_country": "Riverside County, California", "episode": 8, "slot": "B", "kevin_views": 2800, "youtube_url": None},
    {"title": "Exploring San Francisco's ABANDONED Sutro Baths", "location_name": "Sutro Baths", "state_country": "San Francisco, California", "episode": 7, "slot": "A", "kevin_views": 575, "youtube_url": None},
    {"title": "Alcatraz Prison Hospital at Night", "location_name": "Alcatraz Prison Hospital", "state_country": "San Francisco Bay, California", "episode": 7, "slot": "C", "kevin_views": 928, "youtube_url": None},
    {"title": "ABANDONED ZOO - The Old Los Angeles Zoo", "location_name": "Old Los Angeles Zoo", "state_country": "Griffith Park, Los Angeles", "episode": 2, "slot": "C", "kevin_views": 809, "youtube_url": None},
    {"title": "Inside Helena Ghost Town – Gold Rush Settlement", "location_name": "Helena Ghost Town", "state_country": "Trinity County, California", "episode": 11, "slot": "C", "kevin_views": 868, "youtube_url": None},
    {"title": "ABANDONED ASYLUM - Camarillo State Mental Hospital", "location_name": "Camarillo State Hospital", "state_country": "Camarillo, California", "episode": 2, "slot": "D", "kevin_views": 1100, "youtube_url": None},
    {"title": "CALIFORNIA ABANDONED Mine and Cabin - Burro Shmidt Tunnel", "location_name": "Burro Schmidt Tunnel", "state_country": "El Paso Mountains, California", "episode": 3, "slot": "C", "kevin_views": 1100, "youtube_url": None},
    # ARIZONA / ROUTE 66
    {"title": "Two Guns Arizona and Apache Death Cave - Route 66", "location_name": "Two Guns / Apache Death Cave", "state_country": "Coconino County, Arizona", "episode": 2, "slot": "A", "kevin_views": 568, "youtube_url": None},
    {"title": "Meteor City Trading Post - Abandoned Arizona Route 66", "location_name": "Meteor City Trading Post", "state_country": "Arizona, Route 66", "episode": 2, "slot": "B", "kevin_views": 1000, "youtube_url": None},
    # UTAH
    {"title": "Abandoned Utah Pioneer homes - Paria Townsite", "location_name": "Paria Ghost Town", "state_country": "Kane County, Utah", "episode": 1, "slot": "C", "kevin_views": 1000, "youtube_url": None},
    # NEW YORK
    {"title": "ABANDONED HOSPITAL for Mentally and Physically Disabled", "location_name": "Letchworth Village", "state_country": "Thiells, New York", "episode": 2, "slot": "B", "kevin_views": 410, "youtube_url": None},
    # SCOTLAND
    {"title": "Abandoned Scottish Castle Turned Mental Asylum", "location_name": "Lennox Castle Hospital", "state_country": "Lennoxtown, Scotland", "episode": 2, "slot": "C", "kevin_views": 400, "youtube_url": None},
    {"title": "Abandoned Shipwreck in the Scottish Highlands", "location_name": "Scottish Highlands Shipwreck", "state_country": "Scotland", "episode": 24, "slot": "B", "kevin_views": 4900, "youtube_url": None},
    # INTERNATIONAL
    {"title": "Exploring an Abandoned Italian Ghost Town", "location_name": "Poggioreale (Italian Ghost Town)", "state_country": "Sicily, Italy", "episode": 22, "slot": "B", "kevin_views": 1000, "youtube_url": None},
    {"title": "ABANDONED ELEVEN STORY HOTEL IN COSTA RICA", "location_name": "Costa Rica Abandoned Hotel", "state_country": "Costa Rica", "episode": 25, "slot": "A", "kevin_views": 53000, "youtube_url": None},
    {"title": "EXPLORING ABANDONED IGLOO HOTEL - Alaska", "location_name": "Igloo Hotel (Never-Finished)", "state_country": "Alaska, USA", "episode": 23, "slot": "B", "kevin_views": 5700, "youtube_url": None},
]


@app.route("/api/projects/<project_id>/seed-kevin-catalog", methods=["POST"])
def seed_kevin_catalog(project_id):
    """
    Pre-populate Kevin's YouTube library with catalog entries (pending analysis).
    Each entry shows the video title, story context, and episode assignment.
    The user can then find the URL on Kevin's channel and add it to trigger Gemini analysis.
    Skips videos that have already been added (by title match).
    """
    # Get existing clip titles to avoid duplicates
    existing_docs = list(
        db.collection(COLLECTIONS["youtube_clips"])
        .where("project_id", "==", project_id)
        .stream()
    )
    existing_titles = {d.to_dict().get("kevin_title", "").lower() for d in existing_docs}

    created = 0
    skipped = 0

    for video in KEVIN_LACEY_CATALOG:
        if video["title"].lower() in existing_titles:
            skipped += 1
            continue

        channel_search_term = video["title"].replace(" ", "+")
        channel_search_url = f"https://www.youtube.com/@theabandonedproject/search?query={channel_search_term}"

        catalog_doc = {
            "project_id": project_id,
            "status": "pending_url",          # Waiting for user to paste the URL
            "kevin_title": video["title"],
            "location_name": video["location_name"],
            "state_country": video.get("state_country", ""),
            "location_type": "abandoned",
            "episode_assignment": video["episode"],
            "slot_assignment": video["slot"],
            "kevin_views": video["kevin_views"],
            "youtube_url": video.get("youtube_url"),  # None until user adds it
            "clip_name": f"ABDN_{str(video['episode']).zfill(2)}{video['slot']}_MAS_KevinLacy_{video['location_name'].replace(' ', '_')[:30]}",
            "channel_url": KEVIN_CHANNEL_URL,
            "channel_search_url": channel_search_url,
            "video_summary": f"Kevin Lacey explores {video['location_name']} — planned for Episode {video['episode']} ({video['slot']}-Story)",
            "story_potential": f"Episode {video['episode']} {video['slot']}-Story: {video['location_name']}",
            "abc_suitability": f"{video['slot']}-story",
            "abc_reason": f"Assigned as {video['slot']}-Story for Episode {video['episode']} of the Abandoned series",
            "presenter_on_screen": True,
            "total_shots": 0,
            "shots": [],
            "createdAt": datetime.utcnow().isoformat(),
        }
        create_doc("youtube_clips", catalog_doc)
        created += 1

    return jsonify({
        "message": "Kevin's catalog seeded",
        "created": created,
        "skipped": skipped,
        "total_catalog": len(KEVIN_LACEY_CATALOG),
        "channel_url": KEVIN_CHANNEL_URL,
    }), 201


YOUTUBE_SHOT_PROMPT = """You are an archive producer cataloguing footage for a documentary series.

Analyze this YouTube video shot by shot. For EVERY distinct shot or scene change, provide a JSON entry.

Return a JSON object with two keys:
1. "shots" — array of shot objects
2. "meta" — overall video metadata

Each shot object:
{
  "timecode": "HH:MM:SS",
  "duration_seconds": <integer>,
  "shot_type": "<wide|medium|close|cutaway|PTC|aerial|B-roll|title_card>",
  "location_description": "<physical description of the place shown>",
  "subject": "<main subject, person, or object>",
  "action": "<what is happening in this shot>",
  "camera_movement": "<static|pan|tilt|zoom|tracking|handheld|drone>",
  "usability": "<high|medium|low>",
  "tags": "<comma-separated search keywords>"
}

Meta object:
{
  "video_summary": "<2-3 sentence description of the full video>",
  "location_name": "<name of the primary location>",
  "location_type": "<factory|hospital|school|hotel|prison|military|residential|industrial|other>",
  "total_shots": <integer>,
  "presenter_on_screen": <true|false>,
  "story_potential": "<what documentary story does this footage support?>",
  "abc_suitability": "<A-story|B-story|C-story|contextual>",
  "abc_reason": "<why this footage suits that story role>"
}

Be thorough — capture every shot. Return only valid JSON, no markdown wrapping."""


def _run_youtube_analysis(url: str, project_id: str, batch_id: str = None, batch_index: int = 0):
    """Core Gemini YouTube analysis — runs synchronously, safe to call in thread."""
    try:
        video_part = Part.from_uri(uri=url, mime_type="video/mp4")
        response = model.generate_content(
            [video_part, YOUTUBE_SHOT_PROMPT],
            generation_config={"max_output_tokens": 8192, "temperature": 0.1}
        )
        raw = clean_ai_response(response.text)
        data = json.loads(raw)
    except json.JSONDecodeError:
        data = {"shots": [], "meta": {"video_summary": "Parse error", "location_name": "Unknown",
                                       "location_type": "other", "total_shots": 0,
                                       "presenter_on_screen": False, "story_potential": "",
                                       "abc_suitability": "B-story", "abc_reason": ""}}
    except Exception as e:
        data = {"shots": [], "meta": {"video_summary": f"Analysis failed: {str(e)}",
                                       "location_name": "Unknown", "location_type": "other",
                                       "total_shots": 0, "presenter_on_screen": False,
                                       "story_potential": "", "abc_suitability": "B-story", "abc_reason": ""}}

    shots = data.get("shots", [])
    meta = data.get("meta", {})

    # Derive clip name from location
    location = meta.get("location_name", "Unknown Location")
    clip_name = f"ABDN_{str(batch_index + 1).zfill(4)}_MAS_KevinLacy_{location.replace(' ', '_')}"

    # Save to Firestore
    clip_doc = {
        "projectId": project_id,
        "youtube_url": url,
        "clip_name": clip_name,
        "location_name": location,
        "location_type": meta.get("location_type", "other"),
        "video_summary": meta.get("video_summary", ""),
        "story_potential": meta.get("story_potential", ""),
        "abc_suitability": meta.get("abc_suitability", "B-story"),
        "abc_reason": meta.get("abc_reason", ""),
        "presenter_on_screen": meta.get("presenter_on_screen", False),
        "total_shots": len(shots),
        "shots": shots,
        "status": "analyzed",
        "batch_id": batch_id,
    }
    saved = create_doc("youtube_clips", clip_doc)

    # Update batch progress if part of a batch
    if batch_id:
        batch_ref = db.collection(COLLECTIONS["youtube_batches"]).document(batch_id)
        batch_doc = batch_ref.get()
        if batch_doc.exists:
            bdata = batch_doc.to_dict()
            completed = bdata.get("completed", 0) + 1
            total = bdata.get("total", 1)
            batch_ref.update({
                "completed": completed,
                "status": "complete" if completed >= total else "running",
                "updatedAt": datetime.utcnow().isoformat(),
            })

    return saved


@app.route("/api/youtube/analyze", methods=["POST"])
def youtube_analyze_single():
    """Analyze a single YouTube video — shot by shot — and save to Firestore."""
    data = request.get_json()
    url = data.get("url", "").strip()
    project_id = data.get("projectId", "")
    if not url or not project_id:
        return jsonify({"error": "url and projectId required"}), 400

    try:
        result = _run_youtube_analysis(url, project_id)
        return jsonify(result), 200
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/youtube/batch", methods=["POST"])
def youtube_batch_analyze():
    """Kick off batch analysis of multiple YouTube URLs in a background thread."""
    data = request.get_json()
    urls = data.get("urls", [])
    project_id = data.get("projectId", "")
    if not urls or not project_id:
        return jsonify({"error": "urls and projectId required"}), 400

    # Create batch record
    batch_doc = create_doc("youtube_batches", {
        "projectId": project_id,
        "urls": urls,
        "total": len(urls),
        "completed": 0,
        "status": "running",
    })
    batch_id = batch_doc["id"]

    def run_batch():
        for i, url in enumerate(urls):
            try:
                _run_youtube_analysis(url.strip(), project_id, batch_id=batch_id, batch_index=i)
            except Exception as e:
                print(f"[youtube-batch] Failed {url}: {e}")

    thread = threading.Thread(target=run_batch, daemon=True)
    thread.start()

    return jsonify({"batchId": batch_id, "total": len(urls), "status": "running"}), 202


@app.route("/api/youtube/batch/<batch_id>", methods=["GET"])
def youtube_batch_status(batch_id):
    """Get batch analysis progress."""
    doc = get_doc("youtube_batches", batch_id)
    if not doc:
        return jsonify({"error": "Batch not found"}), 404
    return jsonify(doc), 200


@app.route("/api/projects/<project_id>/youtube-clips", methods=["GET"])
def get_youtube_clips(project_id):
    """Get all analyzed YouTube clips for a project."""
    docs = db.collection(COLLECTIONS["youtube_clips"]) \
              .where("projectId", "==", project_id) \
              .order_by("createdAt") \
              .stream()
    return jsonify([{**d.to_dict(), "id": d.id} for d in docs]), 200


@app.route("/api/youtube/score-episodes", methods=["POST"])
def youtube_score_episodes():
    """
    Second-pass analysis: given all analyzed clips for a project,
    score each location's A/B/C episode strength and suggest episode structures.
    """
    data = request.get_json()
    project_id = data.get("projectId", "")
    if not project_id:
        return jsonify({"error": "projectId required"}), 400

    # Load all analyzed clips
    clips_docs = db.collection(COLLECTIONS["youtube_clips"]) \
                   .where("projectId", "==", project_id) \
                   .where("status", "==", "analyzed") \
                   .stream()
    clips = [{**d.to_dict(), "id": d.id} for d in clips_docs]

    if not clips:
        return jsonify({"error": "No analyzed clips found. Run YouTube analysis first."}), 400

    # Build a summary of all footage for Gemini
    footage_summary = "\n\n".join([
        f"VIDEO {i+1}: {c.get('location_name', 'Unknown')} ({c.get('location_type', '')})\n"
        f"URL: {c.get('youtube_url', '')}\n"
        f"Summary: {c.get('video_summary', '')}\n"
        f"Story potential: {c.get('story_potential', '')}\n"
        f"Shots: {c.get('total_shots', 0)} | Presenter on screen: {c.get('presenter_on_screen', False)}\n"
        f"Initial ABC assessment: {c.get('abc_suitability', '')} — {c.get('abc_reason', '')}"
        for i, c in enumerate(clips)
    ])

    scoring_prompt = f"""You are a senior documentary series producer for "What the Hell Happened Here!?" —
a 30-episode presenter-led factual entertainment series about abandoned locations.

Each episode has an A/B/C story structure:
- A Story: Primary location mystery — must carry the full episode (strongest narrative, 25+ minutes of footage needed)
- B Story: Secondary thread from same location or era (supporting, 10-15 mins)
- C Story: Broader historical/contextual thread (adds depth, can use external archive)

Here is the footage that exists across {len(clips)} YouTube videos from the presenter Kevin Lacey:

{footage_summary}

Based on this footage inventory, provide:

1. EPISODE RECOMMENDATIONS — suggest which locations have enough for a full A-story episode (rank top 10)
2. B-STORY PAIRINGS — which locations work well as B-stories alongside stronger A-stories
3. WEAK LOCATIONS — locations with insufficient footage for even a B-story
4. MISSING GAPS — what types of locations are absent that would round out a 30-episode series

Return as JSON:
{{
  "episode_recommendations": [
    {{
      "rank": 1,
      "location_name": "...",
      "youtube_url": "...",
      "episode_title_suggestion": "...",
      "a_story": "...",
      "b_story_suggestion": "...",
      "c_story_suggestion": "...",
      "footage_strength": "strong|moderate|thin",
      "reason": "..."
    }}
  ],
  "b_story_only": [
    {{"location_name": "...", "youtube_url": "...", "best_paired_with": "...", "reason": "..."}}
  ],
  "weak_locations": [
    {{"location_name": "...", "youtube_url": "...", "issue": "..."}}
  ],
  "series_gaps": ["...", "..."]
}}"""

    try:
        response = model.generate_content(
            scoring_prompt,
            generation_config={"max_output_tokens": 8192, "temperature": 0.2}
        )
        raw = clean_ai_response(response.text)
        result = json.loads(raw)
        return jsonify(result), 200
    except json.JSONDecodeError as e:
        return jsonify({"error": f"Failed to parse AI response: {str(e)}", "raw": response.text[:500]}), 500
    except Exception as e:
        return jsonify({"error": str(e)}), 500


# ============== Social Media Importer (yt-dlp + GCS + Gemini) ==============

def _detect_platform(url: str) -> str:
    u = url.lower()
    if "instagram.com" in u: return "instagram"
    if "tiktok.com" in u: return "tiktok"
    if "twitter.com" in u or "x.com" in u: return "twitter"
    if "facebook.com" in u or "fb.watch" in u: return "facebook"
    if "vimeo.com" in u: return "vimeo"
    if "youtube.com" in u or "youtu.be" in u: return "youtube"
    return "unknown"


def _ytdlp_info(url: str) -> dict:
    """Get video metadata without downloading."""
    try:
        import yt_dlp
        with yt_dlp.YoutubeDL({"quiet": True, "no_warnings": True, "skip_download": True}) as ydl:
            return ydl.extract_info(url, download=False) or {}
    except Exception as e:
        raise Exception(f"Could not retrieve info: {str(e)}")


def _download_and_upload_to_gcs(url: str, project_id: str, platform: str):
    """Download video with yt-dlp, upload to GCS, return (gcs_uri, info_dict). Cleans up temp files."""
    import yt_dlp, tempfile, os, uuid, shutil
    tmp_dir = tempfile.mkdtemp()
    try:
        ydl_opts = {
            "format": "best[ext=mp4][filesize<200M]/best[ext=mp4]/best[filesize<200M]/best",
            "outtmpl": os.path.join(tmp_dir, "%(id)s.%(ext)s"),
            "quiet": True, "no_warnings": True,
            "merge_output_format": "mp4",
        }
        with yt_dlp.YoutubeDL(ydl_opts) as ydl:
            info = ydl.extract_info(url, download=True)
            local_path = ydl.prepare_filename(info)

        # yt-dlp sometimes changes ext during merge — find the actual file
        if not os.path.exists(local_path):
            files = os.listdir(tmp_dir)
            if not files:
                raise Exception("yt-dlp produced no output file")
            local_path = os.path.join(tmp_dir, sorted(files)[0])

        size_mb = os.path.getsize(local_path) / (1024 * 1024)
        if size_mb > 250:
            raise Exception(f"Video too large ({size_mb:.0f}MB — max 200MB). Try a shorter clip.")

        ext = os.path.splitext(local_path)[1] or ".mp4"
        gcs_path = f"social-media/{project_id}/{platform}/{uuid.uuid4().hex}{ext}"
        blob = storage_client.bucket(STORAGE_BUCKET).blob(gcs_path)
        blob.upload_from_filename(local_path, content_type="video/mp4")
        return f"gs://{STORAGE_BUCKET}/{gcs_path}", info or {}
    finally:
        shutil.rmtree(tmp_dir, ignore_errors=True)


def _run_social_analysis(url: str, project_id: str, platform: str, catalog_doc_id: str = None) -> dict:
    """Full pipeline: download → GCS → Gemini analysis → Firestore. YouTube delegates to existing function."""
    if platform == "youtube":
        return _run_youtube_analysis(url, project_id)

    gcs_uri, info = _download_and_upload_to_gcs(url, project_id, platform)

    try:
        video_part = Part.from_uri(uri=gcs_uri, mime_type="video/mp4")
        response = model.generate_content(
            [video_part, YOUTUBE_SHOT_PROMPT],
            generation_config={"max_output_tokens": 8192, "temperature": 0.1}
        )
        data = json.loads(clean_ai_response(response.text))
    except Exception as e:
        data = {"shots": [], "meta": {"video_summary": f"Analysis failed: {e}",
            "location_name": info.get("title", "Unknown"), "location_type": "other",
            "total_shots": 0, "presenter_on_screen": False,
            "story_potential": "", "abc_suitability": "B-story", "abc_reason": ""}}

    shots = data.get("shots", [])
    meta = data.get("meta", {})
    location = meta.get("location_name", info.get("title", "Unknown")[:50])

    PLATFORM_ICONS = {"instagram": "📸", "tiktok": "🎵", "twitter": "🐦", "facebook": "👤", "vimeo": "🎬"}
    icon = PLATFORM_ICONS.get(platform, "📹")

    clip_doc = {
        "projectId": project_id,
        "youtube_url": url,
        "gcs_uri": gcs_uri,
        "source_platform": platform,
        "clip_name": f"{icon} [{platform.upper()}] {location[:40]}",
        "location_name": location,
        "location_type": meta.get("location_type", "other"),
        "video_summary": meta.get("video_summary", (info.get("description") or "")[:300]),
        "story_potential": meta.get("story_potential", ""),
        "abc_suitability": meta.get("abc_suitability", "B-story"),
        "abc_reason": meta.get("abc_reason", ""),
        "presenter_on_screen": meta.get("presenter_on_screen", False),
        "total_shots": len(shots),
        "shots": shots,
        "status": "analyzed",
        "social_title": info.get("title", ""),
        "social_uploader": info.get("uploader", ""),
        "social_duration": info.get("duration", 0),
        "social_view_count": info.get("view_count", 0),
        "social_thumbnail": info.get("thumbnail", ""),
        "createdAt": datetime.utcnow().isoformat(),
    }

    if catalog_doc_id:
        db.collection(COLLECTIONS["youtube_clips"]).document(catalog_doc_id).update(clip_doc)
        return {**clip_doc, "id": catalog_doc_id}

    return create_doc("youtube_clips", clip_doc)


@app.route("/api/social/analyze", methods=["POST"])
def social_analyze():
    """Analyze a video from any social platform. yt-dlp downloads it, GCS stores it, Gemini analyses it."""
    data = request.get_json()
    url = data.get("url", "").strip()
    project_id = data.get("projectId", "")
    catalog_doc_id = data.get("catalogDocId")
    if not url or not project_id:
        return jsonify({"error": "url and projectId required"}), 400

    platform = _detect_platform(url)
    try:
        result = _run_social_analysis(url, project_id, platform, catalog_doc_id)
        return jsonify({**result, "platform": platform}), 200
    except Exception as e:
        err = str(e)
        if "instagram" in platform and ("login" in err.lower() or "403" in err or "private" in err.lower()):
            hint = "Instagram requires browser cookies for download. Export cookies from a logged-in browser session and add them to the server."
        elif "login" in err.lower() or "auth" in err.lower():
            hint = "This content requires authentication — make sure it's a public post."
        elif "too large" in err.lower():
            hint = err
        else:
            hint = "Check the URL is correct and the post is publicly visible."
        return jsonify({"error": err, "hint": hint, "platform": platform}), 500


@app.route("/api/social/info", methods=["POST"])
def social_get_info():
    """Get metadata for a social media URL without downloading (fast preview)."""
    data = request.get_json()
    url = data.get("url", "").strip()
    if not url:
        return jsonify({"error": "url required"}), 400
    platform = _detect_platform(url)
    try:
        info = _ytdlp_info(url)
        return jsonify({
            "platform": platform, "title": info.get("title", ""),
            "uploader": info.get("uploader", ""), "duration": info.get("duration", 0),
            "view_count": info.get("view_count", 0), "thumbnail": info.get("thumbnail", ""),
            "description": (info.get("description") or "")[:400],
            "upload_date": info.get("upload_date", ""), "url": url,
        }), 200
    except Exception as e:
        return jsonify({"error": str(e), "platform": platform}), 500


# ============== Creator Profile Bulk Scraper ==========================================

# Kevin's known profiles — used as defaults when no URL is supplied
KEVIN_PLATFORM_PROFILES = {
    "instagram": "https://www.instagram.com/theabandonedproject/",
    "tiktok": "https://www.tiktok.com/@theabandonedproject",
    "youtube": "https://www.youtube.com/@theabandonedproject/videos",
}


def _scrape_profile_metadata(profile_url: str, max_items: int = 300) -> list:
    """
    Use yt-dlp flat-playlist extraction to list all videos on a creator's profile page.
    Returns a list of entry dicts with id, title, url, thumbnail, duration, upload_date.
    Does NOT download any video files.
    """
    import yt_dlp
    ydl_opts = {
        "extract_flat": "in_playlist",
        "quiet": True,
        "no_warnings": True,
        "ignoreerrors": True,
        "playlist_items": f"1:{max_items}",
    }
    try:
        with yt_dlp.YoutubeDL(ydl_opts) as ydl:
            info = ydl.extract_info(profile_url, download=False) or {}
        raw = info.get("entries", [])
        return [e for e in raw if e and (e.get("url") or e.get("webpage_url") or e.get("id"))]
    except Exception as e:
        raise Exception(f"Profile scrape failed: {str(e)}")


def _entry_to_url(entry: dict, platform: str) -> str:
    """Reconstruct a full video URL from a flat-playlist entry."""
    direct = entry.get("webpage_url") or entry.get("url", "")
    if direct.startswith("http"):
        return direct
    vid_id = entry.get("id", "")
    if platform == "tiktok":
        return f"https://www.tiktok.com/@theabandonedproject/video/{vid_id}"
    if platform == "instagram":
        return f"https://www.instagram.com/p/{vid_id}/"
    if platform == "youtube":
        return f"https://www.youtube.com/watch?v={vid_id}"
    return direct


def _process_creator_batch(batch_ref, entries: list, project_id: str, platform: str):
    """
    Background thread: for each scraped entry, download (if needed) → GCS → Gemini analysis → Firestore.
    YouTube videos are analysed via URL directly (no download). Instagram/TikTok use yt-dlp + GCS.
    Updates the batch Firestore doc as it goes so the frontend can poll progress.
    """
    from google.cloud import firestore as fs
    completed = 0
    failed = 0

    for entry in entries:
        video_url = _entry_to_url(entry, platform)
        title = (entry.get("title") or f"{platform}-{entry.get('id', '?')}")[:80]

        if not video_url:
            failed += 1
            batch_ref.update({"failed": fs.Increment(1)})
            continue

        try:
            if platform == "youtube":
                _run_youtube_analysis(video_url, project_id)
            else:
                _run_social_analysis(video_url, project_id, platform)

            completed += 1
            batch_ref.update({
                "completed": fs.Increment(1),
                "last_title": title,
            })
        except Exception as e:
            failed += 1
            batch_ref.update({
                "failed": fs.Increment(1),
                "last_error": f"{title}: {str(e)[:120]}",
            })

    batch_ref.update({"status": "complete"})


@app.route("/api/creator/scrape", methods=["POST"])
def creator_scrape():
    """
    Bulk-scrape and analyse an entire creator profile.

    Phase 1 (synchronous, fast): yt-dlp flat-playlist extraction → get all video URLs + titles.
    Phase 2 (background thread): download each video (low-res) → upload to GCS → Gemini shot analysis → Firestore.

    Returns batchId immediately for the frontend to poll.
    """
    data = request.get_json()
    platform = (data.get("platform") or "").lower().strip()
    project_id = data.get("projectId", "").strip()
    custom_url = (data.get("url") or "").strip()

    if not platform or not project_id:
        return jsonify({"error": "platform and projectId required"}), 400

    profile_url = custom_url or KEVIN_PLATFORM_PROFILES.get(platform)
    if not profile_url:
        return jsonify({"error": f"Unknown platform '{platform}'. Supported: instagram, tiktok, youtube"}), 400

    # Phase 1 — metadata only (synchronous)
    try:
        entries = _scrape_profile_metadata(profile_url)
    except Exception as e:
        err = str(e)
        if "login" in err.lower() or "private" in err.lower() or "403" in err:
            return jsonify({
                "error": "Instagram requires browser cookies to scrape profiles.",
                "hint": "TikTok and YouTube work without auth. For Instagram, export cookies from a logged-in browser session.",
                "platform": platform,
            }), 403
        return jsonify({"error": err, "platform": platform}), 500

    if not entries:
        return jsonify({
            "error": f"No videos found at {profile_url}.",
            "hint": "The profile may be private, the account may not exist, or yt-dlp may need updating.",
            "platform": platform,
        }), 404

    # Create a Firestore batch doc so the frontend can track progress
    batch_ref = db.collection(f"{COLLECTION_PREFIX}doc_creator_batches").document()
    batch_ref.set({
        "projectId": project_id,
        "platform": platform,
        "profileUrl": profile_url,
        "total": len(entries),
        "completed": 0,
        "failed": 0,
        "status": "running",
        "last_title": "",
        "last_error": "",
        "createdAt": datetime.utcnow().isoformat(),
    })

    # Phase 2 — background download + analysis
    thread = threading.Thread(
        target=_process_creator_batch,
        args=(batch_ref, entries, project_id, platform),
        daemon=True,
    )
    thread.start()

    return jsonify({
        "batchId": batch_ref.id,
        "platform": platform,
        "total": len(entries),
        "profileUrl": profile_url,
    }), 202


@app.route("/api/creator/scrape-status/<batch_id>", methods=["GET"])
def creator_scrape_status(batch_id):
    """Poll the progress of a creator profile scrape + analysis batch."""
    doc = db.collection(f"{COLLECTION_PREFIX}doc_creator_batches").document(batch_id).get()
    if not doc.exists:
        return jsonify({"error": "Batch not found"}), 404
    return jsonify({**doc.to_dict(), "id": doc.id}), 200


# ============== Templates (Firestore-backed) ==============

@app.route("/api/templates", methods=["GET"])
def get_templates():
    """Get all planning templates."""
    docs = db.collection(f'{COLLECTION_PREFIX}doc_templates').order_by('created_at').stream()
    return jsonify([{**doc.to_dict(), 'id': doc.id} for doc in docs])


@app.route("/api/templates", methods=["POST"])
def create_template():
    """Create a planning template."""
    data = request.get_json()
    data['created_at'] = datetime.utcnow().isoformat()
    doc_ref = db.collection(f'{COLLECTION_PREFIX}doc_templates').document()
    doc_ref.set(data)
    return jsonify({**data, 'id': doc_ref.id}), 201


# ============== Episode Planner — Seed Data ==============

ABANDONED_EPISODES_SEED = [
    {
        "episode_number": 1, "title": "The Dam That Could Not Break",
        "theme": "Hubris, greed, and the lies powerful men tell",
        "arc": "Three places destroyed by human greed — cost-cutting, gold fever, and hubris. Each got a strange second life: memorial park, Manson pilgrimage, movie set.",
        "score_visual": 8, "score_history": 10, "score_emotion": 9, "score_wtf": 9, "score_total": 36, "locked": True,
        "stories": [
            {"slot": "A", "location_name": "Austin Dam Collapse", "state_country": "Austin, Pennsylvania", "kevin_video_title": "ABANDONED Pennsylvania Bayless Paper Mill and Austin Dam", "kevin_views": 7000, "has_kevin_footage": True, "hook": "A millionaire built a dam 10 feet thinner than designed to save money, told the town it 'could not break,' and when it killed 78 people, he rebuilt it somewhere else. The town brothel madam ran downhill screaming warnings while the rich men ran the other way."},
            {"slot": "B", "location_name": "Ballarat Ghost Town / Manson Intro", "state_country": "Death Valley, California", "kevin_video_title": "WE FOUND CHARLES MANSON'S TRUCK IN A CALIFORNIA GHOST TOWN", "kevin_views": 5500, "has_kevin_footage": True, "hook": "A gold rush ghost town where the last resident didn't bathe for 20 years, Peter Fonda threw away his Rolex for Easy Rider, and Charles Manson hid from the law. Tex Watson's getaway truck is STILL parked there."},
            {"slot": "C", "location_name": "Paria Ghost Town", "state_country": "Kane County, Utah", "kevin_video_title": "Abandoned Utah Pioneer homes – Paria Townsite", "kevin_views": 1000, "has_kevin_footage": True, "hook": "A Mormon town founded by a mass murderer fleeing justice was destroyed by floods, then rebuilt as a Hollywood movie set. Tourists visit the FAKE town and ignore the real one."},
        ],
    },
    {
        "episode_number": 2, "title": "Death Cave",
        "theme": "Murder, myth-making, and the American roadside",
        "arc": "Route 66 promised the American Dream. Two Guns delivered mass murder and con artistry. Meteor City delivered loneliness. The Old Zoo delivered cruelty.",
        "score_visual": 9, "score_history": 8, "score_emotion": 8, "score_wtf": 10, "score_total": 35, "locked": True,
        "stories": [
            {"slot": "A", "location_name": "Two Guns / Apache Death Cave", "state_country": "Coconino County, Arizona", "kevin_video_title": "Two Guns Arizona and Apache Death Cave – Route 66", "kevin_views": 568, "has_kevin_footage": True, "hook": "A white con artist pretended to be an Apache chief, sold skulls of 42 massacred Apache as souvenirs, murdered his business partner, was acquitted — and every subsequent owner died broke or by suicide."},
            {"slot": "B", "location_name": "Meteor City Trading Post", "state_country": "Arizona", "kevin_video_title": "Meteor City Trading Post – Abandoned Arizona Route 66", "kevin_views": 1000, "has_kevin_footage": True, "hook": "'Population: 1.' A man so lonely in the desert he put up a highway sign as a cry for help. Tourists came. Then the highway moved and he was alone again."},
            {"slot": "C", "location_name": "Old Los Angeles Zoo", "state_country": "Griffith Park, Los Angeles", "kevin_video_title": "ABANDONED ZOO – The Old Los Angeles Zoo", "kevin_views": 809, "has_kevin_footage": True, "hook": "LA's original zoo was so cruel to its animals they had to build a new one. The cages are still there — empty, overgrown. People have birthday parties in them now."},
        ],
    },
    {
        "episode_number": 3, "title": "Toxic Paradise",
        "theme": "Man-made environmental disasters and institutional betrayal",
        "arc": "Three paradises poisoned by the people who built them — business, military, industry. The water will kill you in all three.",
        "score_visual": 10, "score_history": 9, "score_emotion": 9, "score_wtf": 9, "score_total": 37, "locked": True,
        "stories": [
            {"slot": "A", "location_name": "Bombay Beach / Salton Sea", "state_country": "Imperial County, California", "kevin_video_title": "Abandoned California City: Bombay Beach", "kevin_views": 3900, "has_kevin_footage": True, "hook": "California's largest lake was created BY ACCIDENT when an irrigation canal broke. Frank Sinatra water-skied here. Now the beach is made of fish bones, the air gives children asthma, and California is letting it die so San Diego can water its lawns."},
            {"slot": "B", "location_name": "George AFB — Toxic Angle", "state_country": "Victorville, California", "kevin_video_title": "Abandoned Military Base Hospital in California", "kevin_views": 10600, "has_kevin_footage": True, "hook": "Women stationed at George AFB were told 'don't get pregnant' — because the water was so poisoned that babies were born with birth defects and dying. Seven infants dead in one year. The Air Force told Congress there were 'no completed exposure pathways.' They lied."},
            {"slot": "C", "location_name": "Trona", "state_country": "San Bernardino County, California", "kevin_video_title": "Trona, California Ghost Town", "kevin_views": 26000, "has_kevin_footage": True, "hook": "A town built to make gunpowder, owned by Kerr-McGee (the Karen Silkwood company), who tried to evict the entire population when the wars ended."},
        ],
    },
    {
        "episode_number": 4, "title": "The Base They Don't Talk About",
        "theme": "Military infrastructure that protected America — and destroyed Americans",
        "arc": "Built to protect, designed to destroy. George AFB poisoned its own. Nike missiles aimed nukes at LA. Fort McArthur imprisoned Americans. The biggest threats weren't foreign — they were us.",
        "score_visual": 8, "score_history": 10, "score_emotion": 10, "score_wtf": 10, "score_total": 38, "locked": True,
        "stories": [
            {"slot": "A", "location_name": "George AFB — Military / Nuclear Angle", "state_country": "Victorville, California", "kevin_video_title": "Exploring an Abandoned Military Base in California", "kevin_views": 10600, "has_kevin_footage": True, "hook": "They flew fighter jets through nuclear mushroom clouds, brought them back to George AFB for decontamination, and let the radioactive wash water drain into the ground where military families lived. Children played on contaminated playgrounds. The Air Force knew. They said nothing."},
            {"slot": "B", "location_name": "LA-88 Nike Missile Silo", "state_country": "Chatsworth, Los Angeles", "kevin_video_title": "EXPLORING UNDERGROUND – LOS ANGELES MISSILE SILO", "kevin_views": 3200, "has_kevin_footage": True, "hook": "For 16 years, there were nuclear warheads in the hills above suburban LA. The 'defence plan' was to detonate atomic bombs in the sky over Los Angeles. The neighbours had no idea."},
            {"slot": "C", "location_name": "Fort McArthur + Golden Gate Batteries", "state_country": "San Pedro & Marin, California", "kevin_video_title": "ABANDONED MILITARY BUNKERS // Fort MacArthur", "kevin_views": 7100, "has_kevin_footage": True, "hook": "Massive coastal guns that never fired a shot — built to stop an invasion that never came. Fort McArthur's darkest chapter: processing Japanese Americans for internment."},
        ],
    },
    {
        "episode_number": 5, "title": "Drowned in Molasses",
        "theme": "Industrial catastrophe and forbidden islands",
        "arc": "Industrial disasters that nobody learned from. The Molasses Flood (corporate negligence), the Slocum disaster (1,021 dead from rotten safety equipment), and Letchworth (institutional violence).",
        "score_visual": 8, "score_history": 10, "score_emotion": 9, "score_wtf": 9, "score_total": 36, "locked": True,
        "stories": [
            {"slot": "A", "location_name": "Great Molasses Flood + North Brother Island", "state_country": "Boston, MA / New York, NY", "kevin_video_title": None, "kevin_views": 0, "has_kevin_footage": False, "hook": "A 25-foot wall of molasses at 35mph killed 21 people in downtown Boston. It moved faster than you can run. The island where they treated survivors is STILL abandoned — hiding the ghosts of 1,021 people who burned alive and Typhoid Mary, imprisoned there for 23 years."},
            {"slot": "B", "location_name": "Letchworth Village", "state_country": "Thiells, New York", "kevin_video_title": "ABANDONED HOSPITAL for Mentally and Physically Disabled", "kevin_views": 410, "has_kevin_footage": True, "hook": "Disabled children were used as guinea pigs for the polio vaccine — without consent. When Geraldo Rivera filmed conditions in 1972, America vomited. 900 graves marked with numbers, not names."},
            {"slot": "C", "location_name": "Helena Ghost Town", "state_country": "Trinity County, California", "kevin_video_title": "Inside Helena Ghost Town – Gold Rush Settlement", "kevin_views": 868, "has_kevin_footage": True, "hook": "They called it 'Baghdad' — the wildest gold rush town on the Trinity River, built on a destroyed Native American village. Now you can buy it for the price of a used car."},
        ],
    },
    {
        "episode_number": 6, "title": "The Impossible Railroad",
        "theme": "Engineering marvels defeated by nature",
        "arc": "Three California desert stories about human ambition vs natural forces — the impossible railroad, the sand that buried a town, the highway that killed a town when it moved.",
        "score_visual": 10, "score_history": 7, "score_emotion": 8, "score_wtf": 9, "score_total": 34, "locked": False,
        "stories": [
            {"slot": "A", "location_name": "Goat Canyon Trestle Bridge", "state_country": "Carrizo Gorge, San Diego County, CA", "kevin_video_title": "ABANDONED CALIFORNIA RAILROAD // Goat Canyon Trestle Bridge", "kevin_views": 31000, "has_kevin_footage": True, "hook": "They spent 12 years building a railroad everyone said was impossible. When the tunnels started collapsing, they built the world's largest wooden bridge — 200 feet high, made of WOOD, in 120°F desert heat. It's still standing. No train has crossed it in 40 years."},
            {"slot": "B", "location_name": "Newberry Springs Buried Town", "state_country": "San Bernardino County, California", "kevin_video_title": "This California Neighborhood was Covered in Sand", "kevin_views": 11000, "has_kevin_footage": True, "hook": "Road crews dumped sand on the wrong side of the highway for 40 years. Now entire houses are buried under sand dunes — rooftops and chimneys poking through. Nobody was ever held responsible."},
            {"slot": "C", "location_name": "Amboy Route 66", "state_country": "San Bernardino County, California", "kevin_video_title": "Amboy California: Route 66 Ghost town", "kevin_views": 13000, "has_kevin_footage": True, "hook": "In 1940, Amboy had 65 people and a future. Today it has zero people and a chicken magnate's son trying to keep the lights on at a gas station nobody needs. He bought the entire town for $425,000."},
        ],
    },
    {
        "episode_number": 7, "title": "The Asylum",
        "theme": "Institutional horror — locking away the unwanted",
        "arc": "Three institutions, three ways society locks away people it doesn't want. An asylum, a castle-hospital, a prison. Different purposes, same cruelty.",
        "score_visual": 8, "score_history": 9, "score_emotion": 10, "score_wtf": 9, "score_total": 36, "locked": False,
        "stories": [
            {"slot": "A", "location_name": "Northern State Hospital", "state_country": "Sedro-Woolley, Washington", "kevin_video_title": "ABANDONED STATE HOSPITAL That you can LEGALLY EXPLORE", "kevin_views": 908, "has_kevin_footage": True, "hook": "Patients went from 100 to over 1,000 in a decade. 'Treatments' included lobotomies without anaesthesia and forced labour. 1,500 patients buried in numbered graves. 200 cans of cremated remains found forgotten in a storage room a decade after closure."},
            {"slot": "B", "location_name": "Lennox Castle Asylum", "state_country": "Lennoxtown, Scotland", "kevin_video_title": "Abandoned Scottish Castle Turned Mental Asylum", "kevin_views": 400, "has_kevin_footage": True, "hook": "A beautiful Scottish castle turned into a mental hospital for 75 years. Hughie McIntyre was abandoned there as a teenager and beaten for 16 years. 'I didn't know why I was there.' He returned to the ruins in 2021 and broke down on camera."},
            {"slot": "C", "location_name": "Alcatraz Prison Hospital", "state_country": "San Francisco Bay, California", "kevin_video_title": "Alcatraz Prison Hospital at Night", "kevin_views": 928, "has_kevin_footage": True, "hook": "In 2013, the SFPD received a letter: 'My name is John Anglin. I escaped from Alcatraz in June 1962. I'm 83 years old and in bad shape.' The FBI still hasn't proved it was real. They still haven't proved it wasn't."},
        ],
    },
    {
        "episode_number": 8, "title": "Swim With the Dead",
        "theme": "Cursed ground and San Francisco's dark secrets",
        "arc": "San Francisco coast — supernatural, toxic, military. All within easy driving distance. Different moods, one haunted region.",
        "score_visual": 9, "score_history": 7, "score_emotion": 8, "score_wtf": 9, "score_total": 33, "locked": False,
        "stories": [
            {"slot": "A", "location_name": "Sutro Baths", "state_country": "Lands End, San Francisco", "kevin_video_title": "Exploring San Francisco's ABANDONED Sutro Baths", "kevin_views": 575, "has_kevin_footage": True, "hook": "The world's largest swimming pool was built on a Native American burial site. Inside: Egyptian mummies, 3,500-year-old mummified heads. Four fires destroyed four structures on this cliff. Someone moved the mummies out before the last fire. Insurance fraud or curse?"},
            {"slot": "B", "location_name": "Rinconada Quicksilver Mine", "state_country": "San Luis Obispo County, CA", "kevin_video_title": "Rinconada Quicksilver Mine – Massive Abandoned Site", "kevin_views": 3000, "has_kevin_footage": True, "hook": "They mined mercury here to extract gold. Mercury is one of the most toxic substances on Earth — 'mad as a hatter' wasn't a joke. The mines are now a public hiking park. Through contaminated ground."},
            {"slot": "C", "location_name": "Golden Gate Batteries", "state_country": "Marin Headlands, San Francisco", "kevin_video_title": "ABANDONED SAN FRANCISCO – Gold Gate Battery Spencer/Wagner", "kevin_views": 2400, "has_kevin_footage": True, "hook": "Millions of tourists take selfies at Battery Spencer — standing in a gun emplacement designed to sink battleships. The Golden Gate Bridge is directly in the line of fire."},
        ],
    },
    {
        "episode_number": 9, "title": "Blood in the Coal",
        "theme": "Labour wars, underground fire, and forgotten people",
        "arc": "Pennsylvania coal country binds the A and B stories — hangings, fire, floods. Dunalastair adds wartime refugees. Theme: forgotten people.",
        "score_visual": 8, "score_history": 10, "score_emotion": 9, "score_wtf": 8, "score_total": 35, "locked": False,
        "stories": [
            {"slot": "A", "location_name": "Molly Maguires & Centralia", "state_country": "Pennsylvania Coal Region", "kevin_video_title": None, "kevin_views": 0, "has_kevin_footage": False, "hook": "A corporation hired Pinkerton spies to hang 20 men in the largest mass execution in Pennsylvania history. The mine owner ran the trial, provided the witnesses, and picked the jury. A handprint pressed into a jail cell wall always comes back. And the town 50 miles away has been on fire underground since 1962 — will burn for 250 more years."},
            {"slot": "B", "location_name": "Austin Dam — Aftermath", "state_country": "Austin, Pennsylvania", "kevin_video_title": "ABANDONED Pennsylvania Bayless Paper Mill and Austin Dam", "kevin_views": 7000, "has_kevin_footage": True, "hook": "George Bayless faced zero consequences. He rebuilt. The second dam also failed. Austin never recovered — 2,300 people became 560. We're back."},
            {"slot": "C", "location_name": "Dunalastair House", "state_country": "Perthshire, Scotland", "kevin_video_title": "Abandoned Mansion in the Scottish Highlands", "kevin_views": 4300, "has_kevin_footage": True, "hook": "A grand Highland mansion became a school for Polish children fleeing the Nazis. Now even the walls are coming down, and nobody remembers the children who lived here."},
        ],
    },
    {
        "episode_number": 10, "title": "The Devil's Playground",
        "theme": "Crime, myth, and the stories we tell about evil places",
        "arc": "Three stories about myths and reality. Manson's fantasy led to real murder. Atolia was built on the fantasy war would last forever. Escobar's mansion is a lie that became truth because it made better business.",
        "score_visual": 8, "score_history": 9, "score_emotion": 9, "score_wtf": 10, "score_total": 37, "locked": False,
        "stories": [
            {"slot": "A", "location_name": "Ballarat / Barker Ranch — Manson Deep Dive", "state_country": "Inyo County, California", "kevin_video_title": "WE FOUND CHARLES MANSON'S TRUCK IN A CALIFORNIA GHOST TOWN", "kevin_views": 5500, "has_kevin_footage": True, "hook": "Charles Manson was found hiding in a bathroom cabinet in a desert ranch. The cops were investigating car theft — they didn't know he was America's most wanted. The desert hermit Paul Crockett tried to deprogram Family members and is history's forgotten hero."},
            {"slot": "B", "location_name": "Atolia Tungsten Mine", "state_country": "San Bernardino County, California", "kevin_video_title": "Exploring ATOLIA – Mojave Desert Mine", "kevin_views": 4300, "has_kevin_footage": True, "hook": "The richest tungsten mine on Earth — fuelling two World Wars — turned into a ghost town the moment peace broke out. A town that could only exist during wartime."},
            {"slot": "C", "location_name": "Pablo Escobar's Mansion (The Lie)", "state_country": "Guatapé, Colombia", "kevin_video_title": "Inside Pablo Escobar Abandoned Mansion in Colombia", "kevin_views": 5400, "has_kevin_footage": True, "hook": "The most famous drug lord's mansion in Colombia… ISN'T HIS. The tour guides made it up because Escobar's name sells better. And everyone fell for it."},
        ],
    },
    {
        "episode_number": 11, "title": "Nazis in the Canyon",
        "theme": "Hidden enemies on American soil — ideological, racial, military",
        "arc": "America's paranoia made physical. Nazis hiding in LA canyons. Chinese immigrants locked in island prisons. Military bunkers watching for submarines.",
        "score_visual": 9, "score_history": 10, "score_emotion": 9, "score_wtf": 8, "score_total": 36, "locked": False,
        "stories": [
            {"slot": "A", "location_name": "Murphy's Ranch — Nazi Camp in LA", "state_country": "Pacific Palisades, Los Angeles", "kevin_video_title": "Exploring an Abandoned Nazi Camp in LA // Murphy's Ranch", "kevin_views": 7500, "has_kevin_footage": True, "hook": "In 1933, a mysterious 'Herr Schmidt' convinced a wealthy couple to build a self-sufficient Nazi holdout compound in a hidden LA canyon — designed for when Germany conquered America. After Pearl Harbor, the FBI raided it. The compound has been decaying in Rustic Canyon for 90 years, reachable only by a 500-step staircase."},
            {"slot": "B", "location_name": "Angel Island Immigration Station", "state_country": "San Francisco Bay", "kevin_video_title": "Abandoned Military Base // Angel Island San Francisco", "kevin_views": 2400, "has_kevin_footage": True, "hook": "The 'Ellis Island of the West' — except Ellis Island let people IN. Between 1910-1940, Chinese immigrants were detained for years. They carved desperate poems into the walls in Classical Chinese. Those poems are still there."},
            {"slot": "C", "location_name": "Devils Slide Bunker", "state_country": "Pacifica, San Francisco", "kevin_video_title": "DEVILS SLIDE BUNKER // ABANDONED SAN FRANCISCO", "kevin_views": 3300, "has_kevin_footage": True, "hook": "A WWII observation bunker perched on a cliff south of San Francisco, used to spot Japanese submarines. A tech tycoon bought it in 2023. Nobody knows why."},
        ],
    },
    {
        "episode_number": 12, "title": "Dreams That Died",
        "theme": "Failed utopias and abandoned ambitions",
        "arc": "Three massive structures that were supposed to change everything — a luxury hotel, a desert waterpark, an igloo resort. If you build it, they'll come. They didn't.",
        "score_visual": 9, "score_history": 6, "score_emotion": 8, "score_wtf": 10, "score_total": 33, "locked": False,
        "stories": [
            {"slot": "A", "location_name": "El Miro Hotel, Costa Rica", "state_country": "Jacó, Costa Rica", "kevin_video_title": "ABANDONED ELEVEN STORY HOTEL IN COSTA RICA", "kevin_views": 53000, "has_kevin_footage": True, "hook": "Kevin's most-viewed video. An 11-story concrete skeleton rises from a Costa Rican beach town — never completed, never opened. The tallest building in Jacó and the emptiest."},
            {"slot": "B", "location_name": "Lake Dolores Waterpark", "state_country": "Newberry Springs, California", "kevin_video_title": "ABANDONED Lake Dolores WATERPARK", "kevin_views": 904, "has_kevin_footage": True, "hook": "A businessman built a private waterpark in the Mojave Desert in the 1950s. Each rebranding lasted less than the one before. Empty cement pools and bleached waterslides bake in the desert sun."},
            {"slot": "C", "location_name": "Igloo Hotel, Alaska", "state_country": "Cantwell, Alaska", "kevin_video_title": "EXPLORING ABANDONED IGLOO HOTEL – Alaska", "kevin_views": 5700, "has_kevin_footage": True, "hook": "A man started building a giant igloo-shaped hotel. Four stories, igloo design, roadside landmark. Never completed. Never opened. Never inhabited. It's been for sale for decades. Nobody wants a giant igloo that doesn't work."},
        ],
    },
    {
        "episode_number": 13, "title": "The Human Mole",
        "theme": "Obsession — one person against the mountain",
        "arc": "California desert obsessives. Schmidt dug for 38 years and it was pointless. Kaiser built a utopia and it evaporated. Randsburg just refused to die.",
        "score_visual": 8, "score_history": 8, "score_emotion": 9, "score_wtf": 9, "score_total": 34, "locked": False,
        "stories": [
            {"slot": "A", "location_name": "Burro Schmidt Tunnel", "state_country": "Kern County, California", "kevin_video_title": "CALIFORNIA ABANDONED Mine and Cabin – Burro Shmidt Tunnel", "kevin_views": 1100, "has_kevin_footage": True, "hook": "In 1902, a prospector named William 'Burro' Schmidt started hand-digging a tunnel through a mountain — using only hand tools and dynamite — to create a shortcut. He dug for 38 YEARS. When he finally broke through in 1938, the road he was trying to avoid had already been built on the other side. Ripley's called him 'The Human Mole.'"},
            {"slot": "B", "location_name": "Eagle Mountain", "state_country": "Riverside County, California", "kevin_video_title": "Eagle Mountain SOLD for $22.5 MILLION", "kevin_views": 2800, "has_kevin_footage": True, "hook": "Henry Kaiser built a complete American town — tennis courts, bowling alley, a dozen churches, 4,000 residents — around an iron mine in the desert. When the mine closed, everyone left in one day. Now someone paid $22.5 million for a ghost town — and nobody knows why."},
            {"slot": "C", "location_name": "Randsburg", "state_country": "San Bernardino County, California", "kevin_video_title": "Abandoned Buildings in Randsburg, California Ghost town", "kevin_views": 1000, "has_kevin_footage": True, "hook": "A gold rush town that never quite died — Randsburg clings to life with a handful of residents and a general store. The gateway to every desert obsession."},
        ],
    },
    {
        "episode_number": 14, "title": "Going on the Water",
        "theme": "Shipwrecks and coastal abandonment",
        "arc": "Three bodies of water, three abandoned vessels. LA coast, Scottish Highlands, Fijian tropics. The sea gives and the sea takes — what it leaves behind is haunting.",
        "score_visual": 8, "score_history": 6, "score_emotion": 8, "score_wtf": 8, "score_total": 30, "locked": False,
        "stories": [
            {"slot": "A", "location_name": "SS Dominator Shipwreck", "state_country": "Palos Verdes, Los Angeles", "kevin_video_title": "ABANDONED L.A. SHIPWRECK // S.S. Dominator", "kevin_views": 466, "has_kevin_footage": True, "hook": "In 1961, a Greek cargo ship ran aground off the LA coast. The Coast Guard tried for two days. Heavy seas pushed her higher onto the rocks. She was stripped and abandoned — her rusting skeleton visible from the cliffs until the sea finally consumed her."},
            {"slot": "B", "location_name": "Corpach Shipwreck (MV Dayspring)", "state_country": "Fort William, Scottish Highlands", "kevin_video_title": "Abandoned Shipwreck in the Scottish Highlands", "kevin_views": 4900, "has_kevin_footage": True, "hook": "A fishing vessel built in 1975. A storm snapped her chain in 2011 and she ran aground. Now she rusts at the foot of Ben Nevis — one of the most photographed shipwrecks in the UK."},
            {"slot": "C", "location_name": "Fiji Resort", "state_country": "Korolevu Beach, Fiji", "kevin_video_title": "ABANDONED FIJI RESORT – Korolevu Beach hotel", "kevin_views": 2000, "has_kevin_footage": True, "hook": "A tropical paradise resort left to rot on one of the most beautiful coastlines in the world. The jungle is eating the buildings. Palm trees grow through the lobby."},
        ],
    },
    {
        "episode_number": 15, "title": "The 27 Club House",
        "theme": "Rock and roll, madness, and California counterculture",
        "arc": "California counterculture and its casualties. The rock star who died at 27. The mental hospital that inspired 'Hotel California.' The ski resort that wagered on nature and lost.",
        "score_visual": 7, "score_history": 8, "score_emotion": 9, "score_wtf": 8, "score_total": 32, "locked": False,
        "stories": [
            {"slot": "A", "location_name": "Canned Heat House, Topanga Canyon", "state_country": "Topanga Canyon, Los Angeles", "kevin_video_title": "Canned Heat Abandoned House In Topanga Canyon", "kevin_views": 5300, "has_kevin_footage": True, "hook": "An abandoned house in Topanga Canyon once belonged to Canned Heat's Bob Hite. Guitarist Alan 'Blind Owl' Wilson died behind the house — age 27 — joining Jimi Hendrix and Janis Joplin in the 27 Club within weeks. The house has been abandoned ever since."},
            {"slot": "B", "location_name": "Camarillo State Mental Hospital", "state_country": "Camarillo, California", "kevin_video_title": "ABANDONED ASYLUM – Camarillo State Mental Hospital", "kevin_views": 1100, "has_kevin_footage": True, "hook": "Camarillo devolved into a house of horrors. Charlie Parker wrote 'Relaxin' at Camarillo' after being committed there. The Eagles referenced it in 'Hotel California.' Now it's a university campus — students walk halls where patients were tortured."},
            {"slot": "C", "location_name": "Abandoned California Ski Resort", "state_country": "California", "kevin_video_title": "Abandoned Ski Resort Drone", "kevin_views": 1500, "has_kevin_footage": True, "hook": "A California ski resort that bet on snow and lost. The lifts stopped, the lodge emptied, the mountain reclaimed it."},
        ],
    },
    {
        "episode_number": 16, "title": "Wheel of Misfortune",
        "theme": "Sin, gambling, and desert secrets outside Las Vegas",
        "arc": "The desert hides everything — toxic mines, ghost towns, car graveyards. Just beyond the Las Vegas glitter, the real Nevada is a landscape of abandonment and contamination.",
        "score_visual": 8, "score_history": 7, "score_emotion": 7, "score_wtf": 9, "score_total": 31, "locked": False,
        "stories": [
            {"slot": "A", "location_name": "Three Kids Mine / Wheel of Misfortune", "state_country": "Henderson, Nevada", "kevin_video_title": "Wheel of Misfortune and Three Kids Mine – ABANDONED LAS VEGAS", "kevin_views": 3900, "has_kevin_footage": True, "hook": "Twenty miles from the Las Vegas Strip sits an abandoned manganese mine with 1.8 million cubic tons of toxic sludge. INDECLINE painted a massive 'Wheel of Misfortune' mural with segments reading 'cancer,' 'birth defects,' and 'death.' It went viral."},
            {"slot": "B", "location_name": "Tybo Ghost Town", "state_country": "Nye County, Nevada", "kevin_video_title": "TYBO NEVADA – Ghost Town exploration", "kevin_views": 545, "has_kevin_footage": True, "hook": "A silver mining ghost town in the Nevada desert, abandoned since the 1930s. Remote, untouched, barely explored."},
            {"slot": "C", "location_name": "Car Canyon", "state_country": "Near Page, Arizona", "kevin_video_title": "ABANDONED CAR CANYON – Full of Vintage Cars in desert", "kevin_views": 32000, "has_kevin_footage": True, "hook": "A hidden canyon in the Arizona desert absolutely STUFFED with vintage cars — rusting Chevys, Fords, and Dodges dumped decades ago. Nobody knows exactly who put them there or why."},
        ],
    },
    {
        "episode_number": 17, "title": "Sand, Wind & Silence",
        "theme": "The Mojave is eating California",
        "arc": "California's coast and desert, both losing the fight. Sand buries houses, resorts freeze in time, motels just fade away.",
        "score_visual": 8, "score_history": 7, "score_emotion": 8, "score_wtf": 7, "score_total": 30, "locked": False,
        "stories": [
            {"slot": "A", "location_name": "Newberry Springs — Buried Town Deep Dive", "state_country": "San Bernardino County, California", "kevin_video_title": "This California Neighborhood was Covered in Sand", "kevin_views": 11000, "has_kevin_footage": True, "hook": "Deep dive into the buried town — the families who watched helplessly. The LA Times investigation. The county's refusal to take responsibility. Extended from the Ep 6 B story into a full exploration."},
            {"slot": "B", "location_name": "Oxnard Marina Hotel (Casa Sirena)", "state_country": "Oxnard, California", "kevin_video_title": "ABANDONED CALIFORNIA OXNARD MARINA HOTEL", "kevin_views": 14000, "has_kevin_footage": True, "hook": "The Casa Sirena seaside resort, abandoned in 2009, frozen in time. The Lobster Trap restaurant still has tablecloths laid, salt shakers filled, a bottle of Tabasco waiting. Mermaid statues guard dry fountains."},
            {"slot": "C", "location_name": "Desert Motel", "state_country": "Mojave Desert, California", "kevin_video_title": "California URBEX – EXPLORING ABANDONED MOTEL in DESERT", "kevin_views": 1400, "has_kevin_footage": True, "hook": "A nameless motel on a desert highway. No history, no story — just empty rooms, a dry pool, and a sign nobody reads. Sometimes abandonment doesn't have a dramatic reason. Sometimes people just stop coming."},
        ],
    },
    {
        "episode_number": 18, "title": "Ellis Island of the West",
        "theme": "Immigration, exclusion, and the lies nations tell about freedom",
        "arc": "Three ways nations process 'the other' — immigration detention, frontier mythology, and war. Who gets in? Who gets kept out? Who gets killed?",
        "score_visual": 8, "score_history": 10, "score_emotion": 9, "score_wtf": 7, "score_total": 34, "locked": False,
        "stories": [
            {"slot": "A", "location_name": "Angel Island — Full Immigration Story", "state_country": "San Francisco Bay", "kevin_video_title": "Abandoned Military Base // Angel Island San Francisco", "kevin_views": 2400, "has_kevin_footage": True, "hook": "Full deep dive: Chinese immigrants detained for months under the Chinese Exclusion Act carved poems of despair into the wooden barrack walls in Classical Chinese. When the building was scheduled for demolition, a ranger discovered the poems. The 'Ellis Island of the West' was America's front door with a 'NO CHINESE' sign."},
            {"slot": "B", "location_name": "Fort Courage Route 66", "state_country": "Arizona", "kevin_video_title": "ABANDONED Arizona Fort Courage and Pancake House Route 66", "kevin_views": 4900, "has_kevin_footage": True, "hook": "A roadside attraction themed after the TV show F Troop — a fake Western fort with a pancake house. The American frontier myth commodified, then abandoned when the highway moved."},
            {"slot": "C", "location_name": "Cambodia War Museum", "state_country": "Siem Reap, Cambodia", "kevin_video_title": "Urban Exploring Cambodia: War Tanks & Helicopters", "kevin_views": 403, "has_kevin_footage": True, "hook": "Tanks, helicopters, and landmines from the Khmer Rouge era and Vietnam War — rusting in a field in Cambodia. Kevin walks among machines that killed millions."},
        ],
    },
    {
        "episode_number": 19, "title": "The Mother Road",
        "theme": "Route 66 — the full journey",
        "arc": "A full Route 66 road trip episode. California to Texas via Arizona. The physical road, the trading posts, the art installations. The Mother Road in all her ruined glory.",
        "score_visual": 8, "score_history": 7, "score_emotion": 7, "score_wtf": 7, "score_total": 29, "locked": False,
        "stories": [
            {"slot": "A", "location_name": "Twin Arrows Trading Post & Route 66 Arizona", "state_country": "Arizona", "kevin_video_title": "Route 66 Abandoned – Twin Arrows Trading Post Ruins", "kevin_views": 2000, "has_kevin_footage": True, "hook": "The Twin Arrows — two massive arrow-shaped signs planted in the desert — marked a Hopi-run trading post. When the interstate came, the arrows stayed but the people left."},
            {"slot": "B", "location_name": "Driving Route 66 to Arizona", "state_country": "California / Arizona", "kevin_video_title": "Driving ABANDONED CA Route 66 to Arizona", "kevin_views": 9700, "has_kevin_footage": True, "hook": "Kevin drives the surviving stretches of Route 66 from California into Arizona — passing through the corpses of a dozen dead towns. The road itself is the story."},
            {"slot": "C", "location_name": "Slug Bug Ranch", "state_country": "Conway, Texas", "kevin_video_title": "SLUG BUG RANCH // Route 66 Texas", "kevin_views": 254, "has_kevin_footage": True, "hook": "Cadillac Ranch's lesser-known cousin — VW Beetles buried nose-down in a Texas field along Route 66. A folk-art tribute to the highway, slowly rusting into the earth."},
        ],
    },
    {
        "episode_number": 20, "title": "Company Towns",
        "theme": "What happens when the corporation IS the town",
        "arc": "Three forms of corporate America's footprint. When capital flows somewhere else, the people who stayed become ghosts.",
        "score_visual": 8, "score_history": 8, "score_emotion": 8, "score_wtf": 8, "score_total": 32, "locked": False,
        "stories": [
            {"slot": "A", "location_name": "Eagle Mountain — Full Kaiser Story", "state_country": "Riverside County, California", "kevin_video_title": "Eagle Mountain SOLD for $22.5 MILLION", "kevin_views": 2800, "has_kevin_footage": True, "hook": "Full Kaiser story — the man who founded Kaiser Permanente and created a utopia in the desert. Eagle Mountain was a triple-bust: mining, prison, landfill — all failed. America's healthcare system literally grew out of this ghost town."},
            {"slot": "B", "location_name": "Trona — Kerr-McGee Connection", "state_country": "San Bernardino County, California", "kevin_video_title": "Trona, California Ghost Town", "kevin_views": 26000, "has_kevin_footage": True, "hook": "The Kerr-McGee/Silkwood connection. Company scrip. The 1974 eviction attempt. The community's resistance. Trona as a study in corporate control and worker resilience."},
            {"slot": "C", "location_name": "Abandoned Car Dealership", "state_country": "New York", "kevin_video_title": "ABANDONED DEALERSHIP // LEFT FOR 30 YEARS", "kevin_views": 552, "has_kevin_footage": True, "hook": "A car dealership in New York left untouched for 30 years. Cars on the showroom floor, keys in the ignition, paperwork on desks. A time capsule of American commerce frozen in place."},
        ],
    },
    {
        "episode_number": 21, "title": "The Cover-Up",
        "theme": "Government lies and contaminated land",
        "arc": "When the government lies: about nuclear contamination, about water safety, about its obligations to the vulnerable. Three cover-ups, three groups of Americans betrayed.",
        "score_visual": 7, "score_history": 10, "score_emotion": 10, "score_wtf": 8, "score_total": 35, "locked": False,
        "stories": [
            {"slot": "A", "location_name": "George AFB — ATSDR Scandal", "state_country": "Victorville, California", "kevin_video_title": "Exploring an Abandoned Military Base in California", "kevin_views": 10600, "has_kevin_footage": True, "hook": "The deepest dive — focusing on how the Air Force deliberately misled the ATSDR, Congress, and the public. 67 toxins identified. Completed exposure pathways concealed. The 2022 lawsuit dismissed on technicalities. The families still fighting."},
            {"slot": "B", "location_name": "LA-88 Nike — Santa Susana Connection", "state_country": "Chatsworth, Los Angeles", "kevin_video_title": "EXPLORING UNDERGROUND – LOS ANGELES MISSILE SILO", "kevin_views": 3200, "has_kevin_footage": True, "hook": "Nearby Santa Susana Field Laboratory had a partial nuclear meltdown in 1959 — WORSE than Three Mile Island — covered up for decades. The Nike site and the meltdown site are neighbours. The entire Chatsworth hillside was an undisclosed nuclear zone."},
            {"slot": "C", "location_name": "Abandoned State Hospital — Connecting Thread", "state_country": "Location TBC", "kevin_video_title": "ABANDONED STATE HOSPITAL That you can LEGALLY EXPLORE", "kevin_views": 908, "has_kevin_footage": True, "hook": "Connecting the government's treatment of the mentally ill with its treatment of military families. Both disposable. Different agencies, same indifference."},
        ],
    },
    {
        "episode_number": 22, "title": "Gold Fever",
        "theme": "The California Gold Rush's toxic legacy",
        "arc": "The Gold Rush killed people, destroyed cultures, poisoned waterways. 175 years later, California is still dealing with the consequences.",
        "score_visual": 7, "score_history": 8, "score_emotion": 8, "score_wtf": 7, "score_total": 30, "locked": False,
        "stories": [
            {"slot": "A", "location_name": "Helena Ghost Town — Full Story", "state_country": "Trinity County, California", "kevin_video_title": "Inside Helena Ghost Town – Gold Rush Settlement", "kevin_views": 868, "has_kevin_footage": True, "hook": "'Baghdad of the West' — deep dive into the Chimariko genocide, the wild mining era, and the town now owned by a dead man's estate. National Register property that nobody claims."},
            {"slot": "B", "location_name": "Rinconada Mine — Mercury Legacy", "state_country": "San Luis Obispo County, California", "kevin_video_title": "Rinconada Quicksilver Mine – Massive Abandoned Site", "kevin_views": 3000, "has_kevin_footage": True, "hook": "The mercury connection to gold — miners dying of poisoning so other miners could extract gold. The environmental contamination that persists today in California waterways."},
            {"slot": "C", "location_name": "Modern Era Gold Mine", "state_country": "California Desert", "kevin_video_title": "ABANDONED MODERN ERA GOLD MINE in California Desert", "kevin_views": 3100, "has_kevin_footage": True, "hook": "Gold mining continued in the California desert into the modern era. This abandoned mine shows the industry never really changed."},
        ],
    },
    {
        "episode_number": 23, "title": "Earthquake Island",
        "theme": "Natural disasters and rebuilding in the wrong place",
        "arc": "Natural disasters and what comes after — earthquake, ecological collapse, economic death. Rebuilding sometimes creates worse ruins than the disaster itself.",
        "score_visual": 9, "score_history": 8, "score_emotion": 8, "score_wtf": 7, "score_total": 32, "locked": False,
        "stories": [
            {"slot": "A", "location_name": "Poggioreale", "state_country": "Belìce Valley, Sicily", "kevin_video_title": "Exploring an Abandoned Italian Ghost Town", "kevin_views": 1000, "has_kevin_footage": True, "hook": "The 1968 earthquake destroyed the town. The government hired famous architects to rebuild it 4km away. The Mafia got the construction contracts. The old town is frozen in 1968. The new town was designed for 10,000 — population 1,400."},
            {"slot": "B", "location_name": "Bombay Beach — Art & Holdouts", "state_country": "Imperial County, California", "kevin_video_title": "Abandoned California City: Bombay Beach", "kevin_views": 3900, "has_kevin_footage": True, "hook": "Return to Bombay Beach — this time the Biennale art movement, the holdout residents, the lithium gold rush beneath the toxic lake. A disaster zone being gentrified by artists."},
            {"slot": "C", "location_name": "Mexico Beach Resort", "state_country": "Mexico", "kevin_video_title": "ABANDONED Mexico Beach Resort – Caught by Security", "kevin_views": 909, "has_kevin_footage": True, "hook": "A Mexican coastal resort abandoned and guarded by security — Kevin gets caught. The tension of trespassing adds drama to an already eerie location."},
        ],
    },
    {
        "episode_number": 24, "title": "The Tunnel",
        "theme": "Underground — what's hidden beneath the surface",
        "arc": "Things buried underground — tunnels, missile silos, abandoned railways. What we hide, what we dig, and what we leave behind when we stop digging.",
        "score_visual": 8, "score_history": 7, "score_emotion": 8, "score_wtf": 7, "score_total": 30, "locked": False,
        "stories": [
            {"slot": "A", "location_name": "Burro Schmidt Tunnel — Full Walk", "state_country": "Kern County, California", "kevin_video_title": "CALIFORNIA ABANDONED Mine and Cabin – Burro Shmidt Tunnel", "kevin_views": 1100, "has_kevin_footage": True, "hook": "Full exploration of the half-mile hand-dug tunnel. Kevin walks through it. The 38 years of obsession. The cabin where Schmidt lived. The ultimate monument to futile human determination."},
            {"slot": "B", "location_name": "LA-88 Nike Underground Magazines", "state_country": "Chatsworth, Los Angeles", "kevin_video_title": "EXPLORING UNDERGROUND – LOS ANGELES MISSILE SILO", "kevin_views": 3200, "has_kevin_footage": True, "hook": "The underground missile magazines — blast doors, elevator shafts, the spaces where nuclear warheads waited to be launched at LA's own sky. Cold War architecture buried in suburban hills."},
            {"slot": "C", "location_name": "Abandoned Train", "state_country": "California Desert", "kevin_video_title": "ABANDONED TRAIN in the California Desert", "kevin_views": 2600, "has_kevin_footage": True, "hook": "A train left in the desert. No engine, no destination, no explanation. Just carriages rusting in the sun."},
        ],
    },
    {
        "episode_number": 25, "title": "Fire on the Mountain",
        "theme": "Places consumed by flame",
        "arc": "Three kinds of fire — underground, arson, and the constant threat. Fire as destroyer, fire as concealer, fire as inescapable reality.",
        "score_visual": 8, "score_history": 8, "score_emotion": 8, "score_wtf": 7, "score_total": 31, "locked": False,
        "stories": [
            {"slot": "A", "location_name": "Centralia Underground Fire — Full Story", "state_country": "Columbia County, Pennsylvania", "kevin_video_title": None, "kevin_views": 0, "has_kevin_footage": False, "hook": "The full Centralia story — Todd Domboski falling into a sinkhole, the $42 million relocation, the 5 holdouts granted permission to stay until death, the Graffiti Highway buried under dirt. A town on fire for 64 years with 250 more to go."},
            {"slot": "B", "location_name": "Lennox Castle Fire", "state_country": "Lennoxtown, Scotland", "kevin_video_title": "Abandoned Scottish Castle Turned Mental Asylum", "kevin_views": 400, "has_kevin_footage": True, "hook": "The 2008 arson that gutted the castle — suspected to clear the site for development. A building that survived 170 years as a home and 66 years as an asylum, destroyed by someone with a match and a motive."},
            {"slot": "C", "location_name": "Topanga Lookout Fire Tower", "state_country": "Topanga Canyon, California", "kevin_video_title": "Hiking to the Topanga lookout – Fire Tower", "kevin_views": 740, "has_kevin_footage": True, "hook": "A fire lookout tower in the Santa Monica Mountains — built to watch for fires that regularly devastate Southern California. Now abandoned itself. The irony of a fire-watching station left to decay."},
        ],
    },
    {
        "episode_number": 26, "title": "The Last Resident",
        "theme": "The people who refused to leave",
        "arc": "The human dimension of abandonment. Some people stay. Some towns have no one left to stay. What makes someone the last person to turn off the lights?",
        "score_visual": 8, "score_history": 7, "score_emotion": 9, "score_wtf": 7, "score_total": 31, "locked": False,
        "stories": [
            {"slot": "A", "location_name": "Amboy — The People Who Stayed", "state_country": "San Bernardino County, California", "kevin_video_title": "Amboy California: Route 66 Ghost town", "kevin_views": 13000, "has_kevin_footage": True, "hook": "Deep dive into Buster Burris (ran the town for 57 years), Albert Okura (chicken magnate who bought it), and Kyle Okura (son, now guardian of a town with zero residents). The economics and emotion of preserving somewhere nobody lives."},
            {"slot": "B", "location_name": "Bombay Beach — The Holdouts", "state_country": "Imperial County, California", "kevin_video_title": "Abandoned California City: Bombay Beach", "kevin_views": 3900, "has_kevin_footage": True, "hook": "The ~250 people who stayed in Bombay Beach — the artists, the elderly who can't afford to leave, the eccentrics who chose to. Life in a dying town, surrounded by toxic dust and fish bones."},
            {"slot": "C", "location_name": "Dunalastair — Last Days", "state_country": "Perthshire, Scotland", "kevin_video_title": "Abandoned Mansion in the Scottish Highlands", "kevin_views": 4300, "has_kevin_footage": True, "hook": "The mansion's last days — who left last? The Polish school closed, the house emptied, the Highlands reclaimed it. Nobody stayed to fight. Sometimes the last resident just isn't there."},
        ],
    },
    {
        "episode_number": 27, "title": "Across the Pacific",
        "theme": "International ruins — paradise found and lost",
        "arc": "Three Latin American / Pacific locations where tourism failed, myths were built, and paradise rotted.",
        "score_visual": 8, "score_history": 6, "score_emotion": 7, "score_wtf": 8, "score_total": 29, "locked": False,
        "stories": [
            {"slot": "A", "location_name": "Fiji Resort — Full Story", "state_country": "Korolevu Beach, Fiji", "kevin_video_title": "ABANDONED FIJI RESORT – Korolevu Beach hotel", "kevin_views": 2000, "has_kevin_footage": True, "hook": "A tropical resort complex on Fiji's Coral Coast, left to the jungle. Palm trees through lobbies, vines through windows, the swimming pool a swamp. Paradise abandoned is somehow more disturbing than a desert ruin."},
            {"slot": "B", "location_name": "Costa Rica Hotel — Economic Story", "state_country": "Jacó, Costa Rica", "kevin_video_title": "ABANDONED ELEVEN STORY HOTEL IN COSTA RICA", "kevin_views": 53000, "has_kevin_footage": True, "hook": "Return to El Miro for the bigger story — the economics of failed development in Central America, the construction boom and bust, the skeleton buildings dotting the coast."},
            {"slot": "C", "location_name": "Pablo Escobar Mansion — Drug Tourism", "state_country": "Guatapé, Colombia", "kevin_video_title": "Inside Pablo Escobar Abandoned Mansion in Colombia", "kevin_views": 5400, "has_kevin_footage": True, "hook": "The myth-busting story expanded — drug tourism in Colombia, the ethics of profiting from narco-history, the real story of the building."},
        ],
    },
    {
        "episode_number": 28, "title": "The Road Trip",
        "theme": "Kevin's journey across America",
        "arc": "The meta-episode. Kevin reflects on the series, the journey, the pattern. Why do we explore ruins? What are we looking for? The road is the story.",
        "score_visual": 7, "score_history": 6, "score_emotion": 8, "score_wtf": 7, "score_total": 28, "locked": False,
        "stories": [
            {"slot": "A", "location_name": "5 Abandoned Places in Southern California", "state_country": "Various, California", "kevin_video_title": "5 Abandoned Places in Southern California – LEGAL URBEX", "kevin_views": 14000, "has_kevin_footage": True, "hook": "Kevin's compilation video — multiple locations in one trip. The episode frames this as Kevin's personal journey: why he explores, what he finds, what it means to him. The most personal episode of the series."},
            {"slot": "B", "location_name": "Route 66 Road Trip — The Drive", "state_country": "California to Arizona", "kevin_video_title": "ABANDONED PLACES ROAD TRIP – Two weeks across USA", "kevin_views": 10000, "has_kevin_footage": True, "hook": "Two weeks, thousands of miles, dozens of abandoned places. The connective tissue of the series."},
            {"slot": "C", "location_name": "Abandoned Airplanes Arizona", "state_country": "Arizona", "kevin_video_title": "3 PLACES TO SEE ABANDONED AIRPLANES IN ARIZONA", "kevin_views": 1300, "has_kevin_footage": True, "hook": "Aircraft boneyards in the Arizona desert — planes from every era, grounded forever. The desert preserves what humans discard."},
        ],
    },
    {
        "episode_number": 29, "title": "The Highland Darkness",
        "theme": "Scotland's abandoned heritage",
        "arc": "Scotland's three abandoned stories — institutional, maritime, aristocratic. The Highland Clearances echo through all of them.",
        "score_visual": 9, "score_history": 7, "score_emotion": 8, "score_wtf": 6, "score_total": 30, "locked": False,
        "stories": [
            {"slot": "A", "location_name": "Lennox Castle — Full Hughie Story", "state_country": "Lennoxtown, Scotland", "kevin_video_title": "Abandoned Scottish Castle Turned Mental Asylum", "kevin_views": 400, "has_kevin_footage": True, "hook": "The full Hughie McIntyre story. 16 years locked inside. His return to the ruins. The 2008 fire. The housing development on the grounds where thousands suffered. Scotland's institutional shame."},
            {"slot": "B", "location_name": "Corpach Shipwreck — Fishing Industry", "state_country": "Fort William, Scotland", "kevin_video_title": "Abandoned Shipwreck in the Scottish Highlands", "kevin_views": 4900, "has_kevin_footage": True, "hook": "The photogenic wreck at the foot of Ben Nevis. The story of a working fishing boat — the last of its kind, a symbol of Scotland's dying fishing industry."},
            {"slot": "C", "location_name": "Scottish Highlands Mansion", "state_country": "Perthshire, Scotland", "kevin_video_title": "Abandoned Mansion in the Scottish Highlands", "kevin_views": 4300, "has_kevin_footage": True, "hook": "The architecture, the landscape, the atmosphere. The Highlands as a landscape of abandonment — castles, estates, communities emptied by clearances and time."},
        ],
    },
    {
        "episode_number": 30, "title": "Where the Bodies Are",
        "theme": "The dead remain — cemeteries, graves, and the forgotten",
        "arc": "The series finale. Three cemeteries — an asylum, a hospital, a ghost town. Numbers instead of names, cans instead of coffins, boot hills in the desert. The show was always about the people who never got out.",
        "score_visual": 7, "score_history": 9, "score_emotion": 10, "score_wtf": 8, "score_total": 34, "locked": False,
        "stories": [
            {"slot": "A", "location_name": "Northern State Hospital Cemetery", "state_country": "Sedro-Woolley, Washington", "kevin_video_title": "ABANDONED STATE HOSPITAL That you can LEGALLY EXPLORE", "kevin_views": 908, "has_kevin_footage": True, "hook": "Return to Northern State for the cemetery story. 1,500 graves, most marked with numbers. 200 cans of cremated remains. The ongoing effort to identify the 'lost patients.' Can we give them their names back?"},
            {"slot": "B", "location_name": "Letchworth Village Cemetery", "state_country": "Thiells, New York", "kevin_video_title": "ABANDONED HOSPITAL for Mentally and Physically Disabled", "kevin_views": 410, "has_kevin_footage": True, "hook": "900 graves marked with numbers. The plaque: 'To Those Who Shall Not Be Forgotten.' The polio vaccine children. The brains in jars. The cost of institutional medicine."},
            {"slot": "C", "location_name": "Ballarat Boot Hill", "state_country": "Death Valley, California", "kevin_video_title": "WE FOUND CHARLES MANSON'S TRUCK IN A CALIFORNIA GHOST TOWN", "kevin_views": 5500, "has_kevin_footage": True, "hook": "Seldom Seen Slim's grave — 'Just bury me where the digging's easy.' The 28th and final burial in Ballarat's cemetery. A gentle ending to a dark series."},
        ],
    },
]


def _seed_episode_planner(project_id: str) -> dict:
    """Create story cards + episode arrangements from ABANDONED_EPISODES_SEED."""
    story_card_ids = {}  # slot_key -> firestore_id
    episode_refs = []

    for ep in ABANDONED_EPISODES_SEED:
        slot_ids = {}
        for story in ep["stories"]:
            card = {
                "project_id": project_id,
                "location_name": story["location_name"],
                "state_country": story["state_country"],
                "kevin_video_title": story.get("kevin_video_title"),
                "kevin_views": story.get("kevin_views", 0),
                "has_kevin_footage": story.get("has_kevin_footage", False),
                "hook": story["hook"],
                "natural_role": story["slot"],
                "themes": [],
                "current_episode": ep["episode_number"],
                "current_slot": story["slot"],
                "kevin_url": None,  # linked when YouTube clips are analyzed
                "created_at": datetime.utcnow().isoformat(),
            }
            doc_ref = db.collection(COLLECTIONS["story_cards"]).document()
            doc_ref.set(card)
            slot_ids[story["slot"]] = doc_ref.id

        arr = {
            "project_id": project_id,
            "episode_number": ep["episode_number"],
            "title": ep["title"],
            "theme": ep["theme"],
            "arc": ep["arc"],
            "locked": ep["locked"],
            "score_visual": ep["score_visual"],
            "score_history": ep["score_history"],
            "score_emotion": ep["score_emotion"],
            "score_wtf": ep["score_wtf"],
            "score_total": ep["score_total"],
            "slot_a": slot_ids.get("A"),
            "slot_b": slot_ids.get("B"),
            "slot_c": slot_ids.get("C"),
            "slot_d": slot_ids.get("D"),
            "created_at": datetime.utcnow().isoformat(),
            "updated_at": datetime.utcnow().isoformat(),
        }
        doc_ref = db.collection(COLLECTIONS["episode_arrangements"]).document()
        doc_ref.set(arr)
        episode_refs.append({**arr, "id": doc_ref.id})

    return {"episodes_created": len(episode_refs), "story_cards_created": sum(len(ep["stories"]) for ep in ABANDONED_EPISODES_SEED)}


def _get_story_card(card_id: str) -> dict | None:
    if not card_id:
        return None
    doc = db.collection(COLLECTIONS["story_cards"]).document(card_id).get()
    if doc.exists:
        return {**doc.to_dict(), "id": doc.id}
    return None


@app.route("/api/projects/<project_id>/seed-episode-planner", methods=["POST"])
def seed_episode_planner(project_id):
    """Seed 30-episode plan + story cards for the Abandoned series template."""
    # Check if already seeded
    existing = list(
        db.collection(COLLECTIONS["episode_arrangements"])
        .where("project_id", "==", project_id)
        .limit(1)
        .stream()
    )
    if existing:
        return jsonify({"message": "Already seeded", "seeded": False}), 200

    result = _seed_episode_planner(project_id)
    return jsonify({"message": "Seeded successfully", "seeded": True, **result}), 201


@app.route("/api/projects/<project_id>/episode-arrangements", methods=["GET"])
def get_episode_arrangements(project_id):
    """Return all episode arrangements with populated story card data."""
    docs = (
        db.collection(COLLECTIONS["episode_arrangements"])
        .where("project_id", "==", project_id)
        .order_by("episode_number")
        .stream()
    )
    result = []
    for doc in docs:
        arr = {**doc.to_dict(), "id": doc.id}
        # Populate story card data for each slot
        for slot in ["slot_a", "slot_b", "slot_c", "slot_d"]:
            if arr.get(slot):
                card = _get_story_card(arr[slot])
                arr[f"{slot}_data"] = card
        result.append(arr)
    return jsonify(result)


@app.route("/api/projects/<project_id>/episode-arrangements/<int:ep_num>", methods=["PUT"])
def update_episode_arrangement(project_id, ep_num):
    """Update an episode arrangement (lock/unlock, title, etc.)."""
    data = request.get_json()
    docs = list(
        db.collection(COLLECTIONS["episode_arrangements"])
        .where("project_id", "==", project_id)
        .where("episode_number", "==", ep_num)
        .limit(1)
        .stream()
    )
    if not docs:
        return jsonify({"error": "Episode arrangement not found"}), 404

    doc_ref = docs[0].reference
    update_data = {k: v for k, v in data.items() if k in ["locked", "title", "theme", "arc", "notes"]}
    update_data["updated_at"] = datetime.utcnow().isoformat()
    doc_ref.update(update_data)
    updated = {**docs[0].to_dict(), **update_data, "id": docs[0].id}
    return jsonify(updated)


@app.route("/api/projects/<project_id>/episode-arrangements/swap", methods=["POST"])
def swap_story_cards(project_id):
    """Swap story cards between two episode slots.
    Body: { from_ep: int, from_slot: 'A'|'B'|'C'|'D', to_ep: int, to_slot: 'A'|'B'|'C'|'D' }
    """
    data = request.get_json()
    from_ep = data.get("from_ep")
    from_slot = data.get("from_slot", "").lower()  # 'a','b','c','d'
    to_ep = data.get("to_ep")
    to_slot = data.get("to_slot", "").lower()

    if not all([from_ep, from_slot, to_ep, to_slot]):
        return jsonify({"error": "from_ep, from_slot, to_ep, to_slot required"}), 400

    slot_field_from = f"slot_{from_slot}"
    slot_field_to = f"slot_{to_slot}"

    # Fetch both episodes
    def fetch_arrangement(ep_num):
        docs = list(
            db.collection(COLLECTIONS["episode_arrangements"])
            .where("project_id", "==", project_id)
            .where("episode_number", "==", ep_num)
            .limit(1)
            .stream()
        )
        return docs[0] if docs else None

    from_doc = fetch_arrangement(from_ep)
    to_doc = fetch_arrangement(to_ep)

    if not from_doc or not to_doc:
        return jsonify({"error": "One or both episodes not found"}), 404

    from_data = from_doc.to_dict()
    to_data = to_doc.to_dict()

    # Check locks
    if from_data.get("locked") or to_data.get("locked"):
        return jsonify({"error": "Cannot swap stories in locked episodes"}), 403

    from_card_id = from_data.get(slot_field_from)
    to_card_id = to_data.get(slot_field_to)

    now = datetime.utcnow().isoformat()

    # Perform swap
    from_doc.reference.update({slot_field_from: to_card_id, "updated_at": now})
    to_doc.reference.update({slot_field_to: from_card_id, "updated_at": now})

    # Update story card's current_episode and current_slot fields
    if from_card_id:
        db.collection(COLLECTIONS["story_cards"]).document(from_card_id).update({
            "current_episode": to_ep,
            "current_slot": to_slot.upper(),
        })
    if to_card_id:
        db.collection(COLLECTIONS["story_cards"]).document(to_card_id).update({
            "current_episode": from_ep,
            "current_slot": from_slot.upper(),
        })

    return jsonify({"swapped": True, "from_card": from_card_id, "to_card": to_card_id})


@app.route("/api/projects/<project_id>/episode-arrangements/assess", methods=["POST"])
def assess_arrangements(project_id):
    """AI-powered assessment of the current episode arrangement."""
    # Fetch all arrangements
    docs = list(
        db.collection(COLLECTIONS["episode_arrangements"])
        .where("project_id", "==", project_id)
        .order_by("episode_number")
        .stream()
    )
    if not docs:
        return jsonify({"error": "No arrangements found. Seed the planner first."}), 404

    # Build a compact summary for the AI
    ep_summaries = []
    for doc in docs:
        arr = doc.to_dict()
        slots = []
        for s in ["slot_a", "slot_b", "slot_c", "slot_d"]:
            if arr.get(s):
                card = _get_story_card(arr[s])
                if card:
                    slots.append(f"{s[-1].upper()}: {card['location_name']} ({card['state_country']})")
        ep_summaries.append(
            f"Ep{arr['episode_number']} '{arr['title']}' [Score:{arr['score_total']}/40, Locked:{arr['locked']}] — {' | '.join(slots)}"
        )

    prompt = f"""You are the story editor for 'What the Hell Happened Here!?' — a 30-episode Fremantle/Arrow documentary series.

Review this current episode arrangement and provide editorial feedback:

{chr(10).join(ep_summaries)}

Analyse and provide:
1. **Arc Assessment**: Does the series build emotionally across 30 episodes? Are there pacing issues?
2. **Thematic Clustering**: Are similar themes too close together? What moves should be made?
3. **Score Analysis**: Identify the weakest episodes (under 32/40) and suggest stronger story swaps.
4. **Geographic Variety**: Is there good geographic spread or are too many similar locations consecutive?
5. **Top 3 Swap Suggestions**: Specific story moves with rationale (format: 'Move [Story] from Ep X slot Y to Ep Z slot W because...')
6. **Series Finale Check**: Does Ep 30 land emotionally? Is it the right closer?

Be direct, editorial, and specific. Reference episode numbers and story names."""

    resp = model.generate_content(prompt)
    assessment = resp.text if hasattr(resp, "text") else str(resp)

    return jsonify({
        "assessment": assessment,
        "episodes_reviewed": len(docs),
        "generated_at": datetime.utcnow().isoformat(),
    })


@app.route("/api/projects/<project_id>/import-script", methods=["POST"])
def import_script(project_id):
    """Fast-track: import a pre-written script and extract research, archive needs,
    expert types, and visual ideas — populating all phases in one pass."""
    data = request.get_json()
    script_text = data.get("script_text", "").strip()
    episode_title = data.get("episode_title", "Untitled Episode")

    if len(script_text) < 100:
        return jsonify({"error": "Script text too short"}), 400

    project = get_doc("projects", project_id)
    if not project:
        return jsonify({"error": "Project not found"}), 404

    prompt = f"""You are a documentary research assistant. A producer has submitted a pre-written script for '{episode_title}'.

Your job is to extract everything needed to populate the production workflow — research context, archive requirements, expert contributors, and visual ideas.

SCRIPT:
{script_text[:50000]}

Return a JSON object with EXACTLY this structure:
{{
  "episode_summary": "2-3 sentence summary of what this episode covers",
  "key_facts": ["fact 1", "fact 2", ...],
  "research_topics": [
    {{"topic": "...", "why_needed": "...", "suggested_sources": ["source1", "source2"]}}
  ],
  "archive_needs": [
    {{"description": "...", "type": "historical_photo|footage|document|map", "era": "...", "suggested_sources": ["..."]}}
  ],
  "expert_types": [
    {{"role": "...", "reason": "...", "ideal_soundbite": "...", "suggested_questions": ["q1", "q2", "q3"]}}
  ],
  "visual_ideas": [
    {{"type": "ai_reconstruction|ai_drone|archive_transition|talking_head|broll", "description": "...", "when_in_script": "..."}}
  ],
  "locations_mentioned": ["location 1", "location 2"],
  "themes": ["theme 1", "theme 2"]
}}

Extract AT LEAST 5 research_topics, 5 archive_needs, 3 expert_types, and 5 visual_ideas."""

    resp = model.generate_content(
        prompt,
        generation_config={"response_mime_type": "application/json"},
    )
    try:
        extracted = json.loads(resp.text)
    except Exception:
        extracted = {"error": "Failed to parse script analysis", "raw": resp.text[:500]}
        return jsonify(extracted), 500

    # Save research topics as research documents
    research_saved = []
    for rt in extracted.get("research_topics", []):
        doc = create_doc("research_documents", {
            "project_id": project_id,
            "topic": rt.get("topic"),
            "why_needed": rt.get("why_needed"),
            "suggested_sources": rt.get("suggested_sources", []),
            "source": "script_import",
            "episode_title": episode_title,
            "created_at": datetime.utcnow().isoformat(),
        })
        research_saved.append(doc)

    # Save archive needs as assets
    assets_saved = []
    for an in extracted.get("archive_needs", []):
        doc = create_doc("assets", {
            "project_id": project_id,
            "description": an.get("description"),
            "asset_type": an.get("type"),
            "era": an.get("era"),
            "suggested_sources": an.get("suggested_sources", []),
            "source": "script_import",
            "status": "needed",
            "created_at": datetime.utcnow().isoformat(),
        })
        assets_saved.append(doc)

    # Save expert types as interview plans
    interviews_saved = []
    for et in extracted.get("expert_types", []):
        doc = create_doc("interviews", {
            "project_id": project_id,
            "topic": et.get("role"),
            "scene_context": et.get("reason"),
            "ideal_soundbite": et.get("ideal_soundbite"),
            "questions": et.get("suggested_questions", []),
            "candidates": [],
            "status": "planning",
            "production_status": "pending_audio",
            "source": "script_import",
            "created_at": datetime.utcnow().isoformat(),
        })
        interviews_saved.append(doc)

    return jsonify({
        "episode_summary": extracted.get("episode_summary"),
        "key_facts": extracted.get("key_facts", []),
        "locations_mentioned": extracted.get("locations_mentioned", []),
        "themes": extracted.get("themes", []),
        "visual_ideas": extracted.get("visual_ideas", []),
        "research_topics_saved": len(research_saved),
        "archive_needs_saved": len(assets_saved),
        "experts_saved": len(interviews_saved),
        "next_phase": "archive",
    })


# ============== Production Episode Override (Eps 1-3, Abandoned Series) ==============
# These replace the generic 30-episode seed data for the three episodes actively in production.
PRODUCTION_EPISODES_OVERRIDE = [
    {
        "episode_number": 1,
        "title": "Military Secrets",
        "theme": "Secret military testing, toxic contamination, and government cover-ups",
        "arc": "Three sites where the US military poisoned its own people and lied about it. From nuclear test pilots drinking radioactive water to the worst nuclear accident in US history happening in a Los Angeles suburb — the desert still holds the evidence.",
        "score_visual": 9, "score_history": 10, "score_emotion": 10, "score_wtf": 10, "score_total": 39,
        "locked": True,
        "stories": [
            {
                "slot": "A",
                "location_name": "George Air Force Base (Toxic/Nuclear)",
                "state_country": "Victorville, California",
                "has_kevin_footage": True,
                "kevin_video_title": "Exploring an Abandoned Military Base in California",
                "kevin_views": 10600,
                "hook": "Pilots were ordered to fly through nuclear mushroom clouds and bring their contaminated jets back to George AFB for 'decontamination' — which meant hosing radioactive water into the ground where military families lived. Children played on contaminated playgrounds for decades. The Air Force knew. They told Congress there were 'no completed exposure pathways.' They lied.",
                "themes": ["nuclear contamination", "government cover-up", "military", "Cold War"],
                "contributors": [
                    {"role": "Nuclear Historian", "name_suggestion": "Alex Wellerstein (Stevens Institute)", "soundbite": "On the scale of the nuclear contamination and what the Air Force knew"},
                    {"role": "Environmental Lawyer", "name_suggestion": "Robert F. Kennedy Jr. / Erin Brockovich", "soundbite": "On the legal battle to force cleanup"},
                    {"role": "Former Base Resident / Veteran", "name_suggestion": "Community member affected", "soundbite": "Personal testimony of health impacts"},
                ],
                "visual_ideas": [
                    {"type": "ai_reconstruction", "description": "1950s-60s pilots in full flight suits walking toward F-86 Sabres, mushroom cloud visible on horizon"},
                    {"type": "ai_drone", "description": "Slow fly-through of the abandoned George AFB base — crumbling barracks, overgrown runways"},
                    {"type": "archive_transition", "description": "Archival photo of bustling base life → same location today, abandoned and fenced off"},
                ],
                "research_summary": "George Air Force Base operated 1941-1992. During atomic testing at Nevada Test Site, pilots at George AFB flew sampling missions through nuclear clouds. Contaminated aircraft returned to base; decontamination wash water contaminated groundwater. The base became an EPA Superfund site in 1990. Multiple studies link base contamination to elevated cancer rates in surrounding communities.",
            },
            {
                "slot": "B",
                "location_name": "Santa Susana Field Laboratory",
                "state_country": "Simi Valley, California",
                "has_kevin_footage": False,
                "kevin_video_title": None,
                "kevin_views": 0,
                "hook": "The worst nuclear accident in US history happened in a suburb of Los Angeles in 1959. It was ten times larger than Three Mile Island. The government and Boeing kept it secret for 20 years. The cleanup still hasn't happened. Children grew up downwind and wondered why so many of them got cancer.",
                "themes": ["nuclear accident", "government cover-up", "environmental contamination", "corporate negligence"],
                "contributors": [
                    {"role": "Nuclear Engineer / Physicist", "name_suggestion": "Arjun Makhijani (IEER)", "soundbite": "On the technical scale of the SSFL accident vs Three Mile Island"},
                    {"role": "Epidemiologist", "name_suggestion": "Researcher who studied cancer rates in Simi Valley", "soundbite": "The data on elevated childhood cancer rates downwind"},
                    {"role": "Local Activist / Survivor", "name_suggestion": "Denise Duffield or Rocketdyne Watch member", "soundbite": "The community fight to force cleanup"},
                ],
                "visual_ideas": [
                    {"type": "ai_reconstruction", "description": "1959 night-time — reactor fire at SSFL, emergency workers in hazmat, the glow visible over the hills"},
                    {"type": "ai_drone", "description": "Aerial view of Santa Susana hillside — abandoned rocket test stands, scrubland, the suburban sprawl of Simi Valley below"},
                    {"type": "archive_transition", "description": "NASA rocket test footage from the 1960s → the same test stand today, rusted and overgrown"},
                ],
                "research_summary": "Santa Susana Field Laboratory (SSFL) in Simi Valley was a Boeing/NASA rocket engine test facility and home to ten nuclear reactors. In July 1959, the SRE partial meltdown released 240 times more radiation than TMI. The accident was concealed for decades. Cleanup was ordered but never fully completed. Studies show elevated cancer rates in communities within 2 miles.",
            },
            {
                "slot": "C",
                "location_name": "George Air Force Base (Hospital / Birth Defects)",
                "state_country": "Victorville, California",
                "has_kevin_footage": True,
                "kevin_video_title": "Abandoned Military Base Hospital in California",
                "kevin_views": 10600,
                "hook": "Women stationed at George AFB were quietly told 'don't get pregnant.' Seven babies died in one year. Children were born with rare cancers and birth defects. The hospital records show the base medical staff suspected the water supply but the Air Force told families the water was safe. The hospital closed. The records were classified.",
                "themes": ["birth defects", "infant mortality", "toxic water", "military cover-up"],
                "contributors": [
                    {"role": "Paediatric Oncologist", "name_suggestion": "Expert in childhood cancer clusters near military bases", "soundbite": "What elevated infant mortality and birth defects indicate about environmental contamination"},
                    {"role": "Former Military Family Member", "name_suggestion": "Family who lost a child while stationed at George AFB", "soundbite": "Personal testimony"},
                    {"role": "Environmental Health Scientist", "name_suggestion": "Expert in TCE and PCE contamination effects", "soundbite": "The chemicals found in the water and their known health effects"},
                ],
                "visual_ideas": [
                    {"type": "ai_reconstruction", "description": "1970s hospital ward — military family with newborn, the clinical sterility contrasting with fear"},
                    {"type": "ai_drone", "description": "The abandoned base hospital building — broken windows, overgrown entrance, medical equipment still inside"},
                    {"type": "talking_head", "description": "AI-generated hospital commander in period uniform explaining water was 'tested and safe'"},
                ],
                "research_summary": "George AFB base housing sat atop a contaminated aquifer containing TCE, PCE, and other industrial solvents from base operations. The base hospital recorded unusual rates of miscarriage, stillbirth, and infant illness. Internal documents obtained under FOIA show base commanders were aware of contamination but did not inform families. The base was placed on the National Priorities List in 1990.",
            },
            {
                "slot": "D",
                "location_name": "Davis-Monthan AFB Airplane Boneyard",
                "state_country": "Tucson, Arizona",
                "has_kevin_footage": False,
                "kevin_video_title": None,
                "kevin_views": 0,
                "hook": "4,400 aircraft parked in the Sonoran Desert — the world's largest aircraft boneyard. The Enola Gay came here. B-52s that dropped bombs over Vietnam sit wingtip-to-wingtip with spy planes and presidential jets. Where American air power goes to die. The perfect ending for a military secrets episode.",
                "themes": ["Cold War", "American military decline", "nuclear history", "aviation history"],
                "contributors": [
                    {"role": "AMARG Historian / Tour Guide", "name_suggestion": "309th AMARG official historian", "soundbite": "On the history of the boneyard and what it represents"},
                    {"role": "Military Aviation Historian", "name_suggestion": "Peter Davies or similar", "soundbite": "On the planes that passed through — especially the Enola Gay"},
                    {"role": "Former Pilot", "name_suggestion": "B-52 veteran", "soundbite": "What it feels like to see your aircraft rotting in the desert"},
                ],
                "visual_ideas": [
                    {"type": "ai_drone", "description": "Sweeping aerial shot of row after row of aircraft stretching to the horizon — scale incomprehensible from ground level"},
                    {"type": "archive_transition", "description": "Archival photo of the Enola Gay in flight → same aircraft in the boneyard, faded and forgotten"},
                    {"type": "ai_reconstruction", "description": "The Enola Gay landing at Davis-Monthan, crew walking away, the plane joining the rows of others"},
                ],
                "research_summary": "Davis-Monthan Air Force Base in Tucson hosts AMARG (Aerospace Maintenance and Regeneration Group), the world's largest aircraft storage and preservation facility. Over 4,400 aircraft are stored in the Sonoran Desert. The dry climate and alkaline soil preserve the aircraft. The Enola Gay was initially stored here before transfer to the Smithsonian. Aircraft range from WWII fighters to Cold War bombers to retired presidential transports.",
            },
        ],
        "episode_structure": {
            "cold_open": "Drone shot over desert at dawn. A vast military base, perfectly preserved — but empty. Rows of identical family homes, grass long dead. A hospital with the lights still off. NARRATOR: 'They told the families the water was safe. They lied.'",
            "act_1": "George AFB: establish the base, its nuclear testing mission, the pilots flying through mushroom clouds. The contamination of the groundwater. The Air Force's internal knowledge.",
            "act_2": "Santa Susana: the secret accident of 1959. The cover-up. The communities downwind. The cancer rates. The cleanup that never happened.",
            "act_3": "George AFB Hospital: the birth defects, the infant deaths, the classified records. The families who still don't have answers.",
            "coda": "Davis-Monthan Boneyard: where the aircraft go to rust in the desert. The Enola Gay, forgotten. A meditation on what American military power costs — and who pays.",
        },
    },
    {
        "episode_number": 2,
        "title": "Horror Hospitals",
        "theme": "Institutional cruelty, the history of mental health 'treatment', and the people we chose to forget",
        "arc": "Four asylums on two continents where the vulnerable were sent away and subjected to horrors in the name of medicine. Lobotomies, slave labour, unconsented vaccine trials, ward fires. The buildings still stand. The graves are still there. We just stopped looking.",
        "score_visual": 9, "score_history": 9, "score_emotion": 10, "score_wtf": 9, "score_total": 37,
        "locked": True,
        "stories": [
            {
                "slot": "A",
                "location_name": "Northern State Hospital",
                "state_country": "Sedro-Woolley, Washington",
                "has_kevin_footage": False,
                "kevin_video_title": None,
                "kevin_views": 0,
                "hook": "2,700 patients are buried in numbered plots in a field behind what is now a farming museum. They were given lobotomies with an icepick through the eye socket, put to work on a 1,000-acre slave-labour farm, and when they died they weren't even given their own names on the headstones. Just numbers.",
                "themes": ["lobotomy", "institutional cruelty", "patient abuse", "forgotten history"],
                "contributors": [
                    {"role": "Psychiatric Historian", "name_suggestion": "Andrew Scull (UC San Diego) or similar", "soundbite": "On the evolution — and atrocities — of asylum-era psychiatry"},
                    {"role": "Genealogist / Descendant Advocate", "name_suggestion": "Researcher working to identify the numbered graves", "soundbite": "On the effort to give names back to the patients buried as numbers"},
                    {"role": "Former State Hospital Employee", "name_suggestion": "Someone who worked at Northern State in its final years", "soundbite": "What the wards were actually like"},
                ],
                "visual_ideas": [
                    {"type": "ai_reconstruction", "description": "1940s ward: rows of metal beds, patients in identical grey pyjamas, a nurse in uniform walking the aisle — clinical, cold, dehumanising"},
                    {"type": "archive_transition", "description": "Archival photo of the grand Northern State campus → same buildings today as a farming museum, families having picnics"},
                    {"type": "ai_drone", "description": "Slow aerial pull-back from the numbered grave markers in the field — 2,700 white markers stretching into the distance"},
                ],
                "research_summary": "Northern State Hospital operated 1912-1973 in Sedro-Woolley, Washington. At peak capacity it held 2,700 patients who were put to work on a 1,000-acre working farm. Walter Freeman performed his icepick lobotomy procedure at Northern State. Upon death, patients were buried in the cemetery as numbered markers rather than named headstones. The site is now a farming museum; the cemetery is maintained by volunteers working to identify the numbered graves.",
            },
            {
                "slot": "B",
                "location_name": "Letchworth Village",
                "state_country": "Thiells, New York",
                "has_kevin_footage": False,
                "kevin_video_title": None,
                "kevin_views": 0,
                "hook": "Letchworth Village was designed as a utopia — a 'village' for disabled children, with cottages and fresh air and gardens. In 1950, doctors secretly used disabled children as test subjects for the first polio vaccine trial without parental consent. The building that housed the vaccine trial is still standing. Nobody put a plaque on it.",
                "themes": ["medical experimentation", "disability rights", "unethical research", "institutional neglect"],
                "contributors": [
                    {"role": "Medical Ethicist", "name_suggestion": "Expert in the history of medical experimentation on institutionalised patients", "soundbite": "On the ethical breach of the Letchworth polio trial and why it set the standard for consent reform"},
                    {"role": "Historian of Disability", "name_suggestion": "Kim Nielsen or similar", "soundbite": "On how disability was treated in mid-20th century America"},
                    {"role": "Descendant / Family Member", "name_suggestion": "Family of a Letchworth resident", "soundbite": "What they were told vs what actually happened"},
                ],
                "visual_ideas": [
                    {"type": "ai_reconstruction", "description": "1950 — a researcher in a white coat, clipboard in hand, approaching a row of children in a ward. The children's faces — not understanding what is about to happen."},
                    {"type": "archive_transition", "description": "1920s promotional photo of Letchworth Village as a garden utopia → the same cottage, now ruined and overgrown"},
                    {"type": "ai_drone", "description": "Aerial shot of the abandoned Letchworth Village campus — dozens of buildings, trees growing through roofs"},
                ],
                "research_summary": "Letchworth Village opened 1911 in Thiells, New York, as a 'planned community' for people with intellectual disabilities. In 1950, Dr Hilary Koprowski administered the first polio vaccine to residents without obtaining parental consent — a landmark in the history of unethical medical experimentation. By the 1970s, journalist Geraldo Rivera exposed horrific overcrowding and neglect. The facility closed in 1996 and has been largely abandoned since.",
            },
            {
                "slot": "C",
                "location_name": "Lennox Castle Hospital",
                "state_country": "Lennoxtown, Scotland",
                "has_kevin_footage": False,
                "kevin_video_title": None,
                "kevin_views": 0,
                "hook": "Patients at Lennox Castle slept 50 to a ward in a castle that hadn't been renovated since the 1930s, wore the same pyjamas every day for years, and in 1995 a nurse was convicted of murdering a patient. The castle sits empty in the Scottish countryside. Local people say they can still hear sounds at night.",
                "themes": ["institutional neglect", "patient murder", "Scotland", "Gothic horror", "NHS history"],
                "contributors": [
                    {"role": "Scottish Psychiatric History Expert", "name_suggestion": "Researcher at University of Glasgow or Edinburgh", "soundbite": "On the state of Scottish mental health institutions in the 20th century"},
                    {"role": "Former Lennox Castle Staff Member or Patient", "name_suggestion": "Someone who lived or worked there", "soundbite": "First-hand account of conditions inside"},
                    {"role": "Local Historian", "name_suggestion": "Lennoxtown community historian", "soundbite": "On the castle's history and its impact on the local community"},
                ],
                "visual_ideas": [
                    {"type": "ai_drone", "description": "Gothic aerial shot of Lennox Castle at dusk — the castle silhouetted against a grey Scottish sky, the empty grounds"},
                    {"type": "ai_reconstruction", "description": "1970s ward: 50 beds in a cavernous Victorian room, minimal heating, patients in identical pyjamas"},
                    {"type": "archive_transition", "description": "Victorian-era photo of the castle as a private residence → the same castle today, windows boarded, ivy-covered"},
                ],
                "research_summary": "Lennox Castle in Lennoxtown, Scotland operated as a psychiatric hospital from 1936 to 2002. At peak it housed over 1,000 patients in grossly overcrowded conditions. A 1995 conviction saw a nurse found guilty of patient murder. Several inquiries identified systematic neglect. The castle itself — a 19th century baronial mansion — is now abandoned and subject to occasional arson attacks. It is considered one of Scotland's most dangerous abandoned buildings.",
            },
            {
                "slot": "D",
                "location_name": "Camarillo State Hospital",
                "state_country": "Camarillo, California",
                "has_kevin_footage": False,
                "kevin_video_title": None,
                "kevin_views": 0,
                "hook": "Charlie Parker was committed here in 1946 and wrote 'Relaxin' at Camarillo' about it. The patients ran the dairy farm — 'the Scary Dairy' — which supplied milk to the rest of California while the patients inside were receiving ECT and insulin coma therapy. The dairy is still there. It's a university now. The ghosts presumably stayed.",
                "themes": ["psychiatry history", "celebrity connection", "shock therapy", "haunting"],
                "contributors": [
                    {"role": "Jazz Historian", "name_suggestion": "Expert on Charlie Parker and the bebop era", "soundbite": "On Parker's time at Camarillo and how it influenced his music"},
                    {"role": "Psychiatric Historian", "name_suggestion": "Expert on ECT and insulin coma therapy", "soundbite": "On the 'treatments' used at Camarillo and what we now know about their effects"},
                    {"role": "CSU Channel Islands Archivist", "name_suggestion": "University staff member with knowledge of the hospital history", "soundbite": "What was found when CSU took over the buildings"},
                ],
                "visual_ideas": [
                    {"type": "ai_reconstruction", "description": "1946 — Charlie Parker in a hospital gown, sitting at a piano in a recreation room, composing. The institutional setting contrasting with his genius."},
                    {"type": "archive_transition", "description": "1950s photo of the dairy farm in full operation with patient-workers → the same building now as a university classroom block"},
                    {"type": "ai_drone", "description": "The Spanish Colonial Revival buildings of Camarillo — beautiful architecture that housed horrors"},
                ],
                "research_summary": "Camarillo State Hospital operated 1936-1997 in Camarillo, California. At peak it housed over 7,000 patients. Jazz musician Charlie Parker was committed in 1946 and wrote 'Relaxin' at Camarillo' during his stay. Patients were used as labour on the hospital's working dairy farm — nicknamed 'the Scary Dairy.' Treatments included ECT and insulin coma therapy. The facility is now California State University Channel Islands; the former dairy buildings remain.",
            },
        ],
        "episode_structure": {
            "cold_open": "A locked ward door. The sound of distant moaning — then silence. A numbered grave marker in an overgrown field. NARRATOR: 'We called it care. We called it treatment. We called it medicine. We sent them away so we wouldn't have to look.'",
            "act_1": "Northern State: establish the 'therapeutic community' ideal of the asylum era. The reality: lobotomies, slave labour, 2,700 numbered graves.",
            "act_2": "Letchworth Village: the utopia that became a testing ground. The children who had no say. The polio vaccine trial that nobody consented to.",
            "act_3": "Lennox Castle: the Scottish angle. The horror of 50-to-a-ward. The murder conviction. The castle that still stands empty.",
            "coda": "Camarillo: a jazz legend committed here wrote one of the most beautiful songs in American music about it. The dairy farm is a university. The patients are gone. But what happened here happened — and most of us never knew.",
        },
    },
    {
        "episode_number": 3,
        "title": "Impossible Engineering",
        "theme": "Man's obsession to conquer nature — and nature's quiet, inevitable victory",
        "arc": "A railroad through a desert fortress. A wooden bridge over an impossibly deep canyon. One man's 38-year tunnel to nowhere. A town that died the day the trains stopped. What connects them: the audacity of human ambition, and the patience of the desert.",
        "score_visual": 10, "score_history": 9, "score_emotion": 9, "score_wtf": 9, "score_total": 37,
        "locked": True,
        "stories": [
            {
                "slot": "A",
                "location_name": "The Impossible Railroad (SD&AE Railway)",
                "state_country": "Carrizo Gorge, California",
                "has_kevin_footage": False,
                "kevin_video_title": None,
                "kevin_views": 0,
                "hook": "Every engineer in America said the railroad through Carrizo Gorge was impossible. John D. Spreckels — sugar baron, owner of the San Diego Union-Tribune — spent $18 million and 20 years to prove them wrong. He succeeded. Then the floods came, the tunnels collapsed, and nature took everything back. The last train ran in 1976.",
                "themes": ["railroad history", "engineering", "ambition", "desert", "California"],
                "contributors": [
                    {"role": "Railroad Historian", "name_suggestion": "Pacific Southwest Railway Museum (PSRM) representative", "soundbite": "On why engineers called it the 'Impossible Railroad' and what it took to build it"},
                    {"role": "Civil/Railway Engineer", "name_suggestion": "Expert on early 20th century railroad construction techniques", "soundbite": "On the specific engineering challenges of Carrizo Gorge"},
                    {"role": "Local Historian", "name_suggestion": "San Diego or Imperial Valley historian", "soundbite": "On what the railroad meant for San Diego's future"},
                ],
                "visual_ideas": [
                    {"type": "ai_reconstruction", "description": "1910s: An army of workers with mule teams and dynamite blasting through the rock walls of Carrizo Gorge. Explosions. Dust. Scale."},
                    {"type": "ai_drone", "description": "The abandoned tracks snaking through the gorge — rusted rails, crumbling roadbed, the desert reclaiming every inch"},
                    {"type": "archive_transition", "description": "Archival photo of a train crossing the finished line → same location today, tracks buried under rockfall"},
                ],
                "research_summary": "The San Diego & Arizona Eastern Railway was built 1907-1919 through Carrizo Gorge in the California desert. Backed by John D. Spreckels, the line crossed 17 tunnels and required the construction of the Goat Canyon Trestle to bypass a collapsed tunnel. Total cost: $18 million (over $500m in today's money). The line suffered repeated flood damage, landslides, and tunnel collapses throughout its operational life. Passenger service ended 1951; freight service continued until 1976. The tracks are largely abandoned.",
            },
            {
                "slot": "B",
                "location_name": "Goat Canyon Trestle Bridge",
                "state_country": "Carrizo Gorge, California",
                "has_kevin_footage": False,
                "kevin_video_title": None,
                "kevin_views": 0,
                "hook": "The largest all-wood trestle bridge in the world was built from California redwood in the middle of an inaccessible desert canyon to replace a tunnel that nature destroyed. A structural engineer explains why it was designed to flex and sway in the desert heat. It still stands. Nobody can easily reach it. The desert keeps its secrets.",
                "themes": ["engineering marvel", "wood construction", "desert", "isolation", "structural engineering"],
                "contributors": [
                    {"role": "Structural Engineer", "name_suggestion": "Expert in historical timber structures", "soundbite": "Why building the world's largest wooden trestle in a desert canyon was an engineering masterstroke"},
                    {"role": "Railway Historian", "name_suggestion": "PSRM or SD&AE specialist", "soundbite": "The decision to build the trestle instead of repairing the tunnel"},
                    {"role": "Wilderness Guide / Historian", "name_suggestion": "Someone who has hiked to the trestle", "soundbite": "The experience of approaching this colossus in the middle of nowhere"},
                ],
                "visual_ideas": [
                    {"type": "ai_reconstruction", "description": "AI timelapse of the trestle being built — redwood beams lowered into the canyon, tiny figures of workers swarming the structure, day-to-night-to-day construction cycle"},
                    {"type": "ai_drone", "description": "The drone approaches the trestle from the canyon floor — pulling back to reveal its impossible height above the gorge"},
                    {"type": "ai_reconstruction", "description": "A steam train crossing the trestle for the first time — the creak of the wood, the sheer drop below, the engineer's white-knuckle grip"},
                ],
                "research_summary": "Goat Canyon Trestle, completed 1933, is the largest curved wooden trestle bridge in the world. Built from California redwood to bypass the collapsed Tunnel 15, it stands 185 feet tall and 630 feet long. Engineered to flex with temperature extremes. Accessible only on foot via a 16-mile round-trip hike. Listed on the National Register of Historic Places. Despite no maintenance for decades, the trestle remains structurally sound — a testament to redwood's durability.",
            },
            {
                "slot": "C",
                "location_name": "Burro Schmidt Tunnel",
                "state_country": "El Paso Mountains, California",
                "has_kevin_footage": False,
                "kevin_video_title": None,
                "kevin_views": 0,
                "hook": "William 'Burro' Schmidt spent 38 years digging half a mile of tunnel through solid granite — alone, by hand, with a pickaxe and a wheelbarrow and his donkeys. He started to create a shortcut for his ore. A proper road was built years before he finished. He kept digging anyway. When he finally broke through into sunlight, there was nobody there to see it.",
                "themes": ["obsession", "solitude", "Mojave Desert", "human perseverance", "pointlessness"],
                "contributors": [
                    {"role": "Mining Historian", "name_suggestion": "Mojave Desert or El Paso Mountains specialist", "soundbite": "On what Schmidt accomplished in raw physical terms — the volume of rock, the primitive tools"},
                    {"role": "Psychologist / Behavioural Expert", "name_suggestion": "Expert on obsession and intrinsic motivation", "soundbite": "On the psychology of continuing when the original purpose is gone"},
                    {"role": "Tunnel / Mining Engineer", "name_suggestion": "Expert who can speak to the technical achievement", "soundbite": "What it takes, physically and technically, to hand-drill half a mile of granite"},
                ],
                "visual_ideas": [
                    {"type": "ai_reconstruction", "description": "Burro Schmidt by candlelight deep inside the tunnel — hand-drilling a hole for dynamite, the only light a flickering candle, his breathing ragged in the silence"},
                    {"type": "ai_drone", "description": "The tunnel entrance from outside — the Mojave Desert stretching away, the dark mouth of the tunnel an incongruous scar in the mountainside"},
                    {"type": "ai_reconstruction", "description": "Schmidt emerging from the tunnel's far end for the first time — blinking in the sunlight, the empty desert before him, no celebration, no audience"},
                ],
                "research_summary": "William 'Burro' Schmidt began digging his tunnel in 1906 in the El Paso Mountains of California's Mojave Desert. His stated goal: a shortcut to haul ore to the smelter at Mojave. He dug entirely alone (with his burros) using hand tools and black powder. The tunnel is 2,087 feet long (nearly half a mile). A wagon road that rendered the tunnel unnecessary was completed in the 1920s. Schmidt broke through in 1938 — 32 years after starting. He sold the tunnel in 1954 for a small sum. It is now a tourist attraction managed by Kern County.",
            },
            {
                "slot": "D",
                "location_name": "Goffs, California",
                "state_country": "San Bernardino County, California",
                "has_kevin_footage": False,
                "kevin_video_title": None,
                "kevin_views": 0,
                "hook": "Goffs was a thriving railroad town on Route 66 — a real community with a school, a post office, a general store. When the Santa Fe Railway rerouted the line, the town was left behind almost overnight. When the new interstate bypassed Route 66, that was it. The Mojave Desert Heritage Association now maintains the old schoolhouse. The desert took everything else.",
                "themes": ["Route 66", "ghost town", "railroad history", "American decline", "desert"],
                "contributors": [
                    {"role": "Route 66 Historian", "name_suggestion": "Mojave Desert Heritage and Cultural Association (MDHCA)", "soundbite": "On the rise and fall of Goffs and what it represents about the American West"},
                    {"role": "Local Historian", "name_suggestion": "Joe de Kehoe or MDHCA representative", "soundbite": "The community effort to preserve what little remains"},
                    {"role": "Railroad Historian", "name_suggestion": "Santa Fe Railway specialist", "soundbite": "Why the railroad was rerouted and what that meant for towns like Goffs"},
                ],
                "visual_ideas": [
                    {"type": "archive_transition", "description": "1930s photo of Goffs — the schoolhouse full of children, the general store busy, trains visible in the background → same view today, the schoolhouse alone in the desert"},
                    {"type": "ai_drone", "description": "Aerial shot of Goffs at golden hour — the schoolhouse, the old station platform, the desert stretching to the horizon in every direction"},
                    {"type": "ai_reconstruction", "description": "The day the railroad left — a final freight train passing through Goffs, townspeople watching from the platform. The train doesn't stop."},
                ],
                "research_summary": "Goffs was established in 1883 as a water stop and railroad junction on the Atlantic & Pacific Railroad (later Santa Fe). During the Route 66 era it was a significant desert community. The Santa Fe rerouted its main line in the 1930s, bypassing Goffs. Interstate 40 later bypassed Route 66 entirely. Population declined from hundreds to single digits. The Mojave Desert Heritage and Cultural Association (MDHCA) rescued and restored the 1914 schoolhouse, which serves as the area's only historical museum.",
            },
        ],
        "episode_structure": {
            "cold_open": "A drone at dawn, flying at speed through the immense and jagged Carrizo Gorge. The rising sun catches the skeletal frame of the Goat Canyon Trestle. It looks alien. Impossible. NARRATOR: 'They said it couldn't be built. A railroad through a fortress of rock. So we built it anyway.' Cut to pitch black. The sound of a pickaxe. A flicker of candlelight.",
            "act_1": "The Impossible Railroad: San Diego's desperation for an eastern rail link. Spreckels' audacity. The engineers who said it couldn't be done. The army of men who proved them wrong. And then the floods.",
            "act_2": "The Goat Canyon Trestle: the solution to a collapsed tunnel. The world's largest wooden trestle built in a place nobody could reach. The structural genius of building in wood, not steel, in a desert.",
            "act_3": "Burro Schmidt: the singular, private obsession. The tunnel that outlasted its purpose. The 38-year negotiation with solid rock. The morning he finally broke through.",
            "coda": "Goffs: the answer to the episode's question. When man battles nature — sometimes he wins for a while, sometimes he loses spectacularly, but sometimes he just gets distracted and walks away. And nature quietly takes everything back.",
        },
    },
]


@app.route("/api/projects/<project_id>/import-production-episodes", methods=["POST"])
def import_production_episodes(project_id):
    """
    Replace episodes 1-3 with the actively developed production episodes.
    Moves displaced seed story cards to an unassigned pool (current_episode=None, current_slot='POOL').
    Also imports research summaries, contributors, and visual ideas into the project's collections.
    """
    # --- Step 1: Find and move displaced story cards for eps 1-3 to pool ---
    displaced_count = 0
    for ep_num in [1, 2, 3]:
        arr_docs = list(
            db.collection(COLLECTIONS["episode_arrangements"])
            .where("project_id", "==", project_id)
            .where("episode_number", "==", ep_num)
            .limit(1)
            .stream()
        )
        if arr_docs:
            arr = arr_docs[0].to_dict()
            arr_id = arr_docs[0].id
            for slot_key in ["slot_a", "slot_b", "slot_c", "slot_d"]:
                card_id = arr.get(slot_key)
                if card_id:
                    db.collection(COLLECTIONS["story_cards"]).document(card_id).update({
                        "current_episode": None,
                        "current_slot": "POOL",
                        "updated_at": datetime.utcnow().isoformat(),
                    })
                    displaced_count += 1
            # Delete the old arrangement
            db.collection(COLLECTIONS["episode_arrangements"]).document(arr_id).delete()

    # --- Step 2: Create new story cards and episode arrangements from production data ---
    story_cards_created = 0
    episodes_created = 0
    research_saved = 0
    interviews_saved = 0
    assets_saved = 0

    for ep in PRODUCTION_EPISODES_OVERRIDE:
        slot_ids = {}
        for story in ep["stories"]:
            card = {
                "project_id": project_id,
                "location_name": story["location_name"],
                "state_country": story["state_country"],
                "kevin_video_title": story.get("kevin_video_title"),
                "kevin_views": story.get("kevin_views", 0),
                "has_kevin_footage": story.get("has_kevin_footage", False),
                "hook": story["hook"],
                "natural_role": story["slot"],
                "themes": story.get("themes", []),
                "current_episode": ep["episode_number"],
                "current_slot": story["slot"],
                "kevin_url": None,
                "research_summary": story.get("research_summary", ""),
                "is_production": True,
                "created_at": datetime.utcnow().isoformat(),
                "updated_at": datetime.utcnow().isoformat(),
            }
            doc_ref = db.collection(COLLECTIONS["story_cards"]).document()
            doc_ref.set(card)
            slot_ids[story["slot"]] = doc_ref.id
            story_cards_created += 1

            # Save contributors as interview plans
            for c in story.get("contributors", []):
                create_doc("interviews", {
                    "project_id": project_id,
                    "topic": c.get("role"),
                    "scene_context": f"Ep{ep['episode_number']} {story['location_name']} — {c.get('name_suggestion', '')}",
                    "ideal_soundbite": c.get("soundbite", ""),
                    "questions": [],
                    "candidates": [c.get("name_suggestion", "")] if c.get("name_suggestion") else [],
                    "status": "planning",
                    "production_status": "pending_audio",
                    "source": "production_import",
                    "story_location": story["location_name"],
                    "created_at": datetime.utcnow().isoformat(),
                })
                interviews_saved += 1

            # Save visual ideas as assets
            for vi in story.get("visual_ideas", []):
                create_doc("assets", {
                    "project_id": project_id,
                    "description": vi.get("description"),
                    "asset_type": vi.get("type"),
                    "status": "needed",
                    "source": "production_import",
                    "story_location": story["location_name"],
                    "episode_number": ep["episode_number"],
                    "created_at": datetime.utcnow().isoformat(),
                })
                assets_saved += 1

            # Save research summary as a research document
            if story.get("research_summary"):
                create_doc("research_documents", {
                    "project_id": project_id,
                    "topic": story["location_name"],
                    "why_needed": f"Primary A/B/C/D story for Episode {ep['episode_number']}",
                    "summary": story["research_summary"],
                    "suggested_sources": [],
                    "source": "production_import",
                    "episode_title": ep["title"],
                    "created_at": datetime.utcnow().isoformat(),
                })
                research_saved += 1

        # Save episode structure as a research document
        if ep.get("episode_structure"):
            es = ep["episode_structure"]
            create_doc("research_documents", {
                "project_id": project_id,
                "topic": f"Episode {ep['episode_number']} Structure — {ep['title']}",
                "why_needed": "Full episode structure and act breakdown",
                "summary": f"COLD OPEN: {es.get('cold_open', '')}\n\nACT 1: {es.get('act_1', '')}\n\nACT 2: {es.get('act_2', '')}\n\nACT 3: {es.get('act_3', '')}\n\nCODA: {es.get('coda', '')}",
                "source": "production_import",
                "episode_title": ep["title"],
                "created_at": datetime.utcnow().isoformat(),
            })
            research_saved += 1

        # Create the new episode arrangement
        arr = {
            "project_id": project_id,
            "episode_number": ep["episode_number"],
            "title": ep["title"],
            "theme": ep["theme"],
            "arc": ep["arc"],
            "locked": ep["locked"],
            "score_visual": ep["score_visual"],
            "score_history": ep["score_history"],
            "score_emotion": ep["score_emotion"],
            "score_wtf": ep["score_wtf"],
            "score_total": ep["score_total"],
            "slot_a": slot_ids.get("A"),
            "slot_b": slot_ids.get("B"),
            "slot_c": slot_ids.get("C"),
            "slot_d": slot_ids.get("D"),
            "is_production": True,
            "created_at": datetime.utcnow().isoformat(),
            "updated_at": datetime.utcnow().isoformat(),
        }
        arr_ref = db.collection(COLLECTIONS["episode_arrangements"]).document()
        arr_ref.set(arr)
        episodes_created += 1

    return jsonify({
        "message": "Production episodes imported successfully",
        "episodes_imported": episodes_created,
        "story_cards_created": story_cards_created,
        "displaced_to_pool": displaced_count,
        "research_documents_saved": research_saved,
        "interview_plans_saved": interviews_saved,
        "visual_assets_saved": assets_saved,
    }), 201


# ============== Admin: Sync prod → dev collections ==============
PROD_COLLECTIONS = [
    "doc_projects", "doc_episodes", "doc_series", "doc_research",
    "doc_interviews", "doc_shots", "doc_assets", "doc_scripts",
    "doc_feedback", "doc_research_documents", "doc_archive_logs",
    "doc_interview_transcripts", "doc_script_versions",
    "doc_compliance_items", "doc_agent_tasks", "doc_users", "doc_templates",
]


@app.route("/api/admin/sync-from-prod", methods=["POST"])
def admin_sync_from_prod():
    """Copy all docs from prod collections into dev_ prefixed collections,
    and sync GCS bucket contents from prod to dev.
    Only available when APP_ENV=dev. Skips items that already exist."""
    if APP_ENV != "dev":
        return jsonify({"error": "Sync only available in dev environment"}), 403

    # ---- Firestore sync ----
    firestore_results = []
    total_copied = 0
    total_skipped = 0

    for prod_col in PROD_COLLECTIONS:
        dev_col = f"dev_{prod_col}"
        prod_docs = list(db.collection(prod_col).stream())
        existing_ids = {doc.id for doc in db.collection(dev_col).stream()}

        copied = 0
        skipped = 0
        batch = db.batch()
        batch_count = 0

        for doc in prod_docs:
            if doc.id in existing_ids:
                skipped += 1
                continue
            batch.set(db.collection(dev_col).document(doc.id), doc.to_dict())
            batch_count += 1
            copied += 1
            if batch_count >= 500:
                batch.commit()
                batch = db.batch()
                batch_count = 0

        if batch_count > 0:
            batch.commit()

        firestore_results.append({"collection": prod_col, "target": dev_col,
                        "prod_count": len(prod_docs), "copied": copied, "skipped": skipped})
        total_copied += copied
        total_skipped += skipped

    # ---- GCS bucket sync ----
    prod_bucket_name = f"{PROJECT_ID}-doc-assets"
    dev_bucket_name = f"{PROJECT_ID}-doc-assets-dev"
    gcs_copied = 0
    gcs_skipped = 0
    gcs_error = None

    try:
        prod_bucket = storage_client.bucket(prod_bucket_name)
        dev_bucket = storage_client.bucket(dev_bucket_name)

        # Ensure dev bucket exists
        if not dev_bucket.exists():
            dev_bucket.create(location="us-central1")

        # Get existing dev blob names for skip check
        existing_blobs = {b.name for b in dev_bucket.list_blobs()}

        for blob in prod_bucket.list_blobs():
            if blob.name in existing_blobs:
                gcs_skipped += 1
                continue
            prod_bucket.copy_blob(blob, dev_bucket, blob.name)
            gcs_copied += 1
    except Exception as e:
        gcs_error = str(e)

    return jsonify({
        "firestore": firestore_results,
        "total_copied": total_copied,
        "total_skipped": total_skipped,
        "gcs": {
            "prod_bucket": prod_bucket_name,
            "dev_bucket": dev_bucket_name,
            "copied": gcs_copied,
            "skipped": gcs_skipped,
            "error": gcs_error,
        },
    })


@app.route("/api/admin/promote-to-prod", methods=["POST"])
def admin_promote_to_prod():
    """Trigger a production Cloud Build from the dev branch.
    Only available when APP_ENV=dev. Runs the unified cloudbuild.yaml
    with _APP_ENV=prod using the Cloud Build triggers API."""
    if APP_ENV != "dev":
        return jsonify({"error": "Promote only available in dev environment"}), 403

    import google.auth
    import google.auth.transport.requests as gauth_requests

    trigger_id = "718e788e-df9d-49da-a3ad-be79405bf1a3"  # aim-documentary-studio (^main$ trigger)
    url = f"https://cloudbuild.googleapis.com/v1/projects/{PROJECT_ID}/triggers/{trigger_id}:run"

    credentials, _ = google.auth.default(scopes=["https://www.googleapis.com/auth/cloud-platform"])
    auth_req = gauth_requests.Request()
    credentials.refresh(auth_req)

    # Request body is a RepoSource object for triggers:run
    resp = requests.post(url, headers={
        "Authorization": f"Bearer {credentials.token}",
        "Content-Type": "application/json",
    }, json={
        "projectId": PROJECT_ID,
        "branchName": "dev",
        "substitutions": {
            "_APP_ENV": "prod",
        },
    })

    if resp.status_code >= 400:
        return jsonify({"error": f"Cloud Build API error: {resp.text}"}), resp.status_code

    build_data = resp.json()
    build_id = build_data.get("metadata", {}).get("build", {}).get("id", "unknown")
    log_url = build_data.get("metadata", {}).get("build", {}).get("logUrl", "")

    return jsonify({
        "message": "Production build triggered from dev branch",
        "build_id": build_id,
        "log_url": log_url,
    })


# ============== Style Lab ==============

def _run_style_analysis(reference_id, gcs_uri, mime_type, batch_id=None, series_id=''):
    """Run four-pass Gemini style analysis on a video reference. Safe to call in a thread."""
    try:
        # --- Pass 1: Beat sheet extraction ---
        update_doc('style_references', reference_id, {'analysis_status': 'pass1_running'})

        video_part = Part.from_uri(uri=gcs_uri, mime_type=mime_type)
        pass1_response = model.generate_content(
            [video_part, STYLE_LAB_PASS1_PROMPT],
            generation_config={"max_output_tokens": 8192, "temperature": 0.2}
        )
        beat_sheet_raw = clean_ai_response(pass1_response.text)
        beat_sheet = json.loads(beat_sheet_raw)

        update_doc('style_references', reference_id, {
            'beat_sheet': beat_sheet,
            'analysis_status': 'pass1_complete',
        })

        # Update batch progress after pass 1
        if batch_id:
            batch_ref = db.collection(COLLECTIONS['style_batches']).document(batch_id)
            batch_snap = batch_ref.get()
            if batch_snap.exists:
                bdata = batch_snap.to_dict()
                results = bdata.get('results', [])
                results.append({'referenceId': reference_id, 'pass': 1, 'status': 'pass1_complete'})
                batch_ref.update({
                    'results': results,
                    'updatedAt': datetime.utcnow().isoformat(),
                })

        # --- Pass 2: Five-pillar deep analysis ---
        update_doc('style_references', reference_id, {'analysis_status': 'pass2_running'})

        pass2_prompt = STYLE_LAB_PASS2_PROMPT.format(beat_sheet_json=json.dumps(beat_sheet, indent=2))
        pass2_response = model.generate_content(
            [video_part, pass2_prompt],
            generation_config={"max_output_tokens": 8192, "temperature": 0.2}
        )
        pillars_raw = clean_ai_response(pass2_response.text)
        pillars = json.loads(pillars_raw)

        update_doc('style_references', reference_id, {
            'pillars': pillars,
            'analysis_status': 'pass2_complete',
        })

        # Update batch progress after pass 2
        if batch_id:
            batch_ref = db.collection(COLLECTIONS['style_batches']).document(batch_id)
            batch_snap = batch_ref.get()
            if batch_snap.exists:
                bdata = batch_snap.to_dict()
                results = bdata.get('results', [])
                results.append({'referenceId': reference_id, 'pass': 2, 'status': 'pass2_complete'})
                batch_ref.update({
                    'results': results,
                    'updatedAt': datetime.utcnow().isoformat(),
                })

        # --- Pass 3: Content layer (opening, interviews, archive) ---
        update_doc('style_references', reference_id, {'analysis_status': 'pass3_running'})

        pass3_prompt = STYLE_LAB_PASS3_PROMPT.format(beat_sheet_json=json.dumps(beat_sheet, indent=2))
        pass3_response = model.generate_content(
            [video_part, pass3_prompt],
            generation_config={"max_output_tokens": 8192, "temperature": 0.2}
        )
        pass3_raw = clean_ai_response(pass3_response.text)
        pass3_pillars = json.loads(pass3_raw)

        # Merge pass 3 pillars into existing pillars
        pillars.update(pass3_pillars)
        update_doc('style_references', reference_id, {
            'pillars': pillars,
            'analysis_status': 'pass3_complete',
        })

        # Update batch progress after pass 3
        if batch_id:
            batch_ref = db.collection(COLLECTIONS['style_batches']).document(batch_id)
            batch_snap = batch_ref.get()
            if batch_snap.exists:
                bdata = batch_snap.to_dict()
                results = bdata.get('results', [])
                results.append({'referenceId': reference_id, 'pass': 3, 'status': 'pass3_complete'})
                batch_ref.update({
                    'results': results,
                    'updatedAt': datetime.utcnow().isoformat(),
                })

        # --- Pass 4: Production layer (sound, story engine, ad breaks, fingerprint) ---
        update_doc('style_references', reference_id, {'analysis_status': 'pass4_running'})

        # Look up ad break config from series config
        ad_break_config = "No ad break configuration available."
        if series_id:
            try:
                series_config_doc = db.collection(COLLECTIONS['series_config']).document(series_id).get()
                if series_config_doc.exists:
                    sc = series_config_doc.to_dict()
                    # Check for development brief with ad break info
                    has_ad_breaks = sc.get('has_ad_breaks', False)
                    ad_breaks_per_part = sc.get('ad_breaks_per_part', 0)
                    ad_break_config = f"has_ad_breaks: {has_ad_breaks}, ad_breaks_per_part: {ad_breaks_per_part}"
            except Exception:
                pass

        pass4_prompt = STYLE_LAB_PASS4_PROMPT.format(
            beat_sheet_json=json.dumps(beat_sheet, indent=2),
            pillars_json=json.dumps({k: v for k, v in pillars.items() if k in ('narrative_pacing', 'visual_language', 'editing_rhythms', 'tone_atmosphere', 'reverse_engineering')}, indent=2),
            pass3_json=json.dumps(pass3_pillars, indent=2),
            ad_break_config=ad_break_config,
        )
        pass4_response = model.generate_content(
            [video_part, pass4_prompt],
            generation_config={"max_output_tokens": 8192, "temperature": 0.2}
        )
        pass4_raw = clean_ai_response(pass4_response.text)
        pass4_pillars = json.loads(pass4_raw)

        # Merge pass 4 pillars into existing pillars
        pillars.update(pass4_pillars)
        update_doc('style_references', reference_id, {
            'pillars': pillars,
            'analysis_status': 'complete',
        })

        # Update batch progress after pass 4 (final)
        if batch_id:
            batch_ref = db.collection(COLLECTIONS['style_batches']).document(batch_id)
            batch_snap = batch_ref.get()
            if batch_snap.exists:
                bdata = batch_snap.to_dict()
                completed = bdata.get('completed', 0) + 1
                total = bdata.get('total', 1)
                results = bdata.get('results', [])
                results.append({'referenceId': reference_id, 'pass': 4, 'status': 'complete'})
                batch_ref.update({
                    'completed': completed,
                    'status': 'complete' if completed >= total else 'running',
                    'results': results,
                    'updatedAt': datetime.utcnow().isoformat(),
                })

    except Exception as e:
        print(f"[style-lab] Analysis failed for {reference_id}: {e}")
        try:
            update_doc('style_references', reference_id, {
                'analysis_status': 'error',
                'analysis_error': str(e),
            })
        except Exception:
            pass

        if batch_id:
            try:
                batch_ref = db.collection(COLLECTIONS['style_batches']).document(batch_id)
                batch_snap = batch_ref.get()
                if batch_snap.exists:
                    bdata = batch_snap.to_dict()
                    completed = bdata.get('completed', 0) + 1
                    total = bdata.get('total', 1)
                    results = bdata.get('results', [])
                    results.append({'referenceId': reference_id, 'status': 'error', 'error': str(e)})
                    batch_ref.update({
                        'completed': completed,
                        'status': 'complete' if completed >= total else 'running',
                        'results': results,
                        'updatedAt': datetime.utcnow().isoformat(),
                    })
            except Exception:
                pass


@app.route("/api/style-lab/upload", methods=["POST"])
def style_lab_upload():
    """Upload a video file as a style reference."""
    if 'file' not in request.files:
        return jsonify({"error": "No file provided"}), 400

    file = request.files['file']
    if file.filename == '':
        return jsonify({"error": "No file selected"}), 400

    project_id = request.form.get('projectId')
    series_id = request.form.get('seriesId', '')
    name = request.form.get('name', file.filename)

    if not project_id:
        return jsonify({"error": "projectId is required"}), 400

    try:
        file_content = file.read()
        content_type = file.content_type or 'video/mp4'
        ext = os.path.splitext(file.filename)[1] or '.mp4'
        unique_id = uuid.uuid4().hex
        blob_path = f"style-refs/{project_id}/{unique_id}{ext}"

        ensure_bucket_exists(STORAGE_BUCKET)
        bucket = storage_client.bucket(STORAGE_BUCKET)
        blob = bucket.blob(blob_path)
        blob.upload_from_string(file_content, content_type=content_type)

        gcs_uri = f"gs://{STORAGE_BUCKET}/{blob_path}"

        ref_doc = create_doc('style_references', {
            'projectId': project_id,
            'seriesId': series_id,
            'name': name,
            'source_type': 'upload',
            'gcs_uri': gcs_uri,
            'thumbnail_url': '',
            'duration_seconds': 0,
            'mime_type': content_type,
            'analysis_status': 'pending',
        })

        return jsonify(ref_doc), 201
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/style-lab/youtube", methods=["POST"])
def style_lab_youtube():
    """Add a YouTube video as a style reference."""
    data = request.get_json()
    url = data.get('url', '').strip()
    project_id = data.get('projectId', '')
    series_id = data.get('seriesId', '')

    if not url or not project_id:
        return jsonify({"error": "url and projectId are required"}), 400

    try:
        gcs_uri, info = _download_and_upload_to_gcs(url, project_id, 'style-ref')
        name = info.get('title', url)[:200]
        duration = info.get('duration', 0) or 0
        thumbnail = info.get('thumbnail', '') or ''

        ref_doc = create_doc('style_references', {
            'projectId': project_id,
            'seriesId': series_id,
            'name': name,
            'source_type': 'youtube',
            'youtube_url': url,
            'gcs_uri': gcs_uri,
            'thumbnail_url': thumbnail,
            'duration_seconds': duration,
            'mime_type': 'video/mp4',
            'analysis_status': 'pending',
        })

        return jsonify(ref_doc), 201
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/style-lab/analyze", methods=["POST"])
def style_lab_analyze():
    """Start analysis on one or more style references. Returns 202 with batchId."""
    data = request.get_json()
    reference_id = data.get('referenceId')
    reference_ids = data.get('referenceIds', [])

    # Accept single or multiple
    if reference_id and not reference_ids:
        reference_ids = [reference_id]

    if not reference_ids:
        return jsonify({"error": "referenceId or referenceIds required"}), 400

    # Validate all references exist
    refs = []
    for rid in reference_ids:
        ref = get_doc('style_references', rid)
        if not ref:
            return jsonify({"error": f"Reference {rid} not found"}), 404
        refs.append(ref)

    # Create batch document
    batch_doc = create_doc('style_batches', {
        'referenceIds': reference_ids,
        'total': len(reference_ids),
        'completed': 0,
        'status': 'running',
        'results': [],
    })
    batch_id = batch_doc['id']

    def run_batch():
        for ref in refs:
            _run_style_analysis(
                reference_id=ref['id'],
                gcs_uri=ref['gcs_uri'],
                mime_type=ref.get('mime_type', 'video/mp4'),
                batch_id=batch_id,
                series_id=ref.get('seriesId', ''),
            )

    thread = threading.Thread(target=run_batch, daemon=True)
    thread.start()

    return jsonify({"batchId": batch_id, "total": len(reference_ids), "status": "running"}), 202


@app.route("/api/style-lab/batch/<batch_id>", methods=["GET"])
def style_lab_batch_status(batch_id):
    """Get style analysis batch progress."""
    doc = get_doc('style_batches', batch_id)
    if not doc:
        return jsonify({"error": "Batch not found"}), 404
    return jsonify(doc), 200


@app.route("/api/projects/<project_id>/style-references", methods=["GET"])
def get_style_references(project_id):
    """List all style references for a project."""
    docs = db.collection(COLLECTIONS['style_references']) \
              .where('projectId', '==', project_id) \
              .stream()
    return jsonify([{**d.to_dict(), "id": d.id} for d in docs]), 200


@app.route("/api/style-lab/reference/<reference_id>", methods=["GET"])
def get_style_reference(reference_id):
    """Get a single style reference with full analysis data."""
    doc = get_doc('style_references', reference_id)
    if not doc:
        return jsonify({"error": "Reference not found"}), 404
    return jsonify(doc), 200


@app.route("/api/style-lab/reference/<reference_id>", methods=["DELETE"])
def delete_style_reference(reference_id):
    """Delete a style reference and its GCS file."""
    doc = get_doc('style_references', reference_id)
    if not doc:
        return jsonify({"error": "Reference not found"}), 404

    # Delete GCS file
    gcs_uri = doc.get('gcs_uri', '')
    if gcs_uri.startswith('gs://'):
        try:
            # Parse gs://bucket/path
            parts = gcs_uri.replace('gs://', '').split('/', 1)
            bucket_name = parts[0]
            blob_path = parts[1] if len(parts) > 1 else ''
            bucket = storage_client.bucket(bucket_name)
            blob = bucket.blob(blob_path)
            if blob.exists():
                blob.delete()
        except Exception as e:
            print(f"[style-lab] Error deleting GCS file: {e}")

    delete_doc('style_references', reference_id)
    return jsonify({"success": True}), 200


@app.route("/api/series-config/<series_id>/style-dna", methods=["PUT"])
def put_style_dna(series_id):
    """Save Style DNA to an existing series config."""
    data = request.get_json()
    selections = data.get('selections', {})
    composite_summary = data.get('composite_summary', '')
    script_directives = data.get('script_directives', {})

    # Check that series config exists
    existing = get_doc('series_config', series_id)
    if not existing:
        return jsonify({"error": "Series config not found"}), 404

    style_dna = {
        'selections': selections,
        'composite_summary': composite_summary,
        'script_directives': script_directives,
        'updatedAt': datetime.utcnow().isoformat(),
    }

    update_doc('series_config', series_id, {'styleDna': style_dna})
    return jsonify({"success": True, "styleDna": style_dna}), 200


@app.route("/api/series-config/<series_id>/style-dna", methods=["GET"])
def get_style_dna(series_id):
    """Get Style DNA from a series config."""
    existing = get_doc('series_config', series_id)
    if not existing:
        return jsonify({"error": "Series config not found"}), 404

    style_dna = existing.get('styleDna', {})
    return jsonify(style_dna), 200


@app.route("/api/style-lab/generate-composite", methods=["POST"])
def style_lab_generate_composite():
    """Generate a composite summary and script directives from selected pillar data."""
    data = request.get_json()
    selections = data.get('selections', {})

    if not selections:
        return jsonify({"error": "selections is required"}), 400

    # Build context from selections
    pillar_context_parts = []
    for pillar_key, sel in selections.items():
        pillar_data = sel.get('data', {})
        source_id = sel.get('source_reference_id', 'unknown')
        pillar_context_parts.append(
            f"=== {pillar_key.upper()} (from reference {source_id}) ===\n"
            f"{json.dumps(pillar_data, indent=2)}"
        )
    pillar_context = "\n\n".join(pillar_context_parts)

    composite_prompt = f"""You are a documentary style consultant. You have been given analysis data from reference videos across up to twelve style pillars. Your job is to synthesize these into a unified Style DNA profile.

SELECTED PILLAR DATA:
{pillar_context}

Generate two things:

1. A COMPOSITE SUMMARY — a natural language synthesis (3-5 paragraphs) that describes the overall documentary style defined by these selections. Write it as if briefing a director and editor on the visual and editorial language of the series. Be specific, concrete, and actionable.

2. SCRIPT DIRECTIVES — exactly 12 instruction strings that can be injected into a script generation prompt. Each must be a clear, direct instruction. If a pillar is not present in the selections, write "N/A" for that directive:
   - Pacing directive (how to pace the narrative)
   - Visual directive (shot types, camera work, framing)
   - Editing directive (cut rhythm, transitions, montage)
   - Tone directive (mood, narration style, atmosphere)
   - Production directive (overall production approach, signature techniques)
   - Opening directive (how to construct the opening 90-120 seconds)
   - Interview directive (how to deploy and structure interviews)
   - Archive directive (how to use archive footage as storytelling)
   - Sound & Score directive (sound design layers, music approach)
   - Story Engine directive (narrative engine, tension mechanics)
   - Ad Break directive (cliffhanger construction, bump structure)
   - Comparative Reference (pitch-deck-ready style fingerprint)

Return ONLY valid JSON:
{{
  "composite_summary": "string — the full synthesis",
  "script_directives": {{
    "pacing_instruction": "string — how to pace the narrative",
    "visual_instruction": "string — shot types, camera work, framing",
    "editing_instruction": "string — cut rhythm, transitions, montage",
    "tone_instruction": "string — mood, narration style, atmosphere",
    "production_instruction": "string — overall production approach, signature techniques",
    "opening_instruction": "string — how to construct the opening",
    "interview_instruction": "string — how to deploy interviews",
    "archive_instruction": "string — how to use archive footage",
    "sound_score_instruction": "string — sound design and music approach",
    "story_engine_instruction": "string — narrative engine and tension mechanics",
    "ad_break_instruction": "string — ad break cliffhanger and bump structure",
    "comparative_reference": "string — pitch-deck-ready style fingerprint"
  }}
}}"""

    try:
        response = model.generate_content(
            [composite_prompt],
            generation_config={"max_output_tokens": 4096, "temperature": 0.3}
        )
        raw = clean_ai_response(response.text)
        result = json.loads(raw)
        return jsonify(result), 200
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/style-lab/drill-down", methods=["POST"])
def style_lab_drill_down():
    """Ask a follow-up question about a specific pillar of a reference video."""
    data = request.get_json()
    reference_id = data.get('referenceId', '')
    pillar = data.get('pillar', '')
    question = data.get('question', '')
    timecode = data.get('timecode', '')

    if not reference_id or not question:
        return jsonify({"error": "referenceId and question are required"}), 400

    ref = get_doc('style_references', reference_id)
    if not ref:
        return jsonify({"error": "Reference not found"}), 404

    gcs_uri = ref.get('gcs_uri', '')
    mime_type = ref.get('mime_type', 'video/mp4')

    # Build context from existing analysis
    pillar_context = ""
    if pillar and ref.get('pillars', {}).get(pillar):
        pillar_context = f"\nExisting analysis for {pillar}:\n{json.dumps(ref['pillars'][pillar], indent=2)}\n"

    beat_sheet_context = ""
    if ref.get('beat_sheet'):
        beat_sheet_context = f"\nBeat sheet:\n{json.dumps(ref['beat_sheet'], indent=2)}\n"

    timecode_instruction = ""
    if timecode:
        timecode_instruction = f"\nFocus especially on the section around timecode {timecode}."

    drill_prompt = f"""You are a documentary style analyst. A user is asking a follow-up question about this video.
{pillar_context}{beat_sheet_context}{timecode_instruction}

USER QUESTION: {question}

Provide a detailed, specific answer based on the video content. Reference specific timecodes where relevant. Return plain text (not JSON)."""

    try:
        video_part = Part.from_uri(uri=gcs_uri, mime_type=mime_type)
        response = model.generate_content(
            [video_part, drill_prompt],
            generation_config={"max_output_tokens": 4096, "temperature": 0.3}
        )
        return jsonify({"answer": response.text}), 200
    except Exception as e:
        return jsonify({"error": str(e)}), 500


# ============== YouTube Transcript (Free Captions) ==============

@app.route("/api/transcribe-youtube", methods=["POST"])
def transcribe_youtube():
    """Batch-fetch YouTube auto-generated captions. No API key needed."""
    from youtube_transcript_api import YouTubeTranscriptApi
    from youtube_transcript_api._errors import (
        TranscriptsDisabled,
        NoTranscriptFound,
        VideoUnavailable,
    )

    data = request.get_json() or {}
    urls = data.get("urls", [])
    if not urls or not isinstance(urls, list):
        return jsonify({"error": "urls array required"}), 400
    if len(urls) > 50:
        return jsonify({"error": "Maximum 50 URLs per request"}), 400

    def extract_video_id(url):
        """Extract YouTube video ID from various URL formats."""
        url = url.strip()
        # youtu.be/ID
        m = re.match(r"(?:https?://)?youtu\.be/([A-Za-z0-9_-]{11})", url)
        if m:
            return m.group(1)
        # youtube.com/watch?v=ID
        m = re.match(r"(?:https?://)?(?:www\.)?youtube\.com/watch\?.*v=([A-Za-z0-9_-]{11})", url)
        if m:
            return m.group(1)
        # youtube.com/embed/ID
        m = re.match(r"(?:https?://)?(?:www\.)?youtube\.com/embed/([A-Za-z0-9_-]{11})", url)
        if m:
            return m.group(1)
        # youtube.com/shorts/ID
        m = re.match(r"(?:https?://)?(?:www\.)?youtube\.com/shorts/([A-Za-z0-9_-]{11})", url)
        if m:
            return m.group(1)
        # Bare video ID
        if re.match(r"^[A-Za-z0-9_-]{11}$", url):
            return url
        return None

    def get_video_title(video_id):
        """Get video title via YouTube oEmbed (free, no API key)."""
        try:
            resp = requests.get(
                f"https://www.youtube.com/oembed?url=https://www.youtube.com/watch?v={video_id}&format=json",
                timeout=10,
            )
            if resp.ok:
                return resp.json().get("title", f"YouTube Video {video_id}")
        except Exception:
            pass
        return f"YouTube Video {video_id}"

    # Deduplicate by video ID
    seen_ids = set()
    unique_entries = []
    for url in urls:
        vid = extract_video_id(url)
        if vid and vid not in seen_ids:
            seen_ids.add(vid)
            unique_entries.append((url, vid))

    api = YouTubeTranscriptApi()
    results = []
    for original_url, video_id in unique_entries:
        try:
            fetched = api.fetch(video_id)
            raw_data = fetched.to_raw_data()
            title = get_video_title(video_id)
            thumbnail_url = f"https://img.youtube.com/vi/{video_id}/hqdefault.jpg"

            # Build segments matching ClipTranscript type
            segments = []
            for i, entry in enumerate(raw_data):
                segments.append({
                    "id": f"seg-{i+1}",
                    "start_time": round(entry["start"], 2),
                    "end_time": round(entry["start"] + entry["duration"], 2),
                    "text": entry["text"],
                })

            full_text = " ".join(e["text"] for e in raw_data)
            total_duration = 0
            if raw_data:
                last = raw_data[-1]
                total_duration = round(last["start"] + last["duration"], 2)

            results.append({
                "video_id": video_id,
                "url": original_url,
                "title": title,
                "thumbnail_url": thumbnail_url,
                "duration_seconds": total_duration,
                "transcript": {
                    "segments": segments,
                    "full_text": full_text,
                    "language": "en",
                    "status": "complete",
                    "created_at": datetime.utcnow().isoformat(),
                },
            })
        except TranscriptsDisabled:
            results.append({
                "video_id": video_id,
                "url": original_url,
                "error": "Transcripts are disabled for this video",
            })
        except NoTranscriptFound:
            results.append({
                "video_id": video_id,
                "url": original_url,
                "error": "No transcript available for this video",
            })
        except VideoUnavailable:
            results.append({
                "video_id": video_id,
                "url": original_url,
                "error": "Video is unavailable or private",
            })
        except Exception as e:
            results.append({
                "video_id": video_id,
                "url": original_url,
                "error": str(e),
            })

    return jsonify({"results": results}), 200


# ============== Episode Fast Track Pipeline ==============

@app.route("/api/ai/extract-stories", methods=["POST"])
def extract_stories():
    """Extract A/B/C story structure from an episode brief using Gemini 3 Flash."""
    try:
        data = request.get_json() or {}
        brief_text = data.get("brief_text", "").strip()
        if not brief_text or len(brief_text) < 50:
            return jsonify({"error": "brief_text must be at least 50 characters"}), 400

        system_prompt = """You are a documentary series development AI for the show ABANDONED PLACES: UNCOVERED.
Each episode has 3 stories: A (centrepiece, 20-25 min), B (counterpoint, 15-20 min), C (gut punch, 5-10 min).
The three stories connect through THEME, not geography.

Extract the story structure from the brief below. Return valid JSON only, no markdown."""

        prompt = f"""From this episode brief, extract the story structure.

BRIEF:
{brief_text}

Return this exact JSON structure:
{{
  "stories": [
    {{
      "slot": "A",
      "title": "Short title of the A story",
      "description": "2-3 sentence description of the story",
      "themes": ["theme1", "theme2"],
      "locations": ["location1"]
    }},
    {{
      "slot": "B",
      "title": "Short title of the B story",
      "description": "2-3 sentence description",
      "themes": ["theme1", "theme2"],
      "locations": ["location1"]
    }},
    {{
      "slot": "C",
      "title": "Short title of the C story",
      "description": "2-3 sentence description",
      "themes": ["theme1"],
      "locations": ["location1"]
    }}
  ],
  "episode_theme": "The overarching thematic connection between all three stories",
  "episode_title_suggestion": "Suggested episode title"
}}

If the brief only describes fewer than 3 stories, extract what you can and leave remaining slots with descriptive placeholders based on the theme."""

        flash_model = GenerativeModel("gemini-2.0-flash-001")
        response = flash_model.generate_content(
            f"{system_prompt}\n\n{prompt}",
            generation_config={"response_mime_type": "application/json"}
        )
        raw = clean_ai_response(response.text)
        result = json.loads(raw)
        return jsonify(result), 200

    except json.JSONDecodeError as e:
        return jsonify({"error": f"Failed to parse AI response as JSON: {str(e)}"}), 500
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/ai/deep-research-fast-track", methods=["POST"])
def deep_research_fast_track():
    """Deep research for a single story using Gemini + Google Search grounding."""
    try:
        data = request.get_json() or {}
        episode_id = data.get("episode_id")
        story = data.get("story", {})
        transcript_text = data.get("transcript_text", "")
        brief_text = data.get("brief_text", "")
        beat_sheet = data.get("beat_sheet", "")

        if not episode_id:
            return jsonify({"error": "episode_id is required"}), 400
        if not story.get("title"):
            return jsonify({"error": "story.title is required"}), 400

        story_slot = story.get("slot", "A")
        story_title = story.get("title", "Untitled")
        story_desc = story.get("description", "")
        story_themes = ", ".join(story.get("themes", []))

        system_prompt = f"""You are a documentary research specialist for the show ABANDONED PLACES: UNCOVERED.
You are generating a deep research document for the {story_slot} Story: "{story_title}".

Your research must be thorough, factual, and production-ready. Use web search to find current, verified information.
Structure your output using the Research Depth Checklist below.

RESEARCH DEPTH CHECKLIST — generate each section:
1. EXECUTIVE SUMMARY (200 words)
2. HISTORICAL TIMELINE (chronological key events with dates)
3. KEY CHARACTERS (named individuals — victims, villains, heroes — with roles and fates)
4. THE DISASTER/SCANDAL (what went wrong, detailed account)
5. THE COVER-UP OR AFTERMATH (institutional response, legal outcomes)
6. HUMAN COST (named victims where possible, statistics, lasting impact)
7. EXPERT VOICES (real experts who could speak on this — names, affiliations, publications)
8. ARCHIVE SOURCES (specific archives, libraries, footage collections to pursue)
9. MODERN PARALLELS (how this connects to current events)
10. THE VILLAIN (who is responsible, what did they do, were they held accountable)
11. VISUAL OPPORTUNITIES (locations to film, objects to show, AI reconstruction possibilities)
12. UNANSWERED QUESTIONS (mysteries, disputed facts, ongoing investigations)
13. WTF MOMENTS (the details that make people say "no way" — for audience retention)
14. PRODUCTION NOTES (access issues, sensitivity considerations, legal risks)

Write with authority. Name names. Cite sources. This document will be used by producers and researchers."""

        prompt = f"""Generate a comprehensive deep research document for this story.

STORY: {story_slot} — {story_title}
DESCRIPTION: {story_desc}
THEMES: {story_themes}

EPISODE BRIEF CONTEXT:
{brief_text[:3000] if brief_text else 'No brief provided.'}

{'TRANSCRIPT CONTEXT (from related YouTube video):' + chr(10) + transcript_text[:5000] if transcript_text else ''}

{'BEAT SHEET REFERENCE:' + chr(10) + beat_sheet[:2000] if beat_sheet else ''}

Generate the full 14-section research document now. Be thorough and specific."""

        # Use Gemini 3.1 Pro with Google Search grounding for live web research
        research_model = GenerativeModel("gemini-3.1-pro-preview")
        search_tool = Tool._from_gapic(
            raw_tool=GapicTool(google_search=GapicTool.GoogleSearch())
        )
        response = research_model.generate_content(
            f"{system_prompt}\n\n{prompt}",
            tools=[search_tool],
        )
        research_text = response.text

        # Save to Firestore
        research_doc = create_doc('research_documents', {
            'episodeId': episode_id,
            'title': f'Deep Research — {story_slot}: {story_title}',
            'content': research_text,
            'documentType': 'fast_track_research',
            'storySlot': story_slot,
            'storyTitle': story_title,
            'wordCount': len(research_text.split()),
            'confidenceLevel': 'web_grounded',
            'source': 'fast_track_pipeline',
        })

        return jsonify({
            "document_id": research_doc['id'],
            "title": research_doc['title'],
            "word_count": research_doc['wordCount'],
            "preview": research_text[:500],
        }), 200

    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/ai/master-script-fast-track", methods=["POST"])
def master_script_fast_track():
    """Generate a master script from research docs + transcripts + beat sheet."""
    try:
        data = request.get_json() or {}
        episode_id = data.get("episode_id")
        episode_title = data.get("episode_title", "Untitled Episode")
        episode_theme = data.get("episode_theme", "")

        if not episode_id:
            return jsonify({"error": "episode_id is required"}), 400

        # Load research documents for this episode
        research_docs = get_docs_by_episode('research_documents', episode_id)
        research_context = ""
        for doc in research_docs:
            research_context += f"\n\n--- {doc.get('title', 'Research')} ---\n"
            research_context += doc.get('content', '')[:8000]

        if not research_context.strip():
            return jsonify({"error": "No research documents found for this episode. Run deep research first."}), 400

        # Load beat sheet from series config (Abandoned)
        beat_sheet = ""
        try:
            config_doc = db.collection(COLLECTIONS['series_config']).document('abandoned').get()
            if config_doc.exists:
                config_data = config_doc.to_dict()
                beat_sheet = config_data.get('beat_sheet', '')
        except Exception:
            pass

        # Load any existing transcripts linked to this episode
        transcript_context = ""
        try:
            transcripts = db.collection(COLLECTIONS['interview_transcripts']).where(
                'episodeId', '==', episode_id
            ).stream()
            for t in transcripts:
                td = t.to_dict()
                transcript_context += f"\n\n--- Transcript: {td.get('title', 'Unknown')} ---\n"
                transcript_context += td.get('content', '')[:5000]
        except Exception:
            pass

        system_prompt = f"""You are a master script writer for the documentary series ABANDONED PLACES: UNCOVERED.

You write in the show's distinctive voice: conversational, darkly humorous, emotionally powerful.
Short sentences for impact. Name names. Never be cruel about victims — irreverence is for the powerful.

{'BEAT SHEET (follow this structure):' + chr(10) + beat_sheet[:4000] if beat_sheet else 'Write a 45-minute documentary script with cold open, 3 acts, and synthesis.'}

SCRIPT FORMAT:
- Write in 2-column format: [VISUALS] and [SCRIPT]
- COMMENTARY (voiceover) in CAPS
- SYNC (interview bites) in lowercase with speaker name
- Include timing markers: [00:00], [02:30], etc.
- Include visual direction: AI RECONSTRUCTION, ARCHIVE, DRONE, EXPLORER FOOTAGE, etc.
- Mark each story transition clearly: === A STORY ===, === B STORY ===, === C STORY ==="""

        prompt = f"""Write a complete master script for this episode.

EPISODE TITLE: {episode_title}
EPISODE THEME: {episode_theme}

RESEARCH MATERIAL:
{research_context[:20000]}

{'TRANSCRIPT MATERIAL:' + chr(10) + transcript_context[:10000] if transcript_context else ''}

Generate the FULL master script now. Include all beats from cold open through to the next episode tease.
Make it production-ready — a producer should be able to hand this to an editor."""

        # Use the configured script model (defaults to gemini-2.5-pro)
        script_model_name = AGENT_MODELS.get('script_writer', 'gemini-2.5-pro')
        script_model = GenerativeModel(script_model_name)
        response = script_model.generate_content(
            f"{system_prompt}\n\n{prompt}",
        )
        script_text = response.text

        # Save to Firestore
        script_doc = create_doc('script_versions', {
            'episodeId': episode_id,
            'title': f'Master Script — {episode_title}',
            'content': script_text,
            'versionNumber': 1,
            'status': 'draft',
            'wordCount': len(script_text.split()),
            'source': 'fast_track_pipeline',
            'model': script_model_name,
        })

        return jsonify({
            "script_id": script_doc['id'],
            "title": script_doc['title'],
            "word_count": script_doc['wordCount'],
            "preview": script_text[:800],
        }), 200

    except Exception as e:
        return jsonify({"error": str(e)}), 500


# ============== Notification Routes ==============

@app.route("/api/notifications", methods=["POST"])
def create_notification():
    """Create a new notification."""
    data = request.get_json()
    required = ['recipientId', 'senderId', 'senderName', 'type', 'title', 'message']
    for field in required:
        if field not in data:
            return jsonify({"error": f"Missing required field: {field}"}), 400

    now = datetime.utcnow().isoformat()
    doc_data = {
        'recipientId': data['recipientId'],
        'senderId': data['senderId'],
        'senderName': data['senderName'],
        'type': data['type'],  # 'mention' | 'message' | 'system'
        'title': data['title'],
        'message': data['message'],
        'read': False,
        'context': data.get('context', {}),
        'createdAt': now,
        'updatedAt': now,
    }
    doc_ref = db.collection(COLLECTIONS['notifications']).document()
    doc_ref.set(doc_data)
    doc_data['id'] = doc_ref.id
    return jsonify(doc_data), 201


@app.route("/api/notifications/<user_id>", methods=["GET"])
def get_notifications(user_id):
    """Get notifications for a user, ordered by createdAt DESC, limit 50."""
    docs = db.collection(COLLECTIONS['notifications']) \
        .where('recipientId', '==', user_id) \
        .order_by('createdAt', direction=firestore.Query.DESCENDING) \
        .limit(50) \
        .stream()

    results = []
    for doc in docs:
        d = doc.to_dict()
        d['id'] = doc.id
        results.append(d)
    return jsonify(results)


@app.route("/api/notifications/<notif_id>/read", methods=["PUT"])
def mark_notification_read(notif_id):
    """Mark a single notification as read."""
    doc_ref = db.collection(COLLECTIONS['notifications']).document(notif_id)
    doc_ref.update({
        'read': True,
        'updatedAt': datetime.utcnow().isoformat(),
    })
    return jsonify({"success": True})


@app.route("/api/notifications/<user_id>/read-all", methods=["PUT"])
def mark_all_notifications_read(user_id):
    """Batch mark all notifications as read for a user."""
    docs = db.collection(COLLECTIONS['notifications']) \
        .where('recipientId', '==', user_id) \
        .where('read', '==', False) \
        .stream()

    batch = db.batch()
    count = 0
    now = datetime.utcnow().isoformat()
    for doc in docs:
        batch.update(doc.reference, {'read': True, 'updatedAt': now})
        count += 1
        if count % 500 == 0:  # Firestore batch limit
            batch.commit()
            batch = db.batch()
    if count % 500 != 0:
        batch.commit()
    return jsonify({"success": True, "updated": count})


# ============== Message Routes ==============

@app.route("/api/messages", methods=["POST"])
def send_message():
    """Send a direct message."""
    data = request.get_json()
    required = ['senderId', 'senderName', 'recipientId', 'recipientName', 'text', 'threadId']
    for field in required:
        if field not in data:
            return jsonify({"error": f"Missing required field: {field}"}), 400

    now = datetime.utcnow().isoformat()
    doc_data = {
        'senderId': data['senderId'],
        'senderName': data['senderName'],
        'senderAvatar': data.get('senderAvatar', ''),
        'recipientId': data['recipientId'],
        'recipientName': data['recipientName'],
        'text': data['text'],
        'read': False,
        'threadId': data['threadId'],
        'createdAt': now,
        'updatedAt': now,
    }
    doc_ref = db.collection(COLLECTIONS['messages']).document()
    doc_ref.set(doc_data)
    doc_data['id'] = doc_ref.id
    return jsonify(doc_data), 201


@app.route("/api/messages/threads/<user_id>", methods=["GET"])
def get_threads(user_id):
    """Get all message threads for a user (latest message per thread)."""
    # Get messages where user is sender
    sent = db.collection(COLLECTIONS['messages']) \
        .where('senderId', '==', user_id) \
        .stream()
    # Get messages where user is recipient
    received = db.collection(COLLECTIONS['messages']) \
        .where('recipientId', '==', user_id) \
        .stream()

    all_messages = []
    for doc in sent:
        d = doc.to_dict()
        d['id'] = doc.id
        all_messages.append(d)
    for doc in received:
        d = doc.to_dict()
        d['id'] = doc.id
        all_messages.append(d)

    # Group by threadId, keep latest message
    threads = {}
    for msg in all_messages:
        tid = msg.get('threadId', '')
        if tid not in threads or msg['createdAt'] > threads[tid]['createdAt']:
            threads[tid] = msg

    # Sort by most recent
    result = sorted(threads.values(), key=lambda m: m['createdAt'], reverse=True)

    # Add unread count per thread
    for thread in result:
        tid = thread['threadId']
        thread['unreadCount'] = sum(
            1 for m in all_messages
            if m['threadId'] == tid and m['recipientId'] == user_id and not m.get('read', True)
        )

    return jsonify(result)


@app.route("/api/messages/thread/<thread_id>", methods=["GET"])
def get_thread_messages(thread_id):
    """Get all messages in a thread, ordered chronologically."""
    docs = db.collection(COLLECTIONS['messages']) \
        .where('threadId', '==', thread_id) \
        .order_by('createdAt') \
        .stream()

    results = []
    for doc in docs:
        d = doc.to_dict()
        d['id'] = doc.id
        results.append(d)
    return jsonify(results)


@app.route("/api/messages/thread/<thread_id>/read", methods=["PUT"])
def mark_thread_read(thread_id):
    """Mark all messages in a thread as read for a specific user."""
    data = request.get_json()
    user_id = data.get('userId')
    if not user_id:
        return jsonify({"error": "userId is required"}), 400

    docs = db.collection(COLLECTIONS['messages']) \
        .where('threadId', '==', thread_id) \
        .where('recipientId', '==', user_id) \
        .where('read', '==', False) \
        .stream()

    batch = db.batch()
    count = 0
    now = datetime.utcnow().isoformat()
    for doc in docs:
        batch.update(doc.reference, {'read': True, 'updatedAt': now})
        count += 1
        if count % 500 == 0:
            batch.commit()
            batch = db.batch()
    if count % 500 != 0:
        batch.commit()
    return jsonify({"success": True, "updated": count})


# ============== Golden Script Analysis Engine ==============


@app.route("/api/golden-scripts/upload", methods=["POST"])
def golden_script_upload():
    """Upload a .docx golden script file, store in GCS, and create a record."""
    if 'file' not in request.files:
        return jsonify({"error": "No file provided"}), 400

    file = request.files['file']
    if file.filename == '':
        return jsonify({"error": "No file selected"}), 400

    project_id = request.form.get('projectId', '')
    series_id = request.form.get('seriesId', '')
    series_name = request.form.get('seriesName', '')
    producer_name = request.form.get('producerName', '')

    if not project_id:
        return jsonify({"error": "projectId is required"}), 400

    try:
        file_content = file.read()
        content_type = file.content_type or 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        ext = os.path.splitext(file.filename)[1] or '.docx'
        unique_id = uuid.uuid4().hex
        blob_path = f"golden-scripts/{project_id}/{unique_id}{ext}"

        ensure_bucket_exists(STORAGE_BUCKET)
        bucket = storage_client.bucket(STORAGE_BUCKET)
        blob = bucket.blob(blob_path)
        blob.upload_from_string(file_content, content_type=content_type)

        gcs_uri = f"gs://{STORAGE_BUCKET}/{blob_path}"

        # Extract text from .docx for analysis
        script_text = ""
        if ext.lower() == '.docx':
            try:
                import io
                from docx import Document as DocxDocument
                doc = DocxDocument(io.BytesIO(file_content))
                paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
                # Also check tables (Mark Carter uses 2-column tables)
                for table in doc.tables:
                    for row in table.rows:
                        cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                        if cells:
                            paragraphs.append(' | '.join(cells))
                script_text = '\n'.join(paragraphs)
            except Exception as parse_err:
                script_text = f"[docx parse error: {parse_err}]"

        word_count = len(script_text.split()) if script_text else 0

        golden_doc = create_doc('golden_scripts', {
            'projectId': project_id,
            'seriesId': series_id,
            'seriesName': series_name,
            'producerName': producer_name,
            'filename': file.filename,
            'gcsUri': gcs_uri,
            'blobPath': blob_path,
            'mimeType': content_type,
            'wordCount': word_count,
            'scriptText': script_text[:100000],  # Cap at 100k chars
            'analysisStatus': 'pending',
            'analysisResult': None,
        })

        return jsonify(golden_doc), 201
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/golden-scripts/<script_id>/analyze", methods=["POST"])
def golden_script_analyze(script_id):
    """Run Gemini analysis on a golden script to extract structural DNA."""
    golden = get_doc('golden_scripts', script_id)
    if not golden:
        return jsonify({"error": "Golden script not found"}), 404

    script_text = golden.get('scriptText', '')
    if len(script_text) < 100:
        return jsonify({"error": "Script text too short for analysis"}), 400

    series_name = golden.get('seriesName', 'Unknown Series')
    producer_name = golden.get('producerName', 'Unknown Producer')
    filename = golden.get('filename', '')

    # Update status to analyzing
    update_doc('golden_scripts', script_id, {'analysisStatus': 'analyzing'})

    # Use Gemini Pro for deep analysis
    analysis_model = GenerativeModel(AGENT_MODELS.get('script_writer', 'gemini-2.5-pro'))

    prompt = f"""You are an expert script analyst for documentary television. You are analyzing a "Golden Script" — the reference episode that defines the style, structure, and format for an entire series.

SERIES: {series_name}
PRODUCER: {producer_name}
FILENAME: {filename}

SCRIPT TEXT:
{script_text[:80000]}

Analyze this script deeply and return a JSON object with EXACTLY this structure:

{{
  "format_analysis": {{
    "document_format": "How the script is physically formatted (e.g., 2-column table, flowing paragraphs, numbered scenes)",
    "column_layout": "Description of columns if table format, or paragraph structure if flowing",
    "scene_marking_style": "How scenes are marked/numbered (e.g., SCENE 1, bold headers, numbered paragraphs)",
    "timing_notation": "How timing/duration is indicated if at all"
  }},
  "structure_analysis": {{
    "total_parts": 1,
    "parts": [
      {{
        "part_number": 1,
        "title": "Part title if any",
        "estimated_duration_minutes": 22,
        "scene_count": 10,
        "scenes": [
          {{
            "scene_number": 1,
            "title": "Scene title",
            "type": "opening|narrative|interview|archive_montage|transition|closing",
            "estimated_duration_minutes": 2.5,
            "beat_types": ["voice_over", "archive", "expert", "ai_visual"],
            "summary": "Brief description of what happens"
          }}
        ]
      }}
    ],
    "total_scenes": 20,
    "total_estimated_duration_minutes": 44,
    "story_threads": ["A story description", "B story description"],
    "has_ad_breaks": false,
    "ad_break_positions": []
  }},
  "voice_over_analysis": {{
    "narrator_style": "How the narrator/presenter sounds (tone, register, vocabulary)",
    "narrator_name": "Name of narrator/presenter if identified",
    "vo_density": "How much of the episode is voice-over vs other elements (percentage estimate)",
    "typical_vo_block_length": "Average length of a VO block in words or sentences",
    "capitalization_style": "Whether VO is written in CAPS, Title Case, or lowercase",
    "examples": ["Example VO line 1", "Example VO line 2", "Example VO line 3"]
  }},
  "expert_analysis": {{
    "expert_format": "How experts are introduced and formatted in the script (e.g., SYNC:, bold name, etc.)",
    "typical_soundbite_length": "Average expert quote length in words",
    "expert_frequency": "How often experts appear (every X scenes, etc.)",
    "expert_types": ["Type 1", "Type 2"],
    "examples": ["Example expert line 1", "Example expert line 2"]
  }},
  "archive_analysis": {{
    "archive_density": "How much of the episode relies on archive footage (percentage estimate)",
    "archive_marking_style": "How archive references are marked in the script",
    "archive_types": ["Type 1", "Type 2"],
    "visual_direction_style": "How visual directions are written",
    "examples": ["Example visual direction 1", "Example visual direction 2"]
  }},
  "pacing_analysis": {{
    "opening_strategy": "How the episode opens (cold open, teaser, title sequence, etc.)",
    "tension_curve": "How tension builds throughout the episode",
    "scene_transition_style": "How the script transitions between scenes",
    "closing_strategy": "How the episode ends",
    "pacing_notes": "General pacing observations"
  }},
  "distinctive_features": [
    "Feature 1 that makes this script style unique",
    "Feature 2",
    "Feature 3"
  ],
  "template_instructions": {{
    "for_new_episodes": "Clear instructions a producer could follow to write a new episode in this exact style",
    "must_include": ["Element that every episode must include"],
    "must_avoid": ["Element that should never appear"],
    "tone_keywords": ["keyword1", "keyword2", "keyword3"]
  }}
}}

Be thorough and specific. Extract real examples from the script text. Count actual scenes, don't estimate."""

    try:
        resp = analysis_model.generate_content(
            prompt,
            generation_config={"response_mime_type": "application/json"},
        )
        analysis = json.loads(resp.text)

        update_doc('golden_scripts', script_id, {
            'analysisStatus': 'complete',
            'analysisResult': analysis,
        })

        return jsonify({
            "id": script_id,
            "status": "complete",
            "analysis": analysis,
        })
    except Exception as e:
        update_doc('golden_scripts', script_id, {
            'analysisStatus': 'error',
            'analysisError': str(e),
        })
        return jsonify({"error": str(e)}), 500


@app.route("/api/golden-scripts/by-project/<project_id>", methods=["GET"])
def golden_scripts_by_project(project_id):
    """Get all golden scripts for a project."""
    docs = db.collection(COLLECTIONS['golden_scripts']) \
        .where('projectId', '==', project_id) \
        .order_by('createdAt', direction=firestore.Query.DESCENDING) \
        .stream()
    results = [doc_to_dict(d) for d in docs]
    return jsonify(results)


@app.route("/api/golden-scripts/by-series/<series_id>", methods=["GET"])
def golden_scripts_by_series(series_id):
    """Get all golden scripts for a series."""
    docs = db.collection(COLLECTIONS['golden_scripts']) \
        .where('seriesId', '==', series_id) \
        .order_by('createdAt', direction=firestore.Query.DESCENDING) \
        .stream()
    results = [doc_to_dict(d) for d in docs]
    return jsonify(results)


@app.route("/api/golden-scripts/<script_id>", methods=["GET"])
def golden_script_get(script_id):
    """Get a single golden script by ID."""
    doc = get_doc('golden_scripts', script_id)
    if not doc:
        return jsonify({"error": "Not found"}), 404
    return jsonify(doc)


@app.route("/api/golden-scripts/<script_id>", methods=["DELETE"])
def golden_script_delete(script_id):
    """Delete a golden script."""
    doc = get_doc('golden_scripts', script_id)
    if not doc:
        return jsonify({"error": "Not found"}), 404

    # Delete from GCS if exists
    blob_path = doc.get('blobPath')
    if blob_path:
        try:
            bucket = storage_client.bucket(STORAGE_BUCKET)
            bucket.blob(blob_path).delete()
        except Exception:
            pass

    delete_doc('golden_scripts', script_id)
    return jsonify({"success": True})


@app.route("/api/series-config/<series_id>/script-profile", methods=["PUT"])
def save_script_profile(series_id):
    """Save or update the series script profile (derived from golden script analysis)."""
    data = request.get_json()
    config_ref = db.collection(COLLECTIONS['series_config']).document(series_id)
    config_doc = config_ref.get()

    if config_doc.exists:
        config_ref.update({
            'scriptProfile': data,
            'scriptProfileUpdatedAt': datetime.utcnow().isoformat(),
            'updatedAt': datetime.utcnow().isoformat(),
        })
    else:
        config_ref.set({
            'scriptProfile': data,
            'scriptProfileUpdatedAt': datetime.utcnow().isoformat(),
            'createdAt': datetime.utcnow().isoformat(),
            'updatedAt': datetime.utcnow().isoformat(),
        })

    return jsonify({"success": True, "seriesId": series_id})


@app.route("/api/series-config/<series_id>/script-profile", methods=["GET"])
def get_script_profile(series_id):
    """Get the series script profile."""
    config_ref = db.collection(COLLECTIONS['series_config']).document(series_id)
    config_doc = config_ref.get()
    if not config_doc.exists:
        return jsonify({"scriptProfile": None})
    data = config_doc.to_dict()
    return jsonify({
        "scriptProfile": data.get('scriptProfile'),
        "updatedAt": data.get('scriptProfileUpdatedAt'),
    })


# ============== Beat Sheet Routes (Episode Development) ==============

@app.route("/api/beat-sheets", methods=["POST"])
def save_beat_sheet():
    """Save or update a beat sheet (upsert by id)."""
    try:
        data = request.get_json() or {}
        sheet_id = data.get('id')
        if sheet_id:
            existing = get_doc('beat_sheets', sheet_id)
            if existing:
                return jsonify(update_doc('beat_sheets', sheet_id, data))
        doc = create_doc('beat_sheets', data)
        return jsonify(doc), 201
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/beat-sheets/by-project/<project_id>", methods=["GET"])
def beat_sheets_by_project(project_id):
    """Get all beat sheets for a project."""
    docs = db.collection(COLLECTIONS['beat_sheets']) \
        .where('projectId', '==', project_id) \
        .order_by('createdAt', direction=firestore.Query.DESCENDING) \
        .stream()
    results = [doc_to_dict(d) for d in docs]
    return jsonify(results)


@app.route("/api/beat-sheets/<sheet_id>", methods=["GET"])
def beat_sheet_get(sheet_id):
    """Get a single beat sheet by ID."""
    doc = get_doc('beat_sheets', sheet_id)
    if not doc:
        return jsonify({"error": "Not found"}), 404
    return jsonify(doc)


@app.route("/api/ai/generate-beat-sheet", methods=["POST"])
def generate_beat_sheet():
    """Generate or refine a beat sheet using Gemini 2.5 Pro with Google Search grounding."""
    try:
        data = request.get_json() or {}
        topic = data.get("topic", "")
        theme = data.get("theme", "")
        subjects = data.get("subjects", [])
        project_id = data.get("projectId", "")
        series_id = data.get("seriesId", "")
        episode_number = data.get("episodeNumber")
        golden_profile = data.get("goldenScriptProfile")
        existing_beats = data.get("existingBeats", [])
        user_message = data.get("userMessage", "")
        target_duration = data.get("targetDuration", 2640)  # 44 min default

        if not topic:
            return jsonify({"error": "topic is required"}), 400

        # Build golden script context
        golden_context = ""
        if golden_profile:
            golden_context = f"""
== GOLDEN SCRIPT STANDARD ==
Format: {golden_profile.get('format_analysis', {}).get('document_format', 'Unknown')}
Column Layout: {golden_profile.get('format_analysis', {}).get('column_layout', 'Unknown')}
Structure: {json.dumps(golden_profile.get('structure_analysis', {}), indent=2)[:2000] if golden_profile.get('structure_analysis') else 'Not available'}
Voice Over Style: {json.dumps(golden_profile.get('voice_over_analysis', {}), indent=2)[:1000] if golden_profile.get('voice_over_analysis') else 'Not available'}
Expert Format: {json.dumps(golden_profile.get('expert_analysis', {}), indent=2)[:1000] if golden_profile.get('expert_analysis') else 'Not available'}
Archive Usage: {json.dumps(golden_profile.get('archive_analysis', {}), indent=2)[:1000] if golden_profile.get('archive_analysis') else 'Not available'}
Pacing: {json.dumps(golden_profile.get('pacing_analysis', {}), indent=2)[:800] if golden_profile.get('pacing_analysis') else 'Not available'}

This is the benchmark. Every beat sheet you create must match this standard.
"""

        # Build series context
        series_context = ""
        if series_id:
            try:
                config_doc = db.collection(COLLECTIONS['series_config']).document(series_id).get()
                if config_doc.exists:
                    config_data = config_doc.to_dict()
                    series_name = config_data.get('name', series_id)
                    series_context = f"""
== SERIES BIBLE ==
Series: {series_name}
Workflow: {json.dumps(config_data.get('workflow', {}), indent=2)[:500]}
Script Style: {json.dumps(config_data.get('scriptStyle', {}), indent=2)[:500]}
"""
            except Exception:
                pass

        # Build existing beats context for refinement
        existing_context = ""
        if existing_beats:
            beats_text = "\n".join([
                f"Beat {b.get('order', i+1)}: [{b.get('type', 'narrative')}] {b.get('title', 'Untitled')} — {b.get('description', '')} ({b.get('duration_estimate_seconds', 0)}s)"
                for i, b in enumerate(existing_beats)
            ])
            existing_context = f"""
== CURRENT BEAT SHEET (refine this) ==
{beats_text}

The producer wants you to refine the above beat sheet based on their feedback below.
"""

        ep_label = f"Episode {episode_number}" if episode_number else "this episode"
        subjects_text = ", ".join(subjects) if subjects else "not specified"
        target_minutes = target_duration // 60

        system_prompt = f"""You are a senior documentary showrunner with 20 years experience at Netflix, BBC, and Channel 4.
You are building a beat sheet for {ep_label}: "{topic}".
{golden_context}
{series_context}

== YOUR TASK ==
Build a production-ready beat sheet. Not a summary — a structural blueprint that a producer can hand to their team.

For each beat:
- SPECIFIC visual direction (not "show archive" — specify WHAT archive: "NASA mission control footage from Apollo 13, the moment Jim Lovell says 'Houston, we have a problem'")
- SPECIFIC narrative purpose (not "introduce the topic" — "establish the stakes: 3 astronauts, 200,000 miles from Earth, and the oxygen tank just exploded")
- NAMED experts where possible (not "an engineer" — "Dr. Emily Lakdawalla, planetary scientist, explains why...")
- PRECISE timing (not "a few minutes" — "4:30 — this is the audience's first emotional anchor")
- ARCHIVE FEASIBILITY notes (is this footage likely to exist? public domain? Getty? needs licensing?)

Think like a showrunner, not a chatbot. Every beat must earn its place.

After generating, assess:
- Is there enough here for {target_minutes} minutes?
- Where are the research gaps? What does the team need to investigate further?
- What's the weakest beat and how could it be strengthened?
- What questions should the producer be asking that they haven't thought of?

IMPORTANT: Return ONLY valid JSON in this exact format:
{{
  "beats": [
    {{
      "id": "beat-1",
      "reference": "EP{episode_number or '01'}-BEAT-01",
      "title": "string",
      "description": "detailed description with specific visual direction",
      "type": "cold_open|narrative|expert|archive_montage|transition|climax|resolution",
      "duration_estimate_seconds": 180,
      "research_notes": "what needs researching for this beat",
      "archive_reference": "specific archive sources or feasibility notes",
      "order": 1
    }}
  ],
  "commentary": "Your showrunner analysis of the beat sheet — what works, what needs attention, structural notes",
  "total_duration_estimate": 2640,
  "viability_note": "Production readiness assessment — can this episode actually be made?",
  "suggested_followups": [
    "Context-aware follow-up question 1",
    "Context-aware follow-up question 2",
    "Context-aware follow-up question 3",
    "Context-aware follow-up question 4"
  ],
  "research_gaps": [
    "Gap 1: what needs investigating",
    "Gap 2: what needs investigating"
  ]
}}"""

        prompt_parts = []
        if existing_context:
            prompt_parts.append(existing_context)

        prompt_parts.append(f"""
EPISODE TOPIC: {topic}
THEME: {theme or 'Not specified'}
KEY SUBJECTS: {subjects_text}
TARGET DURATION: {target_minutes} minutes ({target_duration} seconds)
""")

        if user_message:
            prompt_parts.append(f"""
PRODUCER'S REQUEST:
{user_message}
""")
        elif not existing_beats:
            prompt_parts.append("""
Generate the initial beat sheet now. Create 8-12 beats that cover the full episode arc from cold open through resolution.
""")

        full_prompt = "\n".join(prompt_parts)

        # Use Gemini 3.1 Pro with Google Search grounding
        beat_model = GenerativeModel("gemini-3.1-pro-preview")
        search_tool = Tool._from_gapic(
            raw_tool=GapicTool(google_search=GapicTool.GoogleSearch())
        )
        response = beat_model.generate_content(
            f"{system_prompt}\n\n{full_prompt}",
            tools=[search_tool],
        )
        response_text = response.text

        # Clean and parse JSON response
        cleaned = response_text.strip()
        if cleaned.startswith("```json"):
            cleaned = cleaned[7:]
        if cleaned.startswith("```"):
            cleaned = cleaned[3:]
        if cleaned.endswith("```"):
            cleaned = cleaned[:-3]
        cleaned = cleaned.strip()

        try:
            result = json.loads(cleaned)
        except json.JSONDecodeError:
            # Try to extract JSON from the response
            json_match = re.search(r'\{[\s\S]*\}', cleaned)
            if json_match:
                result = json.loads(json_match.group())
            else:
                return jsonify({
                    "error": "Failed to parse AI response as JSON",
                    "raw_response": response_text[:2000]
                }), 500

        # Ensure required fields exist
        if 'beats' not in result:
            result['beats'] = []
        if 'commentary' not in result:
            result['commentary'] = ''
        if 'total_duration_estimate' not in result:
            result['total_duration_estimate'] = sum(b.get('duration_estimate_seconds', 0) for b in result['beats'])
        if 'viability_note' not in result:
            result['viability_note'] = ''
        if 'suggested_followups' not in result:
            result['suggested_followups'] = []
        if 'research_gaps' not in result:
            result['research_gaps'] = []

        return jsonify(result), 200

    except Exception as e:
        return jsonify({"error": str(e)}), 500


# ============== Auto-seed users on startup ==============
try:
    if not list(db.collection(COLLECTIONS['users']).limit(1).stream()):
        for _ud in SEED_USERS:
            create_doc('users', dict(_ud))
        print(f"[startup] Seeded {len(SEED_USERS)} default users into {COLLECTIONS['users']}")
    else:
        print(f"[startup] Users already exist in {COLLECTIONS['users']}")
except Exception as _e:
    print(f"[startup] Could not auto-seed users: {_e}")


if __name__ == "__main__":
    port = int(os.environ.get("PORT", 5000))
    app.run(host="0.0.0.0", port=port, debug=True)
